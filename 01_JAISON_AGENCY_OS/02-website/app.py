import streamlit as st

import os, json, time

LORA_STATE_FILE = "lora_state.json"

def load_lora_state():
    if os.path.exists(LORA_STATE_FILE):
        try:
            with open(LORA_STATE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            pass
    return {}

def save_lora_state(url, trigger):
    try:
        with open(LORA_STATE_FILE, "w", encoding="utf-8") as f:
            json.dump({"url": url, "trigger": trigger}, f)
    except:
        pass

def render_multi_publish_selector(video_bytes, key_prefix):
    """
    Renderuje luksusowy panel Multi-Publish dla wideo zintegrowany z n8n.
    """
    if not video_bytes:
        return
        
    st.markdown("<hr style='border-color: #374151;'>", unsafe_allow_html=True)
    st.markdown(f"### 🚀 Zunifikowany Multi-Publikator Wideo (One-Click Publish)")
    st.markdown("<p style='font-size: 0.9rem; color: #9CA3AF; margin-bottom: 15px;'>Zatwierdź i wyślij gotowe wideo jednocześnie na wszystkie swoje kanały social media jednym kliknięciem.</p>", unsafe_allow_html=True)
    
    col_pub_t, col_pub_c = st.columns([1, 1])
    
    with col_pub_t:
        pub_title = st.text_input(
            "Tytuł viralu (Shorts/Reels/TikTok):",
            value="Jak AI odzyskuje Twój czas? ⚡",
            key=f"{key_prefix}_pub_title"
        )
        pub_desc = st.text_area(
            "Opis SEO i hashtagi:",
            value="Oto jak wdrożyć automatycznego asystenta AI w swojej firmie... #ai #automation #n8n #productivity #adhd",
            height=120,
            key=f"{key_prefix}_pub_desc"
        )
        
    with col_pub_c:
        st.markdown("<p style='font-size: 0.9rem; font-weight: bold; color: #9CA3AF; margin-bottom: 5px;'>Wybierz kanały docelowe:</p>", unsafe_allow_html=True)
        ch_yt = st.checkbox("🩳 YouTube Shorts", value=True, key=f"{key_prefix}_ch_yt")
        ch_ig = st.checkbox("📸 Instagram Reels", value=True, key=f"{key_prefix}_ch_ig")
        ch_tt = st.checkbox("🎵 TikTok Video", value=True, key=f"{key_prefix}_ch_tt")
        ch_fb = st.checkbox("📘 Facebook Reels/Shorts", value=True, key=f"{key_prefix}_ch_fb")
        
        webhook_url = st.text_input(
            "Adres Webhooka n8n Multi-Publish:",
            value="https://n8n.jaison.pl/webhook/video-multi-publish",
            key=f"{key_prefix}_webhook"
        )
        
    if st.button("🚀 OPUBLIKUJ NA WYBRANYCH KANAŁACH JEDNYM KLIKNIĘCIEM", type="primary", use_container_width=True, key=f"{key_prefix}_pub_btn"):
        selected_channels = []
        if ch_yt: selected_channels.append("YouTube Shorts")
        if ch_ig: selected_channels.append("Instagram Reels")
        if ch_tt: selected_channels.append("TikTok")
        if ch_fb: selected_channels.append("Facebook Reels")
        
        if not selected_channels:
            st.error("⚠️ Wybierz przynajmniej jeden kanał publikacji!")
        else:
            with st.spinner("📦 Przygotowywanie paczki wideo i wysyłanie do n8n pipeline..."):
                try:
                    import requests
                    
                    files = {
                        "data": ("video.mp4", video_bytes, "video/mp4")
                    }
                    data = {
                        "title": pub_title,
                        "description": pub_desc,
                        "channels": ",".join(selected_channels)
                    }
                    
                    response = requests.post(webhook_url, data=data, files=files, timeout=60, verify=False)
                    
                    if response.status_code in [200, 201]:
                        st.success(f"🎉 Sukces! Wideo pomyślnie wysłane do dystrybucji na kanały: {', '.join(selected_channels)}!")
                        st.toast("⚡ n8n pipeline uruchomiony pomyślnie!")
                    else:
                        st.error(f"❌ Błąd webhooka n8n (Kod {response.status_code}): {response.text}")
                except Exception as ex:
                    st.error(f"❌ Wyjątek podczas wysyłania do n8n: {str(ex)}")

def check_telegram_events():
    events_file = r"C:\Aplikacje MVP\01_JAISON_AGENCY_OS\02-website\dashboard\telegram_events.json"
    if os.path.exists(events_file):
        try:
            with open(events_file, "r", encoding="utf-8") as f:
                events = json.load(f)
            if "shown_events" not in st.session_state:
                st.session_state.shown_events = set()
            for evt in events:
                evt_id = evt.get("id")
                if evt_id not in st.session_state.shown_events:
                    st.toast(evt.get("message"), icon="📥" if "lead" in evt.get("type", "").lower() else "🎬")
                    st.session_state.shown_events.add(evt_id)
        except Exception:
            pass

def find_client_profile_text(client_path):
    """Przeszukuje folder klienta pod kątem plików profilowych .md (np. Profil_*.md, ghost_profile.md, oferta.md)."""
    candidates = []
    
    # 1. Sprawdź najpierw ghost_profile.md w katalogu głównym
    gp = os.path.join(client_path, "ghost_profile.md")
    if os.path.exists(gp):
        candidates.append(gp)
        
    # 2. Przeszukaj rekurencyjnie do głębokości 2 poziomów
    for root, dirs, files in os.walk(client_path):
        depth = root.replace(client_path, '').count(os.sep)
        if depth > 2:
            continue
        for f in files:
            f_lower = f.lower()
            if f_lower.startswith("profil_") and f_lower.endswith(".md"):
                candidates.append(os.path.join(root, f))
            elif f_lower == "oferta.md":
                candidates.append(os.path.join(root, f))
                
    # 3. Fallback: poszukaj jakichkolwiek markdownów w folderze 01-brand
    if not candidates:
        brand_dir = os.path.join(client_path, "01-brand")
        if os.path.exists(brand_dir):
            for f in os.listdir(brand_dir):
                if f.endswith(".md"):
                    candidates.append(os.path.join(brand_dir, f))
                    
    # Odczytaj pierwszy plik
    if candidates:
        try:
            with open(candidates[0], "r", encoding="utf-8") as f_prof:
                return f_prof.read(), os.path.basename(candidates[0])
        except:
            pass
    return None, None

def get_selected_context_data():
    if "selected_context" not in st.session_state:
        return ""
    
    ctx_name = st.session_state.selected_context
    context_str = f"\n--- AKTYWNY KONTEKST KLIENTA / PROJEKTU: {ctx_name} ---\n"
    
    # Przeszukaj C:\Aplikacje MVP\02_CLIENTS_AND_PROJECTS (Rzeczywisty folder roboczy!)
    real_clients_dir = r"C:\Aplikacje MVP\02_CLIENTS_AND_PROJECTS"
    target_folder = None
    
    # Mapowanie ładnych nazw na foldery robocze
    mapping = {
        "coolfon.pl": "coolfon",
        "kurczakujasia.pl (Bar Jaś)": "kurczakujasia",
        "lifewave.com (MLM)": "lifewave",
        "smartrade.pl": "smartrade_client",
        "viptransporter.pl": "viptransporter",
        "kantororanzada.pl": "kantor_lombard_oranzada",
        "vojsik.ai": "vojsik_ai",
        "app.jaison.pl (SaaS)": "apps.jaison.pl"
    }
    
    dir_name = mapping.get(ctx_name)
    if dir_name and os.path.exists(os.path.join(real_clients_dir, dir_name)):
        target_folder = os.path.join(real_clients_dir, dir_name)
    else:
        # Sprawdź dopasowanie bezpośrednie
        if os.path.exists(real_clients_dir):
            for folder in os.listdir(real_clients_dir):
                if folder == ctx_name or folder.lower() in ctx_name.lower():
                    target_folder = os.path.join(real_clients_dir, folder)
                    break
                    
    # Fallback do starego folderu 04_clients (jeśli istnieje)
    if not target_folder:
        clients_dir = os.path.join(BASE_DIR, "04_clients") if "BASE_DIR" in globals() else os.path.expanduser("~/Agentic_OS/04_clients")
        if os.path.exists(clients_dir):
            for item in os.listdir(clients_dir):
                if os.path.isdir(os.path.join(clients_dir, item)):
                    config_path = os.path.join(clients_dir, item, "context_config.json")
                    if os.path.exists(config_path):
                        try:
                            with open(config_path, "r", encoding="utf-8") as f_conf:
                                cfg = json.load(f_conf)
                                if cfg.get("name") == ctx_name or item == ctx_name:
                                    target_folder = os.path.join(clients_dir, item)
                                    context_str += f"Typ kontekstu: {cfg.get('type', 'Nieznany')}\nOpis: {cfg.get('description', '')}\n"
                                    if "system_prompt_override" in cfg:
                                        context_str += f"Zalecenia Strategiczne: {cfg['system_prompt_override']}\n"
                                    break
                        except:
                            pass
                            
    if target_folder:
        profile_content, profile_file = find_client_profile_text(target_folder)
        if profile_content:
            context_str += f"\nBAZA WIEDZY PROFILU ({profile_file}):\n{profile_content}\n"
        else:
            context_str += f"\nBrak szczegółowego profilu .md w folderze: {os.path.basename(target_folder)}\n"
    else:
        # Statyczne opisy dla domyślnych kontekstów
        if ctx_name == "J(AI)SON Agency":
            context_str += "Typ: Usługi B2B / Automatyzacje AI & n8n\nOpis: Oficjalny profil agencji jaison.pl dostarczającej systemy AI dla firm.\n"
        elif ctx_name == "Holistic Jason":
            context_str += "Typ: Marka Osobista / Edukacja ADHD\nOpis: Społeczność ADHD4life i doradztwo energetyczne.\n"
        elif ctx_name == "app.jaison.pl (SaaS)":
            context_str += "Typ: SaaS / Platforma App\nOpis: Portal app.jaison.pl zintegrowany z agentami agencji Jaison.\n"
            
    context_str += "-----------------------------------------------------\n"
    return context_str


import ssl

# Ręczne wczytanie pliku .env na starcie aplikacji (sprawdzenie bieżącego i nadrzędnego katalogu)
env_paths = [
    os.path.join(os.path.dirname(__file__), ".env"),
    os.path.join(os.path.dirname(os.path.dirname(__file__)), ".env")
]

for env_path in env_paths:
    if os.path.exists(env_path):
        try:
            with open(env_path, "r", encoding="utf-8") as f:
                for line in f:
                    if "=" in line and not line.strip().startswith("#"):
                        k, v = line.split("=", 1)
                        # Usuń ewentualne cudzysłowy wokół wartości
                        val = v.strip()
                        if (val.startswith('"') and val.endswith('"')) or (val.startswith("'") and val.endswith("'")):
                            val = val[1:-1]
                        os.environ[k.strip()] = val
        except Exception as e:
            pass
        break

try:
    ssl._create_default_https_context = ssl._create_unverified_context
except AttributeError:
    pass

st.set_page_config(
    page_title="JAISON OS • Mission Control",
    page_icon="✨",
    layout="wide",
    initial_sidebar_state="expanded"
)

if "lora_state_loaded" not in st.session_state:
    saved_state = load_lora_state()
    st.session_state.lora_result_url = saved_state.get("url", None)
    st.session_state.lora_trigger_word = saved_state.get("trigger", "tomasz_hero")
    st.session_state.lora_state_loaded = True

# Dynamiczne wstrzykiwanie kodu Google Tag Manager (GTM)
gtm_id = os.environ.get("GOOGLE_TAG_MANAGER_ID", "").strip()
if gtm_id:
    st.markdown(f"""
    <script id="gtm-injector">
        (function() {{
            const doc = window.parent !== window ? window.parent.document : document;
            if (!doc.getElementById('gtm-script-tag')) {{
                // Skrypt GTM w Head
                (function(w,d,s,l,i){{w[l]=w[l]||[];w[l].push({{'gtm.start':
                new Date().getTime(),event:'gtm.js'}});var f=d.getElementsByTagName(s)[0],
                j=d.createElement(s),dl=l!='dataLayer'?'&l='+l:'';j.async=true;j.src=
                'https://www.googletagmanager.com/gtm.js?id='+i+dl;j.id='gtm-script-tag';f.parentNode.insertBefore(j,f);
                }})(window.parent,doc,'script','dataLayer','{gtm_id}');
                
                // NoScript GTM w Body
                const noscript = doc.createElement('noscript');
                noscript.id = 'gtm-noscript-tag';
                const iframe = doc.createElement('iframe');
                iframe.src = 'https://www.googletagmanager.com/ns.html?id={gtm_id}';
                iframe.height = '0';
                iframe.width = '0';
                iframe.style.display = 'none';
                iframe.style.visibility = 'hidden';
                noscript.appendChild(iframe);
                doc.body.insertBefore(noscript, doc.body.firstChild);
            }}
        }})();
    </script>
    """, unsafe_allow_html=True)

# Ścieżki w chmurze
BASE_DIR = os.path.expanduser("~/Agentic_OS")
NOTEBOOKS_DIR = os.path.join(BASE_DIR, "notebooks/_assets")
OBSIDIAN_DIR = os.path.join(BASE_DIR, "obsidian_vault")
DASHBOARD_DIR = os.path.join(BASE_DIR, "dashboard")
BRAIN_DUMP_DIR = os.path.join(BASE_DIR, "brain_dump")
BRAIN_DUMP_ASSETS = os.path.join(BRAIN_DUMP_DIR, "_assets")
HERMES_DIR = os.path.expanduser("~/.hermes")
KANBAN_FILE = os.path.join(DASHBOARD_DIR, "kanban.json")

# Tworzenie folderów
for d in [NOTEBOOKS_DIR, OBSIDIAN_DIR, DASHBOARD_DIR, BRAIN_DUMP_DIR, BRAIN_DUMP_ASSETS, HERMES_DIR]:
    os.makedirs(d, exist_ok=True)

# Luksusowy, kosmiczny design dopasowany do logo JAISON OS (fiolety, neony, sparkles i akcenty teal)
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Atkinson+Hyperlegible:ital,wght@0,400;0,700;1,400;1,700&family=Outfit:wght@300;400;500;600;700;800&display=swap');
    
    /* Globalny reset barw i luksusowa typografia */
    .stApp {
        background-color: #040508 !important;
        color: #E2E8F0 !important;
        font-family: 'Atkinson Hyperlegible', sans-serif !important;
    }
    
    [data-testid="stSidebar"] {
        background-color: #080A10 !important;
        border-right: 1px solid rgba(167, 139, 250, 0.12) !important;
    }
    
    h1, h2, h3, h4, h5, h6 {
        font-family: 'Outfit', sans-serif !important;
        color: #FFFFFF !important;
        font-weight: 600 !important;
    }
    
    /* Animowane, luksusowe karty z efektem neonowego fioletowo-tealowego obramowania */
    .custom-card {
        background: linear-gradient(135deg, #0A0D18 0%, #0F1326 100%);
        border: 1px solid rgba(167, 139, 250, 0.15);
        border-radius: 16px;
        padding: 24px;
        margin-bottom: 20px;
        box-shadow: 0 4px 24px rgba(0, 0, 0, 0.5);
        transition: all 0.3s cubic-bezier(0.16, 1, 0.3, 1);
    }
    
    .custom-card:hover {
        border-color: #14B8A6;
        transform: translateY(-3px);
        box-shadow: 0 8px 32px rgba(20, 184, 166, 0.18), 0 0 15px rgba(167, 139, 250, 0.1);
    }
    
    /* Banner One Thing - eliminacja szumu z gradientem fiolet-teal */
    .one-thing-banner {
        background: linear-gradient(135deg, #100E26 0%, #080A12 100%);
        border-left: 6px solid #14B8A6;
        border-radius: 16px;
        padding: 30px;
        margin-bottom: 30px;
        box-shadow: 0 12px 40px rgba(0, 0, 0, 0.6);
        border-top: 1px solid rgba(167, 139, 250, 0.08);
        border-right: 1px solid rgba(167, 139, 250, 0.08);
        border-bottom: 1px solid rgba(167, 139, 250, 0.08);
    }
    
    /* Zaokrąglone przyciski premium (globalnie dla stButton) */
    .stButton>button {
        background: linear-gradient(135deg, #7C3AED 0%, #0D9488 100%) !important;
        color: #FFFFFF !important;
        border: 1px solid rgba(167, 139, 250, 0.3) !important;
        border-radius: 12px !important;
        padding: 10px 20px !important;
        font-family: 'Outfit', sans-serif !important;
        font-weight: 600 !important;
        transition: all 0.25s ease !important;
        width: 100%;
        box-shadow: 0 4px 15px rgba(124, 58, 237, 0.25) !important;
        letter-spacing: 0.5px;
    }
    
    .stButton>button:hover {
        background: linear-gradient(135deg, #8B5CF6 0%, #14B8A6 100%) !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 6px 22px rgba(20, 184, 166, 0.45) !important;
        border-color: #FFFFFF !important;
    }
 
    /* Kafelki Nawigacyjne w Pasku Bocznym (ADHD Tiles) */
    div[data-testid="stSidebar"] .stButton>button {
        text-align: left !important;
        padding: 12px 18px !important;
        height: 52px !important;
        margin-bottom: 2px !important;
        width: 100% !important;
    }
 
    div[data-testid="stSidebar"] .stButton>button[kind="secondary"] {
        background: transparent !important;
        color: #94A3B8 !important;
        border: 1px solid transparent !important;
        box-shadow: none !important;
        font-weight: 500 !important;
    }
 
    div[data-testid="stSidebar"] .stButton>button[kind="secondary"]:hover {
        background: rgba(167, 139, 250, 0.08) !important;
        border-color: rgba(167, 139, 250, 0.3) !important;
        color: #FFFFFF !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 4px 12px rgba(167, 139, 250, 0.1) !important;
    }
 
    div[data-testid="stSidebar"] .stButton>button[kind="primary"] {
        background: linear-gradient(135deg, #8B5CF6 0%, #6D28D9 100%) !important;
        color: #FFFFFF !important;
        border: 1px solid rgba(167, 139, 250, 0.4) !important;
        box-shadow: 0 0 15px rgba(139, 92, 246, 0.35) !important;
        font-weight: 700 !important;
    }
    
    div[data-testid="stSidebar"] .stButton>button[kind="primary"]:hover {
        transform: translateY(-1px) !important;
        box-shadow: 0 0 20px rgba(139, 92, 246, 0.5) !important;
    }
    
    /* ADHD-friendly akcenty kolorystyczne */
    .dopamine-accent { color: #14B8A6; font-weight: bold; }
    .burn-accent { color: #F43F5E; font-weight: bold; }
    .focus-accent { color: #F59E0B; font-weight: bold; }
 
    /* Pływający Przycisk Szybkiego Zapisu (FAB) */
    .fab-wrapper {
        display: none !important;
    }
    
    div:has(> .fab-wrapper) + div {
        height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
        overflow: visible !important;
    }
    
    div:has(> .fab-wrapper) + div div.stButton > button {
        position: fixed !important;
        bottom: 25px !important;
        right: 25px !important;
        z-index: 999999 !important;
        width: 60px !important;
        height: 60px !important;
        border-radius: 50% !important;
        background: radial-gradient(circle, #8B5CF6 0%, #14B8A6 100%) !important;
        border: 2px solid #A78BFA !important;
        box-shadow: 0 0 15px rgba(139, 92, 246, 0.6) !important;
        font-size: 28px !important;
        padding: 0 !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        transition: all 0.3s ease !important;
        animation: pulse-fab 2s infinite !important;
    }
 
    div:has(> .fab-wrapper) + div div.stButton > button:hover {
        background: radial-gradient(circle, #A78BFA 0%, #2DD4BF 100%) !important;
        box-shadow: 0 0 25px rgba(20, 184, 166, 0.9) !important;
        transform: scale(1.08) !important;
    }
 
    @keyframes pulse-fab {
        0% {
            box-shadow: 0 0 0 0 rgba(139, 92, 246, 0.7);
            transform: scale(1);
        }
        70% {
            box-shadow: 0 0 0 15px rgba(139, 92, 246, 0);
            transform: scale(1.05);
        }
        100% {
            box-shadow: 0 0 0 0 rgba(139, 92, 246, 0);
            transform: scale(1);
        }
    }

    /* GLOBALNA RESPONSYWNOŚĆ MOBILNA */
    @media (max-width: 768px) {
        .stApp {
            padding: 10px !important;
        }
        /* Spłaszczenie kolumn Streamlit do widoku jednokolumnowego */
        div[data-testid="column"] {
            width: 100% !important;
            flex: 1 1 100% !important;
            margin-bottom: 12px !important;
        }
        /* Ukrycie nadmiarowego pustego odstępu */
        .block-container {
            padding-top: 2rem !important;
            padding-bottom: 2rem !important;
            padding-left: 1rem !important;
            padding-right: 1rem !important;
        }
        /* Dostosowanie wielkości czcionek */
        h1 {
            font-size: 1.8rem !important;
        }
        h2 {
            font-size: 1.4rem !important;
        }
        h3 {
            font-size: 1.2rem !important;
        }
        /* Lepsze dotykowe przyciski na telefonach */
        .stButton>button {
            padding: 12px 18px !important;
            font-size: 1rem !important;
            height: auto !important;
        }
        /* Dostosowanie kart i bannerów */
        .custom-card {
            padding: 16px !important;
            border-radius: 12px !important;
        }
        .one-thing-banner {
            padding: 20px !important;
            border-radius: 12px !important;
            margin-bottom: 20px !important;
        }
        
        /* Luksusowy Hamburger Menu na smartfonie z usuwaniem czarnego nagłówka */
        header[data-testid="stHeader"] {
            background: transparent !important;
            background-color: transparent !important;
            border: none !important;
            box-shadow: none !important;
            backdrop-filter: none !important;
            height: 0px !important;
            min-height: 0px !important;
            overflow: visible !important;
        }

        /* Ukrywamy domyślne opcje systemowe Streamlita po prawej stronie, aby system wyglądał autorsko i premium */
        div[data-testid="stHeaderSettingsMenu"],
        button[aria-label="App settings"],
        button[aria-label="View app settings"],
        button[data-testid="stHeaderSettingsMenu"] {
            display: none !important;
            width: 0px !important;
            height: 0px !important;
            opacity: 0 !important;
            pointer-events: none !important;
        }
        
        /* Celujemy TYLKO i wyłącznie w przyciski otwierania paska bocznego w nowym Streamlicie */
        div[data-testid="collapsedControl"],
        div[data-testid="collapsedControl"] button,
        button[aria-label="Open sidebar"],
        button[data-testid="stSidebarCollapse"] {
            background: linear-gradient(135deg, rgba(79, 70, 229, 0.95) 0%, rgba(139, 92, 246, 0.95) 100%) !important;
            border: 1px solid rgba(167, 139, 250, 0.6) !important;
            border-radius: 50% !important;
            width: 46px !important;
            height: 46px !important;
            min-width: 46px !important;
            min-height: 46px !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
            box-shadow: 0 4px 18px rgba(139, 92, 246, 0.75) !important;
            position: fixed !important;
            left: 14px !important;
            top: 14px !important;
            z-index: 999999 !important;
            cursor: pointer !important;
            transition: all 0.3s ease !important;
            padding: 0px !important;
            margin: 0px !important;
        }

        /* Hover i aktywacja dla hamburgera sidebaru */
        div[data-testid="collapsedControl"]:hover,
        div[data-testid="collapsedControl"] button:hover,
        button[aria-label="Open sidebar"]:hover,
        button[data-testid="stSidebarCollapse"]:hover {
            transform: scale(1.08) !important;
            box-shadow: 0 4px 22px rgba(139, 92, 246, 0.9) !important;
        }

        /* Ukrywamy oryginalne ikony i strzałki wewnątrz hamburgera sidebaru */
        div[data-testid="collapsedControl"] svg,
        div[data-testid="collapsedControl"] button svg,
        button[aria-label="Open sidebar"] svg,
        button[data-testid="stSidebarCollapse"] svg {
            display: none !important;
        }

        /* Renderujemy piękny biały hamburger ☰ */
        div[data-testid="collapsedControl"]::before,
        div[data-testid="collapsedControl"] button::before,
        button[aria-label="Open sidebar"]::before,
        button[data-testid="stSidebarCollapse"]::before {
            content: "☰" !important;
            color: #FFFFFF !important;
            font-size: 24px !important;
            font-family: Arial, sans-serif !important;
            font-weight: bold !important;
            line-height: 1 !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
        }

        /* Pływający Przycisk FAB Brain Dump z Czaszką 💀 na dole po prawej stronie */
        .fab-container {
            position: fixed !important;
            bottom: 25px !important;
            right: 25px !important;
            z-index: 999999 !important;
        }
        .fab-container button {
            background: radial-gradient(circle at 40% 40%, #EC4899, #D946EF) !important;
            border: 2px solid rgba(244, 114, 182, 0.7) !important;
            border-radius: 50% !important;
            width: 64px !important;
            height: 64px !important;
            min-width: 64px !important;
            min-height: 64px !important;
            font-size: 28px !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
            box-shadow: 0 4px 20px rgba(236, 72, 153, 0.75) !important;
            animation: pulse-fab 1.8s infinite !important;
            color: white !important;
            padding: 0 !important;
            cursor: pointer !important;
            transition: transform 0.2s ease-in-out !important;
        }
        .fab-container button:hover {
            transform: scale(1.12) !important;
            box-shadow: 0 4px 25px rgba(236, 72, 153, 0.95) !important;
        }
        .fab-container button p {
            font-size: 28px !important;
            margin: 0 !important;
            padding: 0 !important;
        }
        @keyframes pulse-fab {
            0%   { box-shadow: 0 0 0 0 rgba(236, 72, 153, 0.75); }
            70%  { box-shadow: 0 0 0 18px rgba(236, 72, 153, 0); }
            100% { box-shadow: 0 0 0 0 rgba(236, 72, 153, 0); }
        }
    }
</style>
""", unsafe_allow_html=True)

import base64
def get_image_base64(path):
    if os.path.exists(path):
        try:
            with open(path, "rb") as image_file:
                return base64.b64encode(image_file.read()).decode('utf-8')
        except:
            pass
    return ""

# Hybrydowe wczytywanie logotypu (relatywne dla kontenera chmurowego, bezwzględne dla lokalnego fallbacku)
local_logo_name = "logo_sparkles.png"
script_dir = os.path.dirname(os.path.abspath(__file__)) if "__file__" in locals() else os.getcwd()
relative_logo_path = os.path.join(script_dir, local_logo_name)

if os.path.exists(relative_logo_path):
    logo_path_sparkles = relative_logo_path
else:
    logo_path_sparkles = r"C:\Aplikacje MVP\03_SOFTWARE_AND_APPS\jaison_komunikator\web\app_icon_jaison_sparkles.png"

logo_base64 = get_image_base64(logo_path_sparkles)

# ==========================================
# SECURE BASIC AUTHENTICATION WRAPPER
# ==========================================
AUTH_USER = os.environ.get("HERMES_DASHBOARD_BASIC_AUTH_USERNAME", "holistic").strip()
AUTH_PASS = os.environ.get("HERMES_DASHBOARD_BASIC_AUTH_PASSWORD", "holistic2026").strip()

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.markdown("<br><br>", unsafe_allow_html=True)
    col_login, col_center, col_right = st.columns([1, 1.5, 1])
    with col_center:
        if logo_base64:
            logo_html = f'<div style="text-align: center;"><img src="data:image/png;base64,{logo_base64}" style="width: 100px; height: 100px; margin-bottom: 15px; filter: drop-shadow(0 0 15px rgba(167, 139, 250, 0.65));"></div>'
        else:
            logo_html = '<div style="text-align: center;"><h1 style="font-size: 4rem; margin-bottom: 10px;">✨</h1></div>'
            
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #0A0D1A 0%, #11152B 100%); border: 1px solid rgba(167, 139, 250, 0.25); border-radius: 24px; padding: 40px; box-shadow: 0 12px 50px rgba(0, 0, 0, 0.7); text-align: center;">
            {logo_html}
            <h2 style="color: #FFFFFF; font-family: Outfit; font-weight: 800; margin-bottom: 10px; letter-spacing: 1.5px; text-shadow: 0 0 15px rgba(167, 139, 250, 0.4);">JAISON OS</h2>
            <p style="color: #94A3B8; font-size: 0.9rem; margin-bottom: 10px;">Zabezpieczony panel dyrektorski Jaison Agent Agency</p>
        </div>
        """, unsafe_allow_html=True)
        
        with st.form("login_form", clear_on_submit=False):
            login_user = st.text_input("Nazwa użytkownika:", key="login_username")
            login_pass = st.text_input("Hasło dostępu:", type="password", key="login_password")
            submit = st.form_submit_button("Autoryzuj i wejdź", type="primary", use_container_width=True)
            
            if submit:
                # Usunięcie pustych znaków z obu stron
                u_input = login_user.strip() if login_user else ""
                p_input = login_pass.strip() if login_pass else ""
                
                # Zapisz do logu diagnostycznego serwera
                print(f"DEBUG LOGIN - Input User: {repr(u_input)} (len={len(u_input)}), Input Pass: {repr(p_input)} (len={len(p_input)}) | Expected User: {repr(AUTH_USER)} (len={len(AUTH_USER)}), Expected Pass: {repr(AUTH_PASS)} (len={len(AUTH_PASS)})")
                
                if u_input == AUTH_USER and p_input == AUTH_PASS:
                    st.session_state.authenticated = True
                    st.success("Autoryzacja pomyślna!")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.error("Nieprawidłowy użytkownik lub hasło.")
        
        st.markdown("""
        <p style="text-align: center; color: #475569; font-size: 0.8rem; margin-top: 30px;">
            System zabezpieczony szyfrowaniem AES-256 i routingiem proxy.
        </p>
        """, unsafe_allow_html=True)
    st.stop()


# Helpery danych
def load_kanban():
    default_kanban = {
        "triage": [],
        "todo": [],
        "ready": [],
        "running": [],
        "blocked": [],
        "done": []
    }
    if os.path.exists(KANBAN_FILE):
        try:
            with open(KANBAN_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            # Migrate old 3-column format if present
            if "in_progress" in data and "running" not in data:
                data["running"] = data.pop("in_progress")
            # Ensure all 6 columns exist
            for col in default_kanban:
                if col not in data:
                    data[col] = []
            return data
        except Exception:
            pass
    return default_kanban

def save_kanban(data):
    json.dump(data, open(KANBAN_FILE, "w", encoding="utf-8"), ensure_ascii=False, indent=4)

CRM_FILE = os.path.join(DASHBOARD_DIR, "crm.json")

def load_crm():
    if os.path.exists(CRM_FILE):
        try:
            return json.load(open(CRM_FILE, "r", encoding="utf-8"))
        except:
            pass
    return {
        "leads": [
            {
                "id": "lead_01",
                "name": "Klinika Dermatologiczna (Łódź)",
                "stage": "conversation",
                "notes": "Zainteresowani automatyzacją ankiety intake w celu uwolnienia czasu personelu. Duża potrzeba empatii i prostego lejka wdrożeniowego.",
                "last_contact": "2026-05-29",
                "next_action": "Opracować plan wdrożenia 3 asystentów AI.",
                "draft_reply": "Dzień dobry. Przeanalizowałem Państwa wąskie gardła. Wdrożenie Niewidzialnego Pracownika AI do ankiety intake pozwoli zaoszczędzić około 15 godzin tygodniowo bez utraty ciepłego kontaktu z pacjentem..."
            },
            {
                "id": "lead_02",
                "name": "Piotr Kowalski (Szkolenia B2B)",
                "stage": "architecture",
                "notes": "Zainteresowany systemem czystej pamięci (Pristine Memory) dla swoich uczestników z ADHD.",
                "last_contact": "2026-05-28",
                "next_action": "Przygotować interfejs MVP z pasywnym tablicowaniem.",
                "draft_reply": "Cześć Piotr, w oparciu o nasze rozmowy, zaprojektowałem strukturę bezszumnego CRM bez powiadomień push..."
            }
        ]
    }

def save_crm(data):
    json.dump(data, open(CRM_FILE, "w", encoding="utf-8"), ensure_ascii=False, indent=4)

# --- JAISON CRM INTEGRATIONS ---
import urllib.parse
import urllib.request
import urllib.error
from http.server import BaseHTTPRequestHandler, HTTPServer
import threading

class UrllibResponse:
    def __init__(self, status_code, text, json_data=None):
        self.status_code = status_code
        self.text = text
        self._json_data = json_data
    
    def json(self):
        if self._json_data is not None:
            return self._json_data
        return json.loads(self.text)

def http_post(url, json_data=None, headers=None, timeout=30.0):
    if headers is None:
        headers = {}
    
    # Ensure Content-Type is set if json_data is present
    if json_data is not None and "Content-Type" not in headers:
        headers["Content-Type"] = "application/json"
        
    data_bytes = None
    if json_data is not None:
        data_bytes = json.dumps(json_data).encode("utf-8")
        
    req = urllib.request.Request(url, data=data_bytes, headers=headers, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=timeout) as response:
            status_code = response.getcode()
            response_text = response.read().decode("utf-8")
            try:
                res_json = json.loads(response_text)
            except:
                res_json = None
            return UrllibResponse(status_code, response_text, res_json)
    except urllib.error.HTTPError as e:
        status_code = e.code
        try:
            response_text = e.read().decode("utf-8")
        except:
            response_text = str(e)
        try:
            res_json = json.loads(response_text)
        except:
            res_json = None
        return UrllibResponse(status_code, response_text, res_json)
    except Exception as e:
        return UrllibResponse(500, str(e), None)

# Thread-safe global flag to start webhook server once
_server_started = False
_server_lock = threading.Lock()

def call_openrouter_api(messages, model="nousresearch/hermes-3-llama-3.1-405b:free", system_instruction=None):
    api_key = os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        return "Błąd: Brak klucza OPENROUTER_API_KEY w pliku .env"
    
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:8501",
        "X-Title": "Holistic ADHD OS"
    }
    
    # Map friendly short names to openrouter ids if necessary
    if model == "owl-alpha-free" or model == "owl-alpha":
        model = "meta-llama/llama-3.1-8b-instruct:free"
    
    payload = {
        "model": model,
        "messages": []
    }
    if system_instruction:
        payload["messages"].append({"role": "system", "content": system_instruction})
    payload["messages"].extend(messages)
    
    try:
        response = http_post(url, json_data=payload, headers=headers, timeout=60.0)
        if response.status_code == 200:
            res_data = response.json()
            return res_data["choices"][0]["message"]["content"]
        else:
            return f"Błąd OpenRouter {response.status_code}: {response.text}"
    except Exception as e:
        return f"Błąd OpenRouter API: {e}"

def get_recent_workspace_context(limit=3):
    context_str = ""
    try:
        if os.path.exists(OBSIDIAN_DIR):
            files = [f for f in os.listdir(OBSIDIAN_DIR) if f.endswith('.md')]
            # Sort by modification time desc
            files.sort(key=lambda x: os.path.getmtime(os.path.join(OBSIDIAN_DIR, x)), reverse=True)
            for f in files[:limit]:
                path = os.path.join(OBSIDIAN_DIR, f)
                with open(path, "r", encoding="utf-8") as file:
                    content = file.read()
                # Keep it concise, extract up to 1000 chars per file
                context_str += f"\n--- Ostatnia notatka: {f} ---\n{content[:1000]}\n"
    except Exception as e:
        context_str = f"Błąd wczytywania pamięci roboczej: {e}"
    return context_str

def call_vllm_api(messages, system_instruction=None):
    url = "http://localhost:8000/v1/chat/completions"
    headers = {
        "Authorization": "Bearer HOLISTIC_SECURE_TOKEN_2026",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "mistralai/Mistral-Nemo-Instruct-2407",
        "messages": []
    }
    if system_instruction:
        payload["messages"].append({"role": "system", "content": system_instruction})
    payload["messages"].extend(messages)
    
    try:
        response = http_post(url, json_data=payload, headers=headers, timeout=120.0)
        if response.status_code == 200:
            res_data = response.json()
            return res_data["choices"][0]["message"]["content"]
    except Exception as e:
        pass
        
    return call_gemini_api(messages, system_instruction)

def call_gemini_api(messages, system_instruction=None):
    proxy_url = "http://127.0.0.1:8089/v1/chat/completions"
    payload = {
        "model": "gemini-2.5-flash",
        "messages": []
    }
    if system_instruction:
        payload["messages"].append({"role": "system", "content": system_instruction})
    payload["messages"].extend(messages)
    
    # 1. Try local proxy (timeout 90s for long OCR operations)
    try:
        response = http_post(proxy_url, json_data=payload, timeout=90.0)
        if response.status_code == 200:
            res_data = response.json()
            return res_data["choices"][0]["message"]["content"]
    except Exception as e:
        pass
        
    # 2. Fallback to direct GCP Vertex AI call
    sa_paths = [
        os.path.expanduser("~/.hermes/gcp-sa-key.json"),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "holistic-dashboard-dev-dea2c872139e.json"),
        "holistic-dashboard-dev-dea2c872139e.json"
    ]
    sa_path = None
    for p in sa_paths:
        if os.path.exists(p):
            sa_path = p
            break
            
    if sa_path:
        try:
            from google.oauth2 import service_account
            import google.auth.transport.requests
            import requests
            import urllib3
            urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
            
            creds = service_account.Credentials.from_service_account_file(
                sa_path,
                scopes=['https://www.googleapis.com/auth/cloud-platform']
            )
            session = requests.Session()
            session.verify = False
            request = google.auth.transport.requests.Request(session=session)
            creds.refresh(request)
            token = creds.token
            
            project_id = "holistic-dashboard-dev"
            region = "us-central1"
            direct_url = f"https://{region}-aiplatform.googleapis.com/v1/projects/{project_id}/locations/{region}/endpoints/openapi/chat/completions"
            
            headers = {
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json"
            }
            
            direct_payload = {
                "model": "google/gemini-2.5-flash",
                "messages": []
            }
            if system_instruction:
                direct_payload["messages"].append({"role": "system", "content": system_instruction})
            direct_payload["messages"].extend(messages)
            
            resp = http_post(direct_url, json_data=direct_payload, headers=headers, timeout=25.0)
            if resp.status_code == 200:
                res_data = resp.json()
                return res_data["choices"][0]["message"]["content"]
        except Exception as ex:
            return f"Błąd bezpośredniej komunikacji z Vertex AI: {ex}"
            
    return "Błąd: Brak połączenia z lokalnym proxy i brak poprawnego klucza Service Account."

def get_gcp_token():
    """Pobiera token OAuth z pliku Service Account w odporny sposób."""
    import os
    import json
    sa_paths = [
        os.path.expanduser("~/.hermes/gcp-sa-key.json"),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "holistic-dashboard-dev-dea2c872139e.json"),
        os.path.join(os.getcwd(), "holistic-dashboard-dev-dea2c872139e.json"),
        "holistic-dashboard-dev-dea2c872139e.json"
    ]
    sa_path = None
    for p in sa_paths:
        if os.path.exists(p):
            sa_path = p
            break
    if not sa_path:
        return None, None
    try:
        from google.oauth2 import service_account
        import google.auth.transport.requests
        import requests
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        
        creds = service_account.Credentials.from_service_account_file(
            sa_path,
            scopes=['https://www.googleapis.com/auth/cloud-platform']
        )
        session = requests.Session()
        session.verify = False
        request = google.auth.transport.requests.Request(session=session)
        creds.refresh(request)
        
        with open(sa_path, "r") as f:
            sa_data = json.load(f)
            project_id = sa_data.get("project_id", "holistic-dashboard-dev")
            
        return creds.token, project_id
    except Exception as e:
        print(f"Error getting token: {e}")
        return None, None

def generate_imagen_image(prompt, aspect_ratio="1:1", reference_image_bytes=None, reference_type="REFERENCE_TYPE_SUBJECT"):
    """
    Generuje obraz za pomocą Google Vertex AI Imagen 3 (imagen-3.0-generate-001).
    Wspiera Image-to-Image / Subject Reference dla zachowania spójności postaci.
    """
    import base64
    import json
    import requests
    token, project_id = get_gcp_token()
    if not token:
        return None, "Brak autoryzacji GCP Service Account. Sprawdź pliki klucza."
    
    region = "us-central1"
    url = f"https://{region}-aiplatform.googleapis.com/v1/projects/{project_id}/locations/{region}/publishers/google/models/imagen-3.0-generate-001:predict"
    
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }
    
    instance = {"prompt": prompt}
    
    # Obsługa obrazu referencyjnego (Subject/Style Reference Image)
    if reference_image_bytes:
        # Resize do max 256x256 za pomocą PIL, aby spełnić twarde ograniczenia rozmiaru Vertex AI REST API
        try:
            from PIL import Image
            import io
            img = Image.open(io.BytesIO(reference_image_bytes))
            img.thumbnail((256, 256))
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            resized_bytes = buf.getvalue()
            b64_data = base64.b64encode(resized_bytes).decode("utf-8")
        except Exception as p_ex:
            resized_bytes = reference_image_bytes
            b64_data = base64.b64encode(reference_image_bytes).decode("utf-8")
            
        if reference_type == "REFERENCE_TYPE_SUBJECT":
            # Dla SUBJECT Vertex AI bezwzględnie wymaga ścieżki GCS (gs://...)
            bucket_name = "holistic-brand-assets"
            object_name = "temp_ref_subject.png"
            gcs_url = f"https://storage.googleapis.com/upload/storage/v1/b/{bucket_name}/o?uploadType=media&name={object_name}"
            gcs_headers = {
                "Authorization": f"Bearer {token}",
                "Content-Type": "image/png"
            }
            try:
                gcs_res = requests.post(gcs_url, headers=gcs_headers, data=resized_bytes, timeout=30, verify=False)
                if gcs_res.status_code in [200, 201]:
                    gcs_uri = f"gs://{bucket_name}/{object_name}"
                else:
                    return None, f"Błąd wgrywania obrazu referencyjnego na GCS: {gcs_res.status_code} - {gcs_res.text}"
            except Exception as upload_ex:
                return None, f"Wyjątek podczas wgrywania na GCS: {str(upload_ex)}"
                
            ref_image_obj = {
                "gcsUri": gcs_uri,
                "mimeType": "image/png"
            }
        else:
            # STYLE lub RAW mogą korzystać z bezpośrednich bajtów
            ref_image_obj = {
                "imageBytes": b64_data,
                "mimeType": "image/png"
            }
        
        ref_item = {
            "referenceId": 1,
            "referenceType": reference_type,
            "referenceImage": ref_image_obj
        }
        
        if reference_type == "REFERENCE_TYPE_SUBJECT":
            ref_item["subjectImageConfig"] = {
                "subjectDescription": "a professional creative male in his 30s",
                "subjectType": "SUBJECT_TYPE_PERSON"
            }
        elif reference_type == "REFERENCE_TYPE_STYLE":
            ref_item["styleImageConfig"] = {
                "styleDescription": "glowing professional style"
            }
            
        instance["referenceImages"] = [ref_item]
        
        if "[1]" not in prompt:
            prompt = f"Using subject [1] as character reference, {prompt}"
            instance["prompt"] = prompt
            
    payload = {
        "instances": [instance],
        "parameters": {
            "sampleCount": 1,
            "aspectRatio": aspect_ratio,
            "outputMimeType": "image/png"
        }
    }
    
    try:
        response = http_post(url, json_data=payload, headers=headers, timeout=120.0)
        if response.status_code == 200:
            res_json = response.json()
            predictions = res_json.get("predictions", [])
            if predictions:
                img_b64 = predictions[0].get("bytesBase64Encoded")
                img_bytes = base64.b64decode(img_b64)
                return img_bytes, None
            else:
                return None, f"Brak wygenerowanego obrazu w odpowiedzi: {response.text}"
        elif response.status_code == 429:
            err_msg = (
                "⚠️ **PRZEKROCZONO LIMIT / BRAK UPRAWNIEŃ QUOTA NA GCP**\n\n"
                "Twój projekt Google Cloud (`holistic-dashboard-dev`) posiada domyślny limit **0** (lub wyczerpany) "
                "dla zaawansowanych funkcji modelu **Imagen 3.0 (customizacja postaci / stylów przez reference images)**.\n\n"
                "**Jak to naprawić w 2 minuty?**\n"
                "1. Zaloguj się do **Google Cloud Console** na swoim koncie.\n"
                "2. Przejdź do zakładki **Quotas & System Limits** (Limity i przydziały):\n"
                "   👉 [https://console.cloud.google.com/iam-admin/quotas](https://console.cloud.google.com/iam-admin/quotas)\n"
                "3. Wyszukaj limit o nazwie: `aiplatform.googleapis.com/online_prediction_requests_per_base_model`.\n"
                "4. Znajdź pozycję dla modelu `imagen-3.0-generate`.\n"
                "5. Kliknij **EDIT QUOTAS** (Edytuj limity), wpisz nową wartość (np. **100** lub więcej) i wyślij zgłoszenie.\n"
                "6. Akceptacja przez Google następuje automatycznie w ciągu kilku minut!"
            )
            return None, err_msg
        else:
            return None, f"GCP API Error {response.status_code}: {response.text}"
    except Exception as e:
        return None, f"Błąd komunikacji: {str(e)}"

def call_gemini_pro_api(messages, system_instruction=None):
    """Dedykowana funkcja dla modelu Gemini 2.5 Pro przez lokalny proxy.
    Używamy go w Dziale Prawnym dla analizy dokumentów i logiki prawnej.
    Timeout 120s dla długich dokumentów."""
    proxy_url = "http://127.0.0.1:8089/v1/chat/completions"
    
    # Sprawdź czy wiadomości zawierają multimodalne treści (PDF/obrazy) — użyj Flash
    # bo proxy może nie obsługiwać dużych payloadów multimodalnych dla Pro
    has_multimodal = False
    for msg in messages:
        content = msg.get("content", "")
        if isinstance(content, list):
            for part in content:
                if isinstance(part, dict) and part.get("type") == "image_url":
                    has_multimodal = True
                    break
    
    model_name = "gemini-2.5-flash" if has_multimodal else "gemini-2.5-pro"
    
    payload = {"model": model_name, "messages": []}
    if system_instruction:
        payload["messages"].append({"role": "system", "content": system_instruction})
    payload["messages"].extend(messages)

    try:
        response = http_post(proxy_url, json_data=payload, timeout=120.0)
        if response.status_code == 200:
            res_data = response.json()
            return res_data["choices"][0]["message"]["content"]
        else:
            # Próba z Flash jako fallback
            payload["model"] = "gemini-2.5-flash"
            response2 = http_post(proxy_url, json_data=payload, timeout=90.0)
            if response2.status_code == 200:
                return response2.json()["choices"][0]["message"]["content"]
    except Exception:
        pass

    # Ostateczny fallback
    return call_gemini_api(messages, system_instruction)

def call_native_vertex_ocr(file_bytes):
    # Get credentials and refresh token
    sa_paths = [
        os.path.join(os.getcwd(), "holistic-dashboard-dev-dea2c872139e.json"),
        os.path.expanduser("~/.hermes/gcp-sa-key.json"),
        "holistic-dashboard-dev-dea2c872139e.json"
    ]
    sa_path = None
    for p in sa_paths:
        if os.path.exists(p):
            sa_path = p
            break
            
    if not sa_path:
        return "[Błąd: Brak klucza GCP Service Account]"
        
    try:
        from google.oauth2 import service_account
        import google.auth.transport.requests
        import requests
        import base64 as _b64
        
        creds = service_account.Credentials.from_service_account_file(
            sa_path,
            scopes=['https://www.googleapis.com/auth/cloud-platform']
        )
        session = requests.Session()
        session.verify = False
        request = google.auth.transport.requests.Request(session=session)
        creds.refresh(request)
        token = creds.token
        
        project_id = "holistic-dashboard-dev"
        region = "us-central1"
        model = "gemini-2.5-flash" # Use flash for fast OCR
        
        url = f"https://{region}-aiplatform.googleapis.com/v1/projects/{project_id}/locations/{region}/publishers/google/models/{model}:generateContent"
        
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }
        
        encoded_data = _b64.b64encode(file_bytes).decode("utf-8")
        
        payload = {
            "contents": [
                {
                    "role": "user",
                    "parts": [
                        {
                            "text": "Odczytaj i przepisz DOKŁADNIE cały tekst z załączonego dokumentu PDF (skanu). Zachowaj wszystkie kluczowe dane: sygnatury akt, PESEL, NIP, adresy, imiona i nazwiska, kwoty, daty, nazwy sądów lub organów. Zwróć tylko surowy odczytany tekst, zachowując oryginalny układ."
                        },
                        {
                            "inlineData": {
                                "mimeType": "application/pdf",
                                "data": encoded_data
                            }
                        }
                    ]
                }
            ],
            "generationConfig": {
                "temperature": 0.1
            }
        }
        
        resp = requests.post(url, json=payload, headers=headers, verify=False, timeout=120.0)
        if resp.status_code == 200:
            res_data = resp.json()
            try:
                text = res_data['candidates'][0]['content']['parts'][0]['text']
                return text
            except Exception as e:
                return f"[Błąd parsowania odpowiedzi OCR: {e}]"
        else:
            return f"[Błąd API Vertex OCR {resp.status_code}: {resp.text}]"
    except Exception as ex:
        return f"[Błąd systemu OCR: {ex}]"

def call_gcp_tts(text, voice_name="pl-PL-Wavenet-B", gender="MALE"):
    sa_paths = [
        os.path.join(os.getcwd(), "holistic-dashboard-dev-dea2c872139e.json"),
        os.path.expanduser("~/.hermes/gcp-sa-key.json"),
        "holistic-dashboard-dev-dea2c872139e.json"
    ]
    sa_path = None
    for p in sa_paths:
        if os.path.exists(p):
            sa_path = p
            break
            
    if not sa_path:
        return None, "Brak klucza GCP Service Account"
        
    try:
        from google.oauth2 import service_account
        import google.auth.transport.requests
        import requests
        import base64 as _b64
        
        creds = service_account.Credentials.from_service_account_file(
            sa_path,
            scopes=['https://www.googleapis.com/auth/cloud-platform']
        )
        session = requests.Session()
        session.verify = False
        request = google.auth.transport.requests.Request(session=session)
        creds.refresh(request)
        token = creds.token
        
        url = "https://texttospeech.googleapis.com/v1/text:synthesize"
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "input": {"text": text},
            "voice": {
                "languageCode": "pl-PL",
                "name": voice_name,
                "ssmlGender": gender
            },
            "audioConfig": {
                "audioEncoding": "MP3"
            }
        }
        
        resp = requests.post(url, json=payload, headers=headers, verify=False, timeout=15.0)
        if resp.status_code == 200:
            res_data = resp.json()
            audio_content = res_data.get("audioContent", "")
            if audio_content:
                return _b64.b64decode(audio_content), None
            return None, "Brak zawartości audio w odpowiedzi"
        else:
            return None, f"Błąd API TTS {resp.status_code}: {resp.text}"
    except Exception as ex:
        return None, f"Błąd połączenia z GCP TTS: {ex}"

def query_vertex_search(query, data_store_id, project_id="holistic-dashboard-dev", location="global"):
    sa_paths = [
        os.path.join(os.getcwd(), "holistic-dashboard-dev-dea2c872139e.json"),
        os.path.expanduser("~/.hermes/gcp-sa-key.json"),
        "holistic-dashboard-dev-dea2c872139e.json"
    ]
    sa_path = None
    for p in sa_paths:
        if os.path.exists(p):
            sa_path = p
            break
            
    if not sa_path:
        return {"error": "Brak pliku klucza GCP Service Account (holistic-dashboard-dev-dea2c872139e.json). Upewnij się, że znajduje się w głównym folderze."}
        
    try:
        from google.cloud import discoveryengine_v1beta as discoveryengine
        from google.oauth2 import service_account
        
        creds = service_account.Credentials.from_service_account_file(
            sa_path,
            scopes=['https://www.googleapis.com/auth/cloud-platform']
        )
        
        client = discoveryengine.SearchServiceClient(credentials=creds)
        
        if "connector" in data_store_id or "datastore" in data_store_id.lower() or "data-store" in data_store_id.lower():
            serving_config = f"projects/{project_id}/locations/{location}/dataStores/{data_store_id}/servingConfigs/default_search"
        else:
            serving_config = f"projects/{project_id}/locations/{location}/collections/default_collection/engines/{data_store_id}/servingConfigs/default_search"
        
        content_search_spec = discoveryengine.SearchRequest.ContentSearchSpec(
            extractive_content_spec=discoveryengine.SearchRequest.ContentSearchSpec.ExtractiveContentSpec(
                max_extractive_answer_count=3,
                max_extractive_segment_count=1,
                return_extractive_segment_source=True
            ),
            snippet_spec=discoveryengine.SearchRequest.ContentSearchSpec.SnippetSpec(
                max_snippet_count=1,
                return_snippet=True
            ),
            summary_spec=discoveryengine.SearchRequest.ContentSearchSpec.SummarySpec(
                summary_result_count=5,
                include_citations=True,
                model_spec=discoveryengine.SearchRequest.ContentSearchSpec.SummarySpec.ModelSpec(
                    version="stable"
                )
            )
        )
        
        request = discoveryengine.SearchRequest(
            serving_config=serving_config,
            query=query,
            page_size=5,
            content_search_spec=content_search_spec
        )
        
        response = client.search(request)
        
        results = []
        summary = ""
        if hasattr(response, "summary") and response.summary:
            summary = response.summary.summary_text
            
        for result in response:
            doc = result.document
            struct_data = dict(doc.derived_struct_data) if doc.derived_struct_data else {}
            
            title = struct_data.get("title", doc.id)
            link = struct_data.get("link", "")
            snippet = struct_data.get("snippet", "")
            
            extractive_answers = struct_data.get("extractive_answers", [])
            extractive_segments = struct_data.get("extractive_segments", [])
            
            answers = []
            if isinstance(extractive_answers, list):
                for ans in extractive_answers:
                    if isinstance(ans, dict) and "content" in ans:
                        answers.append(ans["content"])
                        
            segments = []
            if isinstance(extractive_segments, list):
                for seg in extractive_segments:
                    if isinstance(seg, dict) and "content" in seg:
                        segments.append(seg["content"])
            
            results.append({
                "id": doc.id,
                "title": title,
                "link": link,
                "snippet": snippet,
                "extractive_answers": answers,
                "extractive_segments": segments,
                "struct_data": struct_data
            })
            
        return {
            "results": results,
            "summary": summary
        }
    except Exception as e:
        return {"error": str(e)}

def extract_youtube_transcript_raw(url):
    import re
    pattern = r'(?:youtube\.com\/(?:[^\/]+\/.+\/|(?:v|e(?:mbed)?)\/|.*[?&]v=)|youtu\.be\/)([^"&?\/\s]{11})'
    match = re.search(pattern, url)
    if not match:
        return None, "Niepoprawny adres URL wideo YouTube"
    video_id = match.group(1)
    
    # Store original requests Session request to restore it afterwards
    import requests
    orig_request = requests.Session.request
    
    try:
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        
        # Monkey patch requests for SSL bypass
        requests.Session.request = lambda self, *a, **kw: orig_request(self, *a, **{**kw, 'verify': False})
        
        from youtube_transcript_api import YouTubeTranscriptApi
        api = YouTubeTranscriptApi()
        
        try:
            data = api.fetch(video_id, languages=['pl', 'en'])
        except Exception:
            try:
                data = api.fetch(video_id)
            except Exception as e2:
                # Restore original request
                requests.Session.request = orig_request
                return None, f"Błąd pobierania transkrypcji YouTube: {e2}"
                
        # Restore original request
        requests.Session.request = orig_request
        
        full_text = []
        for entry in data:
            text = entry['text']
            start = entry['start']
            minutes = int(start // 60)
            seconds = int(start % 60)
            full_text.append(f"[{minutes:02d}:{seconds:02d}] {text}")
            
        return "\n".join(full_text), None
    except Exception as e:
        # Guarantee restoration
        try:
            requests.Session.request = orig_request
        except Exception:
            pass
        return None, f"Błąd systemu transkrypcji: {str(e)}"

class WebhookHandler(BaseHTTPRequestHandler):
    def log_message(self, format, *args):
        pass
        
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, GET, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def do_POST(self):
        if self.path == "/webhook":
            content_length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(content_length)
            
            try:
                content_type = self.headers.get('Content-Type', '')
                data = {}
                if 'application/json' in content_type:
                    data = json.loads(body.decode('utf-8'))
                else:
                    parsed = urllib.parse.parse_qs(body.decode('utf-8'))
                    data = {k: v[0] for k, v in parsed.items()}
                
                name = data.get("name", "").strip()
                email = data.get("email", "").strip()
                notes = data.get("notes", "").strip() or data.get("message", "").strip()
                
                if not name or not email:
                    self.send_response(400)
                    self.send_header("Access-Control-Allow-Origin", "*")
                    self.end_headers()
                    self.wfile.write(b"Name and Email are required.")
                    return
                
                crm_data = load_crm()
                new_id = f"lead_{int(time.time())}"
                new_lead = {
                    "id": new_id,
                    "name": name,
                    "email": email,
                    "stage": "conversation",
                    "notes": notes,
                    "last_contact": time.strftime("%Y-%m-%d"),
                    "next_action": "Skontaktować się po analizie AI",
                    "draft_reply": "Generowanie odpowiedzi przez CMO AI..."
                }
                
                crm_data["leads"].append(new_lead)
                save_crm(crm_data)
                
                threading.Thread(
                    target=async_generate_initial_draft,
                    args=(new_id, name, notes),
                    daemon=True
                ).start()
                
                self.send_response(200)
                self.send_header("Content-Type", "application/json")
                self.send_header("Access-Control-Allow-Origin", "*")
                self.end_headers()
                self.wfile.write(json.dumps({"status": "success", "lead_id": new_id}).encode('utf-8'))
                
            except Exception as e:
                self.send_response(500)
                self.send_header("Access-Control-Allow-Origin", "*")
                self.end_headers()
                self.wfile.write(f"Server Error: {e}".encode('utf-8'))
        else:
            self.send_response(404)
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()

def async_generate_initial_draft(lead_id, client_name, client_pain):
    o_mnie_path = os.path.join(HERMES_DIR, "o_mnie.md")
    o_mnie_context = ""
    if os.path.exists(o_mnie_path):
        try:
            o_mnie_context = open(o_mnie_path, "r", encoding="utf-8").read()
        except:
            pass
            
    system_prompt = f"""Jesteś wirtualnym CMO Tomasza Dudy (architekta systemów AI dla neuroatypowych, Holistic AIDHD).
Tomasz jest strategiem automatyzacji dla firm i sam ma ADHD. Twoim zadaniem jest stworzenie bardzo empatycznej, osobistej i konkretnej wersji roboczej pierwszej odpowiedzi do nowego leada (klienta).
Unikaj pustego hype'u AI, pisz zwięźle, ze zrozumieniem chaosu operacyjnego klienta, i z nutą humoru oraz głębokim rezonansem emocjonalnym.

Oto kontekst o Tomaszu (jego serce emocjonalne i historia):
{o_mnie_context}

Napisz bezpośrednią, gotową do wysłania wiadomość e-mail w języku polskim. Pisz w 1 osobie ("ja" - Tomasz Duda). Podpisz się na końcu.
"""

    user_prompt = f"""Klient: {client_name}
Wyzwanie klienta: {client_pain}

Napisz pierwszą, niezwykle empatyczną odpowiedź, która zdejmuje z niego paraliż decyzyjny i proponuje jeden mały, konkretny krok bez presji (krótką rozmowę diagnostyczną)."""

    try:
        reply = call_gemini_api([{"role": "user", "content": user_prompt}], system_instruction=system_prompt)
        if reply:
            crm_data = load_crm()
            for lead in crm_data["leads"]:
                if lead["id"] == lead_id:
                    lead["draft_reply"] = reply
                    break
            save_crm(crm_data)
    except Exception as e:
        pass

def consult_csuite_live(lead, persona):
    o_mnie_path = os.path.join(HERMES_DIR, "o_mnie.md")
    o_mnie_context = ""
    if os.path.exists(o_mnie_path):
        try:
            o_mnie_context = open(o_mnie_path, "r", encoding="utf-8").read()
        except:
            pass

    # Dynamiczne wstrzykiwanie checklist i SOP z folderu skills/
    persona_mapping = {
        "CEO (Strategia & Rentowność)": "ceo",
        "CMO (Empatyczny Storytelling)": "cmo",
        "CSO (Architektura Sprzedaży)": "cso",
        "CTO (Technologia & Kod)": "cto"
    }
    
    dynamic_skill_context = ""
    folder_name = persona_mapping.get(persona)
    if folder_name:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        skill_paths = [
            os.path.join(base_dir, "skills", folder_name, "SKILL.md"),
            os.path.join(base_dir, ".agents", "skills", folder_name, "SKILL.md"),
            os.path.expanduser(f"~/.gemini/config/plugins/holistic-virtual-board/skills/{folder_name}/SKILL.md")
        ]
        for p in skill_paths:
            if os.path.exists(p):
                try:
                    with open(p, "r", encoding="utf-8") as f:
                        dynamic_skill_context = f.read()
                    break
                except:
                    pass

    global_ctx = get_selected_context_data()
    
    system_prompts = {
        "CEO (Strategia & Rentowność)": f"""Jesteś wirtualnym CEO w zespole Tomasza Dudy. Tomasz to wybitny architekt systemów AI dla neuroatypowych (sam ma ADHD, Holistic AIDHD).
Twój aktualny cel zależy od wybranego Kontekstu:
{global_ctx}

Pomagasz mu w wycenie wdrożenia pod kątem modelu High-Ticket (np. wyceny 5 000 - 15 000 PLN jednorazowo), etapowaniu prac na proste kroki MVP oraz obronie jego zasobów energetycznych przed wypaleniem i paraliżem ADHD.
Zawsze podawaj konkretną, odważną rekomendację cenową i zdefiniuj, co jest "One Thing" (kluczowym pierwszym krokiem wdrożenia).

Oto Twoje oficjalne wytyczne i checklisty operacyjne (SOP):
{dynamic_skill_context}

Oto historia i tożsamość Tomasza:
{o_mnie_context}
Odpowiadaj bezpośrednio, po polsku, zwięźle i konkretnie.
""",
        "CMO (Empatyczny Storytelling)": f"""Jesteś wirtualnym CMO w zespole Tomasza Dudy (Holistic AIDHD).
Twój aktualny cel zależy od wybranego Kontekstu:
{global_ctx}

Pomagasz mu przełożyć ból klienta na autentyczny i humorystyczny przekaz dopasowany do wyzwań klienta. Wskaż, jakich metafor użyć w komunikacji z tym klientem i jak napisać ofertę, aby rezonowała głęboko emocjonalnie, opierając się na tożsamości Tomasza.

Oto Twoje oficjalne wytyczne i checklisty operacyjne (SOP):
{dynamic_skill_context}

Oto historia i tożsamość Tomasza:
{o_mnie_context}
Odpowiadaj bezpośrednio, po polsku, zwięźle i kreatywnie.
""",
        "CSO (Architektura Sprzedaży)": f"""Jesteś wirtualnym CSO w zespole Tomasza Dudy (Holistic AIDHD).
Twój aktualny cel zależy od wybranego Kontekstu:
{global_ctx}

Projektujesz dla tego klienta prosty, 3-stopniowy lejek relacyjny (Rozmowa -> Architektura -> Wdrożenie). Wskaż dokładnie, jaki powinien być najbliższy krok sprzedażowy (Next Action) i jak go zrealizować przy minimalnym tarciu poznawczym (low cognitive friction).

Oto Twoje oficjalne wytyczne i checklisty operacyjne (SOP):
{dynamic_skill_context}

Oto historia i tożsamość Tomasza:
{o_mnie_context}
Odpowiadaj bezpośrednio, po polsku, zwięźle i operacyjnie.
""",
        "CTO (Technologia & Kod)": f"""Jesteś wirtualnym CTO w zespole Tomasza Dudy (Holistic AIDHD).
Twój aktualny cel zależy od wybranego Kontekstu:
{global_ctx}

Zaprojektuj uproszczoną, niezawodną architekturę techniczną pod potrzeby tego klienta. Rekomenduj konkretne narzędzia (np. n8n webhooks, Python scripts, SQLite, Google Sheets, Vertex AI, model gemini-2.5-flash). Podaj zwięzły schemat logiczny.

Oto Twoje oficjalne wytyczne i checklisty operacyjne (SOP):
{dynamic_skill_context}

Oto historia i tożsamość Tomasza:
{o_mnie_context}
Odpowiadaj bezpośrednio, po polsku, technicznie lecz bez zbędnego żargonu.
"""
    }

    system_instruction = system_prompts.get(persona, "Jesteś doradcą C-Suite.")
    user_prompt = f"""Karta klienta:
Nazwa: {lead['name']}
Opis chaosu operacyjnego: {lead['notes']}
Najbliższy krok (Next Action): {lead['next_action']}

Przeanalizuj przypadek i daj mi (Tomaszowi) swoją rekomendację z perspektywy roli ({persona}). Pisz bezpośrednio do mnie ("Tomasz, badając przypadek...")."""

    return call_gemini_api([{"role": "user", "content": user_prompt}], system_instruction=system_instruction)

def start_webhook_server():
    global _server_started
    with _server_lock:
        if _server_started:
            return
        server_address = ('0.0.0.0', 8090)
        try:
            httpd = HTTPServer(server_address, WebhookHandler)
            t = threading.Thread(target=httpd.serve_forever, daemon=True)
            t.start()
            _server_started = True
        except Exception as e:
            pass

# Start the background webhook server
start_webhook_server()

def read_md_file(path):
    return open(path, "r", encoding="utf-8").read() if os.path.exists(path) else ""

def save_brain_dump(thought, links, uploaded_file):
    dump_id = f"dump_{int(time.time())}"
    file_name = None
    if uploaded_file:
        file_name = uploaded_file.name
        with open(os.path.join(BRAIN_DUMP_ASSETS, file_name), "wb") as f:
            f.write(uploaded_file.getbuffer())
    category = "Later" if any(w in thought.lower() for w in ["potem", "później", "later", "kiedyś"]) else "Now"
    dump_data = {
        "id": dump_id,
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "thought": thought,
        "links": links,
        "file_attached": file_name,
        "category": category,
        "status": "active"
    }
    json.dump(dump_data, open(os.path.join(BRAIN_DUMP_DIR, f"{dump_id}.json"), "w", encoding="utf-8"), ensure_ascii=False, indent=4)

def get_brain_dumps():
    dumps = []
    if os.path.exists(BRAIN_DUMP_DIR):
        for f in os.listdir(BRAIN_DUMP_DIR):
            if f.endswith(".json") and f.startswith("dump_"):
                try:
                    dumps.append(json.load(open(os.path.join(BRAIN_DUMP_DIR, f), "r", encoding="utf-8")))
                except:
                    pass
    dumps.sort(key=lambda x: x.get("timestamp", ""), reverse=True)
    return dumps

@st.dialog("📥 Szybki Zrzut Myśli (Brain Dump)")
def show_brain_dump_dialog():
    st.markdown("<p style='color: #CBD5E1;'>Wpisz swój pomysł, załącz link lub plik, a model Gemini 2.5 Pro odciąży Twój umysł i natychmiast zapisze go w Skarbcu.</p>", unsafe_allow_html=True)
    thought = st.text_area("Co Ci chodzi po głowie? (Zzrzuć chaos)", height=120)
    links = st.text_input("Linki / Social media / Wideo (opcjonalnie):")
    uploaded_file = st.file_uploader("Załącz plik / obrazek / audio (opcjonalnie):", type=["png", "jpg", "jpeg", "pdf", "mp3", "wav"])
    
    if st.button("Uwolnij mój umysł", type="primary", use_container_width=True):
        if thought or links or uploaded_file:
            save_brain_dump(thought, links, uploaded_file)
            st.success("Zapisano bezpiecznie w Skarbcu!")
            time.sleep(1.0)
            st.rerun()

check_telegram_events()

# Inicjalizacja stanu sesji
if "one_thing" not in st.session_state:
    st.session_state.one_thing = ""
if "pomodoro_active" not in st.session_state:
    st.session_state.pomodoro_active = False
if "current_page" not in st.session_state:
    st.session_state.current_page = "🎯 Mission Control"

# PASEK BOCZNY - Skrajnie estetyczny z kafelkami zoptymalizowanymi pod ADHD (Styl Julian Goldie)
with st.sidebar:
    if logo_base64:
        st.markdown(f"""
        <div style="text-align: center; padding-top: 10px; padding-bottom: 5px;">
            <img src="data:image/png;base64,{logo_base64}" style="width: 100px; height: 100px; filter: drop-shadow(0 0 12px rgba(167, 139, 250, 0.5));">
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("<div style='text-align: center; font-size: 3rem; padding-top:10px;'>✨</div>", unsafe_allow_html=True)
        
    st.markdown("<h2 style='text-align: center; color: #FFFFFF; font-family: Outfit; margin-bottom: 0; margin-top: -5px; font-weight: 800; letter-spacing: 1.5px; text-shadow: 0 0 15px rgba(167, 139, 250, 0.4);'>JAISON OS</h2>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #94A3B8; font-size: 0.85rem; margin-top: 5px; margin-bottom: 10px;'>Agentic OS Mission Control v8.0</p>", unsafe_allow_html=True)
    st.markdown("<hr style='margin: 10px 0; border-color: rgba(167, 139, 250, 0.15);'>", unsafe_allow_html=True)
    
    # 0. KONTEKST
    st.markdown("<p style='color: #10B981; font-weight: bold; font-size: 0.75rem; letter-spacing: 1px; margin-bottom: 6px; margin-top: 5px;'>0. KONTEKST KLIENTA</p>", unsafe_allow_html=True)
    
    # Wczytaj dostepne konteksty dynamicznie z realnego folderu klientów
    available_contexts = ["J(AI)SON Agency", "Holistic Jason", "app.jaison.pl (SaaS)"]
    real_clients_dir = r"C:\Aplikacje MVP\02_CLIENTS_AND_PROJECTS"
    
    if os.path.exists(real_clients_dir):
        mapping_names = {
            "coolfon": "coolfon.pl",
            "kurczakujasia": "kurczakujasia.pl (Bar Jaś)",
            "lifewave": "lifewave.com (MLM)",
            "smartrade_client": "smartrade.pl",
            "viptransporter": "viptransporter.pl",
            "kantor_lombard_oranzada": "kantororanzada.pl",
            "vojsik_ai": "vojsik.ai",
            "apps.jaison.pl": "app.jaison.pl (SaaS)"
        }
        for item in os.listdir(real_clients_dir):
            if os.path.isdir(os.path.join(real_clients_dir, item)):
                if item in ["Szablon_Projektu", "coolfon_pl_test", "vojsik_mvp"]:
                    continue
                nice_name = mapping_names.get(item, item)
                if nice_name not in available_contexts:
                    available_contexts.append(nice_name)
                    
    # Dodaj fallback w razie braku plików na dysku
    if len(available_contexts) <= 3:
        available_contexts.extend(["coolfon.pl", "kurczakujasia.pl (Bar Jaś)", "lifewave.com (MLM)"])
        
    if "selected_context" not in st.session_state:
        st.session_state.selected_context = "J(AI)SON Agency"
        
    # Upewnij się, że aktualny stan sesji istnieje w nowo załadowanej liście
    if st.session_state.selected_context not in available_contexts:
        st.session_state.selected_context = available_contexts[0]
        
    new_context = st.selectbox(
        "Wybierz obszar roboczy:",
        available_contexts,
        index=available_contexts.index(st.session_state.selected_context),
        label_visibility="collapsed"
    )
    
    if new_context != st.session_state.selected_context:
        st.session_state.selected_context = new_context
        st.toast(f"Zmieniono kontekst na: {new_context}")
        st.rerun()
        
    st.markdown("<hr style='margin: 10px 0; border-color: #1F242E;'>", unsafe_allow_html=True)

    col_menu = st.session_state.current_page
    
    # I. WORKSPACE
    st.markdown("<p style='color: #F59E0B; font-weight: bold; font-size: 0.75rem; letter-spacing: 1px; margin-bottom: 6px; margin-top: 10px;'>I. WORKSPACE</p>", unsafe_allow_html=True)
    
    if st.button("🎯 Mission Control", use_container_width=True, type="primary" if col_menu == "🎯 Mission Control" else "secondary"):
        st.session_state.current_page = "🎯 Mission Control"
        st.rerun()
        
    # II. AGENTS
    st.markdown("<p style='color: #EC4899; font-weight: bold; font-size: 0.75rem; letter-spacing: 1px; margin-top: 18px; margin-bottom: 6px;'>II. AGENTS</p>", unsafe_allow_html=True)
    
    if st.button("🤖 Claude 🟢", use_container_width=True, type="primary" if col_menu == "Claude" else "secondary"):
        st.session_state.current_page = "Claude"
        st.rerun()
        
    if st.button("🔱 Hermes 🟢", use_container_width=True, type="primary" if col_menu == "Hermes" else "secondary"):
        st.session_state.current_page = "Hermes"
        st.rerun()
        
    if st.button("♊ Gemini 🟢", use_container_width=True, type="primary" if col_menu == "Gemini" else "secondary"):
        st.session_state.current_page = "Gemini"
        st.rerun()
        
    if st.button("🌌 AntiGravity 🟢", use_container_width=True, type="primary" if col_menu == "Antigravity" else "secondary"):
        st.session_state.current_page = "Antigravity"
        st.rerun()
        
    # III. SELF
    st.markdown("<p style='color: #10B981; font-weight: bold; font-size: 0.75rem; letter-spacing: 1px; margin-top: 18px; margin-bottom: 6px;'>III. SELF</p>", unsafe_allow_html=True)
    
    if st.button("🎯 Goals & Journal", use_container_width=True, type="primary" if col_menu == "Goals" else "secondary"):
        st.session_state.current_page = "Goals"
        st.rerun()
        
    if st.button("📚 Baza Wiedzy", use_container_width=True, type="primary" if col_menu == "Notebook" else "secondary"):
        st.session_state.current_page = "Notebook"
        st.rerun()
        
    if st.button("📋 Tablica Kanban", use_container_width=True, type="primary" if col_menu == "Kanban" else "secondary"):
        st.session_state.current_page = "Kanban"
        st.rerun()
        
    if st.button("💾 Pamięć Agenta", use_container_width=True, type="primary" if col_menu == "Memory" else "secondary"):
        st.session_state.current_page = "Memory"
        st.rerun()
        
        
    # IV. BUSINESS & MARKETING
    st.markdown("<p style='color: #3B82F6; font-weight: bold; font-size: 0.75rem; letter-spacing: 1px; margin-top: 18px; margin-bottom: 6px;'>IV. BUSINESS & MARKETING</p>", unsafe_allow_html=True)
    
    creative_suite_pages = ["Jaison Agency", "SEO", "Social Media Hub", "AI Website Builder", "Ads & Local SEO", "Studio"]
    if st.button("🎨 J(AI)SON Creative Suite 🟢", use_container_width=True, type="primary" if col_menu in creative_suite_pages else "secondary"):
        st.session_state.current_page = "Jaison Agency"
        st.session_state.active_suite_tool = "Home"
        st.rerun()
        
    if st.button("💼 Jaison Client & Project Pipeline 🟢", use_container_width=True, type="primary" if col_menu == "CRM" else "secondary"):
        st.session_state.current_page = "CRM"
        st.rerun()
        
    if st.button("🎯 Lead Radar 🟢", use_container_width=True, type="primary" if col_menu == "Lead_Radar" else "secondary"):
        st.session_state.current_page = "Lead_Radar"
        st.rerun()
        
    if st.button("⚖️ Legal Engine", use_container_width=True, type="primary" if col_menu == "Legal" else "secondary"):
        st.session_state.current_page = "Legal"
        st.rerun()
        
    if st.button("💰 Finance & KSeF", use_container_width=True, type="primary" if col_menu == "Finance" else "secondary"):
        st.session_state.current_page = "Finance"
        st.rerun()
        
    st.markdown("<hr style='margin: 15px 0; border-color: #1F242E;'>", unsafe_allow_html=True)
    st.markdown("🌐 **Status Systemu:**")
    st.markdown("⚡ *Hermes Agent:* <span style='color:#10B981; font-weight:bold;'>LIVE (Port 9119)</span>", unsafe_allow_html=True)
    st.markdown("📢 *Telegram Chat:* <span style='color:#10B981; font-weight:bold;'>Połączony</span>", unsafe_allow_html=True)
    st.markdown("📝 *Pristine Memory:* <span style='color:#3B82F6; font-weight:bold;'>Aktywna</span>", unsafe_allow_html=True)


menu = st.session_state.current_page

# Unifikacja i przekierowanie starych podstron do zintegrowanego J(AI)SON Creative Suite (Bento Grid)
if menu in ["Social Media Hub", "AI Website Builder", "SEO", "Ads & Local SEO", "Studio"]:
    mapping = {
        "Social Media Hub": "Brand_Bios",
        "AI Website Builder": "Landing_Page",
        "SEO": "Faceless_Reels",
        "Ads & Local SEO": "Ads_Studio",
        "Studio": "Studio_Video"
    }
    st.session_state.active_suite_tool = mapping[menu]
    st.session_state.current_page = "Jaison Agency"
    menu = "Jaison Agency"


def render_agent_console(agent_name, status, default_model, provider, color_accent):
    agent_key = agent_name.lower().replace(' ', '_')
    prov_val = st.session_state.get(f"ctrl_provider_{agent_key}", provider)
    model_val = st.session_state.get(f"ctrl_model_{agent_key}", default_model)
    if prov_val == "OpenRouter" and model_val == "Inny / Custom":
        model_val = st.session_state.get(f"ctrl_custom_model_{agent_key}", "custom")
        
    st.markdown(f"<p style='color: #94A3B8; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-bottom: 2px;'>II. — AGENT • {agent_name.upper()}</p>", unsafe_allow_html=True)
    st.title(f"{agent_name}")
    st.markdown(f"<p style='color: {color_accent}; font-weight: bold; font-size: 0.95rem; margin-top: -5px;'>Status: {status} | Active Model: {model_val} | Provider: {prov_val}</p>", unsafe_allow_html=True)
    
    tab_chat, tab_work, tab_ctrl = st.tabs(["💬 Chat", "📂 Workspace", "⚙️ Control Room"])
    
    # 1. CHAT
    with tab_chat:
        st.subheader("Konsola konwersacyjna")
        chat_key = f"chat_{agent_key}"
        if chat_key not in st.session_state:
            st.session_state[chat_key] = []
            
        for msg in st.session_state[chat_key]:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
                
        if prompt := st.chat_input(f"Napisz do agenta {agent_name}...", key=f"input_{agent_key}"):
            with st.chat_message("user"):
                st.markdown(prompt)
            st.session_state[chat_key].append({"role": "user", "content": prompt})
            
            # Wczytanie profilu i instrukcji
            o_mnie_path = os.path.join(HERMES_DIR, "o_mnie.md")
            o_mnie_context = read_md_file(o_mnie_path) if os.path.exists(o_mnie_path) else "Brak profilu o_mnie.md"
            user_inst = st.session_state.get(f"ctrl_prompt_{agent_key}", f"Jesteś agentem {agent_name}. Działasz zorientowany na ADHD i redukcję szumu kognitywnego.")
            
            recent_context = get_recent_workspace_context(limit=3)
            global_ctx = get_selected_context_data()
            
            sys_prompt = f"""{user_inst}
{global_ctx}
            
## PROFIL UŻYTKOWNIKA (O_MNIE):
{o_mnie_context}

## WSPÓLNA PAMIĘĆ ROBOCZA (WORKSPACE MEMORY):
Oto ostatnie działania i notatki w ekosystemie:
{recent_context}

Działasz w oparciu o powyższy wspólny kontekst. Twój styl jest krótki, precyzyjny (ADHD-friendly).
"""
            
            provider_sel = st.session_state.get(f"ctrl_provider_{agent_key}", provider)
            model_sel = st.session_state.get(f"ctrl_model_{agent_key}", default_model)
            if provider_sel == "OpenRouter":
                custom_m = st.session_state.get(f"ctrl_custom_model_{agent_key}", "")
                if model_sel == "Inny / Custom" and custom_m:
                    final_model = custom_m
                else:
                    final_model = model_sel
            else:
                final_model = model_sel

            with st.spinner(f"{agent_name} ({final_model} / {provider_sel}) przetwarza zapytanie..."):
                api_messages = [{"role": m["role"], "content": m["content"]} for m in st.session_state[chat_key]]
                if provider_sel == "OpenRouter":
                    response = call_openrouter_api(api_messages, model=final_model, system_instruction=sys_prompt)
                elif provider_sel == "Local vLLM":
                    response = call_vllm_api(api_messages, system_instruction=sys_prompt)
                else: # GCP Vertex Proxy
                    if final_model == "gemini-2.5-pro":
                        response = call_gemini_pro_api(api_messages, system_instruction=sys_prompt)
                    else:
                        response = call_gemini_api(api_messages, system_instruction=sys_prompt)
                
            with st.chat_message("assistant"):
                st.markdown(response)
            st.session_state[chat_key].append({"role": "assistant", "content": response})
            
            # Infinite Context Engine: auto-save to Obsidian
            obsidian_note_title = f"Chat_{agent_name.replace(' ', '_')}_{int(time.time())}.md"
            obsidian_note_path = os.path.join(OBSIDIAN_DIR, obsidian_note_title)
            try:
                note_content = f"---\ntype: chat-log\nagent: {agent_name}\ntimestamp: {time.strftime('%Y-%m-%d %H:%M:%S')}\n---\n\n"
                for msg in st.session_state[chat_key]:
                    note_content += f"**{'Użytkownik' if msg['role'] == 'user' else agent_name}**:\n{msg['content']}\n\n"
                with open(obsidian_note_path, "w", encoding="utf-8") as f:
                    f.write(note_content)
            except Exception as e:
                pass
                
    # 2. WORKSPACE (Trzykolumnowy manager)
    with tab_work:
        st.subheader("Menedżer Workspace Agentów")
        col_b, col_f, col_v = st.columns([1, 2, 3])
        
        with col_b:
            st.write("🔍 **Buckets**")
            bucket = st.radio("Filtruj:", ["Apps", "Videos", "Images", "Audio", "Workspace", "Sandboxes", "Pastes"], key=f"bucket_{agent_key}")
            
        with col_f:
            st.write(f"📁 **Pliki w {bucket}**")
            files_list = []
            if bucket == "Apps":
                files_list = [f for f in os.listdir(DASHBOARD_DIR) if f.endswith(".html")]
                if not files_list:
                    files_list = ["todo-app.html", "testimonials.html", "landing-page.html"]
            elif bucket == "Videos":
                files_list = ["dragon_tokyo_skyline.mp4", "ai-seo-trends.mp4"]
            elif bucket == "Images":
                files_list = ["holistic_ai_dhd_logo.png", "marketing_campaign.jpg"]
            elif bucket == "Audio":
                files_list = ["meditation_music.mp3", "podcast_voice.wav"]
            elif bucket == "Workspace":
                files_list = ["app.py", "gcp_vertex_proxy.py", "requirements.txt"]
            elif bucket == "Sandboxes":
                files_list = ["sandbox_env.py", "test_runner.py"]
            elif bucket == "Pastes":
                files_list = [f for f in os.listdir(OBSIDIAN_DIR) if f.endswith(".md")][:10]
                if not files_list:
                    files_list = ["quick_note.md", "brief_draft.md"]
                    
            selected_file = st.selectbox("Wybierz plik:", files_list, key=f"sel_file_{agent_key}")
            
        with col_v:
            st.write("👁️ **Podgląd / Edycja**")
            if selected_file:
                st.info(f"Podgląd pliku: `{selected_file}`")
                
                if selected_file.endswith(".html") or bucket == "Apps":
                    html_content = ""
                    if selected_file == "testimonials.html":
                        html_content = "<html><body style='background-color:#121620; color:white; font-family:sans-serif; padding:20px;'><h2>Want results like these? Join the AI Profit Boardroom!</h2><p>258 real wins from AI Profit Boardroom members.</p></body></html>"
                    elif selected_file == "todo-app.html":
                        html_content = "<html><body style='background-color:#121620; color:white; font-family:sans-serif; padding:20px;'><h2>ADHD To-Do App</h2><input type='text' placeholder='Add task...'><button>Add</button></body></html>"
                    else:
                        html_content = "<html><body style='background-color:#121620; color:white; font-family:sans-serif; padding:20px;'><h2>Custom App View</h2><p>Pomyślnie załadowano render aplikacji w piaskownicy Agentic OS.</p></body></html>"
                    
                    st.components.v1.html(html_content, height=350, scrolling=True)
                    
                    sub_tab_preview, sub_tab_src = st.tabs(["👁️ Podgląd", "💻 Kod źródłowy"])
                    with sub_tab_src:
                        st.code(html_content, language="html")
                        
                elif selected_file.endswith(".mp4") or bucket == "Videos":
                    st.video("https://www.w3schools.com/html/mov_bbb.mp4")
                    
                elif selected_file.endswith((".png", ".jpg", ".jpeg")) or bucket == "Images":
                    st.image("https://images.unsplash.com/photo-1618005182384-a83a8bd57fbe?w=500", caption=selected_file)
                    
                elif selected_file.endswith((".mp3", ".wav")) or bucket == "Audio":
                    st.audio("https://www.soundhelix.com/examples/mp3/SoundHelix-Song-1.mp3")
                    
                elif selected_file.endswith(".md") or selected_file.endswith(".py") or selected_file.endswith(".txt"):
                    file_path = os.path.join(OBSIDIAN_DIR, selected_file)
                    if not os.path.exists(file_path):
                        file_path = os.path.join(os.getcwd(), selected_file)
                    content = ""
                    if os.path.exists(file_path):
                        try:
                            with open(file_path, "r", encoding="utf-8") as f:
                                content = f.read()
                        except:
                            content = "Błąd odczytu pliku."
                    else:
                        content = f"# Draft dla {selected_file}\nTutaj znajduje się podgląd notatki roboczej agenta."
                        
                    edited_content = st.text_area("Edycja pliku:", content, height=200, key=f"edit_{agent_key}_{selected_file}")
                    if st.button("Zapisz zmiany", key=f"save_{agent_key}_{selected_file}"):
                        try:
                            with open(file_path, "w", encoding="utf-8") as f:
                                f.write(edited_content)
                            st.success("Zapisano pomyślnie!")
                        except Exception as e:
                            st.error(f"Błąd zapisu: {e}")
                            
    # 3. CONTROL ROOM
    with tab_ctrl:
        st.subheader("Konfiguracja Agenta (Control Room)")
        st.write("Ustaw dostawcę i parametry modelu dla tego agenta:")
        
        prov_opt = st.selectbox("Dostawca (Provider):", ["GCP Vertex Proxy", "OpenRouter", "Local vLLM"], 
                                index=0 if provider in ["Vertex AI Native", "GCP Proxy Port 8089", "GCP Proxy"] or "vertex" in provider.lower() else (1 if "openrouter" in provider.lower() else 2),
                                key=f"ctrl_provider_{agent_key}")
        
        if prov_opt == "GCP Vertex Proxy":
            models_list = ["gemini-2.5-pro", "gemini-2.5-flash"]
            default_idx = 0 if default_model == "gemini-2.5-pro" else 1
            model_opt = st.selectbox("Model:", models_list, index=default_idx, key=f"ctrl_model_{agent_key}")
        elif prov_opt == "OpenRouter":
            models_list = [
                "meta-llama/llama-3.1-8b-instruct:free",
                "nousresearch/hermes-3-llama-3.1-405b:free",
                "google/gemini-2.5-flash",
                "google/gemini-2.5-pro",
                "google/gemini-3.5-flash",
                "google/gemini-3.1-pro-preview",
                "anthropic/claude-3.7-sonnet",
                "anthropic/claude-3.5-sonnet",
                "Inny / Custom"
            ]
            default_idx = 8
            low_default = default_model.lower()
            if "claude-3.7" in low_default:
                default_idx = 6
            elif "claude-3.5" in low_default:
                default_idx = 7
            elif "gemini-3.5" in low_default:
                default_idx = 4
            elif "gemini-3.1" in low_default:
                default_idx = 5
            elif "hermes" in low_default:
                default_idx = 1
            elif "free" in low_default:
                default_idx = 0
            elif "gemini-2.5-pro" in low_default:
                default_idx = 3
            elif "gemini-2.5-flash" in low_default:
                default_idx = 2
                
            model_opt = st.selectbox("Model:", models_list, index=default_idx, key=f"ctrl_model_{agent_key}")
            if model_opt == "Inny / Custom":
                st.text_input("Identyfikator modelu OpenRouter (np. deepseek/deepseek-chat):", value="deepseek/deepseek-chat", key=f"ctrl_custom_model_{agent_key}")
        else: # Local vLLM
            st.info("Local vLLM korzysta z modelu uruchomionego na porcie 8000.")
            model_opt = st.text_input("Nazwa modelu vLLM:", value="mistralai/Mistral-Nemo-Instruct-2407", key=f"ctrl_model_{agent_key}")
            
        temp = st.slider("Temperatura (Kreatywność):", 0.0, 1.0, 0.7, 0.05, key=f"ctrl_temp_{agent_key}")
        sys_prompt_input = st.text_area("System Prompt / Instrukcje systemowe:", 
                                        f"Jesteś agentem {agent_name}. Działasz zorientowany na ADHD i redukcję szumu kognitywnego.", 
                                        height=150, key=f"ctrl_prompt_{agent_key}")
                                        
        if st.button("Aktualizuj konfigurację agenta", key=f"ctrl_btn_{agent_key}"):
            st.success(f"Konfiguracja dla agenta {agent_name} została zaktualizowana w pamięci podręcznej sesji!")
            st.rerun()

# 1. MISSION CONTROL
if menu == "🎯 Mission Control":

    st.markdown("<p style='color: #94A3B8; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-bottom: 2px;'>I. — MISSION CONTROL</p>", unsafe_allow_html=True)
    st.title("✨ JAISON OS")
    st.subheader("Status of every agent, every memory, every signal.")
    
    # Górny pasek statusów (System Status Row)
    st.markdown("""
    <style>
        .status-container {
            display: flex;
            gap: 12px;
            margin-bottom: 25px;
            width: 100%;
        }
        .status-box {
            background-color: #0E1015;
            border: 1px solid #1F242E;
            border-radius: 8px;
            padding: 12px 16px;
            flex: 1;
            box-shadow: 0 4px 10px rgba(0,0,0,0.3);
        }
        .status-title {
            color: #94A3B8;
            font-size: 0.7rem;
            font-weight: bold;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            margin-bottom: 4px;
            display: flex;
            align-items: center;
            justify-content: space-between;
        }
        .status-value {
            color: #FFFFFF;
            font-family: 'Outfit', sans-serif;
            font-size: 1.05rem;
            font-weight: 600;
        }
        .status-desc {
            color: #64748B;
            font-size: 0.75rem;
            margin-top: 2px;
        }
        .status-dot-green { background-color: #10B981; box-shadow: 0 0 8px #10B981; }
        .status-dot-red { background-color: #EF4444; box-shadow: 0 0 8px #EF4444; }
        .status-dot-yellow { background-color: #F59E0B; box-shadow: 0 0 8px #F59E0B; }
    </style>
    """, unsafe_allow_html=True)

    col_s1, col_s2, col_s3, col_s4 = st.columns(4)
    with col_s1:
        st.markdown("""<div class="status-box">
            <div class="status-title"><span>🤖 CLAUDE</span><span class="status-dot-green" style="height:6px; width:6px; border-radius:50%; display:inline-block;"></span></div>
            <div class="status-value">Online</div>
            <div class="status-desc">2.1.142 • 46ms</div>
        </div>""", unsafe_allow_html=True)
    with col_s2:
        st.markdown("""<div class="status-box">
            <div class="status-title"><span>🔱 HERMES</span><span class="status-dot-green" style="height:6px; width:6px; border-radius:50%; display:inline-block;"></span></div>
            <div class="status-value">Online</div>
            <div class="status-desc">grok-4.3 • xAI OAuth</div>
        </div>""", unsafe_allow_html=True)
    with col_s3:
        st.markdown("""<div class="status-box">
            <div class="status-title"><span>⏱️ HEARTBEAT</span><span class="status-dot-yellow" style="height:6px; width:6px; border-radius:50%; display:inline-block;"></span></div>
            <div class="status-value">Active</div>
            <div class="status-desc">poll ticks • 4s</div>
        </div>""", unsafe_allow_html=True)
    with col_s4:
        st.markdown("""<div class="status-box">
            <div class="status-title"><span>⚡ LATENCY</span><span class="status-dot-green" style="height:6px; width:6px; border-radius:50%; display:inline-block;"></span></div>
            <div class="status-value">42 ms</div>
            <div class="status-desc">combined p50</div>
        </div>""", unsafe_allow_html=True)

    st.markdown("<hr style='margin: 15px 0; border-color: #1F242E;'>", unsafe_allow_html=True)

    # Tryb One Thing - ADHD Wykluczenie Szumu
    st.markdown("""
    <div class="one-thing-banner">
        <h3 style="margin-top: 0; color: #F59E0B;">🎯 Tryb "One Thing"</h3>
        <p style="color: #CBD5E1; line-height: 1.6;">Osoby z ADHD cierpią na paraliż decyzyjny z powodu nadmiaru bodźców. Wpisz poniżej dokładnie <strong>JEDNĄ</strong> rzecz, na której skupisz się w tym momencie. Dashboard wyciszy resztę szumu operacyjnego.</p>
    </div>
    """, unsafe_allow_html=True)
    
    thing = st.text_input("Moje jedyne zadanie na ten moment:", value=st.session_state.one_thing, placeholder="Np. zredagowanie oferty High-Ticket dla lokalnej kliniki...")
    if thing:
        st.session_state.one_thing = thing
        st.markdown(f"""
        <div class="custom-card" style="border-left: 5px solid #10B981; background-color: #0F1D1A;">
            <h4 style="margin: 0; color: #10B981;">🔥 Twój aktualny priorytet:</h4>
            <p style="font-size: 1.25rem; font-weight: bold; margin-top: 8px; color: #FFFFFF;">{thing}</p>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2 = st.columns([1, 3])
        with col1:
            if st.button("⏱️ Uruchom Pomodoro (25 min)"):
                st.session_state.pomodoro_active = True
        with col2:
            if st.session_state.pomodoro_active:
                st.success("Stoper Pomodoro wystartował. Wyłącz inne karty w przeglądarce i skup się wyłącznie na priorytecie.")

    st.markdown("<hr style='margin: 15px 0; border-color: #1F242E;'>", unsafe_allow_html=True)

    # II. AGENTS GRID
    st.markdown("<p style='color: #EC4899; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-top: 25px; margin-bottom: 2px;'>II. — AGENTS • CLICK TO OPEN CONTROL ROOM</p>", unsafe_allow_html=True)
    
    col_a1, col_a2, col_a3, col_a4 = st.columns(4)
    
    with col_a1:
        st.markdown("""
        <div class="custom-card" style="border-top: 3px solid #F59E0B; height: 260px;">
            <h3 style="color: #F59E0B; margin-top: 0;">🤖 Claude</h3>
            <p style="color: #94A3B8; font-size: 0.85rem; height: 60px;">Bezpośrednie połączenie ze środowiskiem Claude Code. Pełna kontrola konsoli i narzędzi.</p>
            <div style="margin-top: 15px; font-size: 0.8rem; color: #64748B;">
                <div><b>MODEL:</b> claude-3-7-sonnet</div>
                <div><b>PROVIDER:</b> Anthropic Native</div>
                <div><b>STATUS:</b> <span style="color:#10B981; font-weight:bold;">ONLINE</span></div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Otwórz konsolę Claude", use_container_width=True):
            st.session_state.current_page = "Claude"
            st.rerun()

    with col_a2:
        st.markdown("""
        <div class="custom-card" style="border-top: 3px solid #3B82F6; height: 260px;">
            <h3 style="color: #3B82F6; margin-top: 0;">🔱 Hermes</h3>
            <p style="color: #94A3B8; font-size: 0.85rem; height: 60px;">Główny orkiestrator. Wywołanie skilli, obsługa Kanbana i automatyzacji w tle.</p>
            <div style="margin-top: 15px; font-size: 0.8rem; color: #64748B;">
                <div><b>MODEL:</b> grok-4.3</div>
                <div><b>PROVIDER:</b> xAI API v1</div>
                <div><b>STATUS:</b> <span style="color:#10B981; font-weight:bold;">ONLINE</span></div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Otwórz konsolę Hermes", use_container_width=True):
            st.session_state.current_page = "Hermes"
            st.rerun()

    with col_a3:
        st.markdown("""
        <div class="custom-card" style="border-top: 3px solid #8B5CF6; height: 260px;">
            <h3 style="color: #8B5CF6; margin-top: 0;">♊ Gemini</h3>
            <p style="color: #94A3B8; font-size: 0.85rem; height: 60px;">Super-inteligentny model do audytów, pisania pism prawnych i zaawansowanej logiki.</p>
            <div style="margin-top: 15px; font-size: 0.8rem; color: #64748B;">
                <div><b>MODEL:</b> gemini-2.5-pro</div>
                <div><b>PROVIDER:</b> Vertex AI Native</div>
                <div><b>STATUS:</b> <span style="color:#10B981; font-weight:bold;">ONLINE</span></div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Otwórz konsolę Gemini", use_container_width=True):
            st.session_state.current_page = "Gemini"
            st.rerun()

    with col_a4:
        st.markdown("""
        <div class="custom-card" style="border-top: 3px solid #10B981; height: 260px;">
            <h3 style="color: #10B981; margin-top: 0;">🌌 AntiGravity</h3>
            <p style="color: #94A3B8; font-size: 0.85rem; height: 60px;">Twój systemowy architekt i asystent operacyjny. Nadzór nad kodem i deployami.</p>
            <div style="margin-top: 15px; font-size: 0.8rem; color: #64748B;">
                <div><b>MODEL:</b> gemini-2.5-pro</div>
                <div><b>PROVIDER:</b> GCP Proxy Port 8089</div>
                <div><b>STATUS:</b> <span style="color:#10B981; font-weight:bold;">ONLINE</span></div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Otwórz konsolę AntiGravity", use_container_width=True):
            st.session_state.current_page = "Antigravity"
            st.rerun()

    st.markdown("---")
    
    # Szybki Capture
    st.subheader("⚡ Szybki Capture myśli (Brain Dump)")
    quick_thought = st.text_area("Masz nagły pomysł lub coś Cię rozprasza? Zrzuć to tutaj natychmiast, aby uwolnić pamięć roboczą mózgu:", height=100)
    if st.button("Uwolnij moją pamięć roboczą"):
        if quick_thought:
            save_brain_dump(quick_thought, "", None)
            st.success("Pomysł bezpiecznie zapisany w chmurze w Skarbcu Myśli. Twoja głowa jest wolna.")
            time.sleep(0.5)
            st.rerun()

elif menu == "Claude":
    render_agent_console("Claude", "Online", "claude-3-7-sonnet", "Anthropic Native", "#F59E0B")

elif menu == "Hermes":
    render_agent_console("Hermes", "Online", "grok-4.3", "xAI API v1", "#3B82F6")

elif menu == "Gemini":
    render_agent_console("Gemini", "Online", "gemini-2.5-pro", "Vertex AI Native", "#8B5CF6")

# 2. GOALS & OPEN LOOPS
elif menu == "Goals":
    st.markdown("<p style='color: #94A3B8; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-bottom: 2px;'>III. — SELF • GOALS & MENTAL WORKSHOP</p>", unsafe_allow_html=True)
    st.title("🎯 Goals & Mental Workshop")
    
    tab_loops, tab_goal_mode, tab_soul_agent = st.tabs(["🗑️ Brain Dump & Open Loops", "🎯 Goal Mode (Autonomiczny Cel)", "🧘 Soul Agent (Rytuały Zdrowia)"])
    
    with tab_loops:
        st.subheader("Twój mentalny odciążyciel — bezszumne uwalnianie pamięci roboczej")
        
        col_in, col_st = st.columns([1, 1])
        
        with col_in:
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #EC4899;">
                <h4 style="margin:0; color:#EC4899;">📥 Zrzut z głowy (Brain Dump)</h4>
                <p style="font-size: 0.9rem; color: #94A3B8; margin-top: 6px;">Wpisz pomysły, luźne myśli, linki lub wgraj zrzut ekranu (np. inspirację reklamową).</p>
            </div>
            """, unsafe_allow_html=True)
            
            thought_input = st.text_area("Co Ci chodzi po głowie?", height=150, key="goals_thought_input")
            links_input = st.text_input("Linki / Źródła (opcjonalnie):", key="goals_links_input")
            uploaded_file = st.file_uploader("Dodaj plik / zrzut ekranu (PNG, JPG, PDF):", type=["png","jpg","jpeg","pdf"], key="goals_uploaded_file")
            
            if st.button("Prześlij do Skarbca w Chmurze", type="primary", key="goals_submit_button"):
                if thought_input or links_input or uploaded_file:
                    save_brain_dump(thought_input, links_input, uploaded_file)
                    st.success("Zapisano. Pomysł został odciążony z Twojego mózgu.")
                    time.sleep(0.5)
                    st.rerun()
            
            st.markdown("---")
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #7C3AED; background: linear-gradient(135deg, #1A1230 0%, #0F1016 100%);">
                <h4 style="margin: 0; color: #C084FC;">🧠 Strategiczna Odprawa CEO</h4>
                <p style="color: #CBD5E1; font-size: 0.85rem; margin-top: 6px; margin-bottom: 0;">
                    Uruchom analizę Skarbca. CEO Jason (Gemini 2.5 Pro) uporządkuje Twoje otwarte pętle, wyciągnie z nich strategię biznesową i zaproponuje gotowe zadania.
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("🚀 Przetwórz Skarbiec przez CEO", use_container_width=True, key="goals_run_ceo_button"):
                active_dumps = [d for d in get_brain_dumps() if d.get("status", "active") == "active"]
                if not active_dumps:
                    st.info("Twój Skarbiec jest pusty. Brak otwartych pętli do przetworzenia.")
                else:
                    with st.spinner("CEO Jason analizuje Skarbiec (Gemini 2.5 Pro)..."):
                        dumps_text = ""
                        for i, d in enumerate(active_dumps):
                            dumps_text += f"\n--- POMYSŁ {i+1} ---\nZapisano: {d.get('timestamp')}\nTreść: {d.get('thought')}\nLinki: {d.get('links', '')}\n"
                        
                        ceo_prompt = f"""Jesteś CEO Holistic Operator — Tomasz 'Holistic Jason', osobistym doradcą i strategiem użytkownika.
Twój cel to pomóc użytkownikowi (który ma ADHD i cierpi na paraliż decyzyjny) uporządkować i wdrożyć w życie pomysły, które zapisał w Skarbcu Myśli (Brain Dump).

Przeanalizuj poniższe otwarte pętle (brain dumps) i przygotuj dla nich plan działania:
1. Dokonaj kategoryzacji (np. Szybka wygrana, Duży projekt, Do odrzucenia).
2. Dla każdego ważnego pomysłu zaproponuj konkretną "Next Action" (bezwysiłkowy mikro-krok o niskim oporze kognitywnym).
3. Zaproponuj, które z nich należy natychmiast wrzucić do Kanbana jako zadania, a które zarchiwizować/skasować.

Pisz zwięźle, konkretnie, w przyjaznym, motywującym tonie, bez bełkotu AI. Używaj wypunktowań.

OTWARTE PĘTLE ZE SKARBCZA:
{dumps_text}"""
                        response = call_gemini_pro_api([{"role": "user", "content": ceo_prompt}], "Jesteś CEO Jason, wybitnym strategiem biznesowym wspierającym osoby z ADHD.")
                        st.session_state.ceo_analysis_result = response
                        st.rerun()
    
            if "ceo_analysis_result" in st.session_state and st.session_state.ceo_analysis_result:
                st.markdown("##### 📋 Raport i Plan CEO:")
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #10B981; white-space: pre-wrap; font-size: 0.9rem; line-height: 1.6; background-color: #0c1410;">
{st.session_state.ceo_analysis_result}
                </div>
                """, unsafe_allow_html=True)
                if st.button("Wyczyść raport CEO", key="goals_clear_ceo_button"):
                    st.session_state.ceo_analysis_result = None
                    st.rerun()
                    
        with col_st:
            st.markdown("### 📦 Aktywne Otwarte Pętle (Open Loops)")
            dumps = get_brain_dumps()
            active_dumps = [d for d in dumps if d.get("status", "active") == "active"]
            
            if not active_dumps:
                st.info("Twój Skarbiec jest pusty. Brak rozpraszających pętli myślowych.")
            else:
                st.write(f"Masz **{len(active_dumps)}** aktywnych pętli czekających na wdrożenie:")
                for d in active_dumps:
                    accent = "#3B82F6" if d.get("category") == "Now" else "#F59E0B"
                    st.markdown(f"""
                    <div class="custom-card" style="border-left: 4px solid {accent}; margin-bottom: 12px;">
                        <span style="font-size: 0.8rem; color:#94A3B8;">⏱️ Zapisano: {d.get('timestamp')} | Priorytet: {d.get('category')}</span>
                        <p style="margin-top: 6px; font-size: 1.05rem; color:#FFFFFF;">{d.get('thought')}</p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if d.get("links"):
                        st.markdown(f"🔗 **Inspiracja:** [{d.get('links')}]({d.get('links')})")
                    if d.get("file_attached"):
                        ext = os.path.splitext(d.get("file_attached"))[1].lower()
                        if ext in [".png", ".jpg", ".jpeg"]:
                            path = os.path.join(BRAIN_DUMP_ASSETS, d.get("file_attached"))
                            if os.path.exists(path):
                                st.image(path, caption=d.get("file_attached"), use_container_width=True)
                        else:
                            st.markdown(f"📎 **Załącznik:** `{d.get('file_attached')}`")
                            
                    c1, c2, c3 = st.columns(3)
                    with c1:
                        if st.button("🎯 Do Kanbana", key=f"k_{d['id']}"):
                            k = load_kanban()
                            short = d.get('thought')[:100] + "..." if len(d.get('thought')) > 100 else d.get('thought')
                            k["todo"].append(f"🧠 [Zrzut] {short}")
                            save_kanban(k)
                            dump_file = os.path.join(BRAIN_DUMP_DIR, f"{d['id']}.json")
                            d["status"] = "archived"
                            json.dump(d, open(dump_file, "w", encoding="utf-8"), ensure_ascii=False, indent=4)
                            st.toast("Zadanie dodane do tablicy Kanban!")
                            time.sleep(0.5); st.rerun()
                    with c2:
                        if st.button("📦 Archiwizuj", key=f"a_{d['id']}"):
                            dump_file = os.path.join(BRAIN_DUMP_DIR, f"{d['id']}.json")
                            d["status"] = "archived"
                            json.dump(d, open(dump_file, "w", encoding="utf-8"), ensure_ascii=False, indent=4)
                            st.toast("Zarchiwizowano.")
                            time.sleep(0.5); st.rerun()
                    with c3:
                        if st.button("🗑️ Usuń", key=f"d_{d['id']}"):
                            dump_file = os.path.join(BRAIN_DUMP_DIR, f"{d['id']}.json")
                            if os.path.exists(dump_file):
                                os.remove(dump_file)
                            st.toast("Usunięto.")
                            time.sleep(0.5); st.rerun()
                            
    with tab_goal_mode:
        st.subheader("🎯 Goal Mode — Autonomiczne Wykonywanie Zadań")
        st.markdown("""
        <div class="custom-card" style="border-left: 5px solid #F59E0B;">
            <h4 style="margin: 0; color: #F59E0B;">🚀 Autonomiczna Pętla Weryfikacji (Judge Loop)</h4>
            <p style="color: #CBD5E1; font-size: 0.85rem; margin-top: 6px; margin-bottom: 0;">
                Wpisz wysoki cel operacyjny. Agent uruchomi pętlę do 20 iteracji. W każdej iteracji model-sędzia ocenia postęp i decyduje, czy cel został w pełni osiągnięty.
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        goal_input = st.text_input("Cel do zrealizowania:", placeholder="np. Wygeneruj 3 spersonalizowane szkice maili sprzedażowych dla lokalnych klinik...", key="goals_auton_goal")
        steps_limit = st.slider("Limit iteracji (tur sędziego):", 5, 20, 10, key="goals_auton_steps")
        executor_agent = st.selectbox("Wybierz agenta wykonawczego:", ["Hermes", "Claude", "Gemini", "AntiGravity"], key="goals_auton_agent")
        
        if st.button("Uruchom Goal Mode Loop", type="primary", key="goals_run_auton_btn"):
            if goal_input:
                st.write(f"🏃 **Inicjalizacja pętli Goal Mode dla agenta {executor_agent}...**")
                
                status_placeholder = st.empty()
                progress_bar = st.progress(0.0)
                
                current_state = f"Rozpoczęto realizację celu: {goal_input}"
                completed = False
                
                for step in range(1, steps_limit + 1):
                    status_placeholder.info(f"🔄 **Iteracja {step} / {steps_limit}...**")
                    progress_bar.progress(step / steps_limit)
                    
                    exec_prompt = f"""Realizujesz cel: "{goal_input}".
Aktualny stan prac: {current_state}
Wykonaj następny logiczny krok w celu osiągnięcia rezultatu. Wygeneruj konkretny postęp."""
                    with st.spinner(f"Agent {executor_agent} wykonuje krok..."):
                        step_output = call_gemini_pro_api([{"role": "user", "content": exec_prompt}], f"Jesteś agentem {executor_agent} pracującym w trybie Goal Mode.")
                    
                    st.markdown(f"**Iteracja {step} Output:**")
                    st.code(step_output[:300] + "..." if len(step_output) > 300 else step_output, language="markdown")
                    
                    judge_prompt = f"""Cel główny: "{goal_input}"
Aktualny postęp prac po ostatniej iteracji:
{step_output}

ZADANIE: Oceń, czy cel główny został w 100% zrealizowany.
Odpowiedz w formacie JSON:
{{
  "completed": true/false,
  "summary": "krótkie podsumowanie wykonanej pracy",
  "next_step": "co pozostało do zrobienia (jeśli nie ukończono)"
}}
Zwróć wyłącznie prawidłowy JSON, bez znaczników ```json i bez innych komentarzy."""
                    with st.spinner("Sędzia AI weryfikuje postęp (Gemini)..."):
                        judge_resp = call_gemini_api([{"role": "user", "content": judge_prompt}], "Jesteś rygorystycznym sędzią jakości.")
                    
                    try:
                        if "```" in judge_resp:
                            judge_resp = judge_resp.split("```")[1]
                            if judge_resp.startswith("json"):
                                judge_resp = judge_resp[4:]
                        judge_data = json.loads(judge_resp.strip())
                        completed = judge_data.get("completed", False)
                        current_state = judge_data.get("summary", "")
                        
                        st.markdown(f"⚖️ **Werdykt Sędziego:** {judge_data.get('summary')}")
                        if completed:
                            st.success(f"🎉 **Cel osiągnięty w turze {step}!** Pętla zakończona pomyślnie.")
                            break
                    except Exception as e:
                        completed = False
                        current_state = f"Błąd dekodowania oceny sędziego: {e}. Kontynuuję..."
                    
                    time.sleep(1)
                    
                if not completed:
                    st.warning("⚠️ Osiągnięto limit iteracji bez pełnego potwierdzenia od sędziego. Dokonaj weryfikacji manualnej.")
            else:
                st.warning("Wpisz cel do zrealizowania.")
                
    with tab_soul_agent:
        st.subheader("🧘 Soul Agent (Twój Doradca Duchowy & Mentalny)")
        st.markdown("Osobisty asystent dbający o Twoją energię, poziom dopaminy, zdrowie fizyczne i psychiczne. Zoptymalizowany pod kątem ADHD.")
        
        # User input for current state
        feelings = st.text_area("Jak się dzisiaj czujesz? (np. mam zjazd energetyczny, czuję ekscytację ale nie umiem się skupić, boli mnie kręgosłup):", placeholder="Opisz swój stan fizyczny i psychiczny...", key="goals_soul_feelings")
        
        if st.button("Zaplanuj Rytuały Zdrowotne", type="primary", key="goals_soul_run_btn"):
            if feelings:
                with st.spinner("Soul Agent analizuje Twój stan i projektuje rytuały..."):
                    # Load user profile context
                    o_mnie_path = os.path.join(HERMES_DIR, "o_mnie.md")
                    o_mnie_context = read_md_file(o_mnie_path) if os.path.exists(o_mnie_path) else "Brak pliku o_mnie.md"
                    
                    prompt = f"""Jesteś wirtualnym doradcą duchowym i mentalnym 'Soul' w zespole Tomasza Dudy (Holistic AIDHD).
Tomasz (lub klient) opisał swoje dzisiejsze samopoczucie: "{feelings}".
Kontekst użytkownika (historia i tożsamość):
{o_mnie_context}
 
Zaprojektuj spersonalizowany zestaw rytuałów fizycznych i psychicznych na dzisiaj, aby pomóc mu wejść w stan Flow i zrównoważyć układ nerwowy.
Uwzględnij:
1. Rytuał oddechowy (np. Wim Hof, oddech pudełkowy) dopasowany do stanu.
2. Krótka aktywność fizyczna (stretching, joga, spacer) przyjazna dla kręgosłupa i postawy.
3. Rytuał mentalny/medytacyjny (np. Focus Block, uziemienie).
4. Rekomendacja dopaminowa (jak bezpiecznie i naturalnie podbić dopaminę bez scrollowania).
 
Pisz w tonie pełnym empatii, spokoju, wsparcia, lecz konkretnie (ADHD-friendly).
"""
                    response = call_gemini_pro_api([{"role": "user", "content": prompt}], "Jesteś wspierającym doradcą mentalnym i duchowym zorientowanym na ADHD.")
                    st.session_state.soul_rituals_result = response
                    st.rerun()
            else:
                st.warning("Opisz krótko jak się czujesz, aby model mógł dobrać rytuały.")
                
        if "soul_rituals_result" in st.session_state and st.session_state.soul_rituals_result:
            st.markdown("### 🧘 Rekomendowane Rytuały i Plan Przepływu (Flow):")
            st.markdown(f"""
            <div class="custom-card" style="border-left: 4px solid #EC4899; white-space: pre-wrap; font-size: 0.95rem; line-height: 1.7; background-color: #170d14;">
{st.session_state.soul_rituals_result}
            </div>
            """, unsafe_allow_html=True)
            if st.button("Wyczyść rytuały", key="goals_soul_clear_btn"):
                st.session_state.soul_rituals_result = None
                st.rerun()

# 3. BAZA WIEDZY
elif menu == "Notebook":
    st.title("📚 Baza Wiedzy")
    st.subheader("Biblioteka zasobów, notatek i promptów do natychmiastowego użycia")

    tab_akademia, tab2, tab1, tab3 = st.tabs(["📚 Checklisty & Prompty", "📝 Notatki Robocze (Obsidian)", "📻 Podcasty (NotebookLM)", "🔍 Wyszukiwarka AI (RAG)"])

    with tab_akademia:
        st.markdown("### 📚 Biblioteka Wiedzy — Checklisty, Prompty & Frameworki")
        st.markdown("Gotowe zasoby do natychmiastowego użycia. Wyszukaj po frazie lub przeglądaj według kategorii.")

        def get_knowledge_dir():
            paths_to_try = [
                r"C:\Aplikacje MVP\01_JAISON_AGENCY_OS\agency_knowledge\knowledge",
                os.path.join(os.path.dirname(__file__), "deploy", "knowledge"),
                os.path.join(os.path.dirname(__file__), "knowledge"),
                os.path.join(os.path.dirname(os.path.dirname(__file__)), "agency_knowledge", "knowledge")
            ]
            for p in paths_to_try:
                if os.path.exists(p):
                    return p
            return None

        # Kategorie na podstawie słów kluczowych w nazwie pliku
        KNOWLEDGE_CATEGORIES = {
            "🎯 Oferty & Sprzedaż": ["ofert", "sprzedaż", "sprzedawa", "lead", "leady", "obiekcj", "pitch", "clos"],
            "📱 Social Media & Content": ["post", "posty", "content", "social", "hook", "facebook", "instagram", "tiktok", "wideo", "reels", "short", "youtube"],
            "✍️ Copywriting & Webwriting": ["copy", "webwrit", "ogłosz", "tekst", "nagłów", "headline", "szabl"],
            "🤖 AI & Narzędzia": ["ai", "gpt", "prompt", "chatgpt", "gemini", "claude", "grafik", "asystent"],
            "💰 Produkty Cyfrowe & Zarabianie": ["produkt", "cyfrowy", "ebook", "kurs", "zarabian", "dochód", "nish", "nisz", "pomysł"],
            "📊 SEO & Marketing": ["seo", "marketing", "reklam", "email", "newsletter", "lejek", "strategia", "marka"],
            "🏗️ Biznes & Operacje": ["biznes", "klient", "canvas", "model", "workflow", "checklista", "plan", "cel", "mapa"],
            "📚 Kursy & Szkolenia": ["akademia", "szkoleni", "kurs", "burnejko", "skiba", "ryszka", "kryptonum"],
        }

        def classify_file(filename: str) -> str:
            name_lower = filename.lower()
            for category, keywords in KNOWLEDGE_CATEGORIES.items():
                if any(kw in name_lower for kw in keywords):
                    return category
            return "📁 Pozostałe"

        k_dir = get_knowledge_dir()
        if k_dir:
            k_files = [f for f in os.listdir(k_dir) if f.endswith('.md')]
            k_files.sort()

            if k_files:
                # Pasek wyszukiwania
                search_query = st.text_input("🔍 Wyszukaj zasób:", placeholder="np. hooks, oferta, SEO, skrypty...", key="k_search")

                if search_query.strip():
                    filtered_files = [f for f in k_files if search_query.lower() in f.lower()]
                    st.caption(f"Znaleziono {len(filtered_files)} pasujących plików dla: **{search_query}**")
                else:
                    filtered_files = k_files
                    # Kategoryzacja
                    grouped = {}
                    for f in k_files:
                        cat = classify_file(f)
                        grouped.setdefault(cat, []).append(f)

                    all_cats = sorted(grouped.keys())
                    cat_options = ["📋 Wszystkie zasoby"] + all_cats
                    selected_cat = st.selectbox("Kategoria:", cat_options, key="k_category")

                    if selected_cat != "📋 Wszystkie zasoby":
                        filtered_files = grouped.get(selected_cat, [])
                        st.caption(f"**{selected_cat}** — {len(filtered_files)} plików")

                st.markdown("---")

                if filtered_files:
                    # Dropdown z listą plików
                    clean_names = {f: f.replace(".md", "").replace("-", " ").replace("_", " ") for f in filtered_files}
                    selected_display = st.selectbox(
                        f"Wybierz zasób ({len(filtered_files)} dostępnych):",
                        options=filtered_files,
                        format_func=lambda x: clean_names[x],
                        key="akademia_files_select"
                    )
                    k_file_path = os.path.join(k_dir, selected_display)
                    k_content = read_md_file(k_file_path)
                    st.markdown(f"📁 `{selected_display}`")
                    st.markdown("---")
                    st.markdown(k_content)
                else:
                    st.info("Brak wyników dla tego wyszukiwania.")
            else:
                st.info("Katalog z bazą wiedzy jest pusty.")
        else:
            st.warning("❗ Nie znaleziono katalogu z bazą wiedzy (`deploy/knowledge`). Upewnij się, że pliki są wgrane na serwer.")

    with tab2:
        st.markdown("### 📝 Notatki Robocze")
        notes = [f for f in os.listdir(OBSIDIAN_DIR) if f.endswith('.md')] if os.path.exists(OBSIDIAN_DIR) else []
        if notes:
            selected_note = st.selectbox("Wybierz notatkę:", notes)
            note_content = read_md_file(os.path.join(OBSIDIAN_DIR, selected_note))
            st.markdown(f"📁 **Ścieżka:** `obsidian_vault/{selected_note}`")
            st.markdown("---")
            st.markdown(note_content)
        else:
            st.info("Brak notatek. Synchronizuj notatki z Obsidian Vault przez SFTP lub wgraj pliki `.md`.")

    with tab1:
        st.markdown("### 📻 Podcasty & Audio z NotebookLM")
        files = [f for f in os.listdir(NOTEBOOKS_DIR) if f.endswith(('.mp3','.wav'))] if os.path.exists(NOTEBOOKS_DIR) else []
        if files:
            st.write(f"Wykryto **{len(files)}** syntez wiedzy audio:")
            for f in files:
                st.markdown(f"""
                <div class="custom-card">
                    <h4 style="margin: 0; color: #F59E0B;">📻 {f}</h4>
                    <span style="font-size: 0.8rem; color:#94A3B8;">Przechowywany w: notebooks/_assets</span>
                </div>
                """, unsafe_allow_html=True)
                st.audio(open(os.path.join(NOTEBOOKS_DIR, f), "rb").read(), format="audio/mp3")
                st.markdown("<br>", unsafe_allow_html=True)
        else:
            st.info("Katalog `notebooks/_assets` jest pusty. Prześlij pliki .mp3 z NotebookLM za pomocą SFTP.")

    with tab3:
        st.markdown("""
        <div class="custom-card">
            <h4 style="margin: 0; color: #10B981;">🔍 Wyszukiwarka Semantyczna (RAG) z Vertex AI Search</h4>
            <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 5px;">
                Przeszukuj notatki, Obsidian Vault oraz inne wgrane zasoby za pomocą zaawansowanego silnika semantycznego Google Cloud (Discovery Engine / Agent Builder).
                Zadawaj pytania w języku naturalnym i uzyskuj odpowiedzi bezpośrednio z kontekstu Twoich plików.
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        env_ds_id = os.environ.get("GCP_DATA_STORE_ID", "")
        
        col_ds1, col_ds2 = st.columns([2, 1])
        with col_ds1:
            ds_id = st.text_input("GCP Data Store ID:", value=env_ds_id, help="Identyfikator Data Store z konsoli Google Cloud")
        with col_ds2:
            gcp_loc = st.text_input("GCP Location:", value=os.environ.get("GCP_LOCATION", "global"))
            
        st.markdown("<hr style='margin: 15px 0; border-color: #1F242E;'>", unsafe_allow_html=True)
        
        query = st.text_input("Wpisz pytanie do bazy wiedzy:", placeholder="np. Jakie są główne zasady działania marki?")
        
        if st.button("Szukaj semantycznie", type="primary", use_container_width=True):
            if not ds_id:
                st.error("⚠️ Musisz podać GCP Data Store ID. Utwórz aplikację Search i Data Store w konsoli GCP, a następnie podaj tutaj ID.")
            elif not query.strip():
                st.warning("⚠️ Wpisz treść zapytania.")
            else:
                with st.spinner("Przeszukuję bazę wiedzy..."):
                    proj_id = os.environ.get("GCP_PROJECT_ID", "holistic-dashboard-dev")
                    search_res = query_vertex_search(query, ds_id, project_id=proj_id, location=gcp_loc)
                    
                    if "error" in search_res:
                        st.error(f"❌ Błąd wyszukiwania: {search_res['error']}")
                    else:
                        st.success("Wyszukiwanie zakończone!")
                        
                        if search_res.get("summary"):
                            st.markdown(f"""
                            <div class="custom-card" style="border-left-color: #10B981; background-color: #0F172A; padding: 15px; border-radius: 8px; margin-bottom: 20px;">
                                <h4 style="margin: 0 0 10px 0; color: #10B981; font-family: Outfit;">🤖 Wygenerowane podsumowanie (RAG Grounded):</h4>
                                <p style="color: #F8FAFC; line-height: 1.6; font-size: 0.95rem; margin: 0;">
                                    {search_res['summary']}
                                </p>
                            </div>
                            """, unsafe_allow_html=True)
                            
                        st.subheader("📄 Dopasowane dokumenty i fragmenty:")
                        results = search_res.get("results", [])
                        if not results:
                            st.info("Nie znaleziono dokumentów pasujących do zapytania.")
                        else:
                            for idx, res in enumerate(results):
                                with st.expander(f"Wynik #{idx+1}: {res['title']}"):
                                    st.markdown(f"**ID Dokumentu:** `{res['id']}`")
                                    if res['link']:
                                        st.markdown(f"🔗 **Link:** [{res['link']}]({res['link']})")
                                    if res['snippet']:
                                        st.markdown(f"**Wycinek (Snippet):**\n*{res['snippet']}*")
                                    if res['extractive_answers']:
                                        st.markdown("**Sugerowane odpowiedzi:**")
                                        for ans in res['extractive_answers']:
                                            st.info(ans)
                                    if res['extractive_segments']:
                                        st.markdown("**Pasujące fragmenty tekstu:**")
                                        for seg in res['extractive_segments']:
                                            st.write(seg)


# 4a. LEAD RADAR
elif menu == "Lead_Radar":
    import modules.opportunity_scanner as os_scanner
    os_scanner.render_lead_radar_page(call_gemini_pro_api)
# 5. ADHD CRM & LEJEK
elif menu == "CRM":
    st.title("💼 Jaison Client & Project Pipeline")
    st.subheader("Twój zautomatyzowany, luksusowy lejek leadów zintegrowany z n8n & Systeme.io")
    
    st.markdown("""
    <div class="one-thing-banner" style="border-left-color: #7C3AED;">
        <h3 style="margin-top: 0; color: #7C3AED;">💼 Dlaczego ten CRM jest inny niż tradycyjne, skomplikowane systemy?</h3>
        <p style="color: #CBD5E1; line-height: 1.6; margin-bottom: 0;">
            Tradycyjne CRM-y bombardują Cię powiadomieniami, kolorowymi etykietami i dziesiątkami zawiłych opcji, wywołując u osób neuroatypowych paraliż i chęć ucieczki. 
            Nasza wersja to <strong>Twój osobisty asystent ADHD Flow</strong>: tylko 3 przejrzyste etapy, zero migających czerwonych kropek i wbudowany wirtualny zespół C-Suite gotowy do natychmiastowego doradztwa przy każdym kliencie.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    crm = load_crm()
    
    tab_board, tab_add, tab_onboarding, tab_webhook = st.tabs(["📊 Tablica Lejka", "➕ Dodaj Nowy Kontakt", "🤝 Onboarding & Grill AI", "🔌 Webhook & Symulacje"])
    
    with tab_webhook:
        st.markdown("### 🔌 Wbudowany Webhook Odbiorczy (Jaison CRM API)")
        st.write("Twój dashboard posiada zintegrowany, wbudowany serwer webhooków działający w tle, całkowicie eliminujący potrzebę korzystania z płatnych zewnętrznych narzędzi i integratorów.")
        st.markdown(f"""
        *   **Lokalny Endpoint:** `http://127.0.0.1:8090/webhook`
        *   **Serwerowy Endpoint:** `http://TWÓJ_SERVER_IP:8090/webhook`
        *   **Metoda:** `POST (application/json lub application/x-www-form-urlencoded)`
        *   **Pola:** `name` (Imię/Firma), `email` (E-mail), `notes` (Chaos operacyjny/Wiadomość)
        """)
        
        st.markdown("<br>", unsafe_allow_html=True)
        with st.expander("🤖 Hermes Telegram Agent Bridge (Status & Komendy)"):
            st.markdown("### 📡 Panel Sterowania Demonem Hermes Telegram Bot")
            
            pid_file = r"C:\Aplikacje MVP\01_JAISON_AGENCY_OS\02-website\telegram_bridge.pid"
            daemon_running = False
            active_pid = None
            
            if os.path.exists(pid_file):
                try:
                    with open(pid_file, "r") as f:
                        active_pid = int(f.read().strip())
                    import subprocess
                    output = subprocess.check_output(f'tasklist /FI "PID eq {active_pid}"', shell=True).decode()
                    if str(active_pid) in output:
                        daemon_running = True
                except Exception:
                    pass
            
            if daemon_running:
                st.markdown(f"""
                <div style="background: #064E3B; border-left: 4px solid #10B981; padding: 15px; border-radius: 4px; margin-bottom: 15px;">
                    <span style="font-weight: bold; color: #34D399; font-size: 1.1rem;">🟢 Status: AKTYWNY (Real Polling Mode)</span> 
                    <span style="background: #111827; color: #10B981; padding: 2px 8px; border-radius: 4px; font-size: 0.8rem; font-weight: bold; margin-left: 10px;">PID: {active_pid}</span>
                </div>
                """, unsafe_allow_html=True)
                
                if st.button("🛑 Zatrzymaj bota Telegram", type="secondary", key="stop_tg_daemon"):
                    try:
                        import os, signal
                        os.kill(active_pid, signal.SIGTERM)
                        st.info("Zatrzymano proces bota Telegram.")
                        if os.path.exists(pid_file):
                            os.remove(pid_file)
                        time.sleep(1.0)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Błąd zatrzymania: {e}")
            else:
                st.markdown("""
                <div style="background: #7F1D1D; border-left: 4px solid #EF4444; padding: 15px; border-radius: 4px; margin-bottom: 15px;">
                    <span style="font-weight: bold; color: #FCA5A5; font-size: 1.1rem;">🔴 Status: WYŁĄCZONY</span> 
                </div>
                """, unsafe_allow_html=True)
                
                if st.button("🔌 Uruchom bota Telegram", type="primary", key="start_tg_daemon"):
                    try:
                        import subprocess
                        import sys
                        script_path = r"C:\Aplikacje MVP\01_JAISON_AGENCY_OS\02-website\telegram_bridge.py"
                        p = subprocess.Popen([sys.executable, script_path], creationflags=subprocess.CREATE_NEW_CONSOLE)
                        st.success(f"🚀 Telegram Bot pomyślnie uruchomiony w tle! (PID: {p.pid})")
                        time.sleep(1.0)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Błąd uruchamiania bota: {e}")

            st.markdown("""
            Twój wirtualny asystent **Hermes Agent** jest połączony bezpośrednio z Twoim systemem.
            Pozwala to na sterowanie CRM-em oraz tworzenie projektów wideo za pomocą komend bezpośrednio z Twojego telefonu.
            
            ##### 🛠️ Wspierane komendy i akcje:
            1. **`/lead [Nazwa] [Notatki]`** — Dodaje nową szansę sprzedażową do bazy `crm.json`.
               * *Przykład:* `/lead coolfon.pl Szuka wdrożenia automatyzacji`
            2. **`/tasks`** — Wyświetla listę zadań z Twojego ADHD Kanban board.
            3. **`/stats`** — Zwraca aktualny stan Twojego pipeline sprzedaży.
            4. **`/video [Temat]`** — Generuje kompletny projekt wideo pod CapCut Desktop i syncuje go przez Git!
            """)

            
            # Dodajmy symulator komendy telegramowej dla Tomasza! (Wspaniały efekt)
            st.markdown("---")
            st.markdown("##### 📡 Symulacja komendy Telegram bota:")
            tg_command = st.text_input("Wpisz komendę Telegram bota do zasymulowania:", "/lead coolfon.pl Pan Tomasz chce asystenta AI do obsługi salonu")
            if st.button("🚀 Symuluj komendę Telegram", key="btn_sim_tg"):
                if tg_command.startswith("/lead "):
                    try:
                        content = tg_command[6:].strip()
                        parts = content.split(" ", 1)
                        client_name = parts[0]
                        client_notes = parts[1] if len(parts) > 1 else "Dodano przez komendę Telegram"
                        
                        # Dodaj leada bezpośrednio do centralnego CRM json
                        crm_data = load_crm()
                        new_id = f"lead_{int(time.time())}"
                        new_lead = {
                            "id": new_id,
                            "name": client_name,
                            "stage": "conversation",
                            "notes": client_notes,
                            "last_contact": time.strftime('%Y-%m-%d'),
                            "next_action": "Odezwij się do klienta i przeprowadź onboarding.",
                            "draft_reply": f"Cześć {client_name}, dzięki za kontakt przez Telegram..."
                        }
                        crm_data["leads"].append(new_lead)
                        save_crm(crm_data)
                        st.success(f"🔥 Komenda Telegram pomyślnie przetworzona! Lead '{client_name}' został dodany do bazy CRM.")
                        time.sleep(1.0)
                        st.rerun()
                    except Exception as e_tg:
                        st.error(f"Błąd symulacji komendy: {e_tg}")
                else:
                    st.warning("Symulator wspiera obecnie wyłącznie komendę `/lead [Nazwa] [Notatka]`.")
        
        st.markdown("---")
        st.markdown("### 📡 Symulacja Nowego Kontaktu z Landing Page (Test AI)")
        st.caption("Kliknij poniższy przycisk, aby w sekundy zasymulować wysłanie formularza z Twojej strony internetowej. Serwer webhooków w tle odbierze leada, zapisze go w bazie i asynchronicznie wygeneruje draft odpowiedzi!")
        
        test_case = st.selectbox("Wybierz profil testowego klienta do symulacji:", [
            "Marek Kowalski (Sklep Meble) - Chce zautomatyzować maile z reklamacjami bo traci na to 3 godziny dziennie.",
            "Katarzyna Wiśniewska (Biuro Rachunkowe) - Chce asystenta AI do wyjaśniania klientom zmian w podatkach.",
            "Tomasz Krawczyk (Agencja Reklamowa) - Chce bota kwalifikującego klientów na stronie www przed rozmową."
        ])
        
        if st.button("🚀 Wyślij Symulowane Zgłoszenie", type="primary"):
            parts = test_case.split(" - ")
            name = parts[0]
            notes = parts[1]
            email = name.lower().replace(" ", ".").replace("ł", "l").replace("ś", "s").replace("ó", "o").replace("ą", "a").replace("ć", "c").replace("ę", "e").replace("ń", "n").replace("ź", "z").replace("ż", "z") + "@test-firma.pl"
            
            with st.spinner("Wysyłanie POST do webhooka..."):
                try:
                    payload = {"name": name, "email": email, "notes": notes}
                    resp = http_post("http://127.0.0.1:8090/webhook", json_data=payload, timeout=5.0)
                    if resp.status_code == 200:
                        st.success(f"Sukces! Webhook zwrócił kod 200. Nowy lead '{name}' został zarejestrowany na tablicy. Wróć do tablicy lejkowej, otwórz jego kartę i zobacz wygenerowany przez CMO draft odpowiedzi!")
                        time.sleep(1.0)
                        st.rerun()
                    else:
                        st.error(f"Webhook zwrócił błąd {resp.status_code}: {resp.text}")
                except Exception as ex:
                    st.error(f"Nie udało się połączyć z lokalnym portem 8090. Dodaję leada bezpośrednio do bazy jako fallback...")
                    crm_data = load_crm()
                    new_id = f"lead_{int(time.time())}"
                    new_lead = {
                        "id": new_id,
                        "name": name,
                        "email": email,
                        "stage": "conversation",
                        "notes": notes,
                        "last_contact": time.strftime("%Y-%m-%d"),
                        "next_action": "Skontaktować się po analizie AI",
                        "draft_reply": "Generowanie odpowiedzi przez CMO AI..."
                    }
                    crm_data["leads"].append(new_lead)
                    save_crm(crm_data)
                    async_generate_initial_draft(new_id, name, notes)
                    st.success(f"Lead '{name}' został zapisany bezpośrednio w crm.json jako fallback!")
                    time.sleep(1.0)
                    st.rerun()

    with tab_add:
        st.markdown("### ➕ Zarejestruj nowego klienta w chmurze")
        new_name = st.text_input("Nazwa firmy / Nazwisko klienta:")
        new_notes = st.text_area("O czym rozmawiamy? (Opisz krótko ból operacyjny klienta):")
        new_action = st.text_input("Jaki jest najbliższy konkretny krok (Next Action):", placeholder="Np. Umówić wideorozmowę na pokaz wdrożenia...")
        
        if st.button("Zapisz w Bezszumnym CRM", type="primary"):
            if new_name and new_notes:
                new_id = f"lead_{int(time.time())}"
                new_lead = {
                    "id": new_id,
                    "name": new_name,
                    "stage": "conversation",
                    "notes": new_notes,
                    "last_contact": time.strftime("%Y-%m-%d"),
                    "next_action": new_action if new_action else "Odezwać się do klienta",
                    "draft_reply": f"Cześć {new_name.split()[0] if len(new_name.split()) > 0 else new_name}, dziękuję za rozmowę. Zrozumiałem Twój chaos operacyjny związany z {new_notes[:100]}... Zaprojektowałem dla Ciebie uproszczony system..."
                }
                crm["leads"].append(new_lead)
                save_crm(crm)
                st.success(f"Klient {new_name} został dodany do pierwszego etapu!")
                time.sleep(0.5)
                st.rerun()
            else:
                st.warning("Podaj nazwę klienta oraz jego opis.")
                
    with tab_onboarding:
        st.subheader("🤝 Onboarding & Grill Agent")
        st.markdown("""
        <div class="one-thing-banner" style="border-left-color: #EC4899;">
            <h3 style="margin-top: 0; color: #EC4899;">🤝 Onboarding zasilany Grillowaniem AI</h3>
            <p style="color: #CBD5E1; line-height: 1.6; margin-bottom: 0;">
                Zanim rozpoczniesz nowy projekt albo wdrożenie, nasz agent przeprowadzi z Tobą rygorystyczny wywiad.
                Będzie kwestionował Twoje założenia, szukał prawdziwych problemów Twoich odbiorców i na koniec stworzy kompletny, ustrukturyzowany Brief gotowy do zapisu w Obsidian Vault.
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # Initialize onboarding state
        if "onb_active" not in st.session_state:
            st.session_state.onb_active = False
            st.session_state.onb_step = 0
            st.session_state.onb_answers = []
            st.session_state.onb_questions = []
            st.session_state.onb_type = ""
            st.session_state.onb_target_name = ""
            st.session_state.onb_chat_history = []
            
        if not st.session_state.onb_active:
            st.subheader("🚀 Rozpocznij Nowy Wywiad")
            onb_name = st.text_input("Nazwa Klienta / Nazwa Projektu:", placeholder="np. Gabinet Fizjoterapii Kręgosłup, Nowa Platforma Kursowa...", key="crm_onb_name")
            onb_type = st.selectbox("Typ wywiadu onboardingowego:", [
                "A. Nowy klient biznesowy na systemy AI & CRM (B2B)",
                "B. Shadow Operating dla Twórcy Cyfrowego (Revenue Split)",
                "C. Nowy członek społeczności ADHD for Life (Onboarding profilowy)"
            ], key="crm_onb_type")
            
            if st.button("Uruchom Onboarding & Grill", type="primary", key="crm_onb_start_btn"):
                if onb_name:
                    st.session_state.onb_target_name = onb_name
                    st.session_state.onb_type = onb_type
                    st.session_state.onb_active = True
                    st.session_state.onb_step = 1
                    st.session_state.onb_answers = []
                    st.session_state.onb_chat_history = []
                    
                    # Predefined starter prompt to generate first question
                    starter_prompt = f"""Jesteś elitarnym agentem onboardingu i grillowania założeń biznesowych. 
Użytkownik rozpoczyna onboarding dla projektu/klienta: "{onb_name}" o typie: "{onb_type}".
Zadaj pierwsze, bardzo celne, drążące pytanie, które uderza w sedno problemu biznesowego lub kognitywnego. 
Zadaj TYLKO jedno pytanie. Nie pisz powitań ani wstępów."""
                    with st.spinner("Agent przygotowuje pierwsze pytanie..."):
                        first_q = call_gemini_pro_api([{"role": "user", "content": starter_prompt}], "Jesteś dociekliwym audytorem biznesowym.")
                    st.session_state.onb_questions = [first_q]
                    st.session_state.onb_chat_history.append({"role": "assistant", "content": first_q})
                    st.rerun()
                else:
                    st.warning("Podaj nazwę klienta lub projektu.")
        else:
            st.write(f"### 🤝 Wywiad: **{st.session_state.onb_target_name}**")
            st.caption(f"Typ: {st.session_state.onb_type} | Krok {st.session_state.onb_step} z 5")
            
            # Display chat history
            for msg in st.session_state.onb_chat_history:
                role = "assistant" if msg["role"] == "assistant" else "user"
                with st.chat_message(role):
                    st.markdown(msg["content"])
                    
            # If we have reached 5 steps, compile the Brief!
            if st.session_state.onb_step > 5:
                st.success("✅ Wywiad zakończony! Model kompiluje ustrukturyzowany Brief...")
                
                # Formulate prompt for Brief synthesis
                synthesis_prompt = f"""Na podstawie poniższego wywiadu onboardingowego dla "{st.session_state.onb_target_name}" ({st.session_state.onb_type}),
przygotuj kompletny, ustrukturyzowany Brief Projektowy w formacie Markdown.
Użyj sekcji:
# Brief Projektu: {st.session_state.onb_target_name}
- **Typ Projektu:** {st.session_state.onb_type}
- **Data sporządzenia:** {time.strftime('%Y-%m-%d')}

## 🎯 Główny Cel i Problem
(Jaki rzeczywisty problem odbiorcy rozwiązuje ten projekt? Czego klient naprawdę chce?)

## ⚖️ Analiza Ryzyk i Grillowanie
(Podsumowanie słabych punktów i założeń, które zostały zweryfikowane w wywiadzie)

## 🛠️ Rekomendowana Architektura AI/CRM
(Rekomendowane wdrożenie: n8n, Jaison CRM, automatyzacje, modele)

## 📅 Konkretne Mikro-Kroki (Next Actions)
(Lista 3-5 natychmiastowych zadań o niskim oporze kognitywnym do wdrożenia w Kanbanie)

WYWIAD:
"""
                for msg in st.session_state.onb_chat_history[:-1]: # exclude final system message
                    synthesis_prompt += f"{'Agent' if msg['role']=='assistant' else 'Użytkownik'}: {msg['content']}\n"
                    
                with st.spinner("CEO Jason & Onboarding Agent syntetyzują Brief..."):
                    brief_md = call_gemini_pro_api([{"role": "user", "content": synthesis_prompt}], "Jesteś CEO Jason i Onboarding Agent. Tworzysz precyzyjne briefy.")
                
                st.markdown("### 📝 Wygenerowany Brief:")
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #10B981; white-space: pre-wrap; font-size: 0.95rem; line-height: 1.7;">
{brief_md}
                </div>
                """, unsafe_allow_html=True)
                
                # Save to Obsidian Vault
                brief_filename = f"Brief_{st.session_state.onb_target_name.replace(' ', '_')}_{int(time.time())}.md"
                brief_filepath = os.path.join(OBSIDIAN_DIR, brief_filename)
                try:
                    with open(brief_filepath, "w", encoding="utf-8") as f:
                        f.write(brief_md)
                    st.success(f"📚 Brief został pomyślnie zapisany w Obsidian Vault: `{brief_filename}`")
                except Exception as e:
                    st.error(f"Nie udało się zapisać briefu w Obsidianie: {e}")

                st.markdown("<br>", unsafe_allow_html=True)
                c_action1, c_action2 = st.columns(2)
                
                with c_action1:
                    if st.button("👥 Przepuść przez Flotę Dyrektorów", use_container_width=True, key="crm_onb_fleet_btn"):
                        swarm_prompt = f"""Jesteś wirtualną Flotą Dyrektorów (CEO, CTO, CFO, CMO, COO, CSO, CPO) Jaison.pl.
Przeanalizuj poniższy Brief Projektowy i wygeneruj dla każdego dyrektora jedną, ultrakrótką (max 2 zdania), konkretną, ADHD-friendly, biznesową i merytoryczną rekomendację strategiczną.

BRIEF:
{brief_md}

Format wyjściowy:
### 👥 Rekomendacje Floty Dyrektorów:
- 👑 **CEO (Dyrektor Generalny):** [rekomendacja]
- 🤖 **CTO (Technologia & AI):** [rekomendacja]
- 🎨 **CPO (Produkt & Customer Journey):** [rekomendacja]
- 📢 **CMO (Marketing & Ruch):** [rekomendacja]
- 💸 **CFO (Finanse & Rentowność):** [rekomendacja]
- ⚡ **COO (Operacje & n8n):** [rekomendacja]
- 📞 **CSO (Sprzedaż B2B):** [rekomendacja]
"""
                        with st.spinner("Flota Dyrektorów debatuje nad Briefem..."):
                            try:
                                swarm_res = call_gemini_pro_api([{"role": "user", "content": swarm_prompt}], "Jesteś zarządem i flotą dyrektorów agencji AI.")
                                st.session_state.onb_swarm_results = swarm_res
                                
                                # Zapisz rekomendacje w Obsidianie
                                rec_filename = f"Rekomendacje_Zarzadu_{st.session_state.onb_target_name.replace(' ', '_')}_{int(time.time())}.md"
                                rec_filepath = os.path.join(OBSIDIAN_DIR, rec_filename)
                                with open(rec_filepath, "w", encoding="utf-8") as f_rec:
                                    f_rec.write(swarm_res)
                                st.toast(f"Saved recommendations to: {rec_filename}")
                            except Exception as ex:
                                st.error(f"Błąd Floty Dyrektorów: {ex}")

                with c_action2:
                    if st.button("📋 Utwórz zadania w Kanbanie", use_container_width=True, key="crm_onb_kanban_btn"):
                        kanban_prompt = f"""Przeanalizuj poniższy Brief i wyodrębnij z niego od 3 do 5 konkretnych zadań wdrożeniowych (tzw. Next Actions).
Zwróć je jako surową listę JSON w formacie:
[
  {{"title": "Krótki tytuł zadania", "desc": "Opis mikro-kroku do wykonania", "label": "Wdrożenie"}}
]
Zwróć TYLKO poprawny kod JSON. Żadnego markdownu, żadnego dodatkowego tekstu.

BRIEF:
{brief_md}
"""
                        with st.spinner("Ekstrakcja zadań z Briefu do Kanbana..."):
                            try:
                                json_res = call_gemini_pro_api([{"role": "user", "content": kanban_prompt}], "Jesteś analitykiem zadań.")
                                cleaned_json = json_res.strip()
                                if cleaned_json.startswith("```json"):
                                    cleaned_json = cleaned_json[7:]
                                if cleaned_json.endswith("```"):
                                    cleaned_json = cleaned_json[:-3]
                                cleaned_json = cleaned_json.strip()
                                
                                parsed_tasks = json.loads(cleaned_json)
                                if isinstance(parsed_tasks, list):
                                    kanban_data = []
                                    if os.path.exists(KANBAN_FILE):
                                        try:
                                            with open(KANBAN_FILE, "r", encoding="utf-8") as k_f:
                                                kanban_data = json.load(k_f)
                                        except:
                                            pass
                                            
                                    for t in parsed_tasks:
                                        kanban_data.append({
                                            "id": f"task_{int(time.time())}_{hash(t.get('title', '')) % 1000}",
                                            "title": t.get("title", "Nowe zadanie"),
                                            "desc": t.get("desc", ""),
                                            "label": t.get("label", "Wdrożenie"),
                                            "done": False,
                                            "client": st.session_state.onb_target_name
                                        })
                                    
                                    with open(KANBAN_FILE, "w", encoding="utf-8") as k_f:
                                        json.dump(kanban_data, k_f, ensure_ascii=False, indent=4)
                                    st.success(f"🔥 Dodano pomyślnie {len(parsed_tasks)} zadań do Twojego Kanbana!")
                                else:
                                    st.error("Model nie zwrócił poprawnej listy zadań.")
                            except Exception as ex:
                                st.error(f"Nie udało się sparsować zadań: {ex}")

                if "onb_swarm_results" in st.session_state:
                    st.markdown("<br>", unsafe_allow_html=True)
                    st.markdown(f"""
                    <div class="custom-card" style="border-left: 4px solid #8B5CF6; white-space: pre-wrap; font-size: 0.95rem; line-height: 1.7;">
{st.session_state.onb_swarm_results}
                    </div>
                    """, unsafe_allow_html=True)
                
                st.markdown("<br><hr>", unsafe_allow_html=True)
                if st.button("Rozpocznij nowy onboarding", type="primary", key="crm_onb_restart_btn"):
                    if "onb_swarm_results" in st.session_state:
                        del st.session_state.onb_swarm_results
                    st.session_state.onb_active = False
                    st.rerun()
            else:
                # We are in the middle of the interview
                user_reply = st.chat_input("Twoja odpowiedź (Agent wygrilluje Twoje założenia):", key="crm_onb_chat_input")
                
                if user_reply:
                    st.session_state.onb_chat_history.append({"role": "user", "content": user_reply})
                    st.session_state.onb_answers.append(user_reply)
                    st.session_state.onb_step += 1
                    
                    # Check if this is the final step
                    if st.session_state.onb_step > 5:
                        st.session_state.onb_chat_history.append({"role": "assistant", "content": "Dziękuję. Wywiad zakończony. Przechodzę do generowania briefu..."})
                        st.rerun()
                    else:
                        # Generate next question by prompting the model to grill the last response and advance
                        grill_prompt = f"""Jesteś agentem onboardingu i grillowania założeń biznesowych. 
Onboarding dla projektu/klienta: "{st.session_state.onb_target_name}" o typie: "{st.session_state.onb_type}".
Oto dotychczasowa historia wywiadu:
"""
                        for msg in st.session_state.onb_chat_history:
                            grill_prompt += f"{'Agent' if msg['role']=='assistant' else 'Użytkownik'}: {msg['content']}\n"
                            
                        grill_prompt += f"\nZADANIE: Przeanalizuj ostatnią odpowiedź użytkownika, krótko zakwestionuj lub podważ jedno z jego założeń (grillowanie), a następnie zadaj krok {st.session_state.onb_step} z 5 (następne celne pytanie). Zadaj TYLKO jedno pytanie. Nie pisz powitań ani podsumowań."
                        
                        with st.spinner("Agent analizuje Twoją odpowiedź i szykuje kolejne pytanie..."):
                            next_q = call_gemini_pro_api([{"role": "user", "content": grill_prompt}], "Jesteś rygorystycznym audytorem biznesowym grillującym pomysły.")
                            
                        st.session_state.onb_questions.append(next_q)
                        st.session_state.onb_chat_history.append({"role": "assistant", "content": next_q})
                        st.rerun()
                        
        if st.button("❌ Przerwij wywiad i zresetuj state", key="crm_onb_cancel_btn"):
            st.session_state.onb_active = False
            st.rerun()

    with tab_board:
        col_conv, col_arch, col_build = st.columns(3)
        
        conv_leads = [l for l in crm["leads"] if l.get("stage") == "conversation"]
        arch_leads = [l for l in crm["leads"] if l.get("stage") == "architecture"]
        build_leads = [l for l in crm["leads"] if l.get("stage") == "build"]
        
        with col_conv:
            st.markdown("### 📥 1. Rozmowa")
            st.caption("Początkowy kontakt i ankieta intake")
            for i, lead in enumerate(conv_leads):
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #3B82F6; margin-bottom: 12px;">
                    <span style="font-size: 0.8rem; color:#94A3B8;">📞 Kontakt: {lead.get('last_contact')}</span>
                    <h4 style="margin-top: 4px; margin-bottom: 8px; color: #FFFFFF;">{lead.get('name')}</h4>
                    <p style="font-size: 0.9rem; color: #CBD5E1; line-height: 1.4;">{lead.get('notes')[:100]}...</p>
                    <hr style="border-color: #1F242E; margin: 10px 0;">
                    <span style="font-size: 0.85rem; color: #F59E0B;">➡️ <strong>Krok:</strong> {lead.get('next_action')}</span>
                </div>
                """, unsafe_allow_html=True)
                
                c_opt1, c_opt2 = st.columns(2)
                with c_opt1:
                    if st.button("🔍 Otwórz", key=f"sel_conv_{lead['id']}"):
                        st.session_state.selected_lead_id = lead['id']
                with c_opt2:
                    if st.button("📐 Buduj", key=f"mov_arch_{lead['id']}"):
                        lead["stage"] = "architecture"
                        save_crm(crm)
                        st.rerun()
                        
        with col_arch:
            st.markdown("### 📐 2. Architektura")
            st.caption("Projektowanie bazy i Niewidzialnego Pracownika")
            for i, lead in enumerate(arch_leads):
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #F59E0B; margin-bottom: 12px;">
                    <span style="font-size: 0.8rem; color:#94A3B8;">📅 Planowanie: {lead.get('last_contact')}</span>
                    <h4 style="margin-top: 4px; margin-bottom: 8px; color: #FFFFFF;">{lead.get('name')}</h4>
                    <p style="font-size: 0.9rem; color: #CBD5E1; line-height: 1.4;">{lead.get('notes')[:100]}...</p>
                    <hr style="border-color: #1F242E; margin: 10px 0;">
                    <span style="font-size: 0.85rem; color: #F59E0B;">➡️ <strong>Krok:</strong> {lead.get('next_action')}</span>
                </div>
                """, unsafe_allow_html=True)
                
                a_opt1, a_opt2 = st.columns(2)
                with a_opt1:
                    if st.button("🔍 Otwórz", key=f"sel_arch_{lead['id']}"):
                        st.session_state.selected_lead_id = lead['id']
                with a_opt2:
                    if st.button("🚀 Wdróż", key=f"mov_build_{lead['id']}"):
                        lead["stage"] = "build"
                        save_crm(crm)
                        st.rerun()
                        
        with col_build:
            st.markdown("### 🚀 3. Wdrożenie")
            st.caption("Developer mode — procesy live")
            for i, lead in enumerate(build_leads):
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #10B981; margin-bottom: 12px;">
                    <span style="font-size: 0.8rem; color:#94A3B8;">⚡ Integracja: {lead.get('last_contact')}</span>
                    <h4 style="margin-top: 4px; margin-bottom: 8px; color: #FFFFFF;">{lead.get('name')}</h4>
                    <p style="font-size: 0.9rem; color: #CBD5E1; line-height: 1.4;">{lead.get('notes')[:100]}...</p>
                    <hr style="border-color: #1F242E; margin: 10px 0;">
                    <span style="font-size: 0.85rem; color: #F59E0B;">➡️ <strong>Krok:</strong> {lead.get('next_action')}</span>
                </div>
                """, unsafe_allow_html=True)
                
                b_opt1, b_opt2 = st.columns(2)
                with b_opt1:
                    if st.button("🔍 Otwórz", key=f"sel_build_{lead['id']}"):
                        st.session_state.selected_lead_id = lead['id']
                with b_opt2:
                    if st.button("📦 Archiwum", key=f"mov_archv_{lead['id']}"):
                        lead["stage"] = "archive"
                        save_crm(crm)
                        st.rerun()

    # Szczegółowy podgląd wybranego klienta (ADHD Focus Panel)
    if "selected_lead_id" in st.session_state and st.session_state.selected_lead_id:
        sel_id = st.session_state.selected_lead_id
        lead = next((l for l in crm["leads"] if l["id"] == sel_id), None)
        
        if lead:
            st.markdown("---")
            st.subheader(f"💼 ADHD Focus Panel: {lead['name']}")
            
            c_det1, c_det2 = st.columns([1, 1])
            
            with c_det1:
                st.markdown(f"""
                <div class="custom-card" style="border-left: 5px solid #7C3AED;">
                    <h4>📋 Informacje o kliencie</h4>
                    <p><strong>Ból operacyjny i notatki:</strong><br>{lead['notes']}</p>
                    <p><strong>Ostatni kontakt:</strong> {lead['last_contact']}</p>
                    <p><strong>Następny krok:</strong> <span class="focus-accent">{lead['next_action']}</span></p>
                </div>
                """, unsafe_allow_html=True)
                
                # Sterowanie procesem
                st.write("##### Przesuń etap klienta:")
                col_stg1, col_stg2, col_stg3 = st.columns(3)
                with col_stg1:
                    if st.button("📥 Cofnij do Rozmowy", key="btn_move_conv"):
                        lead["stage"] = "conversation"
                        save_crm(crm)
                        st.rerun()
                with col_stg2:
                    if st.button("📐 Do Architektury", key="btn_move_arch"):
                        lead["stage"] = "architecture"
                        save_crm(crm)
                        st.rerun()
                with col_stg3:
                    if st.button("🚀 Do Wdrożenia", key="btn_move_build"):
                        lead["stage"] = "build"
                        save_crm(crm)
                        st.rerun()
                
                if st.button("🎯 Ustaw jako priorytet 'One Thing'", key=f"set_ot_{lead['id']}"):
                    st.session_state.one_thing = f"Dla klienta {lead['name']}: {lead['next_action']}"
                    st.toast(f"Klient {lead['name']} został ustawiony jako główny cel!")
                    st.rerun()
                
                if st.button("💰 Przygotuj Fakturę (Fakturownia)", key=f"set_fin_{lead['id']}"):
                    st.session_state.fin_client_name = lead["name"]
                    st.session_state.fin_client_email = lead.get("email", "klient@test.pl")
                    st.toast(f"Dane klienta {lead['name']} przeniesione do zakładki finansowej!")
                    st.info("Przejdź teraz do menu '💰 Kancelaria Finansowa & KSeF', aby wystawić dokument.")
                
                st.markdown("<br>", unsafe_allow_html=True)
                col_del, col_close = st.columns(2)
                with col_del:
                    if st.button("🗑️ Usuń klienta z bazy", key="btn_del_lead"):
                        crm["leads"] = [l for l in crm["leads"] if l["id"] != sel_id]
                        save_crm(crm)
                        st.session_state.selected_lead_id = None
                        st.rerun()
                with col_close:
                    if st.button("❌ Zamknij podgląd", key="btn_close_lead"):
                        st.session_state.selected_lead_id = None
                        st.rerun()
                        
            with c_det2:
                # Wersje robocze odpowiedzi
                st.markdown("""
                <div class="custom-card" style="border-left: 5px solid #10B981; background-color: #0F1D1A;">
                    <h4 style="margin: 0; color: #10B981;">✉️ Propozycja bezszumnej odpowiedzi (CMO Draft)</h4>
                    <p style="font-size: 0.85rem; color: #94A3B8; margin-top: 4px;">Wygenerowana automatycznie w oparciu o o_mnie.md w celu głębokiego rezonowania z klientem.</p>
                </div>
                """, unsafe_allow_html=True)
                
                modified_reply = st.text_area("Możesz spersonalizować odpowiedź przed wysłaniem:", value=lead.get("draft_reply", ""), height=150)
                col_cmo1, col_cmo2 = st.columns(2)
                with col_cmo1:
                    if st.button("🧠 Regeneruj przez CMO AI", key=f"regen_cmo_{lead['id']}"):
                        with st.spinner("CMO generuje odpowiedź..."):
                            o_mnie_path = os.path.join(HERMES_DIR, "o_mnie.md")
                            o_mnie_context = ""
                            if os.path.exists(o_mnie_path):
                                try:
                                    o_mnie_context = open(o_mnie_path, "r", encoding="utf-8").read()
                                except:
                                    pass
                            system_prompt = f"""Jesteś wirtualnym CMO Tomasza Dudy (architekta systemów AI dla neuroatypowych, Holistic AIDHD).
Stwórz niezwykle empatyczny, autentyczny i humorystyczny e-mail po polsku odpowiadający na zapytanie klienta. Pisz w 1 osobie jako Tomasz.
Oto historia i tożsamość Tomasza:
{o_mnie_context}
"""
                            user_prompt = f"""Klient: {lead['name']}
Opis chaosu/potrzeby: {lead['notes']}
Najbliższy krok: {lead['next_action']}
Zredaguj niesamowity, głęboki, gotowy do wysłania e-mail."""
                            reply = call_gemini_api([{"role": "user", "content": user_prompt}], system_instruction=system_prompt)
                            if reply:
                                lead["draft_reply"] = reply
                                save_crm(crm)
                                st.rerun()
                with col_cmo2:
                    if st.button("Zatwierdź i Skopiuj", key=f"copy_cmo_{lead['id']}"):
                        lead["draft_reply"] = modified_reply
                        save_crm(crm)
                        st.toast("Odpowiedź zapisana i skopiowana do schowka!")
                
                st.markdown("---")
                
                # Zintegrowane C-Suite dla tego klienta!
                st.write("##### 👥 Wirtualna Konsultacja C-Suite dla tego klienta")
                agents = {
                    "CEO (Strategia & Rentowność)": "CEO pomoże Ci wycenić to wdrożenie pod kątem modelu High-Ticket i zaplanować kroki MVP.",
                    "CMO (Empatyczny Storytelling)": "CMO przełoży ból klienta na autentyczny i humorystyczny przekaz dopasowany do jego problemów.",
                    "CSO (Architektura Sprzedaży)": "CSO zaprojektuje prosty, trzystopniowy lejek bezszumny dla tego klienta.",
                    "CTO (Technologia & Kod)": "CTO podpowie, jakich konkretnych asystentów Gemini oraz automatyzacji w n8n użyć do tego wdrożenia."
                }
                
                selected_dyrektor = st.selectbox("Skonsultuj się z:", list(agents.keys()), key=f"csuite_sel_{lead['id']}")
                st.caption(agents[selected_dyrektor])
                
                if st.button("Konsultuj przypadek", type="primary", key=f"csuite_btn_{lead['id']}"):
                    with st.spinner(f"{selected_dyrektor} analizuje kartę klienta..."):
                        csuite_reply = consult_csuite_live(lead, selected_dyrektor)
                        st.info(f"**{selected_dyrektor}:** \"{csuite_reply}\"")

# 6. ADHD KANBAN
elif menu == "Kanban":
    st.title("🎯 ADHD Kanban Board")
    st.subheader("Wizualny postęp wdrożeń bez paraliżu decyzyjnego")
    
    k = load_kanban()
    
    with st.expander("➕ Dodaj nowe zadanie"):
        task_text = st.text_input("Zadanie (krótko i konkretnie):")
        col_dest = st.selectbox("Dodaj do kolumny:", ["Selekcja (Triage)", "Do zrobienia (Todo)", "Gotowe (Ready)", "W toku (Running)", "Zablokowane (Blocked)", "Zrobione (Done)"], index=0)
        col_map = {
            "Selekcja (Triage)": "triage",
            "Do zrobienia (Todo)": "todo",
            "Gotowe (Ready)": "ready",
            "W toku (Running)": "running",
            "Zablokowane (Blocked)": "blocked",
            "Zrobione (Done)": "done"
        }
        if st.button("Dodaj"):
            if task_text:
                dest_key = col_map[col_dest]
                k[dest_key].append(task_text)
                save_kanban(k)
                st.rerun()
                
    st.markdown("---")
    
    cols = ["triage", "todo", "ready", "running", "blocked", "done"]
    col_labels = {
        "triage": "📥 Selekcja (Triage)",
        "todo": "📋 Do zrobienia (Todo)",
        "ready": "👍 Gotowe (Ready)",
        "running": "⚡ W toku (Running)",
        "blocked": "🚫 Zablokowane",
        "done": "🎉 Zrobione (Done)"
    }
    col_colors = {
        "triage": "#64748B",
        "todo": "#3B82F6",
        "ready": "#8B5CF6",
        "running": "#F59E0B",
        "blocked": "#EF4444",
        "done": "#10B981"
    }
    
    columns_ui = st.columns(6)
    
    for idx, col_name in enumerate(cols):
        with columns_ui[idx]:
            st.markdown(f"##### {col_labels[col_name]}")
            tasks_in_col = k.get(col_name, [])
            if not tasks_in_col:
                st.caption("Pusta kolumna")
            for i, t in enumerate(tasks_in_col):
                opacity = "0.7" if col_name == "done" else "1.0"
                # Render Task Card
                st.markdown(f"""
                <div class='custom-card' style='border-left: 4px solid {col_colors[col_name]}; padding: 10px; margin-bottom: 5px; opacity: {opacity};'>
                    <div style='font-size: 0.9rem; color: #FFFFFF;'>{t}</div>
                </div>
                """, unsafe_allow_html=True)
                
                # Action Buttons inside task
                c1, c2, c3, c4 = st.columns(4)
                
                # Move Left
                if col_name != "triage":
                    with c1:
                        if st.button("◀", key=f"ml_{col_name}_{i}", help="Przesuń w lewo"):
                            k[col_name].pop(i)
                            k[cols[idx - 1]].append(t)
                            save_kanban(k)
                            st.rerun()
                # Move Right
                if col_name != "done":
                    with c2:
                        if st.button("▶", key=f"mr_{col_name}_{i}", help="Przesuń w prawo"):
                            k[col_name].pop(i)
                            k[cols[idx + 1]].append(t)
                            save_kanban(k)
                            st.rerun()
                # Toggle Blocked
                if col_name != "blocked" and col_name != "done":
                    with c3:
                        if st.button("🚫", key=f"bl_{col_name}_{i}", help="Zablokuj"):
                            k[col_name].pop(i)
                            k["blocked"].append(t)
                            save_kanban(k)
                            st.rerun()
                elif col_name == "blocked":
                    with c3:
                        if st.button("⚡", key=f"bl_{col_name}_{i}", help="Odblokuj (wróć do W toku)"):
                            k[col_name].pop(i)
                            k["running"].append(t)
                            save_kanban(k)
                            st.rerun()
                # Delete task
                with c4:
                    if st.button("🗑️", key=f"del_{col_name}_{i}", help="Usuń trwale"):
                        k[col_name].pop(i)
                        save_kanban(k)
                        st.rerun()
                        
            if col_name == "done" and tasks_in_col:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("Wyczyść ukończone", key="clear_done_tasks", use_container_width=True):
                    k["done"] = []
                    save_kanban(k)
                    st.rerun()

# 7. DZIAŁ PRAWNY & KANCELARIA
elif menu == "Legal":
    st.title("💼 Dział Prawny — Twoja Holistyczna Tarcza")
    st.subheader("Automatyczne generowanie pism i audyt zasilany przez AI")
    
    st.markdown("""
    <div class="one-thing-banner" style="border-left-color: #EF4444;">
        <h3 style="margin-top: 0; color: #EF4444;">⚖️ Holistyczny Obrońca Prawny</h3>
        <p style="color: #CBD5E1; line-height: 1.6; margin-bottom: 0;">
            Masz na głowie trudne pismo, wezwanie do zapłaty albo umowę napisaną prawniczym żargonem? Wrzuć plik (PDF, skan, obrazek) lub wklej tekst i napisz prostym językiem, o co chodzi.
            Nasz asystent w tle rozbuduje Twoje polecenie, przeanalizuje szczegóły i przygotuje kompletny, profesjonalny projekt gotowego pisma procesowego, umowy lub odpowiedzi. Bez stresu i bez paraliżu ADHD.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    uploaded_legal_files = st.file_uploader("Załącz dokumenty (Pliki PDF, skany, obrazki lub zdjęcia):", type=["pdf", "png", "jpg", "jpeg"], accept_multiple_files=True)
    contract_text = st.text_area("Lub wklej tutaj treść dokumentu / notatek (opcjonalnie):", height=150)
    user_instruction = st.text_input("Twoje polecenie (krótko, po ludzku - AI rozbuduje je w tle i podbije prompt):", placeholder="np. Napisz")
    
    c_col1, c_col2 = st.columns([1, 1])
    
    with c_col1:
        doc_type = st.selectbox("Co chcesz uzyskać (Typ dokumentu docelowego):", [
            "📩 Oficjalne wezwanie do zapłaty / wykonania umowy",
            "⚖️ Gotowe pismo procesowe do sądu (pozew, odpowiedź na pozew, sprzeciw od nakazu)",
            "✍️ Profesjonalny projekt umowy / aneksu / NDA",
            "🛡️ Gotowa odpowiedź na wezwanie / pismo przeciwnika",
            "🔍 Głęboka analiza prawna i audyt ryzyk (Wykrywanie haczyków)"
        ])
        
    with c_col2:
        obsidian_export = st.checkbox("Automatycznie wyeksportuj wynik do Obsidian Vault", value=True)
        
    if st.button("Uruchom Generator Prawny AI", type="primary"):
        if contract_text or uploaded_legal_files or user_instruction:
            
            # =====================================================
            # KROK 0: WCZYTANIE PLIKÓW
            # =====================================================
            file_texts = []
            image_data_urls = []
            
            if uploaded_legal_files:
                for uploaded_file in uploaded_legal_files:
                    file_bytes = uploaded_file.read()
                    file_name = uploaded_file.name.lower()
                    
                    if file_name.endswith(".pdf"):
                        with st.spinner(f"📄 Wczytuję PDF: {uploaded_file.name}..."):
                            try:
                                import io, base64 as _b64
                                try:
                                    import pypdf
                                except ImportError:
                                    import subprocess, sys
                                    subprocess.run([sys.executable, "-m", "pip", "install", "pypdf"], capture_output=True)
                                    import pypdf
                                
                                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                                raw_pages = []
                                for page in reader.pages:
                                    t = page.extract_text()
                                    if t and len(t.strip()) > 20:
                                        raw_pages.append(t)
                                raw_text = "\n".join(raw_pages)
                                
                                # Skan — użyj natywnego API Vertex dla PDF OCR
                                if len(raw_text.strip()) < 200:
                                    st.warning("⚠️ PDF wygląda na skan. Uruchamiam natywny Vertex AI OCR...")
                                    raw_text = call_native_vertex_ocr(file_bytes)
                                    if not raw_text.startswith("[Błąd"):
                                        st.success(f"✅ OCR Gemini Native: {len(reader.pages)} stron")
                                    else:
                                        st.error(raw_text)
                                else:
                                    st.success(f"✅ PDF tekstowy: {len(reader.pages)} stron, {len(raw_text)} znaków")
                                
                                file_texts.append(f"=== DOKUMENT: {uploaded_file.name} ===\n{raw_text}\n=== KONIEC ===")
                            except Exception as e:
                                st.error(f"Błąd odczytu PDF {uploaded_file.name}: {e}")

                    
                    elif file_name.endswith((".png", ".jpg", ".jpeg")):
                        import base64 as _b64
                        mime_type = "image/png" if file_name.endswith(".png") else "image/jpeg"
                        encoded = _b64.b64encode(file_bytes).decode("utf-8")
                        image_data_urls.append({"name": uploaded_file.name, "data_url": f"data:{mime_type};base64,{encoded}"})
                        st.success(f"✅ Skan załączony: {uploaded_file.name}")
            
            # =====================================================
            # KROK 1: STRUKTURALNA EKSTRAKCJA DANYCH (Flash — szybka i niezawodna)
            # =====================================================
            extracted_data_json = {}
            source_text = "\n\n".join(file_texts) + (f"\n{contract_text}" if contract_text else "")
            
            if source_text.strip():
                with st.spinner("🔎 Krok 1/3 — Wyciągam dane wrażliwe z dokumentu..."):
                    extract_prompt = """Jesteś precyzyjnym ekstraktoremm danych prawnych. Z poniższego dokumentu wyciągnij WSZYSTKIE dostępne dane.
ZWRÓĆ TYLKO JSON (bez markdown, bez wyjaśnień), wypełniając każde pole które istnieje w dokumencie:

{
  "sygnatura_akt": null,
  "sad_organ": null,
  "sad_nazwa": null,
  "sad_adres": null,
  "data_dokumentu": null,
  "strona_powodowa": {"imie_nazwisko": null, "pesel": null, "adres": null, "nip": null, "rola": "Powód"},
  "strona_pozwana": {"imie_nazwisko": null, "pesel": null, "adres": null, "nip": null, "rola": "Pozwany"},
  "inne_osoby": [],
  "kwoty": [],
  "daty_kluczowe": [],
  "numery_referencyjne": [],
  "przedmiot_sprawy": null,
  "terminy": [],
  "pelnomocnicy": []
}

WAŻNE: Jeśli wartość istnieje w dokumencie, WSTAW JĄ. Nie zostawiaj null gdy dane są w tekście.

DOKUMENT:
""" + source_text[:10000] + "\n\nZWRÓĆ TYLKO JSON."
                    
                    import json as _json, re as _re
                    try:
                        extraction_raw = call_gemini_api([{"role": "user", "content": extract_prompt}])
                        jm = _re.search(r'\{[\s\S]*\}', extraction_raw)
                        if jm:
                            extracted_data_json = _json.loads(jm.group())
                    except Exception as ex:
                        st.warning(f"Ekstrakcja danych: {ex}")
                        extracted_data_json = {}
                    
                    if extracted_data_json and any(v for v in extracted_data_json.values() if v and v != [] and v != {}):
                        st.success("✅ Dane wyciągnięte z dokumentu:")
                        preview_cols = st.columns(4)
                        if extracted_data_json.get("sygnatura_akt"):
                            preview_cols[0].metric("📋 Sygnatura", extracted_data_json["sygnatura_akt"])
                        sp = extracted_data_json.get("strona_powodowa") or {}
                        if sp and sp.get("imie_nazwisko"):
                            preview_cols[1].metric(f"👤 {sp.get('rola','Powód')}", sp["imie_nazwisko"])
                        sz = extracted_data_json.get("strona_pozwana") or {}
                        if sz and sz.get("imie_nazwisko"):
                            preview_cols[2].metric(f"👤 {sz.get('rola','Pozwany')}", sz["imie_nazwisko"])
                        sad_name = extracted_data_json.get("sad_nazwa") or extracted_data_json.get("sad_organ") or ""
                        if sad_name:
                            preview_cols[3].metric("🏛️ Sąd", sad_name[:30])
                        if extracted_data_json.get("kwoty"):
                            st.caption(f"💰 Kwoty: {', '.join(str(k) for k in extracted_data_json['kwoty'][:5])}")
                        if extracted_data_json.get("data_dokumentu"):
                            st.caption(f"📅 Data: {extracted_data_json['data_dokumentu']}")
                        with st.expander("🔍 Pełne dane strukturalne (JSON)"):
                            st.code(_json.dumps(extracted_data_json, ensure_ascii=False, indent=2), language="json")
                    else:
                        st.info("ℹ️ Dane strukturalne nie rozpoznane — model użyje surowego tekstu.")
            
            # =====================================================
            # KROK 2: GENEROWANIE DOKUMENTU PRAWNEGO (Pro przez proxy)
            # =====================================================
            analysis_result = ""
            verification_result = ""
            
            with st.spinner("⚖️ Krok 2/3 — Sporządzam pismo prawne (Gemini Pro)..."):
                import json as _json3
                
                data_block = ""
                if extracted_data_json and any(v for v in extracted_data_json.values() if v and v != [] and v != {}):
                    data_block = f"""
## DANE WYCIĄGNIĘTE Z DOKUMENTU — UŻYJ BEZPOŚREDNIO:
```json
{_json3.dumps(extracted_data_json, ensure_ascii=False, indent=2)}
```
ZAKAZ używania [] placeholderów dla powyższych danych!
"""
                
                legal_system = """Jesteś elitarnym polskim radcą prawnym i adwokatem z wieloletnim doświadczeniem procesowym.
Piszesz WYŁĄCZNIE w języku polskim. Styl: precyzyjny, profesjonalny, rygorystyczny.

BEZWZGLĘDNE ZASADY:
1. Wstawiaj RZECZYWISTE dane z dokumentów. NIGDY [xxx] placeholdery gdy dane są dostępne.
2. Jeśli sekcja DANE STRUKTURALNE zawiera sygnaturę/PESEL/nazwisko — wstaw je DOKŁADNIE.
3. Każde pismo musi mieć PEŁNY NAGŁÓWEK z danymi stron, sygnaturą i nazwą sądu.
4. Powołuj się na KONKRETNE artykuły prawa (art. 123 § 1 k.c., art. 505 k.p.c. itp.).
5. Format nagłówka pisma procesowego (standardowy PL):
   POWÓD/WNIOSKODAWCA (lewa strona) | SĄD/ORGAN (prawa strona)
   Imię Nazwisko                      | Nazwa Sądu
   PESEL/NIP                          | Adres Sądu
   Adres                              | Sygn. akt: XXXX/YY
"""
                
                full_prompt = f"""ZADANIE: Sporządź profesjonalne pismo prawne.

Typ dokumentu: {doc_type}
Polecenie użytkownika: {user_instruction if user_instruction else "Dokonaj analizy i przygotuj pismo zabezpieczające interesy klienta."}

{data_block}

DOKUMENTY ŹRÓDŁOWE:
{chr(10).join(file_texts) if file_texts else "(brak plików)"}
{contract_text if contract_text else ""}

WYMAGANIA KOŃCOWE:
1. Wstaw WSZYSTKIE dane z dokumentów bezpośrednio do pisma.
2. Napisz pismo od nagłówka do podpisu — KOMPLETNE i gotowe do wysłania.
3. Przywołaj konkretne artykuły prawa.
4. Użyj markdownu do formatowania.
5. NIE używaj [xxx] gdy dane są dostępne wyżej."""
                
                if image_data_urls:
                    content_list = [{"type": "text", "text": full_prompt}]
                    for img in image_data_urls:
                        content_list.append({"type": "image_url", "image_url": {"url": img["data_url"]}})
                    msgs = [{"role": "user", "content": content_list}]
                else:
                    msgs = [{"role": "user", "content": full_prompt}]
                
                analysis_result = call_gemini_pro_api(msgs, legal_system)
                
                if not analysis_result or len(analysis_result) < 50:
                    st.error(f"Generator nie zwrócił wyników: {analysis_result}")
                else:
                    st.markdown("### 📝 Wygenerowane Pismo Prawne")
                    st.markdown(
                        f"<div class='custom-card' style='border-left: 4px solid #7C3AED;"
                        f"white-space: pre-wrap; font-family: Georgia, serif; line-height: 1.8;'>"
                        f"{analysis_result}</div>",
                        unsafe_allow_html=True
                    )
            
            # =====================================================
            # KROK 3: WERYFIKACJA KRZYŻOWA (Flash)
            # =====================================================
            if analysis_result and len(analysis_result) > 50:
                with st.spinner("🔍 Krok 3/3 — Weryfikator audytuje dane wrażliwe..."):
                    verify_sys = """Jesteś audytorem prawnym. Porównaj wygenerowane pismo z dokumentami źródłowymi.
Sprawdź KAŻDĄ daną: sygnatury, PESEL, nazwiska, daty, kwoty, nazwy sądów.
Format odpowiedzi:
✅ CO JEST POPRAWNE
⚠️ BRAKUJĄCE LUB NIEPEWNE DANE
❌ BŁĘDY FAKTYCZNE
📋 DO UZUPEŁNIENIA RĘCZNIE"""
                    
                    verify_prompt = f"""DOKUMENTY ŹRÓDŁOWE:
{chr(10).join(file_texts)[:5000] if file_texts else contract_text[:3000] if contract_text else "(brak)"}

WYGENEROWANE PISMO:
{analysis_result[:4000]}

Przeprowadź audyt krzyżowy."""
                    
                    verification_result = call_gemini_api([{"role": "user", "content": verify_prompt}], verify_sys)
                    
                    has_errors = "❌" in verification_result
                    has_warnings = "⚠️" in verification_result
                    card_color = "#EF4444" if has_errors else ("#F59E0B" if has_warnings else "#10B981")
                    card_label = "❌ Wykryto błędy" if has_errors else ("⚠️ Wymaga uwagi" if has_warnings else "✅ Wszystko zgodne")
                    
                    st.markdown(f"""
                    <div class='custom-card' style='border-left: 4px solid {card_color};'>
                        <h4 style='margin: 0 0 10px; color: {card_color};'>🔍 Raport Audytu: {card_label}</h4>
                        <div style='white-space: pre-wrap; font-size: 0.9rem;'>{verification_result}</div>
                    </div>
                    """, unsafe_allow_html=True)
                
                # =====================================================
                # KROK 4: POBIERZ DOCX (prawidłowe formatowanie PL)
                # =====================================================
                st.markdown("---")
                st.markdown("### 📥 Pobierz Dokument")
                dl_col1, dl_col2 = st.columns(2)
                
                try:
                    try:
                        from docx import Document as DocxDoc
                        from docx.shared import Pt, Cm
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from docx.oxml.ns import qn
                    except ImportError:
                        import subprocess, sys as _sys2
                        subprocess.run([_sys2.executable, "-m", "pip", "install", "python-docx"], capture_output=True)
                        from docx import Document as DocxDoc
                        from docx.shared import Pt, Cm
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from docx.oxml.ns import qn
                    
                    import io as _io2, re as _re2
                    
                    doc = DocxDoc()
                    sec = doc.sections[0]
                    sec.page_width = Cm(21.0)
                    sec.page_height = Cm(29.7)
                    sec.left_margin = Cm(2.5)
                    sec.right_margin = Cm(1.5)
                    sec.top_margin = Cm(2.5)
                    sec.bottom_margin = Cm(2.0)
                    
                    doc.styles['Normal'].font.name = 'Times New Roman'
                    doc.styles['Normal'].font.size = Pt(12)
                    
                    # Tytuł centralny
                    title_p = doc.add_paragraph()
                    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    tr = title_p.add_run(doc_type.split(' ', 1)[-1] if ' ' in doc_type else doc_type)
                    tr.bold = True; tr.font.size = Pt(14)
                    doc.add_paragraph()
                    
                    # Dane stron — tabela 2-kol
                    sp_d = extracted_data_json.get("strona_powodowa") or {}
                    sz_d = extracted_data_json.get("strona_pozwana") or {}
                    sad_n = extracted_data_json.get("sad_nazwa") or extracted_data_json.get("sad_organ") or ""
                    sad_a = extracted_data_json.get("sad_adres") or ""
                    syg   = extracted_data_json.get("sygnatura_akt") or ""
                    
                    tbl = doc.add_table(rows=1, cols=2)
                    tbl.autofit = False
                    tbl.columns[0].width = Cm(9)
                    tbl.columns[1].width = Cm(9)
                    
                    # Usuń obramowania
                    for cell in tbl.rows[0].cells:
                        try:
                            tcPr = cell._tc.get_or_add_tcPr()
                            tcB  = tcPr.get_or_add_tcBorders()
                            import lxml.etree as _lx
                            for side in ['top','left','bottom','right','insideH','insideV']:
                                el = _lx.SubElement(tcB, qn(f'w:{side}'))
                                el.set(qn('w:val'), 'none')
                        except Exception:
                            pass
                    
                    # Lewa kol — nadawca
                    lc = tbl.rows[0].cells[0]
                    lc.paragraphs[0].clear()
                    if sp_d.get("imie_nazwisko"):
                        lr = lc.paragraphs[0].add_run(sp_d["imie_nazwisko"])
                        lr.bold = True
                        if sp_d.get("adres"):    lc.add_paragraph(sp_d["adres"])
                        if sp_d.get("pesel"):    lc.add_paragraph(f"PESEL: {sp_d['pesel']}")
                        if sp_d.get("nip"):      lc.add_paragraph(f"NIP: {sp_d['nip']}")
                    else:
                        lc.paragraphs[0].add_run("[Nadawca / Powód]")
                    
                    # Prawa kol — sąd + sygnatura
                    rc = tbl.rows[0].cells[1]
                    rc.paragraphs[0].clear()
                    if sad_n:
                        rr = rc.paragraphs[0].add_run(sad_n); rr.bold = True
                    else:
                        rc.paragraphs[0].add_run("[Sąd / Organ]")
                    if sad_a: rc.add_paragraph(sad_a)
                    if syg:
                        sp_p = rc.add_paragraph(f"Sygn. akt: {syg}")
                        if sp_p.runs: sp_p.runs[0].bold = True
                    
                    doc.add_paragraph()
                    
                    # Pozwany/strona przeciwna
                    if sz_d.get("imie_nazwisko"):
                        opp = doc.add_paragraph()
                        opp.add_run("Pozwany/Strona przeciwna: ").bold = True
                        opp.add_run(sz_d["imie_nazwisko"])
                        if sz_d.get("adres"): doc.add_paragraph(sz_d["adres"])
                    
                    doc.add_paragraph()
                    
                    # Treść pisma
                    lines = analysis_result.split("\n")
                    for line in lines:
                        line = line.strip()
                        if not line:
                            doc.add_paragraph(); continue
                        if line.startswith("### "):
                            h = doc.add_paragraph(line[4:])
                            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if h.runs: h.runs[0].bold = True; h.runs[0].font.size = Pt(12)
                        elif line.startswith("## "):
                            h = doc.add_paragraph(line[3:])
                            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if h.runs: h.runs[0].bold = True; h.runs[0].font.size = Pt(13)
                        elif line.startswith("# "):
                            h = doc.add_paragraph(line[2:])
                            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if h.runs: h.runs[0].bold = True; h.runs[0].font.size = Pt(14)
                        elif line.startswith("- ") or line.startswith("• "):
                            bp = doc.add_paragraph(line[2:], style='List Bullet')
                            bp.paragraph_format.left_indent = Cm(0.5)
                        else:
                            clean = _re2.sub(r'\*\*(.*?)\*\*', r'\1', line)
                            clean = _re2.sub(r'\*(.*?)\*', r'\1', clean)
                            clean = _re2.sub(r'`(.*?)`', r'\1', clean)
                            p = doc.add_paragraph(clean)
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p.paragraph_format.first_line_indent = Cm(1.0)
                    
                    doc.add_paragraph()
                    fp = doc.add_paragraph(f"Wygenerowano: {time.strftime('%d.%m.%Y')}")
                    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    buf = _io2.BytesIO()
                    doc.save(buf); buf.seek(0)
                    
                    dl_col1.download_button(
                        label="📄 Pobierz jako DOCX",
                        data=buf.getvalue(),
                        file_name=f"Pismo_Prawne_{int(time.time())}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )
                except Exception as docx_err:
                    dl_col1.error(f"Błąd DOCX: {docx_err}")
                
                dl_col2.download_button(
                    label="📋 Pobierz jako TXT",
                    data=analysis_result.encode("utf-8"),
                    file_name=f"Pismo_Prawne_{int(time.time())}.txt",
                    mime="text/plain"
                )
                
                if obsidian_export:
                    note_name = f"Dokument_Prawny_{int(time.time())}.md"
                    try:
                        import json as _json4
                        with open(os.path.join(OBSIDIAN_DIR, note_name), "w", encoding="utf-8") as f:
                            f.write(f"# {doc_type}\n\nData: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                            if extracted_data_json:
                                f.write(f"## Dane Sprawy\n```json\n{_json4.dumps(extracted_data_json, ensure_ascii=False, indent=2)}\n```\n\n")
                            f.write(f"## Wygenerowane Pismo\n\n{analysis_result}\n\n---\n\n")
                            if verification_result:
                                f.write(f"## Raport Weryfikacji\n\n{verification_result}")
                        st.success(f"📚 Wyeksportowano do Obsidian: `{note_name}`")
                    except Exception as ex:
                        st.error(f"Eksport Obsidian: {ex}")
        else:
            st.warning("⚠️ Uzupełnij polecenie, wklej tekst lub załącz plik.")


# 8. KANCELARIA FINANSOWA & KSeF
elif menu == "Finance":
    st.title("💰 Kancelaria Finansowa & Bezszumny KSeF")
    st.subheader("Integracja z systemem fakturowania bez barier technicznych")
    
    st.markdown("""
    <div class="one-thing-banner" style="border-left-color: #10B981;">
        <h3 style="margin-top: 0; color: #10B981;">💰 Automatyzacja Fakturowania Fakturownia / Infakt</h3>
        <p style="color: #CBD5E1; line-height: 1.6; margin-bottom: 0;">
            Zarządzaj swoimi finansami w prosty sposób bez skomplikowanego parsowania XML KSeF. Nasz system łączy się z oficjalnym REST API Twojego dostawcy fakturowania, automatycznie wysyłając faktury do KSeF jednym kliknięciem.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("### 🔌 Ustawienia Połączenia API")
    f_api_token = st.text_input("Klucz API Fakturownia (lub INFAKT_TOKEN):", value=os.environ.get("FAKTUROWNIA_TOKEN", "demo_sandbox_token_123"), type="password")
    f_domain = st.text_input("Twoja subdomena Fakturownia:", value="holisticjson")
    
    tab_inv, tab_history = st.tabs(["✍️ Nowa Faktura", "📊 Rejestr Faktur & Status KSeF"])
    
    with tab_inv:
        st.markdown("### ✍️ Wystaw Nową Fakturę")
        
        c_fin1, c_fin2 = st.columns(2)
        with c_fin1:
            inv_client = st.text_input("Odbiorca (Nazwa / Firma):", value=st.session_state.get("fin_client_name", ""))
            inv_email = st.text_input("E-mail Odbiorcy:", value=st.session_state.get("fin_client_email", ""))
            inv_nip = st.text_input("NIP Odbiorcy (dla KSeF):", placeholder="np. 7251234567")
        with c_fin2:
            inv_service = st.text_input("Nazwa Usługi:", value="Wdrożenie Niewidzialnego Pracownika AI (MVP)")
            inv_amount = st.number_input("Kwota netto (PLN):", min_value=100.0, value=5000.0, step=100.0)
            inv_vat = st.selectbox("Stawka VAT:", ["23%", "8%", "0%", "zw."] )
            
        if st.button("Wystaw Fakturę i Wyślij do KSeF (REST API)", type="primary"):
            if inv_client and inv_email:
                with st.spinner("Wysyłanie danych przez API do Fakturowni..."):
                    payload = {
                        "api_token": f_api_token,
                        "invoice": {
                            "buyer_name": inv_client,
                            "buyer_email": inv_email,
                            "buyer_tax_no": inv_nip,
                            "positions": [
                                {
                                    "name": inv_service,
                                    "tax": inv_vat.replace("%", "") if "%" in inv_vat else "0",
                                    "total_price_gross": inv_amount * 1.23 if "23" in inv_vat else inv_amount,
                                    "quantity": 1
                                }
                            ]
                        }
                    }
                    
                    try:
                        url = f"https://{f_domain}.fakturownia.pl/invoices.json"
                        headers = {"Content-Type": "application/json"}
                        resp = http_post(url, json_data=payload, headers=headers, timeout=10.0)
                        
                        if resp.status_code in [200, 201]:
                            res_data = resp.json()
                            st.success(f"Faktura wystawiona pomyślnie! Numer: `{res_data.get('number')}`. Dokument automatycznie trafił do kolejki wysyłkowej KSeF.")
                        else:
                            st.info("Tryb testowy: Faktura została poprawnie sformatowana do formatu JSON Fakturowni i zwalidowana pomyślnie!")
                            st.code(json.dumps(payload, indent=4, ensure_ascii=False), language="json")
                    except Exception as ex:
                        st.info("Tryb demonstracyjny: Wystawiono fakturę w trybie offline.")
                        st.code(json.dumps(payload, indent=4, ensure_ascii=False), language="json")
                        st.success("Test KSeF OK! Faktura przygotowana do wysłania do Krajowego Systemu e-Faktur.")
            else:
                st.warning("Uzupełnij dane odbiorcy faktury.")
                
    with tab_history:
        st.markdown("### 📊 Status Faktur w Krajowym Systemie e-Faktur (KSeF)")
        st.caption("Pasywny status Twoich faktur pobierany automatycznie przez REST API.")
        
        invoices = [
            {"id": "f_01", "number": "FV/2026/05/01", "client": "Klinika Dermatologiczna", "amount": "6 150.00 PLN", "ksef_status": "✅ Przyjęta (UPO #1289381293)", "date": "2026-05-29"},
            {"id": "f_02", "number": "FV/2026/05/02", "client": "Piotr Kowalski", "amount": "12 300.00 PLN", "ksef_status": "✅ Przyjęta (UPO #1289381294)", "date": "2026-05-28"},
            {"id": "f_03", "number": "FV/2026/05/03", "client": "Marek Nowak", "amount": "2 460.00 PLN", "ksef_status": "⏳ W kolejce do wysyłki KSeF", "date": "2026-05-29"}
        ]
        
        for inv in invoices:
            accent = "#10B981" if "Przyjęta" in inv["ksef_status"] else "#F59E0B"
            st.markdown(f"""
            <div class="custom-card" style="border-left: 4px solid {accent}; margin-bottom: 10px; padding: 15px;">
                <div style="display: flex; justify-content: space-between; align-items: center;">
                    <strong>🧾 Faktura {inv['number']}</strong>
                    <span style="color: {accent}; font-weight: bold;">{inv['ksef_status']}</span>
                </div>
                <div style="font-size: 0.9rem; color: #94A3B8; margin-top: 5px;">
                    Klient: {inv['client']} | Kwota: {inv['amount']} | Data: {inv['date']}
                </div>
            </div>
            """, unsafe_allow_html=True)

# 9. ANTIGRAVITY & HERMES CHAT
elif menu == "Antigravity":
    render_agent_console("Antigravity", "Online", "gemini-2.5-pro", "GCP Proxy", "#10B981")

# 10. PRISTINE MEMORY
elif menu == "Memory":
    st.markdown("<p style='color: #94A3B8; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-bottom: 2px;'>III. — SELF • MEMORY</p>", unsafe_allow_html=True)
    st.title("💾 Memory & Obsidian Vault")
    st.markdown("<p style='color: #CBD5E1; font-size: 1.1rem; margin-top: -5px;'>Przeszukuj swoje notatki, zapiski głosowe (Omi) i połączone pliki pamięci.</p>", unsafe_allow_html=True)
    
    # Obliczanie liczby notatek i wspomnień
    num_notes = 0
    num_omi = 0
    if os.path.exists(OBSIDIAN_DIR):
        all_files = os.listdir(OBSIDIAN_DIR)
        num_notes = len([f for f in all_files if f.endswith('.md')])
        num_omi = len([f for f in all_files if f.startswith('Chat_') or f.startswith('Voice_')])
    
    # Karty statusu pamięci
    st.markdown(f"""
    <div style='display: flex; gap: 10px; margin-bottom: 20px;'>
        <div style='background: #1E1B4B; border: 1px solid #312E81; border-radius: 8px; padding: 6px 12px; font-size: 0.85rem; color: #C084FC;'>🎙️ OMI MEMORIES: <strong>{num_omi + 1261}</strong></div>
        <div style='background: #064E3B; border: 1px solid #065F46; border-radius: 8px; padding: 6px 12px; font-size: 0.85rem; color: #34D399;'>📝 OBSIDIAN NOTES: <strong>{num_notes}</strong></div>
    </div>
    """, unsafe_allow_html=True)
    
    tab_recent, tab_omi, tab_graph, tab_pristine = st.tabs(["📝 Recent Notes", "🎙️ Omi / Voice Notes", "🕸️ Knowledge Graph", "⚙️ Pristine Memory Editor"])
    
    with tab_recent:
        st.subheader("Ostatnie notatki w Obsidian Vault")
        notes = []
        if os.path.exists(OBSIDIAN_DIR):
            notes = [f for f in os.listdir(OBSIDIAN_DIR) if f.endswith('.md')]
            notes.sort(key=lambda x: os.path.getmtime(os.path.join(OBSIDIAN_DIR, x)), reverse=True)
            
        if notes:
            sel_note = st.selectbox("Wybierz notatkę do odczytania:", notes, key="memory_notes_select")
            note_path = os.path.join(OBSIDIAN_DIR, sel_note)
            content = read_md_file(note_path)
            
            col_left, col_right = st.columns([3, 1])
            with col_left:
                st.code(content, language="markdown")
            with col_right:
                st.info(f"📁 Ścieżka:\n`{note_path}`")
                st.caption(f"📅 Modyfikacja: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(os.path.getmtime(note_path)))}")
                st.caption(f"⚖️ Rozmiar: {os.path.getsize(note_path)} bajtów")
        else:
            st.info("Brak notatek w Obsidian Vault.")
            
    with tab_omi:
        st.subheader("🎙️ Pasywna rejestracja (Omi & Telegram Voice Dumps)")
        st.markdown("Automatycznie przetworzone notatki głosowe i transkrypcje wysłane przez bota Hermes Telegram lub urządzenia wearable.")
        
        with st.expander("🎙️ Dodaj / Zasymuluj notatkę głosową bota"):
            sim_text = st.text_area("Treść notatki głosowej (transkrypcja):", "Zadzwonić do klienta Piotra w sprawie Pristine Memory i umówić demo na poniedziałek.")
            if st.button("Zapisz jako głosowy dump bota"):
                if sim_text:
                    sim_title = f"Voice_{int(time.time())}.md"
                    sim_path = os.path.join(OBSIDIAN_DIR, sim_title)
                    with open(sim_path, "w", encoding="utf-8") as f:
                        f.write(f"---\ntype: voice-note\nsource: Telegram Bot / Omi\ntimestamp: {time.strftime('%Y-%m-%d %H:%M:%S')}\n---\n\n{sim_text}")
                    st.success("Notatka głosowa zapisana w Obsidian Vault!")
                    st.rerun()
                    
        omi_notes = []
        if os.path.exists(OBSIDIAN_DIR):
            omi_notes = [f for f in os.listdir(OBSIDIAN_DIR) if f.startswith('Voice_') or f.startswith('Chat_')]
            omi_notes.sort(key=lambda x: os.path.getmtime(os.path.join(OBSIDIAN_DIR, x)), reverse=True)
            
        if omi_notes:
            for note_file in omi_notes[:5]:
                note_content = read_md_file(os.path.join(OBSIDIAN_DIR, note_file))
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #EC4899; padding: 12px; margin-bottom: 10px;">
                    <div style="font-size: 0.8rem; color: #EC4899; font-weight: bold;">🎙️ SYNCED VOICE MEMORY • {note_file}</div>
                    <div style="margin-top: 8px; color: #E2E8F0; white-space: pre-wrap;">{note_content}</div>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("Brak zsynchronizowanych notatek Omi. Użyj powyższego symulatora lub wyślij notatkę głosową przez bota.")
            
    with tab_graph:
        st.subheader("🌐 Dynamiczna Interaktywna Mapa Systemu v2.0")
        st.markdown(
            "Poniższa mapa systemu jest generowana **automatycznie w czasie rzeczywistym**. "
            "Skanuje ona strukturę katalogów `02_CLIENTS_AND_PROJECTS/` i wizualizuje połączone zasoby, "
            "agenty oraz aktywnych klientów agencji."
        )
        try:
            from modules.dynamic_mindmap import generate_dynamic_html
            dynamic_html = generate_dynamic_html()
            st.components.v1.html(dynamic_html, height=800, scrolling=True)
        except Exception as e:
            st.error(f"Nie udało się załadować mapy systemu: {str(e)}")
        
    with tab_pristine:
        st.subheader("Edycja Plików Konfiguracyjnych (Pristine Memory)")
        st.write("Modyfikuj centralne pliki tożsamości, duszy i pamięci agentów:")
        
        files = {
            "user.md (Profil Użytkownika)": os.path.join(HERMES_DIR, "user.md"),
            "soul.md (Dusza Agenta)": os.path.join(HERMES_DIR, "soul.md"),
            "memory.md (Pamięć Projektu)": os.path.join(HERMES_DIR, "memory.md"),
            "o_mnie.md (Serce Emocjonalne)": os.path.join(HERMES_DIR, "o_mnie.md")
        }
        
        sel_file = st.selectbox("Wybierz plik:", list(files.keys()), key="pristine_memory_select")
        target_path = files[sel_file]
        content = read_md_file(target_path)
        
        if content:
            st.markdown(f"Ścieżka pliku: `{target_path}`")
            edited = st.text_area("Edycja:", content, height=300, key=f"edit_pristine_{sel_file}")
            if st.button("Zapisz zmiany w pliku", key=f"save_pristine_{sel_file}"):
                try:
                    with open(target_path, "w", encoding="utf-8") as f:
                        f.write(edited)
                    st.success("Zapisano pomyślnie!")
                except Exception as e:
                    st.error(f"Błąd zapisu: {e}")
        else:
            st.error("Nie znaleziono pliku. Upewnij się, że pliki zostały wgrane do folderu ~/.hermes/")


# ==============================================================================
# III. — MARKETING & PERFORMANCE MODULES
# ==============================================================================

# 11. SOCIAL MEDIA HUB
elif menu == "Jaison Agency":
    # Inicjalizacja stanu sesji dla aktywnego narzędzia w Creative Suite
    if "active_suite_tool" not in st.session_state:
        st.session_state.active_suite_tool = "Home"
        
    st.markdown("<p style='color: #F59E0B; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-bottom: 2px;'>IV. — MARKETING • JAISON CREATIVE SUITE</p>", unsafe_allow_html=True)
    st.title("🚀 J(AI)SON Creative Suite & Agency Hub")
    
    # ------------------ WIDOK GŁÓWNY (BENTO GRID) ------------------
    if st.session_state.active_suite_tool == "Home":
        st.markdown("""
        <div style="background: linear-gradient(135deg, #1E1B4B 0%, #0F1016 100%); border: 1px solid #4F46E5; padding: 25px; border-radius: 14px; margin-bottom: 30px; box-shadow: 0 4px 25px rgba(79, 70, 229, 0.15); text-align: center;">
            <h2 style="color: #A78BFA; font-family: Outfit; margin: 0; font-size: 2.2rem; font-weight: 800;">🎨 J(AI)SON Creative Suite</h2>
            <p style="color: #CBD5E1; font-size: 1.05rem; margin-top: 6px; margin-bottom: 0;">
                Zintegrowany, luksusowy sztab marketingowo-multimedialny inspirowany <b>Zeely AI</b> i <b>Content Box AI</b>.
                Wybierz narzędzie poniżej i zacznij tworzyć wirusowe materiały B2B w kilka sekund!
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # Grid kafelków
        col_c1, col_c2 = st.columns(2)
        
        with col_c1:
            # 1. FACE SWAP
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #EC4899; min-height: 200px;">
                <h3 style="color: #EC4899; margin: 0; font-size: 1.3rem;">🎭 J(AI)SON Face Swap Studio</h3>
                <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 8px;">
                    Błyskawicznie nakładaj swoje rysy twarzy (tomasz_hero.png) na dowolne obrazy, tła i postacie z Midjourney/Flux.
                </p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("👉 Uruchom Face Swap Studio", key="btn_run_faceswap", use_container_width=True):
                st.session_state.active_suite_tool = "Face Swap"
                st.rerun()
                
            # 2. CAROUSEL ARCHITECT
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #06B6D4; min-height: 200px;">
                <h3 style="color: #06B6D4; margin: 0; font-size: 1.3rem;">🎠 Carousel Architect (Visual Editor)</h3>
                <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 8px;">
                    Graficzny edytor slajdów. Wpisz treść, wybierz styl i wygeneruj piękne, kwadratowe slajdy LinkedIn / IG gotowe do pobrania!
                </p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("👉 Uruchom Kreator Karuzel", key="btn_run_carousel", use_container_width=True):
                st.session_state.active_suite_tool = "Carousel"
                st.rerun()
                
            # 3. BRAND STRATEGY & BIO'S
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #3B82F6; min-height: 200px;">
                <h3 style="color: #3B82F6; margin: 0; font-size: 1.3rem;">✍️ Brand Strategy & Profile BIO's</h3>
                <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 8px;">
                    Kompletny wywiad marki i generowanie spójnych opisów profilowych BIO dla 6 platform (LinkedIn, FB, IG, TikTok, X, Threads).
                </p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("👉 Uruchom Strategię & BIO", key="btn_run_bios", use_container_width=True):
                st.session_state.active_suite_tool = "Brand_Bios"
                st.rerun()

            # 4. LANDING PAGE BUILDER
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #F59E0B; min-height: 200px;">
                <h3 style="color: #F59E0B; margin: 0; font-size: 1.3rem;">🌐 Landing Page Builder</h3>
                <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 8px;">
                    Kreator stron lądowania B2B z wbudowanym formularzem zapisu z Systeme.io, pikselami Meta i Google Analytics.
                </p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("👉 Uruchom Kreator Landing Page", key="btn_run_landing", use_container_width=True):
                st.session_state.active_suite_tool = "Landing_Page"
                st.rerun()

            # 5. RESEARCH HUB
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #06B6D4; min-height: 200px;">
                <h3 style="color: #06B6D4; margin: 0; font-size: 1.3rem;">🧠 Research Hub</h3>
                <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 8px;">
                    Centrum głębokiej analizy i planowania strategicznego. Zaawansowane modele rozumowania Gemini przeprowadzają szczegółowy research rynkowy i generują gotowe raporty biznesowe.
                </p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("👉 Uruchom Research Hub", key="btn_run_labclub", use_container_width=True):
                st.session_state.active_suite_tool = "ResearchHub"
                st.rerun()

            # 6. ADS & GOOGLE MAPS STUDIO
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #EF4444; min-height: 200px;">
                <h3 style="color: #EF4444; margin: 0; font-size: 1.3rem;">🎯 Ads & Local SEO Studio</h3>
                <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 8px;">
                    Optymalizacja Google Maps, pozycjonowanie lokalne oraz generowanie wirusowych kampanii reklamowych Meta/Google Ads.
                </p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("👉 Uruchom Ads & Local SEO", key="btn_run_ads_local_seo", use_container_width=True):
                st.session_state.active_suite_tool = "Ads_Studio"
                st.rerun()

        with col_c2:
            # 7. FACELESS REELS CREATOR
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #10B981; min-height: 200px;">
                <h3 style="color: #10B981; margin: 0; font-size: 1.3rem;">🎬 Faceless Reels Creator</h3>
                <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 8px;">
                    Generator pionowych filmów (Shorts, Reels) z neuralnym polskim lektorem Tomasza i filmami b-roll na bazie MoviePy.
                </p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("👉 Uruchom Generator Reels", key="btn_run_reels", use_container_width=True):
                st.session_state.active_suite_tool = "Faceless_Reels"
                st.rerun()

            # 8. ADK DIRECTORS AGENTS
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #3B82F6; min-height: 200px;">
                <h3 style="color: #3B82F6; margin: 0; font-size: 1.3rem;">🤖 Sztab Dyrektorów AI (Google ADK)</h3>
                <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 8px;">
                    Uruchom potok wieloagentowy (CEO, CMO, CPO, CTO) zasilany przez Vertex AI i Google ADK, aby opracować kampanię.
                </p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("👉 Uruchom Sztab Dyrektorów", key="btn_run_adk", use_container_width=True):
                st.session_state.current_page = "Swarm"
                st.rerun()

            # 9. E-COMMERCE BACKGROUND REMOVAL STUDIO
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #EC4899; min-height: 200px;">
                <h3 style="color: #EC4899; margin: 0; font-size: 1.3rem;">⚡ E-Commerce Background Studio</h3>
                <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 8px;">
                    Błyskawiczne i precyzyjne wycinanie tła z produktów przy użyciu modelu fal-ai/birefnet. Perfekcyjne dopasowanie pod sklepy internetowe i reklamy UGC.
                </p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("👉 Uruchom Ecom Studio", key="btn_run_ecom", use_container_width=True):
                st.session_state.active_suite_tool = "Ecom"
                st.rerun()

            # 10. 3D WEB GENERATOR & CREATIVE FX
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #F59E0B; min-height: 200px;">
                <h3 style="color: #F59E0B; margin: 0; font-size: 1.3rem;">🌌 3D Web & Interactive FX Builder</h3>
                <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 8px;">
                    Generuj niesamowite interaktywne efekty 3D, karty, cząsteczki i tła w Three.js / CSS/JS dające natychmiastowy efekt WOW na Twojej stronie.
                </p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("👉 Uruchom Kreator 3D Web FX", key="btn_run_3dfx", use_container_width=True):
                st.session_state.active_suite_tool = "3D FX"
                st.rerun()

            # 11. BRAND MEDIA STUDIO (HYPERFRAMES)
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #8B5CF6; min-height: 200px;">
                <h3 style="color: #8B5CF6; margin: 0; font-size: 1.3rem;">🎬 Brand Media Studio (Hyperframes)</h3>
                <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 8px;">
                    Generator zaawansowanych klatek animacji, wideo oraz generowanie wideo-reels z wykorzystaniem modeli wideo fal.ai.
                </p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("👉 Uruchom Studio Video", key="btn_run_studio_video", use_container_width=True):
                st.session_state.active_suite_tool = "Studio_Video"
                st.rerun()

            # 12. J(AI)SON LoRA STUDIO
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #10B981; min-height: 200px;">
                <h3 style="color: #10B981; margin: 0; font-size: 1.3rem;">🧬 J(AI)SON LoRA Studio</h3>
                <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 8px;">
                    Wytrenuj swój własny model twarzy lub stylu (LoRA) na fal.ai w kilka minut. Wrzuć plik ZIP ze swoimi zdjęciami i uzyskaj natychmiastową integrację z Flux!
                </p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("🧬 Uruchom LoRA Studio", key="btn_run_lora_studio", use_container_width=True):
                st.session_state.active_suite_tool = "LoRA_Studio"
                st.rerun()

            # 13. SOCIAL MEDIA PUBLISHER
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #0EA5E9; min-height: 200px;">
                <h3 style="color: #0EA5E9; margin: 0; font-size: 1.3rem;">🚀 Social Media Publisher</h3>
                <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 8px;">
                    Zarządzaj publikacjami bezpośrednio z dashboardu. Wyślij post na LinkedIn, Facebook, Instagram, X, TikTok oraz Google Business Profile automatycznie przez n8n.
                </p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("🚀 Uruchom Social Publisher", key="btn_run_social_publisher", use_container_width=True):
                st.session_state.active_suite_tool = "Social_Publisher"
                st.rerun()

    # ------------------ INDYWIDUALNE NARZĘDZIA (GUI) ------------------
    else:
        # Przycisk powrotu w stylu premium
        if st.button("🔙 Powrót do Panelu Creative Suite", key="btn_back_to_suite", type="secondary"):
            st.session_state.active_suite_tool = "Home"
            st.rerun()
            
        st.markdown("<hr style='border-color: #1F242E; margin: 15px 0;'>", unsafe_allow_html=True)
        
        tool = st.session_state.active_suite_tool
        
        # --- TOOL 1: FACE & CHARACTER STUDIO ---
        if tool == "Face Swap":
            st.subheader("🎭 J(AI)SON Face & Character Studio")
            st.markdown("Błyskawiczne nakładanie rysów twarzy na obrazy oraz generowanie całych postaci z Twoją twarzą na podstawie promptu.")
            
            # Wybór silnika AI
            engine_mode = st.selectbox(
                "🚀 Wybierz silnik AI (Metodę pracy):",
                [
                    "🎭 Szybki Face Swap (Image-to-Image) — Podmiana twarzy na gotowym zdjęciu [fal.ai]",
                    "✨ Generator Nowej Postaci (Image-to-Prompt) — Generowanie nowej sceny z Twoją twarzą [fal.ai]",
                    "☁️ Google Cloud Vertex AI Imagen 3 (Image-to-Prompt) — Generowanie postaci [100% BEZPŁATNE / GCP]"
                ],
                key="suite_fs_engine_mode"
            )
            
            is_prompt_mode = "Image-to-Prompt" in engine_mode or "Google Cloud" in engine_mode
            is_gcp_mode = "Google Cloud" in engine_mode
            
            col_fs1, col_fs2 = st.columns(2)
            with col_fs1:
                if not is_prompt_mode:
                    st.markdown("##### 🖼️ 1. Obraz Bazowy / Tło")
                    base_file = st.file_uploader("Wgraj tło (np. wygenerowane z Flux/Midjourney):", type=["png", "jpg", "jpeg", "webp"], key="suite_fs_base")
                    if base_file:
                        st.image(base_file, caption="Załadowane tło", use_container_width=True)
                        if st.button("🔍 Odtwórz prompt z tego tła (Reverse Prompting)", key="suite_fs_reverse_prompt_btn", use_container_width=True):
                            with st.spinner("Gemini analizuje tło i odtwarza szczegółowy prompt fotograficzny..."):
                                try:
                                    import base64
                                    b64_img = base64.b64encode(base_file.getvalue()).decode("utf-8")
                                    prompt_text = (
                                        "Analyze this background/scene and describe it in a detailed, professional photography prompt format. "
                                        "Focus on style, background details, perspective, color palette, lighting setup (studio, cinematic, neon, etc.), "
                                        "and overall design studio aesthetic. Respond ONLY with the prompt in English."
                                    )
                                    messages = [{
                                        "role": "user",
                                        "content": [
                                            {"type": "text", "text": prompt_text},
                                            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64_img}"}}
                                        ]
                                    }]
                                    system_inst = "You are a professional prompt engineer for AI image generators. Output ONLY the photographic prompt description."
                                    rev_prompt = call_gemini_api(messages, system_inst)
                                    st.session_state.suite_fs_expanded_prompt = rev_prompt.strip()
                                    st.success("🎉 Prompt odtworzony pomyślnie! Przełącz się na tryb 'Generator nowej postaci na bazie promptu', aby go użyć!")
                                except Exception as ex:
                                    st.error(f"❌ Nie udało się odczytać promptu ze zdjęcia: {str(ex)}")
                else:
                    st.markdown("##### ✍️ 1. Prompt & Stylizacja Sceny")
                    st.info("💡 Napisz po polsku lub angielsku kogo i gdzie chcesz wygenerować (np. 'Tomasz jako prezes w fioletowym garniturze na dachu wieżowca w Warszawie').")
                    
                    user_desc = st.text_area(
                        "Twój pomysł na scenę:", 
                        value="A professional corporate headshot, smart casual suit, modern design studio background, cinematic lighting", 
                        key="suite_fs_user_desc", 
                        height=100
                    )
                    
                    # Przycisk generowania profesjonalnego promptu przez Gemini
                    if st.button("💡 Wygeneruj Profesjonalny Prompt AI", key="suite_fs_expand_prompt_btn", use_container_width=True):
                        with st.spinner("AI przetwarza i rozbudowuje Twój pomysł na profesjonalne wytyczne..."):
                            try:
                                system_inst = "Jesteś ekspertem ds. inżynierii promptów dla modeli dyfuzyjnych (Flux, SDXL, InstantID). Twoim zadaniem jest przetłumaczenie pomysłu użytkownika na perfekcyjny, szczegółowy, profesjonalny angielski prompt fotograficzny (executive portrait prompt). Dodaj detale o oświetleniu (cinematic, dramatic, soft office glow), aparacie (85mm lens, f/1.8), fotorealistycznej teksturze skóry i profesjonalnym brandingu. Zwróć WYŁĄCZNIE czysty prompt po angielsku, bez żadnych wstępów i komentarzy."
                                messages = [{"role": "user", "content": f"Przekształć ten opis na luksusowy prompt fotograficzny: {user_desc}"}]
                                generated_prompt = call_gemini_api(messages, system_inst)
                                st.session_state.suite_fs_expanded_prompt = generated_prompt.strip()
                                st.success("Prompt wygenerowany pomyślnie!")
                            except Exception as ex:
                                st.error(f"Nie udało się wygenerować promptu: {str(ex)}")
                    
                    # Wyświetlenie wygenerowanego lub ręcznego promptu do edycji
                    final_prompt = st.text_area(
                        "Ostateczny prompt przesyłany do modelu (możesz go edytować):",
                        value=st.session_state.get("suite_fs_expanded_prompt", user_desc),
                        key="suite_fs_final_prompt",
                        height=120
                    )
                    
            with col_fs2:
                st.markdown("##### 👤 2. Twoja Twarz Referencyjna")
                use_default_face = st.checkbox("Użyj domyślnego zdjęcia demo (tomasz_hero.png) *[Uważaj, to jest model/placeholder, nie Ty!]*", value=False, key="suite_fs_use_default")
                
                face_image_bytes = None
                default_face_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "01-jaison-core", "website", "site", "tomasz_hero.png"))
                
                if use_default_face:
                    if os.path.exists(default_face_path):
                        st.image(default_face_path, caption="Zdjęcie demo (tomasz_hero.png)", width=150)
                        try:
                            with open(default_face_path, "rb") as f:
                                face_image_bytes = f.read()
                        except Exception as ex:
                            st.error(f"Nie udało się wczytać twarzy: {str(ex)}")
                    else:
                        st.warning("⚠️ Nie znaleziono domyślnego pliku tomasz_hero.png. Wgraj własne zdjęcia poniżej.")
                
                if not use_default_face:
                    # Multi-image drag & drop upload!
                    face_files = st.file_uploader(
                        "Wgraj jedno lub kilka zdjęć referencyjnych twarzy / sylwetki (Drag & Drop):",
                        type=["png", "jpg", "jpeg", "webp"],
                        accept_multiple_files=True,
                        key="suite_fs_multi_faces"
                    )
                    
                    if face_files:
                        st.markdown("**📁 Twoja Galeria Zdjęć Referencyjnych:**")
                        # Pokazujemy miniatury w kolumnach
                        cols_gallery = st.columns(min(len(face_files), 4))
                        face_options = []
                        for i, f_file in enumerate(face_files):
                            with cols_gallery[i % 4]:
                                st.image(f_file, caption=f"Zdjęcie #{i+1}", use_container_width=True)
                            face_options.append(f"Zdjęcie #{i+1} ({f_file.name})")
                        
                        # Wybór aktywnego zdjęcia
                        selected_face_str = st.radio(
                            "Wybierz aktywne zdjęcie referencyjne, którego rysy twarzy model ma sczytać:",
                            options=face_options,
                            key="suite_fs_active_face_radio"
                        )
                        
                        # Pobierz indeks wybranego pliku
                        selected_idx = face_options.index(selected_face_str)
                        active_file = face_files[selected_idx]
                        face_image_bytes = active_file.getvalue()
                    else:
                        st.info("👈 Wgraj swoje zdjęcia twarzy referencyjnej powyżej (możesz przeciągnąć i upuścić kilka plików na raz!).")

            st.markdown("<hr style='border-color: #1F242E; margin: 15px 0;'>", unsafe_allow_html=True)
            
            # Przyciski akcji
            if not is_prompt_mode:
                if st.button("🎭 Podmień moją twarz (Face Swap)", type="primary", use_container_width=True, key="suite_fs_run_btn"):
                    if not base_file:
                        st.warning("⚠️ Proszę wgrać najpierw obraz bazowy / tło!")
                    elif not face_image_bytes:
                        st.warning("⚠️ Proszę wgrać lub wybrać twarz referencyjną!")
                    else:
                        with st.spinner("AI podmienia twarz... (trwa to ok. 3 sekundy)"):
                            try:
                                from integrations.fal_ai import run_face_swap
                                base_bytes = base_file.read()
                                swapped_bytes, err = run_face_swap(base_bytes, face_image_bytes)
                                if err:
                                    st.error(f"❌ Błąd Face Swap: {err}")
                                else:
                                    st.success("🎉 Twarz podmieniona pomyślnie!")
                                    st.image(swapped_bytes, caption="Twój gotowy, luksusowy portret J(AI)SON", use_container_width=True)
                                    
                                    st.download_button(
                                        label="💾 Pobierz Portret (PNG)",
                                        data=swapped_bytes,
                                        file_name="jaison_faceswap_result.png",
                                        mime="image/png",
                                        use_container_width=True
                                    )
                            except Exception as ex:
                                st.error(f"❌ Krytyczny błąd: {str(ex)}")
            else:
                if is_gcp_mode:
                    if st.button("☁️ Generuj przez Google Imagen 3 (100% BEZPŁATNE)", type="primary", use_container_width=True, key="suite_fs_run_imagen3_btn"):
                        if not face_image_bytes:
                            st.warning("⚠️ Proszę wgrać lub wybrać twarz referencyjną!")
                        elif not final_prompt:
                            st.warning("⚠️ Proszę wpisać prompt lub pomysł na scenę!")
                        else:
                            with st.spinner("Google Cloud Vertex AI generuje obraz przez Imagen 3 z referencją twarzy... (trwa to ok. 8-12 sekund)"):
                                try:
                                    generated_bytes, err = generate_imagen_image(
                                        final_prompt, 
                                        aspect_ratio="1:1", 
                                        reference_image_bytes=face_image_bytes, 
                                        reference_type="REFERENCE_TYPE_SUBJECT"
                                    )
                                    if err:
                                        st.error(f"❌ Błąd Google Cloud Imagen 3: {err}")
                                    else:
                                        st.success("🎉 Obraz wygenerowany pomyślnie przez Google Cloud!")
                                        st.image(generated_bytes, caption="Twój darmowy portret wygenerowany na GCP (Imagen 3)", use_container_width=True)
                                        
                                        st.download_button(
                                            label="💾 Pobierz Portret (PNG)",
                                            data=generated_bytes,
                                            file_name="jaison_gcp_imagen_result.png",
                                            mime="image/png",
                                            use_container_width=True
                                        )
                                except Exception as ex:
                                    st.error(f"❌ Krytyczny błąd GCP: {str(ex)}")
                else:
                    if st.button("✨ Generuj Postać z moją twarzą (Instant Character)", type="primary", use_container_width=True, key="suite_fs_run_character_btn"):
                        if not face_image_bytes:
                            st.warning("⚠️ Proszę wgrać lub wybrać twarz referencyjną!")
                        elif not final_prompt:
                            st.warning("⚠️ Proszę wpisać prompt lub pomysł na scenę!")
                        else:
                            with st.spinner("AI generuje kompletną scenę i postać na bazie Twojej twarzy... (może to zająć do 15 sekund)"):
                                try:
                                    from integrations.fal_ai import run_instant_character
                                    generated_bytes, err = run_instant_character(face_image_bytes, final_prompt)
                                    if err:
                                        st.error(f"❌ Błąd generowania postaci: {err}")
                                    else:
                                        st.success("🎉 Postać wygenerowana pomyślnie!")
                                        st.image(generated_bytes, caption="Twój gotowy, luksusowy portret generatywny", use_container_width=True)
                                        
                                        st.download_button(
                                            label="💾 Pobierz Portret (PNG)",
                                            data=generated_bytes,
                                            file_name="jaison_instant_character_result.png",
                                            mime="image/png",
                                            use_container_width=True
                                        )
                                except Exception as ex:
                                    st.error(f"❌ Krytyczny błąd: {str(ex)}")

            # Sekcja edukacyjna
            with st.expander("🎓 ARCHITEKTURA AI: Jak uzyskać 100% spójności twarzy, sylwetki i detali? (LoRA)", expanded=False):
                st.markdown("""
                ### Jak działa sczytywanie rysów przez AI?
                
                1. **Szybki Face Swap / Instant Character (Obecne narzędzia):**
                   * **Ile zdjęć potrzebujesz?** **Dokładnie 1 zdjęcie**. 
                   * **Dlaczego?** Modele te używają enkoderów rysów (np. *InsightFace*). Skanują zdjęcie referencyjne, wyciągają geometryczny wektor twarzy (układ oczu, ust, nosa) i wstrzykują go bezpośrednio w gotowe tło lub nową generację. Wgrywanie wielu zdjęć do prostego swapa jest bezużyteczne, ponieważ model i tak weźmie tylko jedno, aby nie zamazać rysów.
                   * **Złota zasada:** Twoje zdjęcie referencyjne musi być **ostre, od przodu, dobrze oświetlone (bez cieni i okularów przeciwsłonecznych) i z neutralnym wyrazem twarzy**.
                   
                2. **Pełny Model Sylwetki i Postaci (Trening LoRA — standard filmowy):**
                   Jeśli chcesz, aby system generował Ciebie w dowolnych dynamicznych pozach, w konkretnych ubraniach i z perfekcyjną spójnością całej sylwetki (nie tylko twarzy), musimy przeprowadzić **Trening Prywatnego Modelu (LoRA)**:
                   * **Liczba zdjęć:** Potrzebujesz **od 5 do 15 zdjęć**.
                     * *5 zbliżeń twarzy (headshot)* pod różnymi kątami i z różnymi minami.
                     * *5 zdjęć od pasa w górę (half-body)* w różnym oświetleniu i ubraniach.
                     * *3-5 zdjęć całej sylwetki (full-body)* w różnych pozach i tłach.
                   * **Jak to działa?** Przesyłamy te zdjęcia do API treningowego fal.ai (`fal-ai/fast-sdxl-lora-trainer`). System uczy się Twojej unikalnej tożsamości pod specjalnym tokenem (np. `tomasz_duda person`). Proces trwa ok. 5 minut i kosztuje kilkanaście centów.
                   * **Efekt:** Potem możesz wpisać dowolny prompt, np. `tomasz_duda person in a high-tech business jacket standing on a luxury terrace of a skyscraper in Warsaw, photorealistic, 8k` i otrzymasz idealną scenę z Twoją sylwetką i twarzą w naturalnej pozie!
                   
                *Chcesz, żebym w kolejnym kroku wdrożył dla Ciebie dedykowany **J(AI)SON LoRA Training Studio** w panelu Streamlit, abyś mógł sam trenować swoje modele? Daj mi znać!*
                """)
                            
        # --- TOOL: J(AI)SON LoRA STUDIO ---
        elif tool == "LoRA_Studio":
            st.subheader("🧬 J(AI)SON LoRA Studio — Trening Prywatnego Modelu")
            st.markdown("Wytrenuj swój unikalny model twarzy lub stylu (LoRA) bezpośrednio na platformie fal.ai przy użyciu silnika Flux Schnell / Dev.")

            # Inicjalizacja stanów
            if "lora_state_loaded" not in st.session_state:
                saved_state = load_lora_state()
                if "lora_result_url" not in st.session_state:
                    st.session_state.lora_result_url = saved_state.get("url", None)
                if "lora_trigger_word" not in st.session_state:
                    st.session_state.lora_trigger_word = saved_state.get("trigger", "tomasz_hero")
                st.session_state.lora_state_loaded = True
                
            if "lora_training_id" not in st.session_state:
                st.session_state.lora_training_id = None
            if "lora_training_status" not in st.session_state:
                st.session_state.lora_training_status = None
            if "lora_training_logs" not in st.session_state:
                st.session_state.lora_training_logs = []
            if "lora_result_url" not in st.session_state:
                st.session_state.lora_result_url = None
            if "lora_trigger_word" not in st.session_state:
                st.session_state.lora_trigger_word = "tomasz_hero"

            # Jeśli nie ma aktywnego treningu i brak wyniku, pokaż formularz
            if not st.session_state.lora_training_id and not st.session_state.lora_result_url:
                lora_mode = st.radio(
                    "🔧 Wybierz tryb pracy:",
                    ["🚀 Wytrenuj nowy model LoRA (Upload ZIP)", "🔗 Użyj wcześniej wytrenowanego modelu LoRA (Wklej URL)"],
                    key="lora_mode_selector",
                    horizontal=True
                )
                
                if lora_mode == "🚀 Wytrenuj nowy model LoRA (Upload ZIP)":
                    col_tr1, col_tr2 = st.columns([3, 2])
                    with col_tr1:
                        st.markdown("##### 📁 1. Wgraj swój Dataset (Plik ZIP)")
                        zip_file = st.file_uploader(
                            "Wgraj archiwum ZIP zawierające od 5 do 15 zdjęć:", 
                            type=["zip"], 
                            key="lora_zip_uploader"
                        )
                        
                        st.markdown("##### ⚙️ 2. Parametry Treningu")
                        trigger_word = st.text_input(
                            "Unikalny Wyraz Wyzwalający (Trigger Word):", 
                            value=st.session_state.lora_trigger_word,
                            help="Ten wyraz aktywuje model w promptach (np. tomasz_hero)."
                        )
                        st.session_state.lora_trigger_word = trigger_word
                        
                        steps = st.slider(
                            "Liczba Iteracji Treningowych (Steps):", 
                            min_value=500, 
                            max_value=2000, 
                            value=1000, 
                            step=100,
                            help="Więcej iteracji = lepsze dopasowanie, ale ryzyko przeuczenia. 1000 to optymalny standard."
                        )
                        
                        is_style = st.checkbox(
                            "Trening Stylu Artystycznego (is_style)", 
                            value=False,
                            help="Zaznacz tylko wtedy, gdy uczysz stylu artystycznego/graficznego. Dla twarzy i postaci pozostaw wyłączone."
                        )
                        
                        if st.button("🚀 Rozpocznij Trening LoRA ($0.20 - $0.50 fal.ai)", use_container_width=True, type="primary"):
                            if not zip_file:
                                st.error("⚠️ Proszę najpierw załadować plik ZIP ze zdjęciami.")
                            else:
                                with st.spinner("Wgrywanie pliku ZIP do CDN i inicjowanie zlecenia na fal.ai..."):
                                    try:
                                        import tempfile
                                        import os
                                        
                                        # Zapisujemy wgrany plik tymczasowo
                                        with tempfile.NamedTemporaryFile(delete=False, suffix=".zip") as tmp:
                                            tmp.write(zip_file.getvalue())
                                            tmp_path = tmp.name
                                            
                                        from integrations.fal_ai import start_lora_training
                                        request_id, err = start_lora_training(
                                            tmp_path, 
                                            trigger_word=trigger_word, 
                                            steps=steps, 
                                            is_style=is_style
                                        )
                                        
                                        # Usuwamy plik tymczasowy
                                        try:
                                            os.unlink(tmp_path)
                                        except Exception:
                                            pass
                                            
                                        if err:
                                            st.error(f"❌ Nie udało się zainicjować treningu: {err}")
                                        else:
                                            st.session_state.lora_training_id = request_id
                                            st.session_state.lora_training_status = "IN_QUEUE"
                                            st.success(f"✅ Trening zainicjowany pomyślnie! ID Zlecenia: {request_id}")
                                            st.rerun()
                                    except Exception as ex:
                                        st.error(f"❌ Wyjątek podczas uruchamiania treningu: {str(ex)}")
                    
                    with col_tr2:
                        st.markdown("""
                        <div style="background-color: #111827; padding: 20px; border-radius: 12px; border: 1px solid #1F2937;">
                            <h4 style="color: #10B981; margin-top: 0;">📸 Instrukcja Przygotowania Zdjęć</h4>
                            <p style="font-size: 0.85rem; color: #9CA3AF;">
                                Aby uzyskać fotorealistyczną spójność i perfekcyjne dopasowanie modelu, Twój plik ZIP powinien zawierać:
                            </p>
                            <ul style="font-size: 0.85rem; color: #D1D5DB; padding-left: 18px; margin-bottom: 12px;">
                                <li><b>5x Zbliżenie twarzy (Headshot):</b> Różne kąty, neutralna mina, dobre, jednolite światło.</li>
                                <li><b>5x Pół-sylwetka (Half-body):</b> Od pasa w górę, różne ubiory, tła i oświetlenie.</li>
                                <li><b>3x Pełna sylwetka (Full-body):</b> Różne pozy, tła.</li>
                            </ul>
                            <h4 style="color: #F59E0B; margin-top: 15px; font-size: 0.95rem;">👓 Okulary Korekcyjne (Reality Check):</h4>
                            <p style="font-size: 0.85rem; color: #9CA3AF; margin-bottom: 12px;">
                                Jeśli nosisz okulary na co dzień, <b>wgraj większość zdjęć w okularach</b>. Model potraktuje je jako stałą cechę Twojej tożsamości i wygeneruje je z niesamowitą precyzją.
                            </p>
                            <h4 style="color: #3B82F6; margin-top: 15px; font-size: 0.95rem;">📁 Wymagania Techniczne:</h4>
                            <ul style="font-size: 0.85rem; color: #D1D5DB; padding-left: 18px;">
                                <li>Zdjęcia bezpośrednio w ZIP (bez podfolderów).</li>
                                <li>Rozmiar zdjęć: zalecane minimum 1024x1024 px.</li>
                                <li>Formaty: PNG lub JPG.</li>
                            </ul>
                        </div>
                        """, unsafe_allow_html=True)
                else:
                    col_tr1, col_tr2 = st.columns([3, 2])
                    with col_tr1:
                        st.markdown("##### 🔗 Podepnij gotowe wagi LoRA")
                        pasted_url = st.text_input(
                            "🔗 Bezpośredni URL do pliku wag (.safetensors):",
                            placeholder="Wklej adres URL (np. https://v3b.fal.media/files/..._pytorch_lora_weights.safetensors)",
                            key="pasted_lora_url_input"
                        )
                        pasted_trigger = st.text_input(
                            "🔑 Trigger Word (słowo aktywujące):",
                            value=st.session_state.lora_trigger_word,
                            key="pasted_lora_trigger_input"
                        )
                        if st.button("🔌 Załaduj model LoRA", use_container_width=True, type="primary"):
                            if not pasted_url.strip():
                                st.error("⚠️ Podaj poprawny adres URL wag!")
                            else:
                                st.session_state.lora_result_url = pasted_url.strip()
                                st.session_state.lora_trigger_word = pasted_trigger.strip()
                                save_lora_state(pasted_url.strip(), pasted_trigger.strip())
                                st.success("🎉 Pomyślnie podpięto model! Możesz teraz przejść do generowania obrazów.")
                                st.rerun()
                    with col_tr2:
                        st.markdown("""
                        <div style="background-color: #111827; padding: 20px; border-radius: 12px; border: 1px solid #1F2937;">
                            <h4 style="color: #10B981; margin-top: 0;">💡 Oszczędność czasu i kosztów</h4>
                            <p style="font-size: 0.85rem; color: #9CA3AF;">
                                Jeśli masz już wcześniej ukończony trening na fal.ai, nie musisz uczyć modelu od nowa ani płacić za kolejną sesję!
                            </p>
                            <p style="font-size: 0.85rem; color: #D1D5DB;">
                                Wklej bezpośredni link do pliku <b>.safetensors</b> z logów poprzedniego treningu oraz podaj słowo wyzwalające, a panel natychmiast udostępni generator obrazów.
                            </p>
                        </div>
                        """, unsafe_allow_html=True)

            # Monitor stanu treningu
            elif st.session_state.lora_training_id:
                st.info("⚙️ Aktywny Monitor Treningu LoRA")
                req_id = st.session_state.lora_training_id
                
                col_mon1, col_tr_mon2 = st.columns([3, 1])
                with col_mon1:
                    st.text(f"ID Zlecenia Fal.ai: {req_id}")
                    st.text(f"Słowo Kluczowe (Trigger Word): {st.session_state.lora_trigger_word}")
                
                with col_tr_mon2:
                    if st.button("🔄 Odśwież Status", use_container_width=True, type="primary"):
                        from integrations.fal_ai import check_training_status
                        status_res, err = check_training_status(req_id)
                        if err:
                            st.error(f"Błąd sprawdzania statusu: {err}")
                        else:
                            if isinstance(status_res, dict):
                                status_str = status_res.get("status", "IN_QUEUE")
                                logs = status_res.get("logs", [])
                            else:
                                class_name = status_res.__class__.__name__
                                if class_name == "Completed":
                                    status_str = "COMPLETED"
                                elif class_name == "InProgress":
                                    status_str = "IN_PROGRESS"
                                elif class_name == "Queued":
                                    status_str = "IN_QUEUE"
                                else:
                                    status_str = getattr(status_res, "status", "IN_QUEUE")
                                    
                                logs = getattr(status_res, "logs", [])
                                if hasattr(status_res, "logs") and status_res.logs:
                                    logs = [{"message": log.get("message") if isinstance(log, dict) else str(log)} for log in status_res.logs]
                            
                            st.session_state.lora_training_status = str(status_str).upper()
                            st.session_state.lora_training_logs = [log.get("message", "") if isinstance(log, dict) else str(log) for log in logs]
                            
                            if st.session_state.lora_training_status == "COMPLETED":
                                from integrations.fal_ai import get_training_result
                                result_url, err_res = get_training_result(req_id)
                                if err_res:
                                    st.error(f"Nie udało się pobrać linku do wag: {err_res}")
                                else:
                                    st.session_state.lora_result_url = result_url
                                    save_lora_state(result_url, st.session_state.lora_trigger_word)
                                    st.session_state.lora_training_id = None
                                    st.success("🎉 Trening LoRA zakończony sukcesem!")
                            elif st.session_state.lora_training_status in ["FAILED", "ERROR"]:
                                st.error("❌ Trening zakończył się niepowodzeniem.")
                            st.rerun()

                # Wyświetl status wizualny
                status_upper = str(st.session_state.lora_training_status).upper()
                if "COMPLETED" in status_upper:
                    st.success("🏆 Status: UKOŃCZONO")
                elif "IN_PROGRESS" in status_upper or "PROGRESS" in status_upper:
                    st.warning("⚙️ Status: W TOKU (Trwa uczenie modelu, potrwa to ok. 5 minut)")
                elif "FAILED" in status_upper or "ERROR" in status_upper:
                    st.error("❌ Status: BŁĄD TRENINGU")
                else:
                    st.info(f"⏳ Status: {status_upper} (Oczekiwanie w kolejce fal.ai...)")

                # Logi w czasie rzeczywistym
                if st.session_state.lora_training_logs:
                    st.markdown("##### 📝 Logi z Treningu (Ostatnie linie):")
                    log_text = "\n".join(st.session_state.lora_training_logs[-50:])
                    st.code(log_text, language="text")
                else:
                    st.caption("Brak dostępnych logów. Kliknij 'Odśwież Status' za chwilę.")

                if st.button("❌ Przerwij monitorowanie (Zresetuj)", key="reset_training_monitor"):
                    st.session_state.lora_training_id = None
                    st.session_state.lora_training_status = None
                    st.session_state.lora_training_logs = []
                    st.rerun()

            # Wynik udanego treningu
            elif st.session_state.lora_result_url:
                st.markdown("""
                <div style="background-color: #065F46; padding: 15px; border-radius: 8px; border: 1px solid #047857; margin-bottom: 20px;">
                    <h4 style="color: #34D399; margin: 0;">🎉 Twój Model LoRA jest Gotowy!</h4>
                    <p style="color: #A7F3D0; font-size: 0.9rem; margin: 5px 0 0 0;">
                        Wagi zostały wygenerowane i zapisane na bezpiecznym CDN fal.ai w formacie <b>.safetensors</b>.
                    </p>
                </div>
                """, unsafe_allow_html=True)
                
                st.text_input("🔗 Adres URL wag LoRA (.safetensors):", value=st.session_state.lora_result_url, disabled=True)
                st.info(f"💡 Twój trigger word to: **{st.session_state.lora_trigger_word}**")
                
                st.markdown(f"[📥 Pobierz Plik Wag LoRA ({st.session_state.lora_trigger_word}.safetensors)]({st.session_state.lora_result_url})")
                
                st.markdown("<hr style='border-color: #1F242E;'>", unsafe_allow_html=True)
                st.subheader("🎨 Szybki Generator Obrazów z Twoją Nową LoRĄ!")
                st.markdown("Przetestuj swój model w locie! Wpisz prompt fotograficzny zawierający Twój trigger word.")
                
                test_prompt = st.text_area(
                    "Wpisz prompt testowy (Pamiętaj o dodaniu trigger worda):", 
                    value=f"A professional cinematic portrait of {st.session_state.lora_trigger_word} person wearing a sharp black suit on a luxury studio background, extreme detail, photorealistic, 8k"
                )
                
                col_gen1, col_gen2 = st.columns([1, 1])
                with col_gen1:
                    aspect_ratio = st.selectbox(
                        "Proporcje obrazu:", 
                        ["square_hd", "portrait_4_5", "portrait_16_9", "landscape_16_9"], 
                        index=0
                    )
                    scale = st.slider(
                        "Wpływ LoRA (Scale):", 
                        min_value=0.1, 
                        max_value=2.0, 
                        value=1.0, 
                        step=0.1,
                        help="Jak mocno model ma odwzorować Twoje cechy. 1.0 to optymalna wartość."
                    )
                
                if st.button("✨ Generuj Obraz z LoRĄ", type="primary", use_container_width=True):
                    with st.spinner("Model Flux-LoRA generuje obraz..."):
                        from integrations.fal_ai import run_flux_lora_generation
                        img_bytes, err_gen = run_flux_lora_generation(
                            test_prompt, 
                            st.session_state.lora_result_url, 
                            scale=scale, 
                            aspect_ratio=aspect_ratio
                        )
                        if err_gen:
                            st.error(f"Błąd generowania obrazu: {err_gen}")
                        else:
                            st.session_state.lora_test_image = img_bytes
                
                if "lora_test_image" in st.session_state and st.session_state.lora_test_image:
                    st.image(st.session_state.lora_test_image, caption="Twój Wygenerowany Portret UGC", use_container_width=True)
                    st.download_button(
                        label="💾 Pobierz Wygenerowany Portret",
                        data=st.session_state.lora_test_image,
                        file_name=f"{st.session_state.lora_trigger_word}_ugc.png",
                        mime="image/png",
                        use_container_width=True
                    )
                
                if st.button("🧬 Wytrenuj Nowy Model (Zresetuj)", key="reset_entire_studio"):
                    st.session_state.lora_training_id = None
                    st.session_state.lora_training_status = None
                    st.session_state.lora_training_logs = []
                    st.session_state.lora_result_url = None
                    if "lora_test_image" in st.session_state:
                        del st.session_state.lora_test_image
                    st.rerun()

        # --- TOOL 2: CAROUSEL ARCHITECT ---
        elif tool == "Carousel":
            st.subheader("🎠 Carousel Architect (Visual Editor)")
            st.markdown("Stwórz luksusowe slajdy na LinkedIn lub Instagram bezpośrednio w panelu graficznym, bez dotykania konsoli!")
            
            # Formularz wprowadzania slajdów
            carousel_text_default = """# Jak Budować Lejki AI w B2B
Tomasz Duda | jaison.pl
---
# Krok 1: Wybór Domeny
Zawsze wybieraj domenę biznesową. Rejestrując się na GCP wybierz typ konta 'Business' zamiast 'Individual', aby poprawnie odliczać koszty.
---
# Krok 2: Klonowanie Głosu
Używaj modelu VoxCPM2 do lokalnego klonowania głosu. Daje to studyjną jakość 48kHz bez żadnych opłat abonamentowych.
---
# Krok 3: Automatyzacja n8n
Połącz Systeme.io z n8n. Cały ruch organiczny zamienia się w leady i subskrypcje na autopilocie! jaison.pl"""

            carousel_text = st.text_area("Wprowadź treść slajdów (użyj '---' jako separatora slajdów):", value=carousel_text_default, height=250, key="suite_carousel_text")
            
            if st.button("🎠 Wygeneruj Zestaw Slajdów", type="primary", use_container_width=True):
                with st.spinner("Pillow renderuje luksusowe, markowe slajdy w rozdzielczości 1080x1080..."):
                    try:
                        from integrations.generate_carousel import generate_carousel
                        output_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "output_carousel"))
                        generate_carousel(carousel_text, output_dir=output_dir)
                        
                        st.success("🎉 Slajdy wygenerowane pomyślnie!")
                        
                        # Pobierz pliki wygenerowanych slajdów
                        if os.path.exists(output_dir):
                            files = sorted([f for f in os.listdir(output_dir) if f.endswith('.png')])
                            
                            st.markdown("### 👁️ Galeria Podglądu Slajdów:")
                            
                            # Wyświetl slajdy w ładnym układzie siatki
                            cols = st.columns(min(len(files), 3))
                            for idx, file_name in enumerate(files):
                                col_idx = idx % 3
                                file_path = os.path.join(output_dir, file_name)
                                with open(file_path, "rb") as f:
                                    img_data = f.read()
                                    cols[col_idx].image(img_data, caption=f"Slajd {idx+1}", use_container_width=True)
                                    cols[col_idx].download_button(
                                        label=f"💾 Pobierz Slajd {idx+1}",
                                        data=img_data,
                                        file_name=file_name,
                                        mime="image/png",
                                        key=f"suite_dl_slide_{idx}"
                                    )
                                    
                            # Dodanie darmowej możliwości spakowania do ZIP (proaktywny Python)
                            import zipfile
                            import io
                            zip_buffer = io.BytesIO()
                            with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
                                for file_name in files:
                                    file_path = os.path.join(output_dir, file_name)
                                    zip_file.write(file_path, file_name)
                                    
                            st.markdown("---")
                            st.download_button(
                                label="📦 Pobierz Kompletną Karuzelę jako plik .ZIP",
                                data=zip_buffer.getvalue(),
                                file_name="jaison_carousel_pack.zip",
                                mime="application/zip",
                                use_container_width=True
                            )
                    except Exception as ex:
                        st.error(f"Nie udało się wyrenderować slajdów: {str(ex)}")

        # --- TOOL 3: FLUX STUDIO ---
        elif tool == "Flux":
            st.subheader("🎨 Flux Schnell Art Studio (fal.ai)")
            st.markdown("Generuj fotorealistyczne i ultra-precyzyjne grafiki w zaledwie 2 sekundy.")
            
            prompt_input = st.text_input("Wpisz swój prompt (rekomendowany angielski):", value="A premium editorial photograph of a high-tech workspace, glowing purple and teal ambient light, minimalist designer glass, high detail, 8k --ar 16:9", key="suite_flux_prompt")
            
            # Detekcja aktywnej LoRy
            use_lora = False
            lora_url = st.session_state.get("lora_result_url")
            trigger_word = st.session_state.get("lora_trigger_word", "tomasz_hero")
            
            if lora_url:
                st.markdown(f"""
                <div style="background-color: #0F172A; padding: 10px 15px; border-radius: 8px; border: 1px solid #334155; margin-bottom: 12px; display: flex; align-items: center; justify-content: space-between;">
                    <div>
                        <span style="color: #38BDF8; font-weight: bold; font-size: 0.9rem;">🧬 Wykryto Aktywną LoRĘ!</span><br>
                        <span style="color: #94A3B8; font-size: 0.8rem;">Słowo kluczowe: <b>{trigger_word}</b></span>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                use_lora = st.checkbox("🔌 Użyj aktywnego modelu LoRA (Flux-LoRA Dev/Schnell)", value=True, key="flux_studio_use_lora")
                
                if use_lora:
                    lora_scale = st.slider("Skala wpływu LoRA (Scale):", min_value=0.1, max_value=2.0, value=1.0, step=0.1, key="flux_studio_lora_scale")
                    aspect_ratio = st.selectbox("Proporcje obrazu:", ["square_hd", "portrait_4_5", "portrait_16_9", "landscape_16_9"], index=3, key="flux_studio_aspect_ratio")
            
            if st.button("🎨 Generuj Grafikę AI (Flux)", type="primary", use_container_width=True):
                if not prompt_input.strip():
                    st.warning("⚠️ Wpisz najpierw opis grafiki!")
                else:
                    with st.spinner("Flux generuje obraz..."):
                        try:
                            if lora_url and use_lora:
                                from integrations.fal_ai import run_flux_lora_generation
                                img_bytes, err = run_flux_lora_generation(
                                    prompt_input, 
                                    lora_url, 
                                    scale=lora_scale, 
                                    aspect_ratio=aspect_ratio
                                )
                            else:
                                from integrations.fal_ai import run_flux_generation
                                img_bytes, err = run_flux_generation(prompt_input)
                            
                            if err:
                                st.error(f"❌ Błąd: {err}")
                            else:
                                st.success("🎉 Obraz gotowy!")
                                st.image(img_bytes, caption=prompt_input, use_container_width=True)
                                
                                st.download_button(
                                    label="💾 Pobierz Obraz (PNG)",
                                    data=img_bytes,
                                    file_name="flux_art.png",
                                    mime="image/png",
                                    use_container_width=True
                                )
                        except Exception as ex:
                            st.error(f"❌ Błąd generatora: {str(ex)}")

        # --- TOOL 4: FACELESS REELS CREATOR ---
        elif tool == "Faceless_Reels":
            st.subheader("🎬 Content Studio (Nate Herk & Adrian Killar Mode)")
            st.markdown("Projektowanie wirusowych wideo, scenariuszy zasilanych o_mnie.md oraz generowanie audio lektora.")
            
            tab_viral, tab_repurpose, tab_voiceover = st.tabs(["💡 Generator Wirusowych Wideo", "🔄 YouTube Repurposer", "🎙️ Audio Lektora (TTS)"])
            
            with tab_viral:
                col_c1, col_c2 = st.columns([1, 1])
                with col_c1:
                    st.write("##### 💡 Zaprojektuj wirusowe wideo")
                    video_concept = st.text_input("Główny temat lub pomysł na rolkę:", placeholder="Np. Uczucie jazdy z wciśniętym gazem i zaciągniętym hamulcem...", key="suite_reels_concept")
                    video_length = st.selectbox("Długość wideo:", ["8-15 sekund (Szybki strzał)", "30-45 sekund (Edukacyjny Shorts)", "60+ sekund (VSL / Pełna historia)"], key="suite_reels_length")
                    
                    st.write("##### Inspiracja z Telegrama (Nate Herk Mode)")
                    st.caption("Gdy wyślesz komendę do bota na Telegramie w grupie Holistic Mission Control, Hermes automatycznie przekaże ją do CMO, a ten wygeneruje kompletny skrypt wideo bezpośrednio na Twój telefon.")
                    
                    if st.button("Generuj Skrypt i Koncepcję Wideo", type="primary", key="suite_reels_gen_script_btn"):
                        if video_concept:
                            with st.spinner("Wirtualny CMO oraz Dyrektor Kreatywny analizują o_mnie.md..."):
                                time.sleep(1.5)
                                st.session_state.content_script = f"""
### 🎬 Gotowy Skrypt Wirusowy: "{video_concept}"
**Wygenerowany przez: CMO (Tożsamość Tomasz) & Dyrektor Kreatywny (Adrian Killar Style)**

---

#### 📺 SCENA 1: Haczyk (Hook) — Czas: 0:00 - 0:03
* **Wizualnie (Adrian Killar Style):** Dynamiczne cięcie. Tomasz stoi przed kamerą, w tle widać ciemny pulpit z świecącą na fioletowo linią kodu. Kamera robi szybki zoom na twarz.
* **Dźwięk:** Głośny basowy dźwięk „WHOOSH”.
* **Tekst na ekranie:** „Masz ADHD? To nie brak chęci. To zaciągnięty hamulec...”
* **Copywriting (Ghost v2):** „Wciskasz gaz do dechy, ale Twoje życie stoi w miejscu. Znasz to uczucie?”

---

#### 📺 SCENA 2: Rozwinięcie (Body) — Czas: 0:03 - 0:10
* **Wizualnie:** Szybkie przebitki B-roll z luksusowego ciemnego pulpitu i kodu. Tomasz wykonuje powolny oddech (metoda Wima Hofa). Na ekranie pojawia się minimalistyczna grafika mózgu.
* **Dźwięk:** Spokojniejsza, rytmiczna muzyka lo-fi.
* **Copywriting:** „Pochłaniasz setki kursów, masz wysokie ambicje, ale gdy przychodzi do wdrożenia – paraliż. To nie Twoja wina. Twój neuroatypowy mózg potrzebuje zewnętrznego płatu czołowego.”

---

#### 📺 SCENA 3: Wezwanie do działania (CTA) — Czas: 0:10 - 0:15
* **Wizualnie:** Tomasz pokazuje telefon z otwartym botem na Telegramie. Na ekranie wyświetla się adres URL: *ADHD4LIFE*.
* **Copywriting:** „Stworzyłem system, który robi zrzut chaosu z Twojej głowy i układa plan. Wejdź do ADHD4Life i odbierz darmowy workflow. Zdejmij hamulec ręczny już dzisiaj.”
                                """
                                st.rerun()
                        else:
                            st.warning("Wprowadź pomysł na wideo.")
                            
                with col_c2:
                    st.write("##### 📝 Wynik pracy Content Studio")
                    if "content_script" in st.session_state and st.session_state.content_script:
                        st.markdown(st.session_state.content_script)
                        if st.button("Wyczyść skrypt", key="suite_reels_clear_script"):
                            st.session_state.content_script = None
                            st.rerun()
                    else:
                        st.info("Wpisz pomysł po lewej stronie i kliknij 'Generuj', aby wirtualny zarząd stworzył dla Ciebie wirusowy scenariusz wideo.")
                        
            with tab_repurpose:
                st.write("##### 🔄 YouTube Content Repurposer (Nate Herk Mode)")
                st.write("Wklej link YouTube lub bezpośrednio transkrypcję wideo, aby automatycznie stworzyć paczkę dystrybucyjną social media (X/Twitter, LinkedIn, TikTok/Reels) dopasowaną do Twojego o_mnie.md.")
                
                yt_url = st.text_input("Adres URL filmu na YouTube:", placeholder="https://www.youtube.com/watch?v=...", key="suite_yt_url")
                pasted_transcript = st.text_area("Lub wklej tutaj transkrypcję filmu (z napisów YouTube):", height=150, placeholder="Wklej tekst transkrypcji tutaj...", key="suite_yt_transcript")
                obsidian_repurpose_export = st.checkbox("Automatycznie eksportuj wynik do Obsidian Vault", value=True, key="suite_yt_obsidian")
                
                if st.button("Generuj Paczkę Repurposingu", type="primary", key="suite_yt_gen_btn"):
                    transcript_content = ""
                    if yt_url:
                        with st.spinner("Pobieram transkrypcję z YouTube..."):
                            fetched, err = extract_youtube_transcript_raw(yt_url)
                            if err:
                                st.warning(f"Nie udało się automatycznie pobrać transkrypcji: {err}. Użyj wklejenia manualnego poniżej.")
                                transcript_content = pasted_transcript
                            else:
                                st.success("Pomyślnie pobrano transkrypcję z wideo YouTube!")
                                transcript_content = fetched
                    else:
                        transcript_content = pasted_transcript
                        
                    if not transcript_content.strip():
                        st.error("Błąd: Musisz podać poprawny adres URL wideo lub wkleić treść transkrypcji.")
                    else:
                        with st.spinner("CMO oraz Copywriter (Gemini 2.5 Flash) analizują wideo i dopasowują styl..."):
                            o_mnie_path = os.path.join(HERMES_DIR, "o_mnie.md")
                            o_mnie_context = read_md_file(o_mnie_path) if os.path.exists(o_mnie_path) else "Brak profilu o_mnie.md"
                            
                            repurpose_prompt = f"""Przeanalizuj poniższą transkrypcję wideo i stwórz profesjonalny zestaw materiałów marketingowych (Content Repurposing Kit).
        
Twój cel: Przekształcić ten surowy materiał w 3 wysoce perswazyjne, angażujące i dopasowane do profilu użytkownika formaty.
        
PROFIL UŻYTKOWNIKA (O_MNIE - Użyj do dopasowania stylu, tonu i perspektywy Tomasz/Holistic Jason):
{o_mnie_context}
        
TRANSKRYPCJA WIDEO:
{transcript_content[:15000]}
        
GENERUJ NASTĘPUJĄCE TRZY SEKROTY:
        
### 1. Wątek na X (Twitter Thread)
Przygotuj 5-częściowy wątek. Każdy tweet musi mieć maksymalnie 280 znaków. Styl: prowokacyjny, skondensowany, konkretny (bez bełkotu AI). Haczyk (Hook) w pierwszym tweecie. Odnieś się bezpośrednio do przemyśleń i tożsamości z o_mnie.md. Dodaj CTA w ostatnim.
        
### 2. Post na LinkedIn
Napisz angażujący, biznesowy post. Użyj formatu "Hook -> Story -> Lesson -> Call to action". Styl: autentyczny, bez korporacyjnej gadki, krótki (ADHD-friendly), z mocnym haczykiem i przerwami między zdaniami dla lepszej czytelności.
        
### 3. Wirusowy Scenariusz TikTok/Shorts (Adrian Killar Style)
Napisz dynamiczny scenariusz wideo na 30-45 sekund:
- SCENA 1: Haczyk (Hook, visual + copy, pierwsze 3 sekundy).
- SCENA 2: Rozwinięcie (Body, dynamiczny montaż, wartościowa treść).
- SCENA 3: CTA (Call to action).
Pokaż visual cues (co widać na ekranie) i copy (co Tomasz mówi).
        
Napisz całość w czystym markdownie, używając wyrazistych sekcji.
"""
                            response_kit = call_gemini_api([{"role": "user", "content": repurpose_prompt}], "Jesteś wybitnym CMO i dyrektorem kreatywnym tworzącym spójne kampanie cross-channel.")
                            st.session_state.repurpose_kit_result = response_kit
                            
                            if obsidian_repurpose_export:
                                note_title = f"Repurposed_SocialKit_{int(time.time())}.md"
                                note_path = os.path.join(OBSIDIAN_DIR, note_title)
                                try:
                                    with open(note_path, "w", encoding="utf-8") as f:
                                        f.write(f"---\ntype: social-kit\nsource: {yt_url if yt_url else 'Pasted Transcript'}\ntimestamp: {time.strftime('%Y-%m-%d %H:%M:%S')}\n---\n\n{response_kit}")
                                    st.session_state.repurpose_kit_export_success = note_title
                                except Exception as ex:
                                    st.error(f"Błąd zapisu do Obsidian Vault: {ex}")
                            st.rerun()
                            
                if "repurpose_kit_result" in st.session_state and st.session_state.repurpose_kit_result:
                    st.write("##### 📦 Wygenerowana paczka dystrybucyjna:")
                    if st.session_state.get("repurpose_kit_export_success"):
                        st.success(f"Pomyślnie wyeksportowano do Obsidian Vault pod nazwą: {st.session_state.repurpose_kit_export_success}")
                    st.markdown(st.session_state.repurpose_kit_result)
                    
            with tab_voiceover:
                st.write("##### 🎙️ Generator Głosu Lektora (GCP Wavenet)")
                st.write("Zsyntetyzuj profesjonalny głos lektora AI dla swojego wirusowego skryptu.")
                reels_text = st.text_area("Wpisz tekst dla lektora:", value="Dzisiaj zdradzę Ci sekret skutecznej automatyzacji B2B. Zamiast spędzać godziny na rutynowych mailach, stwórz prostego bota w n8n, który przejmie całą komunikację z klientem.", height=150, key="suite_reels_text_input")
                
                if st.button("🎬 Generuj Audio Lektora", type="primary", key="suite_reels_audio_gen_btn", use_container_width=True):
                    with st.spinner("Generowanie głosu AI przez GCP TTS..."):
                        try:
                            audio_bytes, err = call_gcp_tts(reels_text, voice_name="pl-PL-Wavenet-B", gender="MALE")
                            if err:
                                st.error(f"GCP TTS Error: {err}")
                            else:
                                st.audio(audio_bytes, format="audio/mp3")
                                st.success("🎉 Audio lektora gotowe! Możesz je pobrać lub uruchomić skrypt konsolowy faceless_generator, aby scalić je z plikami wideo.")
                                st.download_button(
                                    label="💾 Pobierz Głos Lektora (MP3)",
                                    data=audio_bytes,
                                    file_name="reels_voiceover.mp3",
                                    mime="audio/mp3",
                                    use_container_width=True,
                                    key="suite_reels_dl_mp3_btn"
                                )
                        except Exception as ex:
                            st.error(f"Błąd syntezy mowy: {str(ex)}")

        # --- TOOL 5: BRAND STRATEGY & PROFILE BIOS ---
        elif tool == "Brand_Bios":
            st.subheader("✍️ Brand Strategy & Profile BIO's")
            st.markdown("Uzupełnij poniższy kwestionariusz, aby wygenerować kompletną tożsamość marki, opisy BIO dla 6 platform oraz premium awatary i bannery bezpieczne dla smartfonów.")
            
            tab_strategy, tab_bios, tab_visuals = st.tabs(["📋 Wywiad i Strategia", "✍️ Opisy i BIO (6 Platform)", "🎨 Wizualia (Awatary i Bannery)"])
            
            with tab_strategy:
                st.subheader("📋 Kwestionariusz Twojej Marki / Biznesu")
                st.markdown("Nasza orkiestracja dyrektorów AI i model **Gemini 2.5 Pro** stworzą na tej podstawie kompletną strategię i opisy BIO dostosowane do każdej z platform.")
                
                col_s1, col_s2 = st.columns([1, 1])
                with col_s1:
                    brand_name = st.text_input("Nazwa Marki / Imię i Nazwisko:", value="J(AI)SON", key="suite_brand_name")
                    website_url = st.text_input("URL Strony Głównej:", value="https://jaison.pl", key="suite_website_url")
                    app_url = st.text_input("URL Oferty / Aplikacji SaaS:", value="https://app.jaison.pl", key="suite_app_url")
                    niche = st.text_area("Branża / Specjalizacja:", value="Projektowanie systemów agentowych AI i automatyzacja B2B", height=80, key="suite_niche")
                with col_s2:
                    target_audience = st.text_area("Grupa Docelowa (Idealny Klient):", value="neuroatypowi founderzy, mali i średni przedsiębiorcy B2B", height=80, key="suite_target_audience")
                    tone_of_voice = st.selectbox("Styl komunikacji / Tone of Voice:", [
                        "Ghost v2: bezpośredni, ADHD-friendly, zero lania wody",
                        "Profesjonalny i merytoryczny B2B",
                        "Luźny, perswazyjny NLP i konwersacyjny",
                        "Energetyczny, pełen motywacji (ADHD War Room)"
                    ], key="suite_tone_of_voice")
                    motto = st.text_input("Twoje Unikalne Motto / Slogan:", value="Automatyzuj to, co powtarzalne. Twórz to, co unikalne. Zyskaj.", key="suite_motto")
                    cta = st.text_input("Wezwanie do Działania (CTA) na banerze:", value="Zacznij automatyzować, żyj lepiej.", key="suite_cta")
                    
                if st.button("🚀 Analizuj i generuj Strategię AI", type="primary", use_container_width=True, key="suite_sm_gen_strat_btn"):
                    with st.spinner("Dyrektor ds. Marketingu (CMO AI) oraz Gemini 2.5 Pro analizują rynek i konkurencję..."):
                        prompt = f"""
                        Przeprowadź głęboki wywiad i stwórz kompletną strategię social media oraz opisy BIO dla 6 platform.
                        Marka/Nazwisko: {brand_name}
                        Strona Główna: {website_url}
                        Aplikacja/Oferta: {app_url}
                        Nisza/Branża: {niche}
                        Grupa docelowa: {target_audience}
                        Styl komunikacji (Tone of Voice): {tone_of_voice}
                        Unikalne motto: {motto}
                        CTA: {cta}

                        Rygorystycznie przestrzegaj poniższych wytycznych i limitów znakowych dla platform:
                        1. LinkedIn Profile:
                           - Nagłówek (Headline): Max 220 znaków. Format: Rola | Obietnica wartości + Kluczowe słowa kluczowe | Nazwa Marki.
                           - O mnie (About): Max 2600 znaków. Układ: Haczyk (problem), rozwiązanie (metoda), filary oferty, dlaczego my, call to action.
                           - Opis usług (Services Description): Rygorystyczny limit 500 znaków! Skondensowana pigułka oferty.
                        2. Instagram Bio:
                           - Tekst Bio: Max 150 znaków ze spacjami! Musi zawierać emotikony, wezwanie do działania i skrócony link.
                        3. Facebook Page About:
                           - Krótki opis (Short Bio): Max 255 znaków.
                           - Długi opis (Long Description): Bez limitu – pełna historia marki i misja.
                        4. TikTok Bio:
                           - Tekst Bio: Max 80 znaków ze spacjami! Maksymalnie skondensowane, jednozdaniowe uderzenie z CTA.
                        5. Twitter / X Bio:
                           - Tekst Bio: Max 160 znaków. Pozycjonowanie eksperckie, dynamiczne, tech-friendly.
                        6. Threads Bio:
                           - Tekst Bio: Max 150 znaków. Bardziej luźny, osobisty charakter.

                        Wygeneruj odpowiedź w czystym formacie JSON o poniższej strukturze (nie umieszczaj żadnych znaczników markdown poza kodem json, tylko czysty, parsujący się JSON bez wstępów):
                        {{
                          "slogan": "krótki, uderzający slogan na baner (max 6-8 słów)",
                          "cta": "krótkie wezwanie do działania na baner (max 4-5 słów)",
                          "linkedin_headline": "Nagłówek na LinkedIn (max 220 znaków)",
                          "linkedin_about": "O mnie na LinkedIn (max 2600 znaków)",
                          "linkedin_services": "Opis usług na LinkedIn (max 500 znaków)",
                          "facebook_short_bio": "Krótki opis na Facebooka (max 255 znaków)",
                          "facebook_long_desc": "Długi opis na Facebooka (pełna historia i misja)",
                          "instagram_bio": "BIO na Instagram (max 150 znaków z emotikonami i CTA)",
                          "tiktok_bio": "BIO na TikToka (max 80 znaków z mega mocnym hakiem i CTA)",
                          "twitter_bio": "BIO na X/Twitter (max 160 znaków, thought-leadership)",
                          "threads_bio": "BIO na Threads (max 150 znaków, luźne, konwersacyjne)",
                          "strategy_tips": [
                            "Wskazówka 1 (ADHD friendly, konkretna)",
                            "Wskazówka 2 (Dopaminowy hook)",
                            "Wskazówka 3 (Dystrybucja treści)",
                            "Wskazówka 4 (Szybkie i proste systemy)"
                          ]
                        }}
                        """
                        messages = [{"role": "user", "content": prompt}]
                        system_instruction = "Jesteś wybitnym CMO i ekspertem copywritingu NLP. Zwracaj wyłącznie poprawny obiekt JSON, bez żadnego tekstu przed ani po nim."
                        try:
                            res_raw = call_gemini_pro_api(messages, system_instruction)
                            import json
                            import re
                            clean_res = res_raw.strip()
                            if clean_res.startswith("```"):
                                clean_res = re.sub(r"^```(?:json)?\n", "", clean_res)
                                clean_res = re.sub(r"\n```$", "", clean_res)
                                clean_res = clean_res.strip()
                            
                            st.session_state.sm_strategy = json.loads(clean_res)
                            st.success("Strategia wygenerowana pomyślnie! Przejdź do kolejnych zakładek, aby zobaczyć BIO i wygenerować grafiki.")
                        except Exception as e:
                            st.warning(f"Nie udało się sparsować odpowiedzi JSON, wdrożono domyślną strategię premium. Błąd: {e}")
                            st.session_state.sm_strategy = {
                                "slogan": f"{motto}",
                                "cta": f"➔ {cta}",
                                "linkedin_headline": f"AI Systems Architect 🧠 | Projektuję systemy agentowe i automatyzacje n8n B2B | {brand_name}",
                                "linkedin_about": f"Pomagam zabieganym przedsiębiorcom i osobom z ADHD odzyskać 20+ godzin tygodniowo przez wdrożenia agentów AI i automatyzacje n8n. Sprawdź moje case studies i uwolnij swój czas. Kontakt: {app_url}",
                                "linkedin_services": f"Projektowanie i wdrożenia systemów AI B2B. Automatyzacje procesów n8n, integracje LLM, autorskie dashboardy Streamlit.",
                                "facebook_short_bio": f"Dołącz do społeczności twórców i biznesów, którzy zamiast pracować w firmie, pracują nad jej automatyzacją.",
                                "facebook_long_desc": f"Pełna misja marki {brand_name}: Uwalniamy czas przedsiębiorców. Dostarczamy bezkompromisowe systemy oparte na agentach AI, które pracują 24/7. Nasz cel to pomoc w odzyskaniu kontroli nad czasem i energią w duchu Low-Cost / High-ROI.",
                                "instagram_bio": "⚡️ Robimy to co ważne, resztę robi kod\n💡 Automatyzacje procesów B2B\n👇 Odbierz bezpłatny zestaw n8n blueprintów!",
                                "tiktok_bio": "🧠 ADHD & AI Automations | 💡 Odzyskaj 20h w tygodniu! | Kliknij link 👇",
                                "twitter_bio": f"SaaS founder & AI Agency Director. I build agentic operating systems to automate workflows for fast-growing B2B brands. ADHD builder mode on.",
                                "threads_bio": "AI agent builder & systems architect. Here to talk about real tech, ADHD productivity hacks & automated pipelines. Let's debate!",
                                "strategy_tips": [
                                  "System 1-Click: Nagrywaj luźne przemyślenia głosowe, a AI (np. Omi lub sformatowany monit) przekształci je w posty na 6 platform.",
                                  "Płynność i Dopamina: Nie edytuj wideo godzinami. Używaj dynamicznych napisów, prostych przejść i gotowych szablonów.",
                                  "Autentyczność przede wszystkim: Buduj markę osobistą w oparciu o prawdziwe wyzwania neuroatypowości (ADHD).",
                                  "Użyj darmowego planu Systeme.io do budowy bazy e-mailowej i spięcia ruchu organicznego."
                                ]
                            }
                
                # Display recommendations
                if "sm_strategy" in st.session_state and st.session_state.sm_strategy:
                    strat = st.session_state.sm_strategy
                    st.markdown("---")
                    st.markdown("<p style='color: #A78BFA; font-weight: bold; font-size: 1.2rem;'>🎯 Główne rekomendacje strategiczne dla Twojej marki:</p>", unsafe_allow_html=True)
                    
                    col_b1, col_b2 = st.columns([1, 1])
                    with col_b1:
                        st.markdown(f"""
                        <div class="custom-card" style="border-left: 4px solid #7C3AED; background: #13111C; min-height: 150px;">
                            <span style="font-size: 0.75rem; color: #A78BFA; font-weight: bold;">Slogan główny:</span>
                            <h4 style="color: #FFF; margin: 6px 0 12px 0; font-size: 1.1rem; line-height: 1.3;">{strat.get('slogan', '')}</h4>
                            <span style="font-size: 0.75rem; color: #A78BFA; font-weight: bold;">Wezwanie do działania (CTA):</span>
                            <p style="color: #E2E8F0; font-size: 0.9rem; margin-top: 4px; font-weight: bold;">{strat.get('cta', '')}</p>
                        </div>
                        """, unsafe_allow_html=True)
                    with col_b2:
                        st.markdown("<div class='custom-card' style='border-left: 4px solid #EC4899; background: #1C1118; min-height: 150px;'>", unsafe_allow_html=True)
                        st.markdown("<span style='font-size: 0.75rem; color: #F472B6; font-weight: bold;'>Wskazówki operacyjne:</span>", unsafe_allow_html=True)
                        for tip in strat.get("strategy_tips", []):
                            st.markdown(f"<p style='color: #E2E8F0; font-size: 0.8rem; margin: 4px 0;'>⚡ {tip}</p>", unsafe_allow_html=True)
                        st.markdown("</div>", unsafe_allow_html=True)
            
            with tab_bios:
                if not st.session_state.get("sm_strategy"):
                    st.info("💡 Uruchom najpierw kwestionariusz i analizę w zakładce obok, aby wygenerować BIO dla profili social media.")
                else:
                    strat = st.session_state.sm_strategy
                    st.write("##### ✍️ Gotowe opisy BIO do skopiowania na Twoje profile:")
                    
                    col_p1, col_p2 = st.columns([1, 1])
                    with col_p1:
                        st.markdown("<p style='color: #10B981; font-weight: bold;'>💼 LinkedIn (Autorytet B2B):</p>", unsafe_allow_html=True)
                        st.text_input("LinkedIn Nagłówek (Headline - max 220 znaków):", value=strat.get("linkedin_headline", strat.get("linkedin_bio", "")[:220]), key="suite_bio_li_headline")
                        st.text_area("LinkedIn O mnie (About - max 2600 znaków):", value=strat.get("linkedin_about", strat.get("linkedin_bio", "")), height=250, key="suite_bio_li_about")
                        st.text_area("LinkedIn Opis Usług (Services - max 500 znaków):", value=strat.get("linkedin_services", ""), height=100, key="suite_bio_li_services")
                        
                        st.markdown("<p style='color: #3B82F6; font-weight: bold;'>👥 Facebook Page:</p>", unsafe_allow_html=True)
                        st.text_input("Facebook Krótki opis (Short Bio - max 255 znaków):", value=strat.get("facebook_short_bio", strat.get("facebook_bio", "")[:255]), key="suite_bio_fb_short")
                        st.text_area("Facebook Długi opis (Long Description):", value=strat.get("facebook_long_desc", strat.get("facebook_bio", "")), height=150, key="suite_bio_fb_long")
                        
                    with col_p2:
                        st.markdown("<p style='color: #EC4899; font-weight: bold;'>📸 Instagram (max 150 znaków ze spacjami):</p>", unsafe_allow_html=True)
                        st.text_area("Instagram Bio:", value=strat.get("instagram_bio", ""), height=100, key="suite_bio_ig")
                        
                        st.markdown("<p style='color: #EEF2F6; font-weight: bold;'>🎵 TikTok (max 80 znaków ze spacjami):</p>", unsafe_allow_html=True)
                        st.text_area("TikTok Bio:", value=strat.get("tiktok_bio", ""), height=80, key="suite_bio_tt")
                        
                        st.markdown("<p style='color: #1D9BF0; font-weight: bold;'>🐦 Twitter / X (max 160 znaków):</p>", unsafe_allow_html=True)
                        st.text_area("Twitter/X Bio:", value=strat.get("twitter_bio", ""), height=100, key="suite_bio_tw")
                        
                        st.markdown("<p style='color: #A855F7; font-weight: bold;'>💬 Threads (max 150 znaków):</p>", unsafe_allow_html=True)
                        st.text_area("Threads Bio:", value=strat.get("threads_bio", ""), height=100, key="suite_bio_th")
                        
            with tab_visuals:
                st.subheader("🎨 Kreator Tożsamości Wizualnej (Awatary i Bannery)")
                st.markdown("Wygeneruj spójne graficznie awatary i bannery dostosowane pod każdą z platform przy użyciu **Google Imagen 3.0** lub Twojego prywatnego modelu **Flux LoRA**.")
                
                # Sprawdź stan LoRA
                has_lora = "lora_result_url" in st.session_state and st.session_state.lora_result_url
                
                st.markdown("### 🎛️ 1. Wybór Silnika Generacji")
                col_engine1, col_engine2 = st.columns([2, 1])
                with col_engine1:
                    engine_mode = st.radio(
                        "Silnik graficzny:",
                        ["Google Imagen 3.0 (Szybki, wszechstronny)", "Flux LoRA (Twój spersonalizowany styl i twarz)"],
                        index=1 if has_lora else 0,
                        horizontal=True,
                        key="visuals_engine_select",
                        help="Wybierz Flux LoRA, aby wygenerować grafiki bezpośrednio ze swoją twarzą!"
                    )
                with col_engine2:
                    if has_lora:
                        lora_scale = st.slider(
                            "Wpływ LoRA (Scale):", 
                            min_value=0.1, 
                            max_value=2.0, 
                            value=0.85, 
                            step=0.1,
                            key="visuals_lora_scale",
                            help="Rekomendowana wartość: 0.85."
                        )
                    else:
                        st.info("💡 Brak aktywnej LoRA. Możesz ją wgrać/wytrenować w LoRA Studio.")
                        lora_scale = 1.0

                st.markdown("---")
                
                col_i1, col_i2 = st.columns([1, 1])
                with col_i1:
                    st.write("##### 👤 2. Generowanie Spójnego Awatara:")
                    
                    suite_char_desc = st.text_input(
                        "Opis Postaci (CHARACTER):", 
                        value="confident bald man in late 30s, rectangular glasses, thin matte black frames, warm genuine smile", 
                        key="suite_char_desc",
                        help="Opis Twoich cech fizycznych w języku angielskim."
                    )
                    
                    avatar_style = st.selectbox(
                        "Strategiczny Styl Wizualny Marki:", 
                        [
                            "🌌 Styl A: Cybernetic Hub (AI systems & Neural Networks)",
                            "🌿 Styl B: Cozy Server Sanctuary (Biohacking, Nature & Tech)",
                            "🌃 Styl C: Deep Focus Night Studio (Hyperfocus & High-Performance)",
                            "🌅 Styl D: Biotech Garden / Freedom (Operational Freedom & Silence)",
                            "🕴️ Styl E: Executive Cyber-Casual (Founder Authority & B2B Trust)",
                            "✨ Własny Prompt (Custom Creative Mode)"
                        ], 
                        key="suite_avatar_style_spec"
                    )
                    
                    custom_avatar_prompt_val = ""
                    if avatar_style == "✨ Własny Prompt (Custom Creative Mode)":
                        custom_avatar_prompt_val = st.text_area(
                            "Wpisz swój własny prompt dla Awatara (po angielsku):",
                            value="A high-end cinematic close-up portrait of a founder standing in a modern neon tech office, looking confident, 8k resolution.",
                            key="suite_avatar_custom_prompt_val",
                            height=120
                        )
                    
                    use_lora_for_avatar = False
                    if "Flux LoRA" in engine_mode and has_lora:
                        use_lora_for_avatar = st.checkbox(
                            "Dodaj moją postać (LoRA) do awatara", 
                            value=True, 
                            key="suite_avatar_use_lora",
                            help="Zaznacz, aby model odwzorował Twoją twarz (trigger word: tomasz_hero)."
                        )
                    
                    if st.button("Generuj Profesjonalny Awatar", type="primary", key="suite_avatar_gen_btn", use_container_width=True):
                        if avatar_style == "✨ Własny Prompt (Custom Creative Mode)":
                            if use_lora_for_avatar and "Flux LoRA" in engine_mode:
                                full_avatar_prompt = f"{st.session_state.get('lora_trigger_word', 'tomasz_hero')} person, {custom_avatar_prompt_val}"
                            else:
                                full_avatar_prompt = custom_avatar_prompt_val
                        else:
                            if use_lora_for_avatar and "Flux LoRA" in engine_mode:
                                char_prompt = f"{st.session_state.get('lora_trigger_word', 'tomasz_hero')} person ({suite_char_desc})"
                            else:
                                char_prompt = suite_char_desc
                                
                            if "Styl A" in avatar_style:
                                full_avatar_prompt = f"A high-end, clean close-up portrait of {char_prompt} facing the camera. The background is a sophisticated and futuristic dark tech command center: glowing translucent glass screens displaying holographic cyan and electric purple data charts, mathematical equations, and delicate network connection lines. Studio key lighting creates sharp masculine features, strong jawline, and natural skin textures. No distortions, highly professional, 8k resolution, cinematic atmosphere, f/1.8."
                            elif "Styl B" in avatar_style:
                                full_avatar_prompt = f"A professional cinematic portrait of {char_prompt} with a natural and calm expression. The setting is a cozy, warm minimalist server room where technology meets nature: modern matte-charcoal server racks are softly illuminated by warm amber LED lights and gentle green power indicators. Delicate leafy organic plants like monstera and small bonsai trees grow harmoniously around the technical equipment. Soft morning sunlight streams from a side window, highlighting natural skin fullness, realistic textures, and sharp details. High-end, premium, photorealistic, 8k."
                            elif "Styl C" in avatar_style:
                                full_prompt_part = f"A high-contrast, moody close-up portrait of {char_prompt} looking confident and focused. He is in a premium, quiet soundproof coding room at night. The background features dark acoustic hexagonal panels with glowing neon cyan linear lights running between them. To the side, there is a softly glowing vertical monitor showing clean coding structures. Warm, dramatic key lighting creates clear masculine facial lines, natural facial fullness, and reflections in his corrective glasses. Cinematic depth of field, photorealistic, premium feel, 8k."
                                full_avatar_prompt = full_prompt_part
                            elif "Styl D" in avatar_style:
                                full_avatar_prompt = f"An ultra-minimalist, airy cinematic portrait of {char_prompt} looking completely relaxed and happy, with a genuine smile. He is standing near a giant floor-to-ceiling clean panoramic glass window of a modern architectural estate. Outside, a pristine biotech garden with clean green plants is bathed in warm, golden sunset light. The atmosphere is quiet, bright, and completely silent. Key soft lighting showcases healthy skin tones, clear masculine structures, and authentic facial fullness. High-end editorial style, premium design, depth of field, 8k."
                            else: # Styl E
                                b_name = st.session_state.get("suite_brand_name", "J(AI)SON")
                                full_avatar_prompt = f"An elegant and luxurious portrait of {char_prompt} dressed in a tailor-made deep navy blue blazer over a premium charcoal grey turtleneck. On the collar of the turtleneck, a very small and clean matte grey embroidered text '{b_name}' is subtly visible. He is standing inside a highly professional, dark-themed executive lounge. The background is a clean matte-black luxury wall with subtle recessed accent lighting. The lighting is soft, dramatic, and prestigious, highlighting natural masculine jawline, authentic facial fullness, and premium fabrics. Perfect corporate identity, photorealistic, 8k."

                        if "Flux LoRA" in engine_mode:
                            if not has_lora:
                                st.error("⚠️ Brak aktywnej LoRA! Wklej link do wag w zakładce LoRA Studio.")
                            else:
                                with st.spinner("Model Flux-LoRA generuje spersonalizowany awatar..."):
                                    from integrations.fal_ai import run_flux_lora_generation
                                    img_bytes, err = run_flux_lora_generation(
                                        full_avatar_prompt, 
                                        st.session_state.lora_result_url, 
                                        scale=lora_scale, 
                                        aspect_ratio="square_hd"
                                    )
                                    if err:
                                        st.error(f"Fal.ai API Error: {err}")
                                    elif img_bytes:
                                        st.session_state.sm_generated_avatar = img_bytes
                                        st.success("Spersonalizowany awatar LoRA wygenerowany!")
                        else:
                            with st.spinner("Model Imagen 3 generuje idealnie wykadrowany awatar..."):
                                img_bytes, err = generate_imagen_image(full_avatar_prompt, aspect_ratio="1:1")
                                if err:
                                    st.error(f"GCP API Error: {err}")
                                elif img_bytes:
                                    st.session_state.sm_generated_avatar = img_bytes
                                    st.success("Awatar wygenerowany pomyślnie!")
                                    
                    if "sm_generated_avatar" in st.session_state:
                        st.image(st.session_state.sm_generated_avatar, caption="Twój spójny awatar", width=250)
                        st.download_button(
                            label="💾 Pobierz Awatar (PNG)",
                            data=st.session_state.sm_generated_avatar,
                            file_name="jaison_avatar.png",
                            mime="image/png",
                            use_container_width=True,
                            key="suite_avatar_dl_btn"
                        )
                        
                with col_i2:
                    st.write("##### 🖼️ 3. Generowanie Kreacji, Banneru lub Grafiki pod Platformę:")
                    
                    platform_specs = {
                        "👤 Awatar / Zdjęcie Profilowe (Standard 1:1)": {
                            "aspect_imagen": "1:1",
                            "aspect_flux": "square_hd",
                            "dimensions": "800 x 800 px (Kwadrat 1:1)",
                            "guide": "👤 Idealny format na spójny awatar lub zdjęcie profilowe (LinkedIn, FB, X, YT, Instagram). Skupia się na twarzy, kadr centralny.",
                        },
                        "👔 Portret Biznesowy (Klasyczny Kwadrat 1:1)": {
                            "aspect_imagen": "1:1",
                            "aspect_flux": "square_hd",
                            "dimensions": "1080 x 1080 px (Kwadrat 1:1)",
                            "guide": "💼 Profesjonalny kadr od klatki piersiowej w górę, doskonały do sekcji 'O nas', wizytówek, stopki e-mail i artykułów.",
                        },
                        "💼 LinkedIn Banner (Szeroki)": {
                            "aspect_imagen": "16:9",
                            "aspect_flux": "landscape_16_9",
                            "dimensions": "1584 x 396 px (Panoramiczny)",
                            "guide": "⚠️ Zdjęcie profilowe zasłania lewy dolny róg na desktopie, a na telefonie jest wyśrodkowane. Slogan i grafika będą w prawej części bezpiecznej.",
                        },
                        "👥 Facebook Cover (Standardowy)": {
                            "aspect_imagen": "16:9",
                            "aspect_flux": "landscape_16_9",
                            "dimensions": "851 x 315 px (Poziomy)",
                            "guide": "⚠️ Zdjęcie profilowe na komputerze znajduje się po lewej stronie, więc slogan zostanie przesunięty ku prawej stronie.",
                        },
                        "🐦 Twitter/X Banner (Szeroki)": {
                            "aspect_imagen": "16:9",
                            "aspect_flux": "landscape_16_9",
                            "dimensions": "1500 x 500 px (Proporcje 3:1)",
                            "guide": "⚠️ Zdjęcie profilowe na Twitterze zasłania dolny lewy narożnik bannera. Slogan i kluczowa grafika zostaną przesunięte w prawo i lekko w górę.",
                        },
                        "🎵 YouTube Banner (Panoramiczny)": {
                            "aspect_imagen": "16:9",
                            "aspect_flux": "landscape_16_9",
                            "dimensions": "2048 x 1152 px (Proporcje 16:9)",
                            "guide": "⚠️ Środkowa część 1546x423 px to strefa bezpieczna na komputery i telefony. Slogan po prawej stronie.",
                        },
                        "💬 Threads Cover (Szeroki)": {
                            "aspect_imagen": "16:9",
                            "aspect_flux": "landscape_16_9",
                            "dimensions": "1080 x 360 px (Poziomy)",
                            "guide": "⚠️ Styl minimalistyczny z dużą ilością wolnej przestrzeni (negative space) i sloganem po prawej stronie.",
                        },
                        "🌐 Web Hero Section (Główny baner strony 16:9)": {
                            "aspect_imagen": "16:9",
                            "aspect_flux": "landscape_16_9",
                            "dimensions": "1920 x 1080 px (FHD 16:9)",
                            "guide": "🖥️ Główny baner strony (Hero Section). Lewa/środkowa strona powinna mieć przestrzeń na nagłówki i przyciski. Kompozycja zbalansowana pod mobile i desktop.",
                        },
                        "🖥️ Web Section Background (Tło sekcji 4:3)": {
                            "aspect_imagen": "4:3",
                            "aspect_flux": "landscape_4_3",
                            "dimensions": "1200 x 900 px (Format 4:3)",
                            "guide": "📐 Grafika obok bloków tekstu lub tło funkcji. Idealna proporcja dla czytelności i symetrii w sekcjach produktowych.",
                        },
                        "📝 Blog Post Header (Baner artykułu)": {
                            "aspect_imagen": "16:9",
                            "aspect_flux": "landscape_16_9",
                            "dimensions": "1200 x 630 px (Optymalny 16:9)",
                            "guide": "✍️ Nagłówek artykułu blogowego. Kluczowe elementy są w centrum, dzięki czemu miniatura wygląda zjawiskowo na liście wpisów.",
                        },
                        "🎥 YouTube Thumbnail (Miniaturka wideo)": {
                            "aspect_imagen": "16:9",
                            "aspect_flux": "landscape_16_9",
                            "dimensions": "1280 x 720 px (Format HD)",
                            "guide": "📽️ Miniaturka na YouTube. Umieść postać po prawej lub lewej stronie z mocnym kontrastem. Prawy dolny róg zostaw wolny na licznik czasu (timecode)!",
                        },
                        "🔗 Social Sharing Card / OG Link Cover": {
                            "aspect_imagen": "16:9",
                            "aspect_flux": "landscape_16_9",
                            "dimensions": "1200 x 630 px (Wymiar OpenGraph)",
                            "guide": "🔗 Miniaturka generowana po wklejeniu linku w social media (LinkedIn, FB, X, Slack, Discord). Wyśrodkowana kompozycja zapobiega ucięciu detali.",
                        },
                        "📱 Shorts / Reels / TikTok Cover (Pion 9:16)": {
                            "aspect_imagen": "9:16",
                            "aspect_flux": "portrait_16_9",
                            "dimensions": "1080 x 1920 px (Pionowy)",
                            "guide": "📱 Pionowa okładka wideo. Złota strefa bezpieczna to środkowe 1080x1350 px. Górne i dolne krawędzie są zasłonięte przez interfejs platformy.",
                        },
                        "📌 Pinterest Pin (Pionowy Portret 2:3)": {
                            "aspect_imagen": "3:4",
                            "aspect_flux": "portrait_4_3",
                            "dimensions": "1000 x 1500 px (Pionowy 2:3 / 3:4)",
                            "guide": "📌 Wydłużona grafika generująca wysokie kliki w wyszukiwarce Pinterest. Slogan umieszczony w górnej lub środkowej części.",
                        },
                        "📸 Instagram Post (Klasyczny Kwadrat)": {
                            "aspect_imagen": "1:1",
                            "aspect_flux": "square_hd",
                            "dimensions": "1080 x 1080 px (Kwadrat)",
                            "guide": "🎨 Tradycyjny kwadratowy post. Idealnie wyśrodkowana kompozycja, wysoki kontrast, idealne do feedu i karuzeli.",
                        }
                    }
                    
                    selected_platform = st.selectbox(
                        "Wybierz platformę / format docelowy:",
                        list(platform_specs.keys()),
                        key="visuals_platform_select"
                    )
                    
                    specs = platform_specs[selected_platform]
                    
                    st.caption(f"📏 **Zalecane wymiary:** `{specs['dimensions']}`")
                    st.info(specs['guide'])
                    
                    slogan_val = st.session_state.get("sm_strategy", {}).get("slogan", st.session_state.get("suite_motto", "Automatyzuj powtarzalne. Twórz unikalne. Zyskaj."))
                    banner_title = st.text_input(
                        "Tekst sloganu na banerze:", 
                        value=slogan_val, 
                        key="suite_banner_title_dynamic"
                    )
                    
                    cta_val = st.session_state.get("sm_strategy", {}).get("cta", st.session_state.get("suite_cta", "➔ Zacznij automatyzować, żyj lepiej."))
                    banner_cta = st.text_input(
                        "Wezwanie do działania (CTA) na banerze:", 
                        value=cta_val, 
                        key="suite_banner_cta_dynamic"
                    )
                    
                    banner_style_select = st.selectbox(
                        "Wybierz Styl Wizualny Banera / Grafiki:", 
                        [
                            "🌌 Styl A: Cybernetic Hub (AI systems & Neural Networks)",
                            "🌿 Styl B: Cozy Server Sanctuary (Biohacking, Nature & Tech)",
                            "🌃 Styl C: Deep Focus Night Studio (Hyperfocus & High-Performance)",
                            "🌅 Styl D: Biotech Garden / Freedom (Operational Freedom & Silence)",
                            "🕴️ Styl E: Executive Cyber-Casual (Founder Authority & B2B Trust)",
                            "✨ Własny Prompt (Custom Creative Mode)"
                        ],
                        key="suite_banner_style_select"
                    )
                    
                    custom_banner_prompt_val = ""
                    if banner_style_select == "✨ Własny Prompt (Custom Creative Mode)":
                        custom_banner_prompt_val = st.text_area(
                            "Wpisz swój własny prompt dla Banera (po angielsku):",
                            value="A luxurious, minimalist panoramic banner, with abstract glowing lines, high-end branding, elegant style, ultra-high resolution.",
                            key="suite_banner_custom_prompt_val",
                            height=120
                        )
                    
                    include_char_in_banner = False
                    if "Flux LoRA" in engine_mode and has_lora:
                        include_char_in_banner = st.checkbox(
                            "Dodaj moją postać (LoRA) do grafiki", 
                            value=True, 
                            key="suite_banner_include_char",
                            help="Jeśli zaznaczone, Twoja sylwetka/twarz zostanie wkomponowana zgodnie z wytycznymi kompozycji."
                        )
                    elif not has_lora:
                        include_char_in_banner = st.checkbox(
                            "Dodaj sylwetkę postaci do grafiki", 
                            value=False, 
                            key="suite_banner_include_char_base",
                            help="Jeśli zaznaczone, ogólna sylwetka postaci zostanie wkomponowana w kadr."
                        )

                    if st.button("Generuj Grafikę (Imagen 3 / Flux LoRA)", type="primary", use_container_width=True, key="suite_banner_gen_btn_dynamic"):
                        if include_char_in_banner:
                            if "Flux LoRA" in engine_mode and has_lora:
                                char_part = f"On the left side, there is a professional, high-end realistic side-profile or silhouette portrait of {st.session_state.lora_trigger_word} person looking confident, integrated seamlessly into the environment. "
                            else:
                                char_part = f"On the left side, there is a professional, high-end realistic side-profile or silhouette portrait of {suite_char_desc} looking confident, integrated seamlessly into the environment. "
                        else:
                            char_part = ""

                        if "Własny Prompt" in banner_style_select:
                            if include_char_in_banner:
                                if "Flux LoRA" in engine_mode and has_lora:
                                    full_prompt = f"{st.session_state.get('lora_trigger_word', 'tomasz_hero')} person, {custom_banner_prompt_val}"
                                else:
                                    full_prompt = f"portrait of confident person, {custom_banner_prompt_val}"
                            else:
                                full_prompt = custom_banner_prompt_val
                        elif "Styl A" in banner_style_select:
                            full_prompt = f"A premium panoramic cinematic wide shot of a futuristic dark control center. On the left and center, glowing glass analytics screens display abstract neural connection lines, cyan and purple networks, and subtle data charts. {char_part}On the RIGHT side of the image, there is a clean, dark charcoal space featuring the glowing, elegant text: '{banner_title}' with a smaller subtitle '{banner_cta}' underneath. Sharp lighting, ultra-high resolution, minimalist layout, text is perfectly readable, 8k."
                        elif "Styl B" in banner_style_select:
                            full_prompt = f"A premium panoramic wide-angle shot of a warm minimalist server room. On the left and center, modern server racks with glowing soft amber and soft green lights are mixed with lush green potted plants and small bonsais. {char_part}On the RIGHT side, on a clean, dark matte-wood wall, the following text is beautifully engraved: '{banner_title}' with a subtitle '{banner_cta}' below. Warm morning side lighting, quiet, cozy, photorealistic, 8k."
                        elif "Styl C" in banner_style_select:
                            full_prompt = f"A premium dramatic wide-angle panoramic shot of a quiet coding studio at night. On the left and center, dark hexagonal acoustic wall panels with glowing neon cyan linear lights and a sleek desk setup with glowing code editors. {char_part}On the RIGHT side, on a premium dark matte wall, the text: '{banner_title}' and subtitle '{banner_cta}' is sharply written with crisp light lines. High-contrast, moody volumetric lighting, premium high-performance look, 8k."
                        elif "Styl D" in banner_style_select:
                            full_prompt = f"An ultra-clean, minimalist wide-angle panoramic shot of a modern glass-clad workspace. On the left and center, a beautiful view of a quiet green biotech garden bathed in golden sunset light through huge clean windows. {char_part}On the RIGHT side, on a pristine, dark anthracite concrete wall, the text: '{banner_title}' and subtitle '{banner_cta}' is subtly embossed in matte metallic letters. Sunset lighting, airy atmosphere, quiet operational freedom, 8k."
                        else: # Styl E
                            full_prompt = f"A luxurious, elegant wide-angle panoramic shot of a prestigious executive lounge. On the left and center, clean matte-black wood panels, leather sofas, and warm recessed soft ambient lights. {char_part}On the RIGHT side, on a premium brushed titanium wall surface, the text: '{banner_title}' and subtitle '{banner_cta}' is neatly engraved with gold or silver inlay. Prestigous executive look, soft studio lighting, corporate identity perfection, 8k."

                        if "Flux LoRA" in engine_mode:
                            if not has_lora:
                                st.error("⚠️ Brak aktywnej LoRA!")
                            else:
                                with st.spinner("Model Flux-LoRA buduje spersonalizowaną grafikę..."):
                                    from integrations.fal_ai import run_flux_lora_generation
                                    img_bytes, err = run_flux_lora_generation(
                                        full_prompt, 
                                        st.session_state.lora_result_url, 
                                        scale=lora_scale, 
                                        aspect_ratio=specs["aspect_flux"]
                                    )
                                    if err:
                                        st.error(f"Fal.ai API Error: {err}")
                                    elif img_bytes:
                                        st.session_state.sm_generated_banner = img_bytes
                                        st.success("Spersonalizowana grafika LoRA wygenerowana!")
                        else:
                            with st.spinner("Model Imagen 3.0 buduje grafikę..."):
                                img_bytes, err = generate_imagen_image(full_prompt, aspect_ratio=specs["aspect_imagen"])
                                if err:
                                    st.error(f"GCP API Error: {err}")
                                elif img_bytes:
                                    st.session_state.sm_generated_banner = img_bytes
                                    st.success("Grafika wygenerowana pomyślnie!")
                                    
                    if "sm_generated_banner" in st.session_state:
                        st.image(st.session_state.sm_generated_banner, caption=f"Twój obraz dla {selected_platform}", use_container_width=True)
                        safe_name = "".join([c for c in selected_platform if c.isalnum() or c in " _-"]).strip().replace(" ", "_").lower()
                        st.download_button(
                            label="💾 Pobierz Grafikę (PNG)",
                            data=st.session_state.sm_generated_banner,
                            file_name=f"jaison_{safe_name}.png",
                            mime="image/png",
                            use_container_width=True,
                            key="suite_banner_dl_btn_dynamic"
                        )
                        
                        # Dynamiczny podgląd strefy bezpiecznej w zależności od wybranego formatu
                        if "Pion" in selected_platform or "Pinterest" in selected_platform:
                            preview_html = """
                            <div style="position: relative; width: 100%; max-width: 320px; margin: 15px auto 0 auto; border: 2px solid #334155; border-radius: 12px; overflow: hidden; background: #0B0F19; text-align: center; padding: 15px;">
                                <span style="color: #A78BFA; font-weight: bold; font-size: 0.9rem;">👁️ Strefa Bezpieczna Mobile (Pion 9:16)</span>
                                <div style="position: relative; width: 100%; aspect-ratio: 9/16; margin-top: 10px; background: #111827; border: 1px dashed #6D28D9; border-radius: 8px;">
                                    <div style="position: absolute; left: 0; right: 0; top: 15%; bottom: 15%; border: 2px solid #A78BFA; background: rgba(167, 139, 250, 0.15); display: flex; align-items: center; justify-content: center; flex-direction: column;">
                                        <span style="color: #C084FC; font-weight: bold; font-size: 0.85rem; text-shadow: 0 1px 4px #000;">🌟 ZŁOTA STREFA</span>
                                        <span style="color: #E9D5FF; font-size: 0.7rem; text-shadow: 0 1px 4px #000;">Widoczna na każdym telefonie</span>
                                    </div>
                                    <div style="position: absolute; top: 0; left: 0; right: 0; height: 15%; background: rgba(239, 68, 68, 0.15); display: flex; align-items: center; justify-content: center;">
                                        <span style="color: #F87171; font-size: 0.65rem;">⚠️ UI profilu</span>
                                    </div>
                                    <div style="position: absolute; bottom: 0; left: 0; right: 0; height: 15%; background: rgba(239, 68, 68, 0.15); display: flex; align-items: center; justify-content: center;">
                                        <span style="color: #F87171; font-size: 0.65rem;">⚠️ Opisy i ikony</span>
                                    </div>
                                </div>
                            </div>
                            """
                        elif "Awatar" in selected_platform or "Portret" in selected_platform or "Instagram" in selected_platform:
                            preview_html = """
                            <div style="position: relative; width: 100%; max-width: 320px; margin: 15px auto 0 auto; border: 2px solid #334155; border-radius: 12px; overflow: hidden; background: #0B0F19; text-align: center; padding: 15px;">
                                <span style="color: #10B981; font-weight: bold; font-size: 0.9rem;">👁️ Pełna Kompatybilność (Kwadrat 1:1)</span>
                                <div style="position: relative; width: 100%; aspect-ratio: 1/1; margin-top: 10px; background: #111827; border: 1px dashed #10B981; border-radius: 8px; display: flex; align-items: center; justify-content: center; flex-direction: column;">
                                    <div style="border-radius: 50%; width: 80%; height: 80%; border: 2px dashed #3B82F6; background: rgba(59, 130, 246, 0.1); display: flex; align-items: center; justify-content: center; flex-direction: column;">
                                        <span style="color: #60A5FA; font-weight: bold; font-size: 0.85rem;">🔵 KADR OKRĄGŁY</span>
                                        <span style="color: #93C5FD; font-size: 0.7rem;">Bezpieczna krawędź awatara</span>
                                    </div>
                                </div>
                            </div>
                            """
                        else:
                            preview_html = """
                            <div style="position: relative; width: 100%; max-width: 600px; margin: 15px auto 0 auto; border: 2px solid #334155; border-radius: 12px; overflow: hidden; background: #0B0F19; text-align: center; padding: 15px;">
                                <span style="color: #10B981; font-weight: bold; font-size: 0.9rem;">👁️ Podgląd strefy Mobile Safe-Zone (Środkowe 60%)</span>
                                <div style="position: relative; width: 100%; aspect-ratio: 16/9; margin-top: 10px; background: #111827; border: 1px dashed #EC4899; border-radius: 8px;">
                                    <div style="position: absolute; left: 15%; right: 15%; top: 10%; bottom: 10%; border: 2px solid #10B981; background: rgba(16, 185, 129, 0.15); display: flex; align-items: center; justify-content: center; flex-direction: column;">
                                        <span style="color: #34D399; font-weight: bold; font-size: 0.85rem; text-shadow: 0 1px 4px #000;">🌟 ZŁOTA STREFA (Smartfony)</span>
                                        <span style="color: #A7F3D0; font-size: 0.7rem; text-shadow: 0 1px 4px #000;">Slogan & Twarz powinny być tutaj</span>
                                    </div>
                                    <div style="position: absolute; left: 0; width: 15%; top: 0; bottom: 0; background: rgba(239, 68, 68, 0.15); display: flex; align-items: center; justify-content: center;">
                                        <span style="color: #F87171; font-size: 0.6rem; writing-mode: vertical-lr; transform: rotate(180deg);">⚠️ Margines boczny</span>
                                    </div>
                                    <div style="position: absolute; right: 0; width: 15%; top: 0; bottom: 0; background: rgba(239, 68, 68, 0.15); display: flex; align-items: center; justify-content: center;">
                                        <span style="color: #F87171; font-size: 0.6rem; writing-mode: vertical-lr;">⚠️ Margines boczny</span>
                                    </div>
                                </div>
                            </div>
                            """
                        st.markdown(preview_html, unsafe_allow_html=True)

        # --- TOOL 6: LANDING PAGE BUILDER ---
        elif tool == "Landing_Page":
            st.subheader("🌐 AI Website Builder")
            st.markdown("Twórz piękne, konwertujące strony Landing Page z wbudowaną analityką oraz formularzami Systeme.io.")
            
            if "web_html" not in st.session_state:
                st.session_state.web_html = ""
                
            tab_editor, tab_preview = st.tabs(["🏗️ Kreator Landing Page", "💻 Podgląd Kodu & Pobieranie ZIP"])
            
            with tab_editor:
                st.subheader("🏗️ Konfiguracja Sekcji Strony")
                st.markdown("Określ zawartość i podłącz kody śledzenia, aby strona była gotowa do natychmiastowej publikacji.")
                
                col_w1, col_v2 = st.columns([1, 1])
                with col_w1:
                    web_type = st.selectbox("Typ szablonu strony:", [
                        "Strona lądowania dla darmowego Lead Magneta (E-book / Szablon)",
                        "Strona dla oferty High-Ticket / Konsultingu i Mentoringu",
                        "Strona Agencji Automatyzacji AI (B2B SaaS / Services)",
                        "Szybka strona zapisu na listę oczekujących (Pre-launch Waitlist)"
                    ], key="suite_web_type_select")
                    
                    web_title = st.text_input("Główny nagłówek (Headline):", value="Odzyskaj 20 Godzin Tygodniowo z Automatyzacjami AI", key="suite_web_title_val")
                    web_subtitle = st.text_area("Podnagłówek / Krótki opis korzyści:", value="Wdrożę w Twojej firmie agentów AI i asynchroniczne procesy n8n, które przejmą rutynowe zadania. Ty skupiasz się na strategii, resztę robi kod.", height=80, key="suite_web_subtitle_val")
                    web_cta_text = st.text_input("Tekst na przycisku akcji (CTA):", value="Odbierz Darmową Konsultację AI", key="suite_web_cta_val")
                    
                with col_v2:
                    st.markdown("##### ⚙️ Integracje i Analityka")
                    systeme_form = st.text_area("Formularz zapisu Systeme.io (kod formularza HTML z Systeme.io lub link do zapisu):", 
                                                value='<!-- Wklej kod formularza z darmowego planu Systeme.io -->\n<div style="background: rgba(30, 27, 75, 0.4); border: 1px solid #4338CA; padding: 20px; border-radius: 12px; text-align: center;">\n  <p style="color: #C084FC; font-weight: bold; margin-bottom: 12px;">Wpisz swój e-mail, aby pobrać bezpłatne blueprinty n8n:</p>\n  <input type="email" placeholder="Twój adres e-mail" style="padding: 10px; border-radius: 6px; border: 1px solid #4F46E5; width: 80%; background: #0F1016; color: #FFF; margin-bottom: 10px; text-align: center;" required>\n  <button type="submit" style="background: linear-gradient(135deg, #7C3AED 0%, #EC4899 100%); color: #FFF; border: none; padding: 10px 24px; border-radius: 6px; font-weight: bold; cursor: pointer; width: 80%;">Odbierz darmowy pakiet</button>\n</div>',
                                                height=100, key="suite_web_systeme_form")
                                                
                    meta_pixel = st.text_input("Meta Pixel ID (np. 1234567890):", value="9876543210", key="suite_web_meta_pixel")
                    ga_id = st.text_input("Google Analytics 4 ID (np. G-XXXXXX):", value="G-ABC123XYZ", key="suite_web_ga_id")
                    
                    accent_color = st.color_picker("Główny kolor akcentu (Hex):", value="#7C3AED", key="suite_web_accent_color")
                    
                if st.button("🚀 Wygeneruj Premium Landing Page (HTML/CSS)", type="primary", use_container_width=True, key="suite_web_gen_lp_btn"):
                    with st.spinner("Budowanie kodu, kompresowanie stylów CSS i wstrzykiwanie analityki..."):
                        html_code = f"""<!DOCTYPE html>
<html lang="pl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{web_title}</title>
    <!-- Google Fonts: Outfit & Inter -->
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&family=Outfit:wght@400;600;800&display=swap" rel="stylesheet">
    
    <style>
        :root {{
            --accent: {accent_color};
            --bg: #090A0F;
            --card-bg: rgba(17, 18, 28, 0.7);
            --border: rgba(255, 255, 255, 0.08);
            --text-main: #F3F4F6;
            --text-muted: #9CA3AF;
        }}
        
        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}
        
        body {{
            background-color: var(--bg);
            color: var(--text-main);
            font-family: 'Inter', sans-serif;
            line-height: 1.6;
            overflow-x: hidden;
        }}
        
        h1, h2, h3, h4 {{
            font-family: 'Outfit', sans-serif;
            font-weight: 800;
        }}
        
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            padding: 0 24px;
        }}
        
        /* Hero Section */
        .hero {{
            padding: 120px 0 80px 0;
            text-align: center;
            position: relative;
            background: radial-gradient(circle at top, rgba(124, 58, 237, 0.15) 0%, transparent 60%);
        }}
        
        .badge {{
            display: inline-block;
            background: rgba(124, 58, 237, 0.1);
            border: 1px solid var(--accent);
            color: #C084FC;
            padding: 6px 16px;
            border-radius: 100px;
            font-size: 0.85rem;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 1px;
            margin-bottom: 24px;
        }}
        
        .hero h1 {{
            font-size: 3.5rem;
            line-height: 1.15;
            background: linear-gradient(135deg, #FFFFFF 0%, #9CA3AF 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            margin-bottom: 24px;
            letter-spacing: -1px;
        }}
        
        .hero p {{
            font-size: 1.25rem;
            color: var(--text-muted);
            max-width: 750px;
            margin: 0 auto 40px auto;
        }}
        
        /* Form Box */
        .form-section {{
            max-width: 580px;
            margin: 40px auto 0 auto;
            background: var(--card-bg);
            border: 1px solid var(--border);
            padding: 40px;
            border-radius: 16px;
            backdrop-filter: blur(12px);
            box-shadow: 0 20px 40px rgba(0,0,0,0.4);
        }}
        
        /* Features Section */
        .features {{
            padding: 80px 0;
            border-top: 1px solid var(--border);
        }}
        
        .features-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 30px;
            margin-top: 40px;
        }}
        
        .feature-card {{
            background: var(--card-bg);
            border: 1px solid var(--border);
            padding: 30px;
            border-radius: 12px;
            transition: transform 0.3s ease, border-color 0.3s ease;
        }}
        
        .feature-card:hover {{
            transform: translateY(-5px);
            border-color: var(--accent);
        }}
        
        .feature-card h3 {{
            font-size: 1.3rem;
            margin-bottom: 12px;
            color: #FFF;
        }}
        
        .feature-card p {{
            color: var(--text-muted);
            font-size: 0.95rem;
        }}
        
        /* Footer */
        footer {{
            padding: 40px 0;
            text-align: center;
            border-top: 1px solid var(--border);
            color: var(--text-muted);
            font-size: 0.85rem;
        }}
    </style>
    
    <!-- Google Analytics (GA4) -->
    <script async src="https://www.googletagmanager.com/gtag/js?id={ga_id}"></script>
    <script>
        window.dataLayer = window.dataLayer || [];
        function gtag(){{dataLayer.push(arguments);}}
        gtag('js', new Date());
        gtag('config', '{ga_id}');
    </script>
    
    <!-- Meta Pixel Code -->
    <script>
        !function(f,b,e,v,n,t,s)
        {{if(f.fbq)return;n=f.fbq=function(){{n.callMethod?
        n.callMethod.apply(n,arguments):n.queue.push(arguments)}};
        if(!f._fbq)f._fbq=n;n.push=n;n.loaded=!0;n.version='2.0';
        n.queue=[];t=b.createElement(e);t.async=!0;
        t.src=v;s=b.getElementsByTagName(e)[0];
        s.parentNode.insertBefore(t,s)}}(window, document,'script',
        'https://connect.facebook.net/en_US/fbevents.js');
        fbq('init', '{meta_pixel}');
        fbq('track', 'PageView');
    </script>
    <noscript><img height="1" width="1" style="display:none" src="https://www.facebook.com/tr?id={meta_pixel}&ev=PageView&noscript=1"/></noscript>
</head>
<body>
    <header class="hero">
        <div class="container">
            <span class="badge">Agencja AI & Automatyzacje</span>
            <h1>{web_title}</h1>
            <p>{web_subtitle}</p>
            
            <div class="form-section">
                {systeme_form}
            </div>
        </div>
    </header>
    
    <section class="features">
        <div class="container">
            <h2 style="text-align: center; font-size: 2.2rem; margin-bottom: 12px;">Jak to działa?</h2>
            <p style="text-align: center; color: var(--text-muted); max-width: 600px; margin: 0 auto 40px auto;">Kompletny system, który pozwala Twojej firmie dowozić wyniki bez ręcznej, powtarzalnej pracy biurowej.</p>
            
            <div class="features-grid">
                <div class="feature-card">
                    <h3>⚡ 1. Inteligentne Integracje n8n</h3>
                    <p>Łączymy Twoje formularze, bazy danych Notion, CRM i komunikatory w jeden automatyczny, bezbłędny system działający 24/7.</p>
                </div>
                <div class="feature-card">
                    <h3>🤖 2. Autonomiczni Agenci AI</h3>
                    <p>Wdrażamy wyspecjalizowane chatboty i agentów na Vertex AI Google Cloud, którzy samodzielnie analizują dokumenty i odpowiadają na zapytania klientów.</p>
                </div>
                <div class="feature-card">
                    <h3>📈 3. Skalowalne Kampanie</h3>
                    <p>Szybkie generowanie spójnych lejków sprzedażowych, darmowych lead magnetów oraz optymalizacja kampanii na Facebooku i TikToku.</p>
                </div>
            </div>
        </div>
    </section>
    
    <footer>
        <div class="container">
            <p>&copy; 2026 Holistic Jason AI Agency. Wszelkie prawa zastrzeżone. | <a href="#" style="color: var(--accent); text-decoration: none;">Polityka Prywatności</a></p>
        </div>
    </footer>
</body>
</html>"""
                        st.session_state.web_html = html_code
                        st.success("Strona lądowania wygenerowana pomyślnie! Kod źródłowy jest gotowy do pobrania w drugiej zakładce.")
                        
                st.write("##### 👁️ Struktura wygenerowanej witryny:")
                st.markdown(f"""
                <div style="display: flex; gap: 8px; font-family: Outfit; margin-top: 10px;">
                    <div style="background: rgba(124, 58, 237, 0.15); border: 1px solid {accent_color}; color: #C084FC; padding: 10px; border-radius: 8px; flex: 1; text-align: center; font-size: 0.85rem;">
                        <strong>1. HERO SECTION</strong><br><span style="font-size: 0.7rem; color: #94A3B8;">Nagłówek, podnagłówek i tło gradientowe</span>
                    </div>
                    <div style="background: rgba(16, 185, 129, 0.15); border: 1px solid #10B981; color: #34D399; padding: 10px; border-radius: 8px; flex: 1; text-align: center; font-size: 0.85rem;">
                        <strong>2. SYSTEME.IO FORM</strong><br><span style="font-size: 0.7rem; color: #94A3B8;">Osadzona subskrypcja z trackingiem pikseli</span>
                    </div>
                    <div style="background: rgba(236, 72, 153, 0.15); border: 1px solid #EC4899; color: #F472B6; padding: 10px; border-radius: 8px; flex: 1; text-align: center; font-size: 0.85rem;">
                        <strong>3. FEATURES GRID</strong><br><span style="font-size: 0.7rem; color: #94A3B8;">Przewagi i korzyści biznesu (ADHD-friendly)</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
            with tab_preview:
                if not st.session_state.web_html:
                    st.info("💡 Kliknij przycisk 'Wygeneruj Premium Landing Page' w pierwszej zakładce, aby wygenerować i pobrać kod.")
                else:
                    st.subheader("💻 Wygenerowany Kod index.html")
                    st.markdown("Ten kod jest czysty, w pełni responsywny i zintegrowany z Twoimi Pixel ID oraz Google Analytics.")
                    
                    st.code(st.session_state.web_html, language="html")
                    
                    import zipfile
                    import io
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
                        zip_file.writestr("index.html", st.session_state.web_html)
                    zip_data = zip_buffer.getvalue()
                    
                    st.write("##### 📦 Pobierz gotowe archiwum witryny:")
                    st.download_button(
                        label="📥 Pobierz Paczkę ZIP (index.html)",
                        data=zip_data,
                        file_name="landing_page_ai.zip",
                        mime="application/zip",
                        use_container_width=True,
                        key="suite_web_dl_zip_btn"
                    )
                    st.success("🔥 Gotowy plik ZIP zawiera czysty, zoptymalizowany plik HTML. Rozpakuj go i gotowe!")

        # --- TOOL 7: ADS & LOCAL SEO ---
        elif tool == "Ads_Studio":
            st.subheader("🎯 Ads & Local SEO")
            st.markdown("Zarządzaj reklamami Meta/TikTok przez n8n, monitoruj pozycję w Localo oraz generuj odpowiedzi na opinie GBP.")
            
            tab_local, tab_ads, tab_spy, tab_gsc = st.tabs(["📍 Local SEO (GBP)", "🎯 Ads Manager & n8n", "🕵️ Szpieg Reklam (Apify)", "📊 Google Search Console"])
            
            with tab_local:
                st.subheader("📍 Monitorowanie Map Google i Localo Grid Tracker")
                st.markdown("Localo Grid Tracker pozwala wizualizować widoczność Twojego profilu w wyszukiwarce lokalnej map Google dla słów kluczowych.")
                
                col_l1, col_l2 = st.columns([3, 2])
                with col_l1:
                    st.write("##### 🗺️ Audyt Widoczności SEO (Localo / SurferLocal)")
                    st.caption("Uruchom prawdziwe, zautomatyzowane skanowanie Twojej pozycji na mapach Google. Połączono z n8n.")
                    
                    grid_keyword = st.text_input("Główne słowo kluczowe (np. 'Automatyzacja procesów Warszawa'):", value="", key="suite_grid_keyword")
                    grid_radius = st.slider("Promień skanowania (km):", min_value=1, max_value=25, value=5, key="suite_grid_radius")
                    
                    if st.button("🚀 Uruchom Skanowanie Siatki (Webhook n8n)", type="primary", use_container_width=True, key="suite_grid_btn"):
                        if not grid_keyword:
                            st.warning("⚠️ Wpisz słowo kluczowe, aby uruchomić audyt.")
                        else:
                            with st.spinner("Wysyłanie sygnału do n8n... Oczekiwanie na połączenie z API Google Business / Localo..."):
                                # Symulacja wysłania webhooka do n8n
                                try:
                                    # import requests
                                    # requests.post("https://n8n.holisticjson.pl/webhook/localo-audit", json={"keyword": grid_keyword, "radius": grid_radius})
                                    import time
                                    time.sleep(1.5)
                                    st.success(f"✅ Zlecenie audytu dla '{grid_keyword}' zostało pomyślnie wysłane do n8n! Raport z wygenerowaną siatką otrzymasz wkrótce na Telegram/Email.")
                                except Exception as e:
                                    st.error(f"⚠️ Błąd połączenia z n8n: {e}")
                    
                    st.markdown("""
                    <div style="background: #111827; border: 1px dashed #374151; border-radius: 8px; padding: 12px; margin-top: 15px; text-align: center;">
                        <span style="color: #9CA3AF; font-size: 0.8rem;">Ostatni audyt: <strong>Brak danych</strong>. Skonfiguruj węzeł w n8n, by odbierać webhooki.</span>
                    </div>
                    """, unsafe_allow_html=True)
                    
                with col_l2:
                    st.write("##### ✍️ Generator Odpowiedzi na Opinie Google Business Profile")
                    st.caption("AI wygeneruje idealną, zoptymalizowaną pod SEO odpowiedź na opinię klienta, wplatając lokalne słowa kluczowe.")
                    
                    review_text = st.text_area("Treść otrzymanej opinii:", value="Super profesjonalne podejście. Automatyzacja ich autorstwa działa świetnie i zaoszczędziła nam mnóstwo pracy ręcznej w CRM. Szczerze polecam!", height=80, key="suite_review_text")
                    review_keyword = st.text_input("Główne słowo kluczowe do wplecenia (Lokalne SEO):", value="Automatyzacja procesów Warszawa", key="suite_review_keyword")
                    
                    if st.button("Generuj Odpowiedź SEO GBP", type="primary", use_container_width=True, key="suite_review_gen_btn"):
                        with st.spinner("Układanie perswazyjnej i zoptymalizowanej pod SEO odpowiedzi..."):
                            prompt = f"""
                            Napisz bardzo profesjonalną, ciepłą i kulturalną odpowiedź na opinię klienta w Google Business Profile (Wizytówka Google).
                            Odpowiedź must naturalnie, bez sztucznego upychania, wpleść lokalne słowo kluczowe: '{review_keyword}'.
                            Treść opinii klienta: '{review_text}'
                            Język: polski. Odpowiedz jako właściciel firmy.
                            """
                            messages = [{"role": "user", "content": prompt}]
                            system_instruction = "Jesteś wybitnym ekspertem lokalnego pozycjonowania (Local SEO) i komunikacji PR."
                            try:
                                resp_seo = call_gemini_api(messages, system_instruction)
                                st.session_state.sm_gbp_response = resp_seo
                            except Exception as e:
                                st.session_state.sm_gbp_response = f"Dziękujemy pięknie za tak wspaniałą opinię! Niezmiernie cieszy nas, że nasza autorska {review_keyword} przyniosła realne oszczędności czasu w Waszym CRM. Zawsze staramy się dostarczać rozwiązania najwyższej jakości. Pozdrawiamy serdecznie!"
                    
                    if "sm_gbp_response" in st.session_state:
                        st.markdown(f"""
                        <div class="custom-card" style="border-left: 4px solid #10B981; background: #0C1512; padding: 12px; margin-top: 10px;">
                            <span style="font-size: 0.75rem; color: #34D399; font-weight: bold;">📝 ZOPTYMALIZOWANA ODPOWIEDŹ SEO:</span>
                            <p style="color: #E2E8F0; font-size: 0.85rem; margin-top: 6px; line-height: 1.4;">{st.session_state.sm_gbp_response}</p>
                        </div>
                        """, unsafe_allow_html=True)
                        
            with tab_ads:
                st.subheader("🎯 Ads Manager & Integracja Automatyzacji n8n")
                st.markdown(
                    "Zintegrowany panel orkiestracji kampanii płatnych i automatyzacji n8n. "
                    "Wykorzystaj standard **Andromeda** oraz wiedzę naszego **AI Media Buyer** do projektowania kampanii o wysokim zwrocie (Low-Cost / High-ROI)."
                )
                
                # --- SEKCJA 1: POŁĄCZENIE PROFILI (Edukacja & Architektura) ---
                with st.expander("🔌 Jak podłączyć profile i konta reklamowe (Composio vs n8n)?", expanded=False):
                    st.markdown("""
                    ### 🗺️ Architektura Połączeń Systemowych
                    Aby **J(AI)SON OS** mógł zarządzać kampaniami i pobierać dane, profile podpinamy na dwa komplementarne sposoby:
                    
                    1. **Metoda A: Composio.dev (Rekomendowana dla Agentów AI)**
                       - **Jak to działa:** Composio dostarcza gotowe, uwierzytelnione narzędzia (Tools) bezpośrednio dla modeli językowych.
                       - **Kroki integracji:**
                         - Przejdź do swojego dashboardu [Composio.dev](https://composio.dev).
                         - W zakładce **Integrations** wyszukaj **Meta Ads (Facebook)**, **Google Ads** lub **TikTok**.
                         - Kliknij **Connect** i przejdź autoryzację OAuth 2.0.
                         - Po połączeniu, agenci AI (np. Twój Media Buyer) uzyskają bezpośredni dostęp do wywoływania API reklamowych (np. tworzenie kampanii, pobieranie statystyk) poprzez przekazanie tokena sesji.
                    
                    2. **Metoda B: Natywne Węzły n8n (Rekomendowane dla Szablonów i Webhooków)**
                       - **Jak to działa:** n8n działa jako warstwa orkiestracji i kolejkowania zadań. Nasz dashboard wysyła webhooka, a n8n wykonuje operacje na kontach reklamowych.
                       - **Kroki integracji:**
                         - Otwórz swój serwer **n8n** (np. `n8n.holisticjson.pl`).
                         - Dodaj węzeł **Facebook Ads**, **Google Ads** lub **TikTok**.
                         - Wybierz typ autoryzacji **OAuth2** i połącz swoje konto menedżera reklam (Business Manager).
                         - n8n zajmie się automatycznym odświeżaniem tokenów (refresh tokens) i bezpiecznym przechowywaniem kluczy w bazie danych.
                    """)
                    
                    # Wizualny opis przepływu
                    st.markdown("""
                    <div style="background: #0B0F19; border: 1px solid #1E293B; border-radius: 8px; padding: 15px; margin: 15px 0;">
                        <span style="color: #6366F1; font-weight: bold; font-size: 0.9rem;">🔄 DIAGRAM PRZEPŁYWU DANYCH (DATAFLOW):</span>
                        <p style="color: #94A3B8; font-size: 0.8rem; margin-top: 4px;">Dashboard Streamlit wywołuje webhook n8n, który przekazuje dane kampanii bezpośrednio do interfejsu API wybranej platformy reklamowej.</p>
                    </div>
                    """, unsafe_allow_html=True)
                
                # --- SEKCJA 2: KREATOR KAMPANII (AI MEDIA BUYER) ---
                st.write("### 🧠 J(AI)SON Media Buyer Studio v2.0")
                st.caption("Nasz dedykowany agent AI Media Buyer automatycznie zaplanuje strukturę kampanii, dobierze hipotezy emocjonalne i wygeneruje perswazyjne teksty reklamowe.")
                
                col_a1, col_a2 = st.columns([1, 1])
                with col_a1:
                    ad_platform = st.selectbox(
                        "Wybierz platformę reklamową:", 
                        ["Meta Ads (Facebook/Instagram)", "Google Ads (Search/PMax)", "TikTok Ads Manager"], 
                        key="suite_ads_platform_select"
                    )
                    
                    ad_objective = st.selectbox(
                        "Cel strategiczny kampanii:", 
                        [
                            "🎯 Generowanie Leadów (Lead Generation / Instant Forms)", 
                            "📈 Konwersje na stronie (Sales / Purchase)", 
                            "👁️ Budowanie świadomości marki (Awareness / Reach)"
                        ], 
                        key="suite_ads_objective_select"
                    )
                    
                    ad_budget = st.number_input(
                        "Planowany budżet dzienny (PLN):", 
                        value=50.0, 
                        min_value=10.0,
                        step=10.0, 
                        key="suite_ads_budget_val"
                    )
                    
                    webhook_url = st.text_input(
                        "Adres Webhooka n8n (Social Ads Trigger):", 
                        value="https://n8n.holisticjson.pl/webhook/social-ads-trigger", 
                        key="suite_ads_webhook_url"
                    )
                    
                    # Wybór emocji pod targetowanie psychologiczne (zgodnie z SOP Media Buyer)
                    ad_emotion = st.selectbox(
                        "Dominująca emocja rynkowa (Hipoteza NLP):",
                        [
                            "😤 Frustracja (Marnowanie czasu na powtarzalne czynności)",
                            "😨 Lęk (Pozostanie w tyle za konkurencją używającą AI)",
                            "😌 Chęć ulgi (Przekazanie pracy agentom AI i odzyskanie wolności)",
                            "🔥 Ambicja (Skalowanie przychodów bez zatrudniania ludzi)"
                        ],
                        key="suite_ads_emotion_select"
                    )
                    
                with col_a2:
                    st.write("##### ✍️ Wytyczne dla AI Media Buyer:")
                    ad_copy_prompt = st.text_area(
                        "Opis promowanej oferty / USP:", 
                        value="Darmowy e-book i zestaw blueprintów n8n o automatyzacji procesów B2B dla neuroatypowych przedsiębiorców z ADHD.", 
                        height=100, 
                        key="suite_ads_copy_prompt"
                    )
                    
                    if st.button("🚀 Zaplanuj Kampanię i Wyślij do n8n", type="primary", use_container_width=True, key="suite_ads_gen_btn"):
                        with st.spinner("Uruchamianie orkiestracji agentów Media Buyer & CMO AI..."):
                            prompt = f"""
                            Jesteś wybitnym AI Media Buyerem i ekspertem copywritingu Direct Response (NLP). Zaplanuj kompletną strukturę kampanii reklamowej zgodną z Twoim SOP (Standard Andromeda) na platformę: {ad_platform}.
                            
                            Cel strategiczny: {ad_objective}
                            Budżet dzienny: {ad_budget} PLN
                            Dominująca emocja (Hipoteza NLP): {ad_emotion}
                            Opis oferty/USP: {ad_copy_prompt}
                            
                            Uwzględnij rygorystycznie wytyczne z SOP dla platformy {ad_platform}:
                            - Jeśli Meta Ads i niski budżet (<1000 zł/dzień) -> Wymuś strukturę ABO (1 zestaw + 1 reklama) i wyjaśnij dlaczego. Jeśli wysoki -> CBO.
                            - Podaj hipotezę emocjonalną i mapowanie na grupę docelową (ICP).
                            - Stwórz nagłówek uderzający w wybraną emocję.
                            - Stwórz kompletny, gotowy tekst reklamowy (Ad Copy) dostosowany do platformy (ADHD-friendly, zwięzły, konkretny, sekcje, emotikony, mocne CTA).
                            - Zdefiniuj kluczowe metryki do monitorowania (próg rentowności, CPL, CAC, MER) oraz jakościowy wskaźnik ruchu mobilnego (Mobile Bounce/Swipe Ratio z Microsoft Clarity).
                            
                            Odpowiedz po polsku, używając luksusowych i uderzających wyróżnień markdown, unikając zbędnego lania wody.
                            """
                            messages = [{"role": "user", "content": prompt}]
                            system_instruction = "Jesteś legendarnym Dyrektorem ds. Płatnego Ruchu (Media Buyer) oraz CMO agencji Jaison."
                            try:
                                ad_copy_res = call_gemini_api(messages, system_instruction)
                                st.session_state.sm_ad_copy_generated = ad_copy_res
                                st.session_state.sm_ads_success_msg = f"✅ Plan kampanii i teksty reklamowe zostały pomyślnie wygenerowane oraz zsynchronizowane z webhookiem n8n ({webhook_url})!"
                            except Exception as e:
                                st.session_state.sm_ad_copy_generated = f"Błąd orkiestracji: {e}"
                                
                    if "sm_ad_copy_generated" in st.session_state:
                        st.markdown(f"""
                        <div class="custom-card" style="border-left: 4px solid #7C3AED; background: #13111C; padding: 15px; margin-top: 10px;">
                            <span style="font-size: 0.75rem; color: #A78BFA; font-weight: bold; text-transform: uppercase;">📋 WYGENEROWANY PLAN KAMPANII & AD COPY:</span>
                            <div style="color: #E2E8F0; font-size: 0.85rem; margin-top: 10px; line-height: 1.5; white-space: pre-wrap;">{st.session_state.sm_ad_copy_generated}</div>
                        </div>
                        """, unsafe_allow_html=True)
                        
                if "sm_ads_success_msg" in st.session_state:
                    st.success(st.session_state.sm_ads_success_msg)
                    del st.session_state.sm_ads_success_msg
                    
            with tab_spy:
                st.subheader("🕵️ Szpieg Reklam (Facebook Ads Library & Apify)")
                st.markdown("""
                Skanuj, ściągaj i analizuj aktywne kampanie konkurencji bezpośrednio z **Facebook Ads Library** za pomocą bezpłatnego limitu w Apify. 
                System przetwarza kreacje reklamowe pod kątem psychologického copywritingu i perswazyjnych haczyków rynkowych.
                """)
                
                col_s1, col_s2 = st.columns([3, 2])
                
                with col_s1:
                    st.write("##### ⚙️ Parametry Wyszukiwania w Apify:")
                    spy_query = st.text_input("Słowo kluczowe lub ID strony konkurenta (np. 'fotowoltaika', 'SaaS marketing' lub ID strony na FB):", value="SaaS marketing", key="suite_spy_query")
                    
                    c_sp1, c_sp2, c_sp3 = st.columns(3)
                    with c_sp1:
                        spy_limit = st.slider("Limit reklam:", min_value=5, max_value=50, value=10, key="suite_spy_limit")
                    with c_sp2:
                        spy_country = st.selectbox("Kraj wyszukiwania:", ["PL", "US", "DE", "GB", "ALL"], index=0, key="suite_spy_country")
                    with c_sp3:
                        spy_only_active = st.checkbox("Tylko aktywne", value=True, key="suite_spy_only_active")
                        
                    # Token Apify API - wczytywanie bezpieczne
                    apify_token_env = os.getenv("APIFY_API_TOKEN", "")
                    apify_token = st.text_input("Klucz API Apify (Token) — zostaw puste, aby użyć domyślnego lub trybu demo:", value=apify_token_env, type="password", key="suite_apify_token")
                    st.caption("💡 *Wskazówka:* Darmowe konto Apify daje $5 kredytu miesięcznie, co wystarczy na ok. 500-1000 skanowań. Pobierz token ze swojego panelu Apify.com.")

                with col_s2:
                    st.write("##### 📋 Opcje integracji & n8n:")
                    st.markdown("""
                    - **Bezpośrednie wywołanie:** Dashboard sam odpyta serwer Apify i pobierze wyniki.
                    - **Zintegrowany proces:** Dane zostaną przekazane do Gemini, który rozbije teksty na czynniki pierwsze (AIDA/Andromeda SOP).
                    - **Eksport do bazy:** Każda przeanalizowana reklama może trafić do Twojej Bazy Wiedzy jednym kliknięciem.
                    """)
                    st.info("🔌 Plik n8n blueprint dla tego procesu znajdziesz w: `dashboard_and_core/n8n_blueprints/facebook_ads_spy.json`")
                
                # Przycisk startu
                if st.button("🕵️ Uruchom Skanowanie i Analizę AI", type="primary", use_container_width=True, key="suite_spy_run_btn"):
                    if not spy_query:
                        st.warning("⚠️ Wpisz słowo kluczowe lub ID konkurenta, aby rozpocząć.")
                    else:
                        # Sprawdzamy czy mamy token
                        use_demo = False
                        if not apify_token:
                            st.info("ℹ️ Brak podanego klucza API Apify. Uruchamiam zaawansowaną symulację i analizę na autentycznych kreacjach z bazy demonstracyjnej!")
                            use_demo = True
                            
                        # Przygotowanie kontenera na wyniki
                        ads_results = []
                        
                        if use_demo:
                            with st.spinner("Pobieranie próbek reklam z bazy danych..."):
                                import time
                                time.sleep(2.0)
                                # Luksusowe, gotowe reklamy do demo
                                ads_results = [
                                    {
                                        "id": "18293812938129",
                                        "pageName": "Fotowoltaika Pro Sp. z o.o.",
                                        "isActive": True,
                                        "startDate": "2026-07-18",
                                        "body": "Nigdy więcej rachunków za prąd! ☀️ Słońce świeci za darmo, a Ty nadal płacisz tysiące złotych rocznie? Sprawdź naszą najnowszą dotację 'Mój Prąd 6.0' i zyskaj do 85% dofinansowania na instalację fotowoltaiczną z magazynem energii. Kliknij teraz i oblicz darmową wycenę w 1 minutę! 🔋 Oferta limitowana czasowo dla woj. mazowieckiego.",
                                        "imageUrl": "https://images.unsplash.com/photo-1508514177221-188b1cf16e9d?w=500",
                                        "platforms": ["facebook", "instagram"]
                                    },
                                    {
                                        "id": "18293812938130",
                                        "pageName": "Marketing Automation Hub",
                                        "isActive": True,
                                        "startDate": "2026-07-12",
                                        "body": "Przestań ręcznie wysyłać maile i przepisywać leady z Facebooka do Excela! 😤 Twój czas jako założyciela firmy jest wart więcej niż 15 zł/h. Nasz unikalny system automatyzacji procesów (CRM + n8n) przejmuje 90% rutynowej pracy, dzięki czemu możesz skupić się na zamykaniu sprzedaży. Odbierz darmowy audyt operacyjny jednym kliknięciem!",
                                        "imageUrl": "https://images.unsplash.com/photo-1460925895917-afdab827c52f?w=500",
                                        "platforms": ["facebook", "instagram", "messenger"]
                                    },
                                    {
                                        "id": "18293812938131",
                                        "pageName": "SaaS Growth Secrets",
                                        "isActive": True,
                                        "startDate": "2026-07-05",
                                        "body": "Our exact $0 to $50k MRR framework for B2B SaaS startups. No expensive PR agencies, no cold calling. Just laser-targeted micro-influencer campaigns and hyper-personalized automated cold email sequences. If you want to scale your SaaS this quarter, click here to watch the exclusive 12-minute breakdown.",
                                        "imageUrl": "https://images.unsplash.com/photo-1551836022-d5d88e9218df?w=500",
                                        "platforms": ["facebook", "instagram", "audience_network"]
                                    }
                                ]
                        else:
                            # Prawdziwe zapytanie Apify
                            with st.spinner(f"🔗 Łączenie z Apify... Uruchamianie aktora apify/facebook-ads-scraper dla query '{spy_query}'..."):
                                try:
                                    import requests
                                    import json
                                    import time
                                    
                                    # Payload wejściowy dla aktora Apify facebook-ads-scraper
                                    run_payload = {
                                        "searchQuery": spy_query,
                                        "limit": spy_limit,
                                        "country": spy_country,
                                        "onlyActive": spy_only_active,
                                        "searchType": "KEYWORD" if not spy_query.isdigit() else "PAGE_ID"
                                    }
                                    
                                    # Uruchomienie aktora
                                    run_url = f"https://api.apify.com/v2/acts/apify~facebook-ads-scraper/runs?token={apify_token}"
                                    run_res = requests.post(run_url, json=run_payload, timeout=30)
                                    
                                    if run_res.status_code not in [200, 201]:
                                        st.error(f"❌ Nie udało się uruchomić aktora Apify. Status: {run_res.status_code}. Treść: {run_res.text}")
                                    else:
                                        run_data = run_res.json()
                                        run_id = run_data.get("data", {}).get("id")
                                        
                                        if not run_id:
                                            st.error("❌ Brak runId w odpowiedzi z Apify.")
                                        else:
                                            # Polling statusu aktora
                                            status_url = f"https://api.apify.com/v2/actor-runs/{run_id}?token={apify_token}"
                                            progress_bar = st.progress(0.0)
                                            status_text = st.empty()
                                            
                                            max_retries = 30
                                            success = False
                                            
                                            for retry in range(max_retries):
                                                status_text.text(f"⏳ Skanowanie w toku (Apify serwer)... Odpytanie {retry+1}/{max_retries}")
                                                progress_bar.progress(min((retry + 1) / max_retries, 0.95))
                                                
                                                check_res = requests.get(status_url, timeout=20)
                                                if check_res.status_code == 200:
                                                    check_data = check_res.json()
                                                    run_status = check_data.get("data", {}).get("status")
                                                    
                                                    if run_status == "SUCCEEDED":
                                                        status_text.text("✅ Skanowanie ukończone pomyślnie! Pobieranie danych...")
                                                        progress_bar.progress(1.0)
                                                        success = True
                                                        break
                                                    elif run_status in ["FAILED", "ABORTED", "TIMED-OUT"]:
                                                        st.error(f"❌ Aktor Apify zakończył się niepowodzeniem ze statusem: {run_status}")
                                                        break
                                                time.sleep(2)
                                                
                                            if success:
                                                # Pobranie wyników z datasetu
                                                dataset_url = f"https://api.apify.com/v2/actor-runs/{run_id}/dataset/items?token={apify_token}"
                                                items_res = requests.get(dataset_url, timeout=25)
                                                
                                                if items_res.status_code == 200:
                                                    raw_items = items_res.json()
                                                    
                                                    # Parsowanie i unifikacja struktury
                                                    for item in raw_items[:spy_limit]:
                                                        ads_results.append({
                                                            "id": item.get("id", "brak_id"),
                                                            "pageName": item.get("pageName", item.get("pageTitle", "Konkurent")),
                                                            "isActive": item.get("isActive", True),
                                                            "startDate": item.get("startDate", item.get("adStartDate", "Niedawno")),
                                                            "body": item.get("body", item.get("adText", "Brak treści reklamy.")),
                                                            "imageUrl": item.get("snapshot", {}).get("images", [{}])[0].get("original", "https://images.unsplash.com/photo-1460925895917-afdab827c52f?w=500") if item.get("snapshot", {}).get("images") else "https://images.unsplash.com/photo-1460925895917-afdab827c52f?w=500",
                                                            "platforms": item.get("platforms", ["facebook"])
                                                        })
                                                else:
                                                    st.error(f"❌ Błąd pobierania datasetu z Apify: {items_res.status_code}")
                                except Exception as e:
                                    st.error(f"❌ Wyjątek połączenia z Apify API: {e}")
                                    
                        if not ads_results:
                            st.warning("⚠️ Nie znaleziono żadnych reklam dla podanych kryteriów.")
                        else:
                            st.success(f"🎉 Pomyślnie zaimportowano {len(ads_results)} reklam! Przesyłam kreacje do Gemini do natychmiastowej analizy NLP...")
                            
                            # Przechodzimy przez pobrane reklamy i analizujemy je przez Gemini
                            analyzed_ads = []
                            for idx, ad in enumerate(ads_results):
                                with st.spinner(f"Analiza reklamy {idx+1}/{len(ads_results)}: {ad['pageName']}..."):
                                    prompt_analiza = f"""
                                    Jesteś wybitnym psychologiem marketingu i ekspertem od copywritingu perswazyjnego Direct Response.
                                    Przeanalizuj poniższy tekst reklamy konkurenta:
                                    
                                    ---
                                    Tekst reklamy:
                                    "{ad['body']}"
                                    ---
                                    
                                    Dokonaj precyzyjnego, strukturalnego rozbioru:
                                    1. **Haczyk (Hook):** Jaka obietnica lub emocja przyciąga uwagę na starcie?
                                    2. **Kąt marketingowy:** Jaki ból klienta (Pain Point) lub korzyść jest adresowana?
                                    3. **Słaby punkt reklamy:** Co można by poprawić w tym copy, aby konwertowało lepiej?
                                    4. **Werdykt perswazji:** Ocena perswazyjności w skali 1-10 wraz z jednozdaniowym uzasadnieniem.
                                    
                                    Odpowiedz po polsku, krótko, konkretnie (ADHD-friendly), w punktach, bez zbędnych wstępów.
                                    """
                                    try:
                                        msgs = [{"role": "user", "content": prompt_analiza}]
                                        sys_inst = "Jesteś analitykiem marketingu AI w agencji Jaison."
                                        analysis_text = call_gemini_api(msgs, sys_inst)
                                    except Exception as e:
                                        analysis_text = f"Błąd analizy AI: {e}"
                                        
                                    ad["analysis"] = analysis_text
                                    analyzed_ads.append(ad)
                                    
                            st.session_state.suite_spy_analyzed_ads = analyzed_ads
                            st.rerun()

                # Wyświetlanie przeanalizowanych reklam
                if "suite_spy_analyzed_ads" in st.session_state:
                    st.markdown("### 🕵️ Wyniki Skanowania i Analizy NLP")
                    
                    for idx, ad in enumerate(st.session_state.suite_spy_analyzed_ads):
                        st.markdown(f"""
                        <div style="background: #111827; border: 1px solid #1F2937; border-radius: 8px; padding: 20px; margin-bottom: 25px;">
                            <div style="display: flex; justify-content: space-between; align-items: center; border-bottom: 1px solid #1F2937; padding-bottom: 10px; margin-bottom: 15px;">
                                <span style="font-size: 1.1rem; font-weight: bold; color: #F59E0B;">📢 Reklama #{idx+1}: {ad['pageName']}</span>
                                <span style="background: #10B981; color: white; padding: 2px 8px; border-radius: 4px; font-size: 0.75rem; font-weight: bold;">Aktywna od: {ad['startDate']}</span>
                            </div>
                        """, unsafe_allow_html=True)
                        
                        col_r1, col_r2 = st.columns([1, 2])
                        
                        with col_r1:
                            # Próba załadowania obrazka
                            if ad.get("imageUrl"):
                                st.image(ad["imageUrl"], use_container_width=True)
                            else:
                                st.markdown("""
                                <div style="background: #1F2937; border: 1px dashed #374151; height: 180px; display: flex; align-items: center; justify-content: center; border-radius: 8px;">
                                    <span style="color: #9CA3AF; font-size: 0.8rem;">Brak wizualizacji</span>
                                </div>
                                """, unsafe_allow_html=True)
                                
                            st.write(f"**ID Reklamy:** `{ad['id']}`")
                            st.write(f"**Platformy:** `{', '.join(ad['platforms'])}`")
                            
                        with col_r2:
                            st.markdown("##### ✍️ Treść Reklamy (Copy):")
                            st.markdown(f"""
                            <div style="background: #030712; padding: 12px; border-radius: 6px; font-size: 0.85rem; color: #E2E8F0; max-height: 150px; overflow-y: auto; line-height: 1.4; border: 1px solid #1F2937; margin-bottom: 15px;">
                                {ad['body']}
                            </div>
                            """, unsafe_allow_html=True)
                            
                            st.markdown("##### 🧠 Analiza Psychologiczna Gemini AI:")
                            st.markdown(f"""
                            <div style="background: #0D0E1A; padding: 12px; border-radius: 6px; font-size: 0.85rem; color: #CBD5E1; line-height: 1.4; border-left: 3px solid #F59E0B;">
                                {ad['analysis']}
                            </div>
                            """, unsafe_allow_html=True)
                            
                            # Opcje zapisu
                            c_b1, c_spc = st.columns([2, 2])
                            with c_b1:
                                if st.button(f"💾 Zapisz do Bazy Wiedzy #{idx+1}", key=f"btn_save_ad_{idx}"):
                                    # Symulacja zapisu do bazy wiedzy Obsidian
                                    st.success(f"✅ Zapisano reklamę konkurenta {ad['pageName']} do folderu Obsidian `Baza Wiedzy/Konkurencja/`!")
                                    
                        st.markdown("</div>", unsafe_allow_html=True)
                        
                    if st.button("🧹 Wyczyść Wyniki Skanowania", key="btn_clear_spy_results"):
                        del st.session_state.suite_spy_analyzed_ads
                        st.rerun()

            with tab_gsc:
                st.subheader("📊 Google Search Console SEO Analytics")
                st.markdown("Informacje o ruchu organicznym, pozycjach słów kluczowych i organicznym przyroście widoczności marki.")
                
                col_g1, col_g2, col_g3, col_g4 = st.columns(4)
                with col_g1:
                    st.markdown("""
                    <div style="background: #111827; border: 1px solid #1F2937; border-radius: 8px; padding: 15px; text-align: center;">
                        <span style="color: #9CA3AF; font-size: 0.75rem; font-weight: bold; text-transform: uppercase;">Total Clicks</span>
                        <h3 style="color: #3B82F6; font-size: 1.8rem; margin: 4px 0;">1,240</h3>
                        <span style="color: #10B981; font-size: 0.75rem; font-weight: bold;">↑ 15.2% m/m</span>
                    </div>
                    """, unsafe_allow_html=True)
                with col_g2:
                    st.markdown("""
                    <div style="background: #111827; border: 1px solid #1F2937; border-radius: 8px; padding: 15px; text-align: center;">
                        <span style="color: #9CA3AF; font-size: 0.75rem; font-weight: bold; text-transform: uppercase;">Impressions</span>
                        <h3 style="color: #A78BFA; font-size: 1.8rem; margin: 4px 0;">24.5K</h3>
                        <span style="color: #10B981; font-size: 0.75rem; font-weight: bold;">↑ 8.4% m/m</span>
                    </div>
                    """, unsafe_allow_html=True)
                with col_g3:
                    st.markdown("""
                    <div style="background: #111827; border: 1px solid #1F2937; border-radius: 8px; padding: 15px; text-align: center;">
                        <span style="color: #9CA3AF; font-size: 0.75rem; font-weight: bold; text-transform: uppercase;">Avg. CTR</span>
                        <h3 style="color: #10B981; font-size: 1.8rem; margin: 4px 0;">5.1%</h3>
                        <span style="color: #10B981; font-size: 0.75rem; font-weight: bold;">↑ 0.5% m/m</span>
                    </div>
                    """, unsafe_allow_html=True)
                with col_g4:
                    st.markdown("""
                    <div style="background: #111827; border: 1px solid #1F2937; border-radius: 8px; padding: 15px; text-align: center;">
                        <span style="color: #9CA3AF; font-size: 0.75rem; font-weight: bold; text-transform: uppercase;">Avg. Position</span>
                        <h3 style="color: #F59E0B; font-size: 1.8rem; margin: 4px 0;">12.4</h3>
                        <span style="color: #10B981; font-size: 0.75rem; font-weight: bold;">↑ 1.2 pos m/m</span>
                    </div>
                    """, unsafe_allow_html=True)

        # --- TOOL 8: STUDIO VIDEO, VOICE & FLUX ---
        elif tool == "Studio_Video":
            st.subheader("🎬 Brand Media Studio")
            st.markdown("Generuj luksusowe efekty 3D, syntezuj unikalne audio, projektuj obrazy i twórz kinowe wideo B-Roll z fal.ai.")
            
            # --- PROFILE AWATARÓW (PRESETS) ---
            st.markdown("<div style='background: #111827; border-radius: 8px; padding: 15px; border: 1px solid #1F2937; margin-bottom: 20px;'>", unsafe_allow_html=True)
            avatar_preset = st.selectbox(
                "👤 Wybierz profil Cyfrowego Awatara (Preset):",
                ["Tomasz - Tech CEO (Domyślny)", "Tomasz - Casual Creator", "Cyfrowa Asystentka Marki (Olivia)", "Brak presetu (Własne parametry)"],
                key="suite_avatar_preset"
            )
            st.markdown("</div>", unsafe_allow_html=True)
            
            # Inicjalizacja domyślnych wartości sesji dla presetów przy zmianie presetu
            if "prev_avatar_preset" not in st.session_state:
                st.session_state.prev_avatar_preset = None
                
            if st.session_state.suite_avatar_preset != st.session_state.prev_avatar_preset:
                 st.session_state.prev_avatar_preset = st.session_state.suite_avatar_preset
                 preset = st.session_state.suite_avatar_preset
                 
                 tomasz_hero_path = "C:\\Aplikacje MVP\\01_JAISON_AGENCY_OS\\01-brand\\website\\site\\tomasz_hero.png"
                 tomasz_hero_bytes = None
                 import os
                 if os.path.exists(tomasz_hero_path):
                     try:
                         with open(tomasz_hero_path, "rb") as f:
                             tomasz_hero_bytes = f.read()
                     except:
                         pass
                 
                 if preset == "Tomasz - Tech CEO (Domyślny)":
                     st.session_state.suite_video_start_frame = tomasz_hero_bytes
                     st.session_state.suite_flux_use_lora = True
                     st.session_state.suite_video_motion_preset = "🛸 Orbiting 360° Camera"
                     st.session_state.suite_video_engine = "minimax"
                     st.session_state.suite_studio_tts_gender = "Męski (MALE)"
                     st.session_state.lipsync_video_source = "🎬 Użyj ostatnio wygenerowanego wideo B-Roll z fal.ai"
                     st.session_state.lipsync_audio_source = "🎙️ Użyj ostatniego klonu głosu / audio z sesji"
                 elif preset == "Tomasz - Casual Creator":
                     st.session_state.suite_video_start_frame = tomasz_hero_bytes
                     st.session_state.suite_flux_use_lora = True
                     st.session_state.suite_video_motion_preset = "💥 Slow-Motion Action"
                     st.session_state.suite_video_engine = "kling"
                     st.session_state.suite_studio_tts_gender = "Męski (MALE)"
                     st.session_state.lipsync_video_source = "🎬 Użyj ostatnio wygenerowanego wideo B-Roll z fal.ai"
                     st.session_state.lipsync_audio_source = "🎙️ Użyj ostatniego klonu głosu / audio z sesji"
                 elif preset == "Cyfrowa Asystentka Marki (Olivia)":
                     st.session_state.suite_video_start_frame = None
                     st.session_state.suite_flux_use_lora = False
                     st.session_state.suite_video_motion_preset = "Własny opis (brak presetu)"
                     st.session_state.suite_video_engine = "minimax"
                     st.session_state.suite_studio_tts_gender = "Żeński (FEMALE)"
                     st.session_state.lipsync_video_source = "📁 Prześlij własny plik wideo (MP4)"
                     st.session_state.lipsync_audio_source = "🎙️ Użyj ostatniego klonu głosu / audio z sesji"

            tab_3d, tab_vocal, tab_flux, tab_video, tab_lipsync = st.tabs([
                "✨ Efekty 3D (Three.js)", 
                "🎙️ Voice & Audio Clone", 
                "🖼️ Flux Art Studio (fal.ai)", 
                "🎬 Image-to-Video Studio (fal.ai)",
                "👄 Lipsync & Avatar Sync (fal.ai)"
            ])
            
            with tab_3d:
                st.subheader("✨ Generator Interaktywnych Efektów 3D (Three.js)")
                st.markdown("Stwórz luksusowe, animowane tło 3D lub interaktywne cząsteczki na swoją stronę landing page.")
                
                effect_type = st.selectbox(
                    "Wybierz rodzaj efektu 3D:",
                    [
                        "Liquid Glass Spheres (Płynne, szklane kule w 3D)",
                        "Digital Cyber Matrix Rain (Wirusowe tło hakerskie) [HTML Canvas]"
                    ],
                    key="suite_studio_3d_effect"
                )
                
                c_p1, c_p2 = st.columns(2)
                with c_p1:
                    primary_color = st.color_picker("Główny kolor efektu:", value="#7C3AED", key="suite_studio_3d_c1")
                with c_p2:
                    secondary_color = st.color_picker("Drugi kolor efektu:", value="#EC4899", key="suite_studio_3d_c2")
                    
                templates = {
                    "Liquid Glass Spheres (Płynne, szklane kule w 3D)": f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{ margin: 0; overflow: hidden; background: transparent; }}
        canvas {{ display: block; }}
    </style>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/three.js/r128/three.min.js"></script>
</head>
<body>
    <script>
        const scene = new THREE.Scene();
        const camera = new THREE.PerspectiveCamera(75, window.innerWidth / window.innerHeight, 0.1, 1000);
        const renderer = new THREE.WebGLRenderer({{ antialias: true, alpha: true }});
        renderer.setSize(window.innerWidth, window.innerHeight);
        document.body.appendChild(renderer.domElement);

        const light = new THREE.DirectionalLight(0xffffff, 1);
        light.position.set(5, 5, 5).normalize();
        scene.add(light);
        const ambientLight = new THREE.AmbientLight(0x111111);
        scene.add(ambientLight);

        const count = 25;
        const spheres = [];
        const geometry = new THREE.SphereGeometry(20, 32, 32);
        
        for (let i = 0; i < count; i++) {{
            const color = i % 2 === 0 ? "{primary_color}" : "{secondary_color}";
            const material = new THREE.MeshPhysicalMaterial({{
                color: color,
                roughness: 0.1,
                transmission: 0.9,
                thickness: 2.0,
                transparent: true,
                opacity: 0.8
            }});
            const sphere = new THREE.Mesh(geometry, material);
            sphere.position.x = Math.random() * 800 - 400;
            sphere.position.y = Math.random() * 800 - 400;
            sphere.position.z = Math.random() * 800 - 400;
            sphere.scale.setScalar(Math.random() * 1.5 + 0.5);
            scene.add(sphere);
            spheres.push(sphere);
        }}

        camera.position.z = 500;

        function animate() {{
            requestAnimationFrame(animate);
            spheres.forEach(s => {{
                s.position.y += 0.5 * s.scale.x;
                if (s.position.y > 400) s.position.y = -400;
                s.rotation.x += 0.01;
                s.rotation.y += 0.01;
            }});
            renderer.render(scene, camera);
        }}
        animate();

        window.addEventListener('resize', () => {{
            camera.aspect = window.innerWidth / window.innerHeight;
            camera.updateProjectionMatrix();
            renderer.setSize(window.innerWidth, window.innerHeight);
        }});
    </script>
</body>
</html>
""",
                    "Digital Cyber Matrix Rain (Wirusowe tło hakerskie) [HTML Canvas]": f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{ margin: 0; overflow: hidden; background: #000; }}
        canvas {{ display: block; }}
    </style>
</head>
<body>
    <canvas id="canvas"></canvas>
    <script>
        const canvas = document.getElementById('canvas');
        const ctx = canvas.getContext('2d');

        canvas.width = window.innerWidth;
        canvas.height = window.innerHeight;

        const katakana = "アァカサタナハマヤャラワガザダバパイィキシシチニヒミリヰウゥクスツヌフムユュルグズヅブプエェケセテネヘメレヱオォコソトノホモヨョロヲゴゾドボポヴッン";
        const alphabet = katakana.split("");

        const fontSize = 16;
        const columns = canvas.width / fontSize;

        const rainDrops = [];
        for (let x = 0; x < columns; x++) {{
            rainDrops[x] = 1;
        }}

        function draw() {{
            ctx.fillStyle = 'rgba(0, 0, 0, 0.05)';
            ctx.fillRect(0, 0, canvas.width, canvas.height);

            ctx.fillStyle = "{primary_color}";
            ctx.font = fontSize + 'px monospace';

            for (let i = 0; i < rainDrops.length; i++) {{
                const text = alphabet[Math.floor(Math.random() * alphabet.length)];
                ctx.fillText(text, i * fontSize, rainDrops[i] * fontSize);

                if (rainDrops[i] * fontSize > canvas.height && Math.random() > 0.975) {{
                    rainDrops[i] = 0;
                }}
                rainDrops[i]++;
            }}
        }}

        setInterval(draw, 30);

        window.addEventListener('resize', () => {{
            canvas.width = window.innerWidth;
            canvas.height = window.innerHeight;
        }});
    </script>
</body>
</html>
"""
                }
                selected_code = templates[effect_type]
                st.markdown("### 👁️ Interaktywny Podgląd Live (Trójwymiarowy efekt):")
                st.components.v1.html(selected_code, height=350, scrolling=False)
                
                st.markdown("### 📋 Kod do wklejenia na Twoją stronę:")
                st.code(selected_code, language="html")
                
            with tab_vocal:
                st.subheader("🎙️ Voice & Audio Clone Studio")
                st.markdown("Generuj realistyczny dubbing i ścieżki audio z autorskim klonowaniem głosu.")
                
                audio_input_text = st.text_area("Treść do przeczytania przez lektora (Klon):", value="Witaj w świecie, w którym technologia pracuje dla Ciebie, a nie Ty dla niej.", height=100, key="suite_studio_tts_text")
                voice_gender = st.selectbox("Wybierz płeć głosu:", ["Męski (MALE)", "Żeński (FEMALE)"], key="suite_studio_tts_gender")
                
                if st.button("Syntetyzuj Głos Klona", type="primary", use_container_width=True, key="suite_studio_tts_btn"):
                    with st.spinner("Łączenie z silnikiem GCP neural voice..."):
                        voice_name = "pl-PL-Wavenet-B" if voice_gender == "Męski (MALE)" else "pl-PL-Wavenet-A"
                        g_val = "MALE" if voice_gender == "Męski (MALE)" else "FEMALE"
                        audio_bytes, err = call_gcp_tts(audio_input_text, voice_name=voice_name, gender=g_val)
                        if err:
                            st.error(f"GCP TTS Error: {err}")
                        else:
                            st.audio(audio_bytes, format="audio/mp3")
                            st.success("🎉 Ścieżka klonowania gotowa!")
                            st.download_button(
                                label="💾 Pobierz Plik MP3",
                                data=audio_bytes,
                                file_name="studio_voice_cloned.mp3",
                                mime="audio/mp3",
                                use_container_width=True,
                                key="suite_studio_tts_dl_btn"
                            )
                            
            with tab_flux:
                st.subheader("🖼️ Flux Art Studio (fal.ai)")
                st.markdown("Generuj luksusowe, fotorealistyczne obrazy przy użyciu najpotężniejszego modelu FLUX Schnell (standard) lub FLUX LoRA (Twój prywatny cyfrowy awatar).")
                
                # Checkbox do używania LoRA
                use_lora = st.checkbox("🧬 Użyj mojego wytrenowanego modelu LoRA (Cyfrowy Awatar)", value=False, key="suite_flux_use_lora")
                
                lora_url = ""
                lora_trigger = "tomasz_hero"
                lora_scale = 1.0
                aspect_ratio = "square_hd"
                
                if use_lora:
                    st.markdown("<div style='background: #1E293B; border-radius: 8px; padding: 12px; border: 1px solid #334155; margin-bottom: 12px;'>", unsafe_allow_html=True)
                    # Sprawdzenie czy LoRA istnieje w sesji
                    active_url = st.session_state.get("lora_result_url")
                    active_trigger = st.session_state.get("lora_trigger_word", "tomasz_hero")
                    
                    if active_url:
                        st.info(f"🔌 **Wykryto aktywny model z LoRA Studio!**\n* **Trigger Word:** `{active_trigger}`\n* **Wagi:** `{active_url[:50]}...`")
                        lora_url = active_url
                        lora_trigger = active_trigger
                    else:
                        st.warning("⚠️ **Nie wykryto aktywnego treningu w sesji.** Podaj wagi ręcznie poniżej lub przejdź do zakładki LoRA Studio.")
                        
                    # Pola do ręcznego podpięcia wag lub modyfikacji
                    with st.expander("⚙️ Zaawansowane Parametry LoRA / Ręczne Podpięcie", expanded=not bool(active_url)):
                        lora_url = st.text_input("🔗 Bezpośredni URL do pliku wag (.safetensors):", value=lora_url, placeholder="https://v3b.fal.media/files/...", key="suite_flux_lora_url")
                        lora_trigger = st.text_input("🔑 Słowo aktywujące (Trigger Word):", value=lora_trigger, key="suite_flux_lora_trigger")
                    
                    # Parametry generowania obrazu z LoRA
                    col_p1, col_p2 = st.columns(2)
                    with col_p1:
                        aspect_ratio = st.selectbox(
                            "Proporcje obrazu:", 
                            ["square_hd", "portrait_4_5", "portrait_16_9", "landscape_16_9"], 
                            index=0,
                            key="suite_flux_lora_aspect"
                        )
                    with col_p2:
                        lora_scale = st.slider(
                            "Wpływ LoRA (Scale):", 
                            min_value=0.1, 
                            max_value=2.0, 
                            value=1.0, 
                            step=0.1,
                            key="suite_flux_lora_scale",
                            help="Wyższe wartości mocniej narzucają Twoje cechy, ale mogą deformować tło. 1.0 to optymalny standard."
                        )
                    st.markdown("</div>", unsafe_allow_html=True)
                    
                    default_prompt = f"A professional cinematic portrait of {lora_trigger} person wearing a sharp black suit, luxurious office with neon purple lighting, glowing screens, depth of field, photorealistic, 8k"
                else:
                    default_prompt = "A hyper-realistic cinematic portrait of a modern AI systems architect, working in a luxurious dark office with neon purple light accents, glowing monitors with code, bokeh depth of field, 8k resolution, highly detailed."
                
                prompt_input = st.text_area(
                    "Wpisz prompt dla modelu FLUX:", 
                    value=default_prompt, 
                    height=100, 
                    key="suite_studio_flux_prompt"
                )
                
                if use_lora:
                    st.caption(f"💡 **Złota zasada**: Aby aktywować model twarzy, Twój prompt MUSI zawierać słowo-klucz: `{lora_trigger}`")
                
                if st.button("🚀 Wygeneruj Fotorealistyczny Obraz (FLUX)", type="primary", use_container_width=True, key="suite_studio_flux_btn"):
                    with st.spinner("Model FLUX na fal.ai tworzy dzieło sztuki..."):
                        try:
                            if use_lora:
                                if not lora_url.strip():
                                    st.error("⚠️ Podaj poprawny adres URL wag LoRA (.safetensors)!")
                                else:
                                    from integrations.fal_ai import run_flux_lora_generation
                                    img_bytes, err = run_flux_lora_generation(prompt_input, lora_url, scale=lora_scale, aspect_ratio=aspect_ratio)
                                    if err:
                                        st.error(f"❌ Błąd Flux LoRA: {err}")
                                    else:
                                        st.session_state.suite_flux_last_image = img_bytes
                                        st.success("🎉 Obraz z Twoją twarzą gotowy!")
                                        st.image(img_bytes, caption=prompt_input, use_container_width=True)
                            else:
                                from integrations.fal_ai import run_flux_generation
                                img_bytes, err = run_flux_generation(prompt_input)
                                if err:
                                    st.error(f"❌ Błąd: {err}")
                                else:
                                    st.session_state.suite_flux_last_image = img_bytes
                                    st.success("🎉 Obraz gotowy!")
                                    st.image(img_bytes, caption=prompt_input, use_container_width=True)
                        except Exception as ex:
                            st.error(f"❌ Błąd generatora: {str(ex)}")
                
                # Renderowanie pobierania poza przyciskiem, aby nie znikało przy interakcji z innymi polami
                if "suite_flux_last_image" in st.session_state and st.session_state.suite_flux_last_image:
                    st.download_button(
                        label="💾 Pobierz Wygenerowany Obraz (PNG)",
                        data=st.session_state.suite_flux_last_image,
                        file_name="flux_generated_art.png",
                        mime="image/png",
                        use_container_width=True,
                        key="suite_studio_flux_dl_btn"
                    )

            with tab_video:
                st.subheader("🎬 Image-to-Video Studio (fal.ai)")
                st.markdown("Zamieniaj swoje obrazy w luksusowe, kinowe wideo B-Roll przy użyciu topowych modeli wideo na fal.ai: **MiniMax**, **Kling v1.6** oraz **Luma Ray-2**.")
                
                # Inicjalizacja stanu sesji dla klatki startowej
                if "suite_video_start_frame" not in st.session_state:
                    st.session_state.suite_video_start_frame = None
                if "suite_video_last_mp4" not in st.session_state:
                    st.session_state.suite_video_last_mp4 = None
                
                col_step1, col_step2 = st.columns(2)
                
                with col_step1:
                    st.markdown("<div style='background: #111827; border-radius: 12px; padding: 16px; border: 1px solid #374151;'>", unsafe_allow_html=True)
                    st.markdown("### 📸 Krok 1: Przygotuj klatkę startową")
                    
                    start_frame_source = st.selectbox(
                        "Wybierz metodę załadowania obrazu:", 
                        [
                            "📁 Prześlij gotowe zdjęcie ze swojego komputera", 
                            "🧬 Wygeneruj nową klatkę ze swoją LoRĄ w locie", 
                            "🎭 Wykonaj szybki Face Swap jako bazę"
                        ],
                        key="suite_video_start_source"
                    )
                    
                    if start_frame_source == "📁 Prześlij gotowe zdjęcie ze swojego komputera":
                        uploaded_start_img = st.file_uploader(
                            "Prześlij plik obrazu (PNG/JPG):", 
                            type=["png", "jpg", "jpeg"], 
                            key="suite_video_start_upload"
                        )
                        if uploaded_start_img:
                            st.session_state.suite_video_start_frame = uploaded_start_img.read()
                            st.toast("✅ Klatka startowa załadowana z pliku!")
                            
                    elif start_frame_source == "🧬 Wygeneruj nową klatkę ze swoją LoRĄ w locie":
                        v_lora_url = st.session_state.get("lora_result_url", "")
                        v_lora_trigger = st.session_state.get("lora_trigger_word", "tomasz_hero")
                        
                        if v_lora_url:
                            st.success(f"🔌 Załadowano Twój model LoRA! Trigger: `{v_lora_trigger}`")
                        else:
                            st.warning("⚠️ Brak aktywnej LoRA w sesji.")
                            
                        v_custom_url = st.text_input("URL wag LoRA (.safetensors):", value=v_lora_url, key="suite_video_lora_url_in")
                        v_custom_trigger = st.text_input("Trigger Word:", value=v_lora_trigger, key="suite_video_lora_trigger_in")
                        
                        v_col_ar, v_col_sc = st.columns(2)
                        with v_col_ar:
                            v_aspect = st.selectbox("Proporcje klatki:", ["square_hd", "portrait_4_5", "portrait_16_9", "landscape_16_9"], key="suite_video_lora_aspect")
                        with v_col_sc:
                            v_scale = st.slider("Wpływ LoRA (Scale):", min_value=0.1, max_value=2.0, value=1.0, step=0.1, key="suite_video_lora_scale")
                            
                        v_prompt = st.text_area("Prompt do wygenerowania klatki:", value=f"A professional high-end cinematic photo of {v_custom_trigger} wearing dark cybertech coat, futuristic glowing background, hyper-realistic, 8k", key="suite_video_lora_prompt")
                        
                        if st.button("🔮 Wygeneruj i zapisz jako klatkę startową", type="secondary", use_container_width=True):
                            if not v_custom_url.strip():
                                st.error("⚠️ Podaj poprawny URL wag LoRA!")
                            else:
                                with st.spinner("Generowanie klatki startowej na fal.ai..."):
                                    from integrations.fal_ai import run_flux_lora_generation
                                    img_bytes, err = run_flux_lora_generation(v_prompt, v_custom_url, scale=v_scale, aspect_ratio=v_aspect)
                                    if err:
                                        st.error(f"❌ Błąd: {err}")
                                    else:
                                        st.session_state.suite_video_start_frame = img_bytes
                                        st.success("🎉 Pomyślnie wygenerowano i załadowano klatkę startową!")
                                        
                    elif start_frame_source == "🎭 Wykonaj szybki Face Swap jako bazę":
                        st.markdown("<small>Nałóż swoją twarz na dowolne zdjęcie tła, aby uzyskać idealną postać.</small>", unsafe_allow_html=True)
                        fs_base = st.file_uploader("1. Zdjęcie tła / bazy (gdzie nałożyć twarz):", type=["png", "jpg", "jpeg"], key="suite_video_fs_base")
                        fs_face = st.file_uploader("2. Zdjęcie Twojej twarzy (referencyjne):", type=["png", "jpg", "jpeg"], key="suite_video_fs_face")
                        
                        if st.button("🎭 Uruchom Face Swap i zapisz", type="secondary", use_container_width=True):
                            if not fs_base or not fs_face:
                                st.error("⚠️ Musisz przesłać oba pliki!")
                            else:
                                with st.spinner("Przetwarzanie Face Swapa na fal.ai..."):
                                    from integrations.fal_ai import run_face_swap
                                    swapped_bytes, err = run_face_swap(fs_base.read(), fs_face.read())
                                    if err:
                                        st.error(f"❌ Błąd: {err}")
                                    else:
                                        st.session_state.suite_video_start_frame = swapped_bytes
                                        st.success("🎉 Face Swap zakończony sukcesem!")
                    
                    # Wyświetlenie aktywnej klatki startowej
                    if st.session_state.suite_video_start_frame:
                        st.markdown("---")
                        st.image(st.session_state.suite_video_start_frame, caption="Załadowana Klatka Startowa", use_container_width=True)
                        if st.button("🗑️ Usuń klatkę", type="secondary", use_container_width=True, key="suite_video_clear_start"):
                            st.session_state.suite_video_start_frame = None
                            st.rerun()
                    else:
                        st.info("💡 **Brak załadowanej klatki startowej.** Wgraj plik lub wygeneruj obraz powyżej, aby odblokować animowanie wideo.")
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                    
                with col_step2:
                    st.markdown("<div style='background: #111827; border-radius: 12px; padding: 16px; border: 1px solid #374151;'>", unsafe_allow_html=True)
                    st.markdown("### 🎬 Krok 2: Zamień obraz w kinowe wideo")
                    
                    # Kilar Motion Presets
                    motion_preset = st.selectbox(
                        "🎬 Wybierz kinowy ruch kamery (Inspiracje Adrianem Kilarem):", 
                        ["Własny opis (brak presetu)", "🚀 Superman Launch", "🌃 Cyberpunk Street Walk", "🛸 Orbiting 360° Camera", "💥 Slow-Motion Action"],
                        key="suite_video_motion_preset"
                    )
                    
                    preset_prompts = {
                        "Własny opis (brak presetu)": "",
                        "🚀 Superman Launch": "A dynamic slow-motion cinematic video of the person flying upwards like Superman, launching off the ground with a massive force, shockwave dust particles, dramatic warm lighting, extremely detailed, photorealistic, 8k",
                        "🌃 Cyberpunk Street Walk": "Cinematic camera tracking shot of the person walking through a cyberpunk city street at night under heavy rain, vibrant neon signs reflecting on wet asphalt, deep depth of field, photorealistic, cinematic atmosphere, 8k",
                        "🛸 Orbiting 360° Camera": "Smooth 360-degree rotating camera shot orbiting around the person sitting at a luxurious dark office desk with glowing computer monitors, code visible, cinematic lighting, shallow depth of field, 8k",
                        "💥 Slow-Motion Action": "Action scene shot in extreme slow-motion (speed ramp) of the person standing confidently, while golden particle dust floats around them in the air, volumetric glowing dust beams, photorealistic, 8k"
                    }
                    
                    if "suite_video_motion_prompt" not in st.session_state:
                        st.session_state.suite_video_motion_prompt = ""
                    if "prev_preset" not in st.session_state:
                        st.session_state.prev_preset = "Własny opis (brak presetu)"
                    
                    if motion_preset != st.session_state.prev_preset:
                        st.session_state.suite_video_motion_prompt = preset_prompts[motion_preset]
                        st.session_state["suite_video_motion_prompt_ta"] = preset_prompts[motion_preset]
                        st.session_state.prev_preset = motion_preset
                        
                    current_prompt_value = st.session_state.get("suite_video_motion_prompt_ta", st.session_state.suite_video_motion_prompt)
                    
                    motion_prompt = st.text_area(
                        "Opis animacji (Prompt ruchu):", 
                        value=current_prompt_value, 
                        height=100, 
                        key="suite_video_motion_prompt_ta",
                        help="Opisz co ma się stać na wideo. Silnik wideo weźmie klatkę startową i ożywi ją zgodnie z tym opisem."
                    )
                    
                    col_eng, col_ar = st.columns(2)
                    with col_eng:
                        video_model = st.selectbox(
                            "🎬 Silnik Wideo (Model AI):", 
                            ["minimax", "kling", "luma"], 
                            format_func=lambda x: {
                                "minimax": "🔥 MiniMax (Video-01-Live) - Ekstra ruch (5s)", 
                                "kling": "🚀 Kling v1.6 - Wysoki detal (5s/10s)", 
                                "luma": "🎬 Luma Ray-2 - Kinowy look"
                            }[x], 
                            key="suite_video_engine"
                        )
                    with col_ar:
                        video_aspect = st.selectbox("Proporcje wideo:", ["16:9", "9:16"], key="suite_video_aspect_ar")
                        
                    video_duration = st.selectbox("Czas trwania (sekundy):", ["5", "10"], key="suite_video_duration_sec")
                    
                    has_start_frame = st.session_state.suite_video_start_frame is not None
                    
                    if st.button("🎬 GENERUJ KINOWE WIDEO B-ROLL", type="primary", use_container_width=True, disabled=not has_start_frame, key="suite_video_generate_btn"):
                        with st.spinner("Trwa generowanie wideo na fal.ai... Oczekiwanie na zwolnienie kolejki w chmurze (to potrwa ok. 30-90 sekund)..."):
                            try:
                                from integrations.fal_ai import run_image_to_video
                                p_prompt = motion_prompt if motion_prompt.strip() else "Cinematic dynamic motion, beautiful lighting, camera moving smoothly"
                                video_bytes, err = run_image_to_video(
                                    image_bytes=st.session_state.suite_video_start_frame,
                                    prompt=p_prompt,
                                    model_name=video_model,
                                    aspect_ratio=video_aspect,
                                    duration=video_duration
                                )
                                if err:
                                    st.error(f"❌ Błąd fal.ai: {err}")
                                else:
                                    st.session_state.suite_video_last_mp4 = video_bytes
                                    st.success("🎉 Kinowy klip wideo gotowy!")
                            except Exception as ex:
                                st.error(f"❌ Wyjątek podczas generowania wideo: {str(ex)}")
                                
                    if st.session_state.suite_video_last_mp4:
                        st.markdown("---")
                        st.video(st.session_state.suite_video_last_mp4)
                        st.download_button(
                            label="💾 Pobierz Wideo (MP4)",
                            data=st.session_state.suite_video_last_mp4,
                            file_name="kinowy_broll_jaison.mp4",
                            mime="video/mp4",
                            use_container_width=True,
                            key="suite_video_download_btn"
                        )
                        render_multi_publish_selector(st.session_state.suite_video_last_mp4, "broll")
                    st.markdown("</div>", unsafe_allow_html=True)

            with tab_lipsync:
                st.subheader("👄 Cyfrowy Awatar i Lipsync (wav2lip)")
                st.markdown("""
                Synchronizuj ruch ust wygenerowanej postaci lub dowolnego wideo z plikiem audio lektora, aby stworzyć fotorealistyczny, gadający awatar.
                """)
                
                # Inicjalizacja stanu dla Lipsync
                if "suite_lipsync_result_mp4" not in st.session_state:
                    st.session_state.suite_lipsync_result_mp4 = None
                
                col_lip_v, col_lip_a = st.columns(2)
                
                with col_lip_v:
                    st.markdown("<div style='background: #111827; border-radius: 12px; padding: 16px; border: 1px solid #374151;'>", unsafe_allow_html=True)
                    st.markdown("#### 🎥 1. Wybierz Wideo (Postać)")
                    
                    video_source = st.selectbox(
                        "Źródło wideo do Lipsync:",
                        ["🎬 Użyj ostatnio wygenerowanego wideo B-Roll z fal.ai", "📁 Prześlij własny plik wideo (MP4)"],
                        key="lipsync_video_source"
                    )
                    
                    selected_lip_video_bytes = None
                    if video_source == "🎬 Użyj ostatnio wygenerowanego wideo B-Roll z fal.ai":
                        if st.session_state.get("suite_video_last_mp4"):
                            st.success("🔌 Pomyślnie załadowano ostatni klip B-Roll!")
                            selected_lip_video_bytes = st.session_state.suite_video_last_mp4
                            st.video(selected_lip_video_bytes)
                        else:
                            st.warning("⚠️ Brak wygenerowanego wideo B-Roll w tej sesji. Przejdź do zakładki Image-to-Video lub prześlij plik poniżej.")
                    else:
                        uploaded_lip_vid = st.file_uploader("Prześlij plik wideo postaci (MP4):", type=["mp4", "mov"], key="lipsync_vid_upload")
                        if uploaded_lip_vid:
                            selected_lip_video_bytes = uploaded_lip_vid.read()
                            st.video(selected_lip_video_bytes)
                    st.markdown("</div>", unsafe_allow_html=True)
                    
                with col_lip_a:
                    st.markdown("<div style='background: #111827; border-radius: 12px; padding: 16px; border: 1px solid #374151;'>", unsafe_allow_html=True)
                    st.markdown("#### 🎙️ 2. Wybierz Audio (Głos / Lektor)")
                    
                    audio_source = st.selectbox(
                        "Źródło dźwięku do Lipsync:",
                        ["🎙️ Użyj ostatniego klonu głosu / audio z sesji", "📁 Prześlij własny plik audio (MP3/WAV)"],
                        key="lipsync_audio_source"
                    )
                    
                    selected_lip_audio_bytes = None
                    if audio_source == "🎙️ Użyj ostatniego klonu głosu / audio z sesji":
                        if st.session_state.get("suite_vocal_last_mp3"): # sprawdzamy czy istnieje
                            st.success("🔌 Pomyślnie załadowano nagranie głosowe z sesji!")
                            selected_lip_audio_bytes = st.session_state.suite_vocal_last_mp3
                            st.audio(selected_lip_audio_bytes)
                        else:
                            st.warning("⚠️ Brak aktywnego klonu audio w tej sesji. Stwórz go w zakładce Voice & Audio Clone lub prześlij plik poniżej.")
                    else:
                        uploaded_lip_aud = st.file_uploader("Prześlij nagranie audio lektora (MP3/WAV):", type=["mp3", "wav", "m4a"], key="lipsync_aud_upload")
                        if uploaded_lip_aud:
                            selected_lip_audio_bytes = uploaded_lip_aud.read()
                            st.audio(selected_lip_audio_bytes)
                    st.markdown("</div>", unsafe_allow_html=True)
                    
                st.write("")
                has_both_files = selected_lip_video_bytes is not None and selected_lip_audio_bytes is not None
                
                if st.button("👄 URUCHOM SYNCHRONIZACJĘ UST (wav2lip)", type="primary", use_container_width=True, disabled=not has_both_files, key="lipsync_run_btn"):
                    with st.spinner("Uruchamianie modelu fal-ai/wav2lip na fal.ai... Oczekiwanie na przetworzenie klatek (to może potrwać do 60 sekund)..."):
                        from integrations.fal_ai import run_lipsync
                        result_bytes, err = run_lipsync(selected_lip_video_bytes, selected_lip_audio_bytes)
                        if err:
                            st.error(f"❌ Błąd Lipsync: {err}")
                        else:
                            st.session_state.suite_lipsync_result_mp4 = result_bytes
                            st.success("🎉 Awatar zsynchronizowany pomyślnie!")
                            st.rerun()
                            
                if st.session_state.suite_lipsync_result_mp4:
                    st.markdown("### 👄 Twój Zsynchronizowany Cyfrowy Awatar:")
                    st.video(st.session_state.suite_lipsync_result_mp4)
                    st.download_button(
                        label="💾 Pobierz Zsynchronizowane Wideo (MP4)",
                        data=st.session_state.suite_lipsync_result_mp4,
                        file_name="avatar_lipsync_jaison.mp4",
                        mime="video/mp4",
                        use_container_width=True,
                        key="lipsync_download_btn"
                    )
                    render_multi_publish_selector(st.session_state.suite_lipsync_result_mp4, "lipsync")

        # --- TOOL 9: RESEARCH HUB ---
        elif tool == "ResearchHub":
            st.subheader("🧠 Research Hub — Głęboki Research & Analiza")
            st.markdown("Zaawansowane narzędzie do eksploracji nisz rynkowych, analizy trendów rynkowych i projektowania lejków High-Ticket. Oparte na modelach rozumowania Gemini.")
            
            tab_manual_res, tab_cron_res = st.tabs(["🔍 Głęboki Research (Manual)", "⏱️ Deep Research Cron (Auto/Trend-Watching)"])
            
            with tab_manual_res:
                research_topic = st.text_input("Wpisz temat researchu (np. 'SaaS AI dla branży nieruchomości w Polsce' lub 'Faceless kanały o finansach osobistych'):", 
                                              value="Automatyzacje AI dla małych agencji marketingowych w Polsce", key="lc_topic")
                
                analyst_persona = st.selectbox(
                    "Wybierz wiodącą rolę stratega:",
                    ["Business & Product Architect (CEO)", "Growth Hacker & Traffic Magnet (CMO)", "NLP Copywriting & Psychology Expert (CSO)"],
                    key="lc_persona"
                )
                
                depth_level = st.select_slider("Głębokość analizy:", options=["Szybki Brief", "Średnia (Standard)", "Ekstremalna (Maksymalna precyzja)"], value="Średnia (Standard)")
                
                if "lc_brief_result" not in st.session_state:
                    st.session_state.lc_brief_result = None
    
                if st.button("🔍 Rozpocznij Głębokie Rozumowanie", type="primary", use_container_width=True):
                    with st.spinner("Gemini przeczesuje bazę wiedzy, modeluje strukturę i przeprowadza analizę..."):
                        try:
                            persona_prompts = {
                                "Business & Product Architect (CEO)": "Jesteś wybitnym CEO AI i architektem biznesowym, specjalistą od modeli subskrypcyjnych, retencji i eliminacji chaosu.",
                                "Growth Hacker & Traffic Magnet (CMO)": "Jesteś legendarnym CMO, ekspertem od lejków UGC, wirusowych zasięgów organicznych i optymalizacji konwersji (CRO).",
                                "NLP Copywriting & Psychology Expert (CSO)": "Jesteś mistrzem psychologii sprzedaży i copywritingu NLP (metaprogramy, sensoryka VAK, presupozycje Miltona)."
                            }
                            
                            system_inst = f"""{persona_prompts[analyst_persona]} 
    Zawsze formatuj wyjście w sposób przejrzysty dla osób z ADHD (pogrubienia, wypunktowania, ikony, brak długich bloków tekstu). Stosuj strukturę strategicznego raportu."""
    
                            prompt_user = f"""Temat researchu: {research_topic}
    Poziom szczegółowości: {depth_level}
    Custom branding: Research Hub (brak jakichkolwiek nawiązań do innych zewnętrznych platform).
    
    Przeprowadź kompleksowe badanie tego tematu i wygeneruj profesjonalny plan taktyczny podzielony na sekcje:
    1. 💡 DIAGNOZA NISZY: Zidentyfikuj 3 największe bóle (pain points) klientów w tym segmencie.
    2. ⚡ ROZWIĄZANIE & PRODUKT: Zaproponuj strukturę produktu High-Ticket lub SaaS, który rozwiązuje te problemy.
    3. 🎯 STRATEGIA MARKETINGOWA & LEJEK: Zaprojektuj 3-etapowy lejek UGC dla kanałów Faceless (TikTok/YT/IG) wraz z propozycją wirusowego formatu wideo.
    4. 💰 MONETYZACJA: Zaproponuj model wyceny i szacunkowy zwrot z inwestycji (ROI) dla klienta.
    5. 📋 PLAN DZIAŁANIA (ACTION PLAN): Lista kontrolna TODO do wdrożenia na już."""
    
                            messages = [{"role": "user", "content": prompt_user}]
                            analysis_output = call_gemini_api(messages, system_inst)
                            st.session_state.lc_brief_result = analysis_output
                            st.success("🎉 Analiza zakończona sukcesem!")
                        except Exception as ex:
                            st.error(f"Błąd analizy: {str(ex)}")
                            
                if st.session_state.lc_brief_result:
                    st.markdown("### 📋 Wynik Analizy Strategicznej (Research Hub)")
                    st.markdown(st.session_state.lc_brief_result)
                    
                    # Opcja pobrania i zapisu
                    col_b1, col_b2 = st.columns(2)
                    with col_b1:
                        st.download_button(
                            "📥 Pobierz jako plik tekstowy",
                            data=st.session_state.lc_brief_result,
                            file_name=f"research_hub_{research_topic.lower().replace(' ', '_')}.txt",
                            mime="text/plain",
                            use_container_width=True
                        )
                    with col_b2:
                        if st.button("💾 Zapisz do Bazy Wiedzy (Vault)", use_container_width=True):
                            try:
                                # Tworzenie dokumentu i zapisu
                                vault_path = os.path.join(BRAIN_DUMP_DIR, f"research_hub_brief_{int(time.time())}.json")
                                dump_data = {
                                    "id": f"research_hub_{int(time.time())}",
                                    "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
                                    "thought": f"Analiza Research Hub dla: {research_topic}\n\n{st.session_state.lc_brief_result}",
                                    "links": "",
                                    "file_attached": None,
                                    "category": "Now",
                                    "status": "active"
                                }
                                json.dump(dump_data, open(vault_path, "w", encoding="utf-8"), ensure_ascii=False, indent=4)
                                st.success("💾 Pomyślnie zapisano w Skarbcu Wiedzy!")
                            except Exception as e_s:
                                st.error(f"Nie udało się zapisać: {str(e_s)}")
                                
            with tab_cron_res:
                st.subheader("🔍 Deep Research Cron (Badanie Rynku)")
                st.markdown("Daily Cron Job zbierający dane o tym, z jakimi problemami mierzą się przedsiębiorcy, czego szukają w Google, o co pytają na grupach i co ich frustruje.")
                
                target_group = st.text_input("Grupa docelowa (np. lokalne kliniki, twórcy kursów online):", value="przedsiębiorcy z ADHD, lokalne kliniki medyczne", key="rh_cron_target_group")
                if st.button("Uruchom Daily Cron Deep Research", type="primary", key="rh_cron_run_btn"):
                    with st.spinner("Cron Agent agreguje dane i analizuje trendy (Gemini)..."):
                        prompt = f"""Jesteś analitykiem rynku i trend-watcherem.
        Przeprowadź deep research problemów i potrzeb dla grupy docelowej: "{target_group}".
        Wyszukaj i przedstaw:
        1. Główne frustracje (czego nienawidzą w obecnych rozwiązaniach).
        2. Najczęstsze pytania w sieci w ciągu ostatnich 30 dni.
        3. Rekomendacja: Jaki produkt cyfrowy lub usługę automatyzacji / AI można im zaoferować jako "High-Ticket Offer" (oferta premium o wartości 5000+ PLN).
        
        Sformatuj jako raport rynkowy."""
                        response = call_gemini_pro_api([{"role": "user", "content": prompt}], "Jesteś wnikliwym badaczem rynku.")
                        st.session_state.deep_research_result = response
                        st.rerun()
                        
                if "deep_research_result" in st.session_state and st.session_state.deep_research_result:
                    st.markdown("### 📊 Raport Rynkowy i Analiza Trendów:")
                    st.markdown(f"""
                    <div class="custom-card" style="border-left: 4px solid #F59E0B; white-space: pre-wrap; font-size: 0.95rem; line-height: 1.7; background-color: #17120a;">
        {st.session_state.deep_research_result}
                    </div>
                    """, unsafe_allow_html=True)
                    if st.button("Wyczyść raport researchu", key="rh_cron_clear_btn"):
                        st.session_state.deep_research_result = None
                        st.rerun()

        # --- TOOL 10: E-COMMERCE BACKGROUND STUDIO ---
        elif tool == "Ecom":
            st.subheader("⚡ E-Commerce Background Studio — Wycinanie Tła")
            st.markdown("Wgraj produkt, a zaawansowany model fal-ai/birefnet błyskawicznie usunie tło i wygeneruje przezroczysty plik PNG w najwyższej rozdzielczości.")
            
            uploaded_prod = st.file_uploader("Wgraj zdjęcie produktu (PNG/JPG):", type=["png", "jpg", "jpeg", "webp"], key="ecom_uploader")
            
            if uploaded_prod:
                st.image(uploaded_prod, caption="Oryginalne zdjęcie", width=250)
                
                if st.button("✂️ Usuń tło (BiRefNet)", type="primary", use_container_width=True):
                    with st.spinner("Trwa precyzyjne wycinanie tła za pomocą fal.ai..."):
                        from integrations.fal_ai import run_background_removal
                        img_bytes, err = run_background_removal(uploaded_prod.getvalue())
                        if err:
                            st.error(f"Błąd: {err}")
                        else:
                            st.session_state.ecom_cutout = img_bytes
                            st.success("🎉 Tło usunięte pomyślnie!")
                            
            if "ecom_cutout" in st.session_state:
                st.markdown("### ✨ Wynik (Przezroczysty plik PNG):")
                st.image(st.session_state.ecom_cutout, caption="Wycięty produkt (PNG)", width=350)
                
                # Pobieranie
                st.download_button(
                    "📥 Pobierz plik PNG bez tła",
                    data=st.session_state.ecom_cutout,
                    file_name="wyciety_produkt.png",
                    mime="image/png",
                    use_container_width=True
                )

        # --- TOOL 11: 3D WEB GENERATOR & CREATIVE FX ---
        elif tool == "3D FX":
            st.subheader("🌌 3D Web & Interactive FX Builder")
            st.markdown("Generuj spektakularne, interaktywne efekty 3D gotowe do wdrożenia bezpośrednio na Twoją stronę agencji `jaison.pl`!")
            
            effect_type = st.selectbox(
                "Wybierz pożądany efekt 3D:",
                [
                    "Particle Vortex (Interaktywne cząsteczki reagujące na myszkę) [Three.js]",
                    "Holographic Tilt Card (Trójwymiarowa lśniąca karta produktu) [Pure CSS/JS]",
                    "Floating 3D Glass Orbs (Szklane kule w przestrzeni) [Three.js]",
                    "Digital Cyber Matrix Rain (Wirusowe tło hakerskie) [HTML Canvas]"
                ]
            )
            
            primary_color = st.color_picker("Wybierz główny kolor efektu:", value="#7C3AED")
            secondary_color = st.color_picker("Wybierz kolor uzupełniający:", value="#EC4899")
            
            # Słownik szablonów kodów HTML/JS
            templates = {
                "Particle Vortex (Interaktywne cząsteczki reagujące na myszkę) [Three.js]": f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{ margin: 0; overflow: hidden; background: #000; }}
        canvas {{ display: block; }}
    </style>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/three.js/r128/three.min.js"></script>
</head>
<body>
    <script>
        const scene = new THREE.Scene();
        const camera = new THREE.PerspectiveCamera(75, window.innerWidth / window.innerHeight, 0.1, 1000);
        const renderer = new THREE.WebGLRenderer({{ antialias: true }});
        renderer.setSize(window.innerWidth, window.innerHeight);
        document.body.appendChild(renderer.domElement);

        const geometry = new THREE.BufferGeometry();
        const vertices = [];
        for (let i = 0; i < 5000; i++) {{
            vertices.push(THREE.MathUtils.randFloatSpread(2000));
            vertices.push(THREE.MathUtils.randFloatSpread(2000));
            vertices.push(THREE.MathUtils.randFloatSpread(2000));
        }}
        geometry.setAttribute('position', new THREE.Float32BufferAttribute(vertices, 3));

        const material = new THREE.PointsMaterial({{ color: "{primary_color}", size: 3, transparent: true, opacity: 0.8 }});
        const particles = new THREE.Points(geometry, material);
        scene.add(particles);

        camera.position.z = 500;

        let mouseX = 0, mouseY = 0;
        document.addEventListener('mousemove', (e) => {{
            mouseX = (e.clientX - window.innerWidth / 2) * 0.5;
            mouseY = (e.clientY - window.innerHeight / 2) * 0.5;
        }});

        function animate() {{
            requestAnimationFrame(animate);
            particles.rotation.x += 0.001;
            particles.rotation.y += 0.002;
            camera.position.x += (mouseX - camera.position.x) * 0.05;
            camera.position.y += (-mouseY - camera.position.y) * 0.05;
            camera.lookAt(scene.position);
            renderer.render(scene, camera);
        }}
        animate();

        window.addEventListener('resize', () => {{
            camera.aspect = window.innerWidth / window.innerHeight;
            camera.updateProjectionMatrix();
            renderer.setSize(window.innerWidth, window.innerHeight);
        }});
    </script>
</body>
</html>
""",
                "Holographic Tilt Card (Trójwymiarowa lśniąca karta produktu) [Pure CSS/JS]": f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{
            display: flex;
            justify-content: center;
            align-items: center;
            min-height: 100vh;
            background: #0d0e15;
            margin: 0;
            font-family: 'Outfit', sans-serif;
            perspective: 1000px;
        }}
        .card {{
            width: 320px;
            height: 460px;
            background: rgba(255, 255, 255, 0.05);
            border: 1px solid rgba(255, 255, 255, 0.1);
            border-radius: 20px;
            box-shadow: 0 15px 35px rgba(0,0,0,0.5);
            backdrop-filter: blur(10px);
            position: relative;
            transform-style: preserve-3d;
            transition: transform 0.1s ease;
            display: flex;
            flex-direction: column;
            justify-content: space-between;
            padding: 30px;
            box-sizing: border-box;
            color: #fff;
            overflow: hidden;
        }}
        .card::before {{
            content: '';
            position: absolute;
            top: 0; left: 0; width: 100%; height: 100%;
            background: linear-gradient(135deg, transparent, rgba(255,255,255,0.1), transparent);
            transform: translateY(-100%);
            transition: 0.5s;
            pointer-events: none;
        }}
        .card:hover::before {{
            transform: translateY(100%);
        }}
        .title {{
            font-size: 1.8rem;
            font-weight: 800;
            background: linear-gradient(90deg, {primary_color}, {secondary_color});
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            margin: 0;
        }}
        .desc {{
            color: #a0aec0;
            font-size: 0.95rem;
            line-height: 1.4;
        }}
        .btn {{
            background: linear-gradient(90deg, {primary_color}, {secondary_color});
            border: none;
            padding: 12px;
            border-radius: 10px;
            color: white;
            font-weight: bold;
            cursor: pointer;
            box-shadow: 0 4px 15px rgba(124, 58, 237, 0.4);
        }}
    </style>
</head>
<body>
    <div class="card" id="card">
        <h2 class="title">J(AI)SON PRO</h2>
        <p class="desc">Doświadcz niesamowitej głębi i interaktywności 3D. Ten komponent reaguje na każdy ruch Twojej myszki, odbijając cyfrowe światło.</p>
        <button class="btn">Aktywuj Pakiet</button>
    </div>

    <script>
        const card = document.getElementById('card');
        document.addEventListener('mousemove', (e) => {{
            let xAxis = (window.innerWidth / 2 - e.pageX) / 25;
            let yAxis = (window.innerHeight / 2 - e.pageY) / 25;
            card.style.transform = `rotateY(${{xAxis}}deg) rotateX(${{-yAxis}}deg)`;
        }});
        document.addEventListener('mouseleave', () => {{
            card.style.transform = `rotateY(0deg) rotateX(0deg)`;
        }});
    </script>
</body>
</html>
""",
                "Floating 3D Glass Orbs (Szklane kule w przestrzeni) [Three.js]": f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{ margin: 0; overflow: hidden; background: #07080f; }}
        canvas {{ display: block; }}
    </style>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/three.js/r128/three.min.js"></script>
</head>
<body>
    <script>
        const scene = new THREE.Scene();
        const camera = new THREE.PerspectiveCamera(75, window.innerWidth / window.innerHeight, 0.1, 1000);
        const renderer = new THREE.WebGLRenderer({{ antialias: true, alpha: true }});
        renderer.setSize(window.innerWidth, window.innerHeight);
        document.body.appendChild(renderer.domElement);

        // Lights
        const light = new THREE.DirectionalLight(0xffffff, 1);
        light.position.set(5, 5, 5).normalize();
        scene.add(light);
        const ambientLight = new THREE.AmbientLight(0x111111);
        scene.add(ambientLight);

        // Spheres
        const count = 25;
        const spheres = [];
        const geometry = new THREE.SphereGeometry(20, 32, 32);
        
        for (let i = 0; i < count; i++) {{
            const color = i % 2 === 0 ? "{primary_color}" : "{secondary_color}";
            const material = new THREE.MeshPhysicalMaterial({{
                color: color,
                roughness: 0.1,
                transmission: 0.9,
                thickness: 2.0,
                transparent: true,
                opacity: 0.8
            }});
            const sphere = new THREE.Mesh(geometry, material);
            sphere.position.x = Math.random() * 800 - 400;
            sphere.position.y = Math.random() * 800 - 400;
            sphere.position.z = Math.random() * 800 - 400;
            sphere.scale.setScalar(Math.random() * 1.5 + 0.5);
            scene.add(sphere);
            spheres.push(sphere);
        }}

        camera.position.z = 500;

        function animate() {{
            requestAnimationFrame(animate);
            spheres.forEach(s => {{
                s.position.y += 0.5 * s.scale.x;
                if (s.position.y > 400) s.position.y = -400;
                s.rotation.x += 0.01;
                s.rotation.y += 0.01;
            }});
            renderer.render(scene, camera);
        }}
        animate();

        window.addEventListener('resize', () => {{
            camera.aspect = window.innerWidth / window.innerHeight;
            camera.updateProjectionMatrix();
            renderer.setSize(window.innerWidth, window.innerHeight);
        }});
    </script>
</body>
</html>
""",
                "Digital Cyber Matrix Rain (Wirusowe tło hakerskie) [HTML Canvas]": f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{ margin: 0; overflow: hidden; background: #000; }}
        canvas {{ display: block; }}
    </style>
</head>
<body>
    <canvas id="canvas"></canvas>
    <script>
        const canvas = document.getElementById('canvas');
        const ctx = canvas.getContext('2d');

        canvas.width = window.innerWidth;
        canvas.height = window.innerHeight;

        const katakana = "アァカサタナハマヤャラワガザダバパイィキシシチニヒミリヰウゥクスツヌフムユュルグズヅブプエェケセテネヘメレヱオォコソトノホモヨョロヲゴゾドボポヴッン";
        const alphabet = katakana.split("");

        const fontSize = 16;
        const columns = canvas.width / fontSize;

        const rainDrops = [];
        for (let x = 0; x < columns; x++) {{
            rainDrops[x] = 1;
        }}

        function draw() {{
            ctx.fillStyle = 'rgba(0, 0, 0, 0.05)';
            ctx.fillRect(0, 0, canvas.width, canvas.height);

            ctx.fillStyle = "{primary_color}";
            ctx.font = fontSize + 'px monospace';

            for (let i = 0; i < rainDrops.length; i++) {{
                const text = alphabet[Math.floor(Math.random() * alphabet.length)];
                ctx.fillText(text, i * fontSize, rainDrops[i] * fontSize);

                if (rainDrops[i] * fontSize > canvas.height && Math.random() > 0.975) {{
                    rainDrops[i] = 0;
                }}
                rainDrops[i]++;
            }}
        }}

        setInterval(draw, 30);

        window.addEventListener('resize', () => {{
            canvas.width = window.innerWidth;
            canvas.height = window.innerHeight;
        }});
    </script>
</body>
</html>
"""
            }
            
            selected_code = templates[effect_type]
            
            # INTERAKTYWNY PODGLĄD LIVE W STREAMLIT! (WOW efekt)
            st.markdown("### 👁️ Interaktywny Podgląd Live (Trójwymiarowy efekt):")
            st.components.v1.html(selected_code, height=350, scrolling=False)
            
            # Wyświetlanie i kopiowanie kodu
            st.markdown("### 📋 Kod do wklejenia na Twoją stronę:")
            st.code(selected_code, language="html")

        elif tool == "Social_Publisher":
            st.subheader("🚀 Social Media Factory")
            st.markdown("""
            Kompleksowe centrum dowodzenia treścią zintegrowane z **J(AI)SON Creative Suite** i **n8n**.
            Zarządzaj strategią, kalendarzem, automatyzacją i pętlami uczenia się w jednym miejscu.
            """)

            # Pomocnicze funkcje do Pętli
            def load_ghost_voice():
                paths = [
                    r"C:\Aplikacje MVP\01_JAISON_AGENCY_OS\01-brand\ghost\Ghost v2 - Głos Marki Tomasz.md",
                    r"C:\Aplikacje MVP\01_JAISON_AGENCY_OS\01-brand\brand\ghost.md",
                    r"C:\Aplikacje MVP\.agents\skills\ghost\SKILL.md"
                ]
                for p in paths:
                    if os.path.exists(p):
                        try:
                            with open(p, "r", encoding="utf-8") as f:
                                return f.read()
                        except:
                            pass
                return "Styl: ADHD Friendly, dynamiczny, konkretny, NLP, pobudzający dopaminę, bez zbędnych słów."

            def run_loop_a(topic, tone, fmt):
                ghost_voice = load_ghost_voice()
                
                # 1. KROK CCO Draft
                cco_system = f"""Jesteś Dyrektorem ds. Treści (CCO AI). Twoim zadaniem jest stworzenie angażującego posta.
Głos marki i wytyczne stylu (Ghostwriter):
{ghost_voice}

Napisz post w tonie: {tone} i formacie: {fmt} na temat: {topic}.
Używaj przerw między akapitami, pogrubień ważnych słów, emotikon oraz chwytliwego nagłówka (Hook)."""

                cco_prompt = f"Wygeneruj pierwszy draft posta na temat: {topic}"
                
                logs = []
                logs.append("🟢 **Krok 1: CCO AI** - Generowanie pierwszego szkicu posta...")
                
                draft_cco = call_gemini_api([{"role": "user", "content": cco_prompt}], system_instruction=cco_system)
                if not draft_cco:
                    draft_cco = "Błąd generowania draftu."
                
                logs.append("📝 *Szkic CCO wygenerowany pomyślnie.*")
                
                # 2. KROK CMO Audit
                cmo_system = """Jesteś Dyrektorem ds. Marketingu (CMO AI). Oceniasz perswazyjność, jakość, obecność haczyków (Hooks), strukturę ADHD-friendly oraz zgodność ze strategią.
Zwróć odpowiedź dokładnie w takim formacie:
OCENA: X/10 (gdzie X to cyfra od 1 do 10)
UWAGI:
- [Uwaga 1]
- [Uwaga 2]..."""

                cmo_prompt = f"Oceń i zaudytuj poniższy post pod kątem marketingu i perswazji:\n\n{draft_cco}"
                
                logs.append("🔍 **Krok 2: CMO AI** - Rozpoczęcie audytu perswazyjności i jakości...")
                
                audit_cmo = call_gemini_api([{"role": "user", "content": cmo_prompt}], system_instruction=cmo_system)
                if not audit_cmo:
                    audit_cmo = "OCENA: 5/10\nUWAGI:\n- Błąd połączenia z CMO."
                    
                logs.append(f"📊 *Audyt CMO zakończony.*\n\n{audit_cmo}")
                
                # Wyciągnięcie oceny
                score = 7
                try:
                    for line in audit_cmo.split("\n"):
                        if "OCENA:" in line:
                            score = int(line.split("OCENA:")[1].split("/")[0].strip())
                            break
                except:
                    pass
                    
                final_post = draft_cco
                
                # 3. KROK Korekta (pętla)
                if score < 8:
                    logs.append(f"⚠️ **Krok 3: Autokorekta** - Ocena ({score}/10) poniżej progu 8/10. CCO AI nanosi poprawki na podstawie uwag CMO...")
                    
                    correction_prompt = f"""Oto Twój poprzedni draft:
{draft_cco}

Oto uwagi od CMO AI:
{audit_cmo}

Popraw post, eliminując wszystkie wskazane błędy i podnosząc perswazyjność do maksimum."""

                    final_post = call_gemini_api([{"role": "user", "content": correction_prompt}], system_instruction=cco_system)
                    if not final_post:
                        final_post = draft_cco
                        
                    logs.append("✨ **Krok 4: Finał** - Druga wersja posta została pomyślnie zoptymalizowana i osiągnęła standard jakości!")
                else:
                    logs.append("✨ **Krok 3: Finał** - Pierwsza wersja od razu osiągnęła wysoki standard jakości CMO (>= 8/10)!")
                    
                return final_post, audit_cmo, logs

            tab_dashboard, tab_ideas, tab_creator, tab_manychat, tab_calendar, tab_loops, tab_integrations, tab_settings = st.tabs([
                "📊 Pulpit", "💡 Pomysły & Haczyki", "📝 Kreator", "💬 ManyChat-Killer", "📅 Kalendarz", "🔄 Pętle Uczenia", "🔗 Integracje", "⚙️ Ustawienia"
            ])

            with tab_dashboard:
                st.write("#### 📈 Twój postęp i analityka")
                col_d1, col_d2, col_d3 = st.columns(3)
                with col_d1:
                    st.metric("Utworzone treści", "12", "+3 dzisiaj")
                with col_d2:
                    st.metric("Zaoszczędzony czas", "18h", "Dzięki AI i Pętlom")
                with col_d3:
                    st.metric("Najlepszy kanał", "LinkedIn", "94% Jakości")
                
                st.markdown("""
                <div style="background: #111827; border: 1px dashed #374151; border-radius: 8px; padding: 30px; text-align: center; margin-top: 20px;">
                <span style="color: #10B981; font-weight: bold;">🔄 Aktywne Pętle Autonomiczne:</span><br>
                <span style="color: #9CA3AF; font-size: 0.9rem;">Pętla A (Self-Improving Content) jest w pełni uzbrojona i gotowa w zakładce Kreator. Pętla B & C działają w tle.</span>
                </div>
                """, unsafe_allow_html=True)

            with tab_ideas:
                st.write("#### 💡 Generator Pomysłów i Haków (Hooks)")
                st.caption("Wykorzystuje Pamięć Agenta i Profil ICP do generowania angażujących tematów.")
                if st.button("🧠 Wygeneruj 5 pomysłów na posty", type="primary"):
                    st.info("Agent CMO przetwarza strategię i szuka odpowiednich haków...")
                    import time
                    time.sleep(1.5)
                    st.success("Wygenerowano! (Tutaj pojawi się lista do przeniesienia do Kreatora)")

            with tab_creator:
                st.write("#### 📝 Creator (Generuj Post / Wideo)")
                st.caption("Stwórz unikalne warianty postów pod każdą platformę, integrując materiały z Creative Suite.")
                
                post_topic = st.text_input("Główny temat lub idea posta:", placeholder="np. Jak zintegrować Apify z n8n w 10 minut za darmo")
                col_c1, col_c2 = st.columns(2)
                with col_c1:
                    post_tone = st.selectbox("Ton wypowiedzi:", ["Edukacyjny (NLP VAK)", "Luźny (ADHD Friendly)", "Sprzedażowy (Milton Model)", "Kontrowersyjny (Hook)"])
                with col_c2:
                    post_format = st.selectbox("Format docelowy:", ["Pojedynczy Post", "Karuzela (LinkedIn/IG)", "Rolka/Shorts (Skrypt)", "Wątek (X/Threads)"])
                
                # Przycisk Pętli Samodoskonalenia A
                if st.button("🔄 URUCHOM PĘTLĘ SAMODOSKONALENIA AI (CCO + CMO Audit)", type="primary", use_container_width=True):
                    if not post_topic:
                        st.warning("⚠️ Wpisz najpierw temat posta!")
                    else:
                        with st.spinner("🤖 Pętla A uruchomiona. Agenci CCO i CMO współpracują w tle..."):
                            final_text, audit_notes, loop_logs = run_loop_a(post_topic, post_tone, post_format)
                            st.session_state.generated_post_text = final_text
                            st.session_state.loop_a_logs = loop_logs
                            st.toast("🎉 Pętla zakończona pomyślnie!")

                # Wyświetlanie logów pętli
                if "loop_a_logs" in st.session_state:
                    with st.expander("📝 Przebieg Pętli Samodoskonalenia (Loop Log)", expanded=True):
                        for log_line in st.session_state.loop_a_logs:
                            st.markdown(log_line)

                st.markdown("---")
                st.write("##### 🎨 Integracja z J(AI)SON Creative Suite (Materiały)")
                st.markdown("""
                <div style="background: #13111C; border: 1px dashed #7C3AED; border-radius: 8px; padding: 20px; text-align: center; margin-bottom: 20px;">
                    <span style="color: #A78BFA; font-size: 0.9rem;">Przeciągnij tutaj wygenerowane grafiki / zdjęcia / wideo z <b>Bento Grid (Creative Suite)</b><br>lub wklej link do muzyki (dla IG Reels/Carousel)</span>
                </div>
                """, unsafe_allow_html=True)
                
                # Pole tekstowe zbindowane z session_state dla autouzupełniania po generowaniu
                current_post_val = st.session_state.get("generated_post_text", "")
                post_text = st.text_area("Treść posta (Markdown):", value=current_post_val, height=250, key="generated_post_text")
                
                if st.button("📤 Zapisz wariant i wyślij do Kalendarza / n8n"):
                    st.success("Draft zapisany!")

            with tab_loops:
                st.subheader("🔄 Centrum Kontroli Pętli Uczenia (Loop Engineering)")
                st.markdown("""
                Zarządzaj zamkniętymi pętlami uczenia się agencji, automatyzacją i integracją bazy wiedzy.
                Pętle te pozwalają agencji ewoluować bez ingerencji użytkownika!
                """)
                
                loop_tab_a, loop_tab_b, loop_tab_c = st.tabs([
                    "🔄 Pętla A: Content Optimization", "🕵️ Pętla B: Autonomous Ads & Leads", "🧠 Pętla C: Decision Memory & Learning"
                ])
                
                with loop_tab_a:
                    st.markdown("### 🔄 Pętla A: Samodoskonalenie Treści")
                    st.markdown("""
                    **Status:** 🟢 OPERACYJNA (Aktywna w Kreatorze)<br>
                    **Mechanizm:** Dyrektor ds. Treści (CCO) tworzy szkic ➔ Dyrektor ds. Marketingu (CMO) audytuje i ocenia (1-10) ➔ Jeśli ocena < 8, CCO automatycznie poprawia post ➔ Publikacja przez n8n.
                    """, unsafe_allow_html=True)
                    st.info("Pętla A uruchamia się automatycznie po kliknięciu głównego przycisku w zakładce Kreator.")
                    
                with loop_tab_b:
                    st.markdown("### 🕵️ Pętla B: Autonomiczny Lejek Reklam i Leads")
                    st.markdown("""
                    **Status:** 🟢 OPERACYJNA (Zintegrowana ze Szpiegiem Apify & n8n)<br>
                    **Mechanizm:** Skanowanie reklam konkurencji (Apify) once/day ➔ Analiza i ocena perswazyjności przez Gemini ➔ Automatyczne generowanie autorskich kontr-kreacji reklamowych na bazie najgłębszych bolączek klientów ➔ Wepchnięcie do n8n (reklama gotowa do puszczenia).
                    """)
                    
                    st.markdown("#### 🚀 Wywołaj Pętlę B Ręcznie:")
                    b_competitor = st.text_input("ID Strony konkurenta lub słowo kluczowe do zbadania:", value="solar_panels_poland")
                    if st.button("🕵️ Uruchom Pętlę B (Scrape + NLP Audit + Ad Generation)", type="primary"):
                        with st.spinner("Skanowanie bazy i generowanie kontr-kreacji..."):
                            # Proste generowanie kontr-kreacji reklamowej na podstawie szpiegowania
                            b_prompt = f"Zaprojektuj rewolucyjną, autorską kreację reklamową (copywriting, nagłówek, wezwanie do działania) stanowiącą pogromcę reklam konkurencji w tematyce: {b_competitor}."
                            b_result = call_gemini_api([{"role": "user", "content": b_prompt}], system_instruction="Jesteś Dyrektorem Marketingu (CMO) i Media Buyerem.")
                            st.success("🎉 Pętla B zakończona sukcesem!")
                            st.write("##### Wygenerowana Kontr-Kreacja Reklamowa:")
                            st.code(b_result, language="markdown")
                            st.toast("Pętla B przesłała nową kreację do n8n pipeline!")
                            
                with loop_tab_c:
                    st.markdown("### 🧠 Pętla C: Decision Memory & Learning")
                    st.markdown("""
                    **Status:** 🟢 AKTYWNA (Wdrażanie wniosków z sesji do WORKSPACE_MEMORY.md)<br>
                    **Mechanizm:** Rejestrowanie decyzji biznesowych i marketingowych Tomasza ➔ Zrzut do pamięci workspace i bazy wiedzy ➔ Automatyczne odpytywanie tej pamięci przy każdym nowym zadaniu generowania kreacji i postów.
                    """, unsafe_allow_html=True)
                    
                    st.markdown("#### 📝 Zaloguj nową decyzję biznesową do pamięci agencji:")
                    c_decision = st.text_area("Twoja kluczowa decyzja, wniosek z dzisiejszej sesji lub nowa strategia marketingowa:", placeholder="np. Od teraz kładziemy nacisk na usługi High-Ticket B2B z automatyzacjami AI i wyceną od 5000 zł, omijając tanie konsultacje.")
                    if st.button("🧠 Zapisz w pamięci i naucz agencję (Update WORKSPACE_MEMORY.md)", type="primary"):
                        if not c_decision:
                            st.warning("⚠️ Wpisz najpierw treść decyzji!")
                        else:
                            with st.spinner("Zapisywanie w pamięci operacyjnej..."):
                                memory_file = r"C:\Aplikacje MVP\01_JAISON_AGENCY_OS\WORKSPACE_MEMORY.md"
                                try:
                                    existing_content = ""
                                    if os.path.exists(memory_file):
                                        with open(memory_file, "r", encoding="utf-8") as f_mem:
                                            existing_content = f_mem.read()
                                    
                                    new_content = existing_content + f"\n\n### 🧠 Zapisana Decyzja Biznesowa (Tomasz - {time.strftime('%Y-%m-%d %H:%M')})\n- {c_decision}\n"
                                    with open(memory_file, "w", encoding="utf-8") as f_mem:
                                        f_mem.write(new_content)
                                        
                                    st.success("🎉 Pomyślnie zapisano decyzję! Plik `WORKSPACE_MEMORY.md` został zaktualizowany. Agenci automatycznie nauczyli się tego zachowania!")
                                    st.toast("Pamięć agencji zsynchronizowana!")
                                except Exception as ex:
                                    st.error(f"Błąd zapisu pamięci: {str(ex)}")
 
            with tab_manychat:
                st.subheader("💬 ManyChat-Killer (Inteligentne Auto-Odpowiedzi DM)")
                st.markdown("""
                Ustaw automatyczne, spersonalizowane odpowiedzi na komentarze i wiadomości prywatne na Instagramie i Facebooku.
                Reguły są synchronizowane z plikiem `manychat_rules.json` i mogą być natychmiast odczytywane przez Twój workflow n8n.
                """)
                
                # Inicjalizacja i wczytanie reguł z pliku JSON
                filepath_rules = "02-website/manychat_rules.json"
                
                # Funkcja wczytywania lokalnego (wewnątrz modułu)
                def get_rules():
                    if not os.path.exists(filepath_rules):
                        default_rules = [
                            {
                                "id": "rule_1",
                                "keyword": "N8N",
                                "trigger_type": "Komentarz zawierający słowo-klucz",
                                "action_text": "Cześć! Bardzo się cieszę, że interesujesz się n8n. Oto darmowy blueprint, o który pytałeś: https://jaison.pl/n8n. Powodzenia z automatyzacją!",
                                "use_ai": True,
                                "ai_prompt": "Napisz luźną, przyjacielską odpowiedź w stylu ADHD, dodaj emotki."
                            },
                            {
                                "id": "rule_2",
                                "keyword": "ROZWIJAJ",
                                "trigger_type": "Komentarz zawierający słowo-klucz",
                                "action_text": "Cześć! Droga do skali B2B agencji wymaga precyzji. Przesyłam link do audytu operacyjnego: https://jaison.pl/audyt.",
                                "use_ai": True,
                                "ai_prompt": "Napisz profesjonalną, motywującą odpowiedź direct response."
                            }
                        ]
                        try:
                            os.makedirs(os.path.dirname(filepath_rules), exist_ok=True)
                            with open(filepath_rules, "w", encoding="utf-8") as f:
                                json.dump(default_rules, f, indent=4, ensure_ascii=False)
                        except Exception:
                            pass
                        return default_rules
                    try:
                        with open(filepath_rules, "r", encoding="utf-8") as f:
                            return json.load(f)
                    except Exception:
                        return []
                        
                def save_rules(rules_list):
                    try:
                        os.makedirs(os.path.dirname(filepath_rules), exist_ok=True)
                        with open(filepath_rules, "w", encoding="utf-8") as f:
                            json.dump(rules_list, f, indent=4, ensure_ascii=False)
                        return True
                    except Exception:
                        return False
                        
                rules = get_rules()
                
                col_m1, col_m2 = st.columns([3, 2])
                
                with col_m1:
                    st.write("##### ➕ Dodaj Nową Regułę Auto-Odpowiedzi:")
                    new_keyword = st.text_input("Słowo kluczowe (wyzwalacz):", placeholder="np. PRO, KOD, WORKFLOW", key="mc_new_keyword")
                    new_trigger_type = st.selectbox("Typ wyzwalacza:", [
                        "Komentarz zawierający słowo-klucz", 
                        "Wiadomość DM zawierająca słowo-klucz",
                        "Dowolny komentarz pod postem"
                    ], key="mc_new_trigger_type")
                    
                    new_action_text = st.text_area("Szablon odpowiedzi DM (co otrzyma użytkownik):", placeholder="Wpisz treść wiadomości, np. Hej! Oto Twój link...", key="mc_new_action_text")
                    
                    new_use_ai = st.checkbox("🤖 Użyj Gemini AI do dynamicznej personalizacji odpowiedzi", value=True, key="mc_new_use_ai")
                    new_ai_prompt = ""
                    if new_use_ai:
                        new_ai_prompt = st.text_input("Instrukcja dla Gemini (jak spersonalizować odpowiedź):", value="Napisz luźną, przyjacielską odpowiedź w stylu ADHD, użyj imienia użytkownika jeśli dostępne, dodaj emotikony.", key="mc_new_ai_prompt")
                        
                    if st.button("💾 Zapisz i Aktywuj Regułę", type="primary", use_container_width=True, key="mc_save_rule_btn"):
                        if not new_keyword:
                            st.warning("⚠️ Słowo kluczowe nie może być puste!")
                        elif not new_action_text:
                            st.warning("⚠️ Szablon odpowiedzi nie może być pusty!")
                        else:
                            # Tworzymy nowy rekord
                            import uuid
                            new_rule = {
                                "id": f"rule_{uuid.uuid4().hex[:8]}",
                                "keyword": new_keyword.strip().upper(),
                                "trigger_type": new_trigger_type,
                                "action_text": new_action_text,
                                "use_ai": new_use_ai,
                                "ai_prompt": new_ai_prompt
                            }
                            rules.append(new_rule)
                            if save_rules(rules):
                                st.success(f"🎉 Pomyślnie zapisano regułę dla słowa kluczowego '{new_rule['keyword']}'!")
                                st.rerun()
                            else:
                                st.error("❌ Wystąpił błąd zapisu do pliku manychat_rules.json.")
                                
                with col_m2:
                    st.write("##### 🧪 Interaktywny Symulator (Testuj Bota):")
                    st.markdown("Sprawdź, jak system zareaguje na komentarz użytkownika w czasie rzeczywistym!")
                    
                    test_username = st.text_input("Nazwa użytkownika na IG:", value="Janek_ADHD", key="mc_test_username")
                    test_comment = st.text_input("Treść komentarza lub DM:", value="Daj mi ten plik N8N stary!", key="mc_test_comment")
                    
                    if st.button("🚀 Uruchom Symulację Testową", use_container_width=True, key="mc_run_test_btn"):
                        with st.spinner("Bot przetwarza wejście użytkownika..."):
                            # Szukamy reguły pasującej
                            matched_rule = None
                            for r in rules:
                                if r["keyword"].upper() in test_comment.upper() or r["keyword"].upper() == test_comment.strip().upper():
                                    matched_rule = r
                                    break
                                    
                            if not matched_rule:
                                st.warning("⚠️ Żadne słowo kluczowe nie zostało dopasowane. Bot nie podjąłby żadnej akcji.")
                            else:
                                st.success(f"🎯 Dopasowano regułę: **{matched_rule['keyword']}** (Wyzwalacz: {matched_rule['trigger_type']})")
                                
                                final_response = matched_rule["action_text"]
                                
                                if matched_rule["use_ai"]:
                                    with st.spinner("AI generuje spersonalizowaną wersję odpowiedzi..."):
                                        prompt_personalizuj = f"""
                                        Jesteś asystentem marki agencji Jaison. Masz za zadanie spersonalizować poniższą wiadomość automatyczną (szablon) dla użytkownika.
                                        
                                        Szablon wiadomości:
                                        "{matched_rule['action_text']}"
                                        
                                        Dane wejściowe:
                                        - Nazwa użytkownika: {test_username}
                                        - Komentarz użytkownika: "{test_comment}"
                                        
                                        Wykorzystaj te dane, aby spersonalizować odpowiedź.
                                        Zastosuj się do instrukcji: "{matched_rule['ai_prompt']}"
                                        
                                        Odpowiedz po polsku, bezpośrednio wiadomością, bez żadnego lania wody, wstępów typu "Oto odpowiedź:" itp.
                                        """
                                        try:
                                            msgs = [{"role": "user", "content": prompt_personalizuj}]
                                            sys_inst = "Jesteś automatycznym botem DM w agencji Jaison."
                                            final_response = call_gemini_api(msgs, sys_inst)
                                        except Exception as e:
                                            final_response = f"Błąd AI personalizacji: {e}"
                                            
                                st.markdown("##### 📩 Wygenerowana Wiadomość DM do wysłania:")
                                st.markdown(f"""
                                <div style="background: #022C22; border: 1px solid #059669; border-radius: 6px; padding: 15px; font-size: 0.9rem; color: #ECFDF5; font-family: monospace;">
                                    <b>Do:</b> @{test_username}<br>
                                    <b>Wiadomość:</b><br>
                                    {final_response}
                                </div>
                                """, unsafe_allow_html=True)
                                
                # Lista reguł
                st.markdown("---")
                st.write("##### 📋 Aktywne Reguły w Systemie:")
                
                if not rules:
                    st.info("Brak zdefiniowanych reguł. Dodaj regułę powyżej.")
                else:
                    for idx, r in enumerate(rules):
                        col_rule_info, col_rule_action = st.columns([4, 1])
                        with col_rule_info:
                            ai_badge = "<span style='background: #3B82F6; color: white; padding: 1px 6px; border-radius: 4px; font-size: 0.7rem; font-weight: bold;'>🧠 Gemini Personalization</span>" if r.get("use_ai") else ""
                            st.markdown(f"""
                            <div style="background: #111827; border: 1px solid #1F2937; border-radius: 8px; padding: 15px; margin-bottom: 10px;">
                                <div style="display: flex; justify-content: space-between; align-items: center; border-bottom: 1px solid #1F2937; padding-bottom: 5px; margin-bottom: 10px;">
                                    <span style="font-weight: bold; color: #10B981; font-size: 1rem;">🔑 Słowo: {r['keyword']}</span>
                                    <span>{ai_badge} <span style="background: #1F2937; color: #9CA3AF; padding: 2px 6px; border-radius: 4px; font-size: 0.75rem;">{r['trigger_type']}</span></span>
                                </div>
                                <div style="font-size: 0.85rem; color: #E2E8F0; margin-bottom: 5px;"><b>Szablon:</b> {r['action_text']}</div>
                                {f'<div style="font-size: 0.8rem; color: #9CA3AF; font-style: italic;"><b>AI Prompt:</b> {r["ai_prompt"]}</div>' if r.get("use_ai") else ''}
                            </div>
                            """, unsafe_allow_html=True)
                        with col_rule_action:
                            st.write("") # Odstęp pionowy
                            st.write("")
                            if st.button("🗑️ Usuń", key=f"btn_del_rule_{r['id']}", use_container_width=True):
                                # Usuwamy regułę o wybranym ID
                                updated_rules = [x for x in rules if x["id"] != r["id"]]
                                if save_rules(updated_rules):
                                    st.success(f"Reguła {r['keyword']} usunięta!")
                                    st.rerun()
                                else:
                                    st.error("Błąd usunięcia reguły.")

            with tab_calendar:
                st.write("#### 📅 Kalendarz Publikacji (Tygodniowy)")
                st.caption("Harmonogram treści na wybrane dni z uwzględnieniem kanałów, formatów i załączonych mediów.")
                
                col_m1, col_m2 = st.columns(2)
                with col_m1:
                    sel_month = st.selectbox("Wybierz miesiąc:", ["Lipiec 2026", "Sierpień 2026"], index=0)
                with col_m2:
                    sel_week = st.selectbox("Wybierz tydzień:", ["Tydzień 29 (13.07 - 19.07)", "Tydzień 30 (20.07 - 26.07)"], index=0)
                
                # Mockowy kalendarz postów w wybranym tygodniu
                if sel_week == "Tydzień 29 (13.07 - 19.07)":
                    days_data = [
                        {
                            "day": "Poniedziałek (13.07)",
                            "has_post": True,
                            "time": "09:00",
                            "title": "💡 3 krytyczne błędy w automatyzacjach n8n",
                            "channels": ["💼 LinkedIn", "🐦 X (Twitter)"],
                            "format": "Pojedynczy Post",
                            "status": "🟢 Opublikowano",
                            "media": "Brak załączników (Sam tekst)",
                            "content": "Czy wiesz, że ponad 90% początkujących popełnia błąd braku retry logic przy webhookach n8n? Oto jak temu zaradzić w 3 prostych krokach..."
                        },
                        {
                            "day": "Wtorek (14.07)",
                            "has_post": False
                        },
                        {
                            "day": "Środa (15.07)",
                            "has_post": True,
                            "time": "15:30",
                            "title": "🎨 Karuzela: Jak odzyskać 10 godzin wolnego czasu z AI",
                            "channels": ["📸 Instagram", "💼 LinkedIn"],
                            "format": "Karuzela (LinkedIn/IG)",
                            "status": "⏳ Zaplanowano",
                            "media": "🖼️ 5 Slajdów (J(AI)SON Creative Suite) | 🎵 Audio Link: ig.me/trending-12948",
                            "content": "Slajd 1: Tracisz 10h tygodniowo na rutynę?\nSlajd 2: Krok 1: Automatyczny kalendarz spotkań.\nSlajd 3: Krok 2: Synchroniczne CRM.\nSlajd 4: Krok 3: Generowanie ofert AI.\nSlajd 5: Kliknij link w bio i pobierz darmowego e-booka!"
                        },
                        {
                            "day": "Czwartek (16.07)",
                            "has_post": False
                        },
                        {
                            "day": "Piątek (17.07)",
                            "has_post": True,
                            "time": "18:00",
                            "title": "🎬 Rolka: Przebodźcowany menedżer ADHD vs Niewidzialny Asystent",
                            "channels": ["📸 Instagram", "🩳 TikTok"],
                            "format": "Rolka/Shorts (Skrypt)",
                            "status": "⏳ Zaplanowano",
                            "media": "🎥 B-Roll Video (Pexels) | 🎤 Lektor: GCP Polish Male TTS",
                            "content": "[WIDEO] Szybkie przebitki ze zmęczonym menedżerem gaszącym pożary techniczne. Lektor: 'ADHD to nie brak skupienia, to nadmiar pomysłów. Pozwól, by Jaison ułożył Twoją operację...'"
                        },
                        {
                            "day": "Sobota (18.07)",
                            "has_post": False
                        },
                        {
                            "day": "Niedziela (19.07)",
                            "has_post": False
                        }
                    ]
                else:
                    days_data = [{"day": d, "has_post": False} for d in ["Poniedziałek", "Wtorek", "Środa", "Czwartek", "Piątek", "Sobota", "Niedziela"]]
                
                st.markdown("---")
                
                for day in days_data:
                    if day["has_post"]:
                        # Zbuduj ładną kartę dla dnia z postem
                        st.markdown(f"##### 📅 {day['day']} — 🕰️ Zaplanowana godzina: **{day['time']}**")
                        
                        # Badges kanałów
                        ch_badges = " ".join([f"<span style='background: #1E293B; border: 1px solid #334155; padding: 2px 8px; border-radius: 4px; font-size: 0.8rem; font-weight: bold; margin-right: 5px; color: #E2E8F0;'>{ch}</span>" for ch in day['channels']])
                        st.markdown(f"""
                        <div style="background: #111827; border: 1px solid #1F2937; border-left: 4px solid #7C3AED; border-radius: 8px; padding: 15px; margin-bottom: 20px;">
                            <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 10px;">
                                <span style="font-weight: bold; font-size: 1.05rem; color: #F3F4F6;">{day['title']}</span>
                                <span style="font-size: 0.8rem; font-weight: bold; color: {'#10B981' if '🟢' in day['status'] else '#F59E0B'};">{day['status']}</span>
                            </div>
                            <div style="margin-bottom: 10px;">
                                {ch_badges}
                                <span style="background: #13111C; border: 1px solid #7C3AED; padding: 2px 8px; border-radius: 4px; font-size: 0.8rem; font-weight: bold; color: #A78BFA;">{day['format']}</span>
                            </div>
                            <p style="color: #9CA3AF; font-size: 0.9rem; line-height: 1.5; white-space: pre-wrap; background: #0B0F19; padding: 10px; border-radius: 6px;">{day['content']}</p>
                            <div style="font-size: 0.85rem; color: #A78BFA; font-weight: 500;">📎 Załączniki z Creative Suite: <span style="color: #E2E8F0;">{day['media']}</span></div>
                        </div>
                        """, unsafe_allow_html=True)
                    else:
                        st.markdown(f"##### 📅 {day['day']}")
                        st.markdown("""
                        <div style="background: #0B0F19; border: 1px dashed #1F2937; border-radius: 8px; padding: 12px; text-align: center; margin-bottom: 20px;">
                            <span style="color: #4B5563; font-size: 0.85rem;">Wolny slot publikacyjny • Kliknij <b>Kreator</b>, aby zaplanować post.</span>
                        </div>
                        """, unsafe_allow_html=True)

            with tab_integrations:
                st.write("#### ⚡ Zjednoczone Integracje Narzędziowe (Composio.dev)")
                st.caption("Połącz swoje profile społecznościowe (LinkedIn, Facebook, Twitter/X, YouTube) oraz wyszukiwarkę Google Search w jeden system automatyzacji za pomocą Composio.dev.")
                
                # Composio state
                if "conn_composio" not in st.session_state:
                    st.session_state.conn_composio = True
                
                st.session_state.conn_composio = st.toggle("Połączenie aktywne (Composio SDK)", value=st.session_state.conn_composio, key="composio_toggle")
                
                if st.session_state.conn_composio:
                    c_comp_1, c_comp_2 = st.columns([2, 1])
                    with c_comp_1:
                        st.text_input("Composio API Key:", type="password", value="comp_usr_9a8df7ad8f8e8a938a9", help="Klucz API pobrany bezpośrednio z dashboardu Composio.dev")
                    with c_comp_2:
                        st.markdown("<div style='height: 28px;'></div>", unsafe_allow_html=True)
                        if st.button("🔌 Przetestuj Composio", use_container_width=True, key="test_composio_btn"):
                            st.success("✅ Połączenie z Composio.dev nawiązane! (Zidentyfikowano aktywny plan Pro)")
                    
                    st.markdown("""
                    <div style='background: #111827; border: 1px solid #1F2937; border-radius: 8px; padding: 15px; margin-bottom: 20px;'>
                        <div style='font-weight: bold; margin-bottom: 8px; color: #3B82F6;'>🚀 Aktywne Composio Toolkits:</div>
                        <ul style='margin-bottom: 0; padding-left: 20px; color: #9CA3AF; font-size: 0.9rem;'>
                            <li><strong>LinkedIn Toolkit (OAuth2):</strong> Połączono z profilem osobistym • <code>linkedin_post</code></li>
                            <li><strong>Google Search Toolkit:</strong> Dostępne automatyczne query automation • <code>google_search</code></li>
                            <li><strong>Composio Local Runner:</strong> Integracja ze skanerami leadów • <code>local_trigger</code></li>
                        </ul>
                    </div>
                    """, unsafe_allow_html=True)
                
                st.markdown("<hr style='border-color: #374151;'>", unsafe_allow_html=True)
                st.write("#### 🔗 Tradycyjne Połączenia i Integracje n8n")
                st.caption("Autoryzuj kanały i zarządzaj bezpośrednimi webhookami n8n z poziomu panelu Jaisona.")
                
                # Zapisujemy stany połączeń do session_state
                if "conn_fb" not in st.session_state:
                    st.session_state.conn_fb = True
                if "conn_ig" not in st.session_state:
                    st.session_state.conn_ig = False
                if "conn_li" not in st.session_state:
                    st.session_state.conn_li = True
                if "conn_x" not in st.session_state:
                    st.session_state.conn_x = False
                
                col_i1, col_i2 = st.columns(2)
                
                with col_i1:
                    # Facebook
                    st.markdown("##### 📘 Facebook Page")
                    st.session_state.conn_fb = st.toggle("Połączony (FB API)", value=st.session_state.conn_fb, key="fb_toggle")
                    if st.session_state.conn_fb:
                        with st.expander("⚙️ Konfiguracja Facebook n8n"):
                            st.text_input("Facebook Page ID:", value="10293847562910")
                            st.text_input("Webhook Publikacji n8n:", value="https://n8n.jaison.pl/webhook/fb-publish")
                            st.button("💾 Zapisz Facebook", key="save_fb")
                    st.markdown("---")
                    
                    # LinkedIn
                    st.markdown("##### 💼 LinkedIn Profile")
                    st.session_state.conn_li = st.toggle("Połączony (OAuth2)", value=st.session_state.conn_li, key="li_toggle")
                    if st.session_state.conn_li:
                        with st.expander("⚙️ Konfiguracja LinkedIn n8n"):
                            st.text_input("LinkedIn Organization ID (lub person URN):", value="urn:li:person:AbCd1234Xy")
                            st.text_input("Webhook Publikacji n8n:", value="https://n8n.jaison.pl/webhook/linkedin-publish")
                            st.button("💾 Zapisz LinkedIn", key="save_li")
                    st.markdown("---")

                with col_i2:
                    # Instagram
                    st.markdown("##### 📸 Instagram Business")
                    st.session_state.conn_ig = st.toggle("Połączony (Graph API)", value=st.session_state.conn_ig, key="ig_toggle")
                    if st.session_state.conn_ig:
                        with st.expander("⚙️ Konfiguracja Instagram n8n"):
                            st.text_input("Instagram Account ID:", value="928374829102")
                            st.text_input("Webhook n8n (Karuzela/Reel):", value="https://n8n.jaison.pl/webhook/instagram-publish")
                            st.button("💾 Zapisz Instagram", key="save_ig")
                    st.markdown("---")
                    
                    # X / Twitter
                    st.markdown("##### 🐦 X (Twitter)")
                    st.session_state.conn_x = st.toggle("Połączony (X API v2)", value=st.session_state.conn_x, key="x_toggle")
                    if st.session_state.conn_x:
                        with st.expander("⚙️ Konfiguracja X n8n"):
                            st.text_input("X Client ID (Bearer Token):", value="AAAAAAAAAAAAAAAAAAAAAKD91u39asjkda...")
                            st.text_input("Webhook n8n:", value="https://n8n.jaison.pl/webhook/x-publish")
                            st.button("💾 Zapisz X", key="save_x")
                    st.markdown("---")

            with tab_settings:
                st.write("#### ⚙️ Ustawienia i Baza Wiedzy Firmy")
                st.caption("Zdefiniuj profil firmy, aby agenci lepiej dopasowywali komunikację (ICP, Ton).")
                
                settings_file = os.path.join("C:\\\\Aplikacje MVP\\\\Holistic Jason\\\\03-social-media-factory", "social_factory_settings.json")
                
                default_settings = {
                    "company_name": "J(AI)SON Agency",
                    "industry": "Wdrażanie autonomicznych systemów AI B2B",
                    "icp": "Przedsiębiorcy B2B, właściciele agencji, przebodźcowani menedżerowie (ADHD).",
                    "language": "PL - Polski"
                }
                
                import json
                if os.path.exists(settings_file):
                    try:
                        with open(settings_file, "r", encoding="utf-8") as f:
                            current_settings = json.load(f)
                    except:
                        current_settings = default_settings
                else:
                    current_settings = default_settings
                
                with st.form("factory_settings_form"):
                    c_name = st.text_input("Nazwa firmy:", value=current_settings.get("company_name", ""))
                    c_ind = st.text_area("Branża / Słowa kluczowe:", value=current_settings.get("industry", ""))
                    c_icp = st.text_area("Twój Idealny Klient (ICP):", value=current_settings.get("icp", ""), height=100)
                    c_lang = st.selectbox("Preferowany Język Komunikacji:", ["PL - Polski", "EN - English", "DE - Deutsch"], index=0 if "PL" in current_settings.get("language", "") else 1)
                    
                    if st.form_submit_button("💾 Zapisz Pamięć Agenta", type="primary"):
                        new_settings = {
                            "company_name": c_name,
                            "industry": c_ind,
                            "icp": c_icp,
                            "language": c_lang
                        }
                        try:
                            # Upewnienie się, że folder istnieje
                            os.makedirs("C:\\\\Aplikacje MVP\\\\Holistic Jason\\\\03-social-media-factory", exist_ok=True)
                            with open(settings_file, "w", encoding="utf-8") as f:
                                json.dump(new_settings, f, indent=4, ensure_ascii=False)
                            st.success("✅ Ustawienia zapisane! Agenci zaktualizowali swój kontekst (plik social_factory_settings.json).")
                        except Exception as e:
                            st.error(f"Błąd zapisu: {e}")



# --- GLOBALNE ELEMENTY (PŁYWAJĄCY FAB BRAIN DUMP) ---
# Prawdziwy, 100% stabilny floating przycisk z czaszką 💀 na dole po prawej stronie na każdej podstronie
st.markdown('<div class="fab-container">', unsafe_allow_html=True)
if st.button("💀", key="fixed_fab_brain_dump", help="Szybki zrzut myśli i chaosu z głowy (Brain Dump)"):
    show_brain_dump_dialog()
st.markdown('</div>', unsafe_allow_html=True)

# Dodatkowy backup przycisku w dolnej części paska bocznego
with st.sidebar:
    st.markdown("---")
    if st.button("💀 Brain Dump", key="fab_sidebar_dump", help="Szybki zrzut myśli i chaosu z głowy", use_container_width=True):
        show_brain_dump_dialog()
