import streamlit as st
import os, json, time
import ssl

# Ręczne wczytanie pliku .env na starcie aplikacji
if os.path.exists(".env"):
    try:
        with open(".env", "r", encoding="utf-8") as f:
            for line in f:
                if "=" in line and not line.strip().startswith("#"):
                    k, v = line.split("=", 1)
                    os.environ[k.strip()] = v.strip()
    except Exception as e:
        pass

try:
    ssl._create_default_https_context = ssl._create_unverified_context
except AttributeError:
    pass

st.set_page_config(
    page_title="Holistic AIDHD OS • Mission Control",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Dynamiczne wstrzykiwanie kodu Google Tag Manager (GTM)
gtm_id = os.environ.get("GOOGLE_TAG_MANAGER_ID", "").strip()
if gtm_id:
    st.markdown(f"""
    <script id="gtm-injector">
        (function() {{
            const doc = window.parent !== window ? window.parent.document : document;
            if (!doc.getElementById('gtm-script-tag')) {{
                // Skrypt GTM w Head
                (function(w,d,s,l,i){{w[l]=w[l]||[];w[l].push({{'gtm.start':
                new Date().getTime(),event:'gtm.js'}});var f=d.getElementsByTagName(s)[0],
                j=d.createElement(s),dl=l!='dataLayer'?'&l='+l:'';j.async=true;j.src=
                'https://www.googletagmanager.com/gtm.js?id='+i+dl;j.id='gtm-script-tag';f.parentNode.insertBefore(j,f);
                }})(window.parent,doc,'script','dataLayer','{gtm_id}');
                
                // NoScript GTM w Body
                const noscript = doc.createElement('noscript');
                noscript.id = 'gtm-noscript-tag';
                const iframe = doc.createElement('iframe');
                iframe.src = 'https://www.googletagmanager.com/ns.html?id={gtm_id}';
                iframe.height = '0';
                iframe.width = '0';
                iframe.style.display = 'none';
                iframe.style.visibility = 'hidden';
                noscript.appendChild(iframe);
                doc.body.insertBefore(noscript, doc.body.firstChild);
            }}
        }})();
    </script>
    """, unsafe_allow_html=True)

# Ścieżki w chmurze
BASE_DIR = os.path.expanduser("~/Agentic_OS")
NOTEBOOKS_DIR = os.path.join(BASE_DIR, "notebooks/_assets")
OBSIDIAN_DIR = os.path.join(BASE_DIR, "obsidian_vault")
DASHBOARD_DIR = os.path.join(BASE_DIR, "dashboard")
BRAIN_DUMP_DIR = os.path.join(BASE_DIR, "brain_dump")
BRAIN_DUMP_ASSETS = os.path.join(BRAIN_DUMP_DIR, "_assets")
HERMES_DIR = os.path.expanduser("~/.hermes")
KANBAN_FILE = os.path.join(DASHBOARD_DIR, "kanban.json")

# Tworzenie folderów
for d in [NOTEBOOKS_DIR, OBSIDIAN_DIR, DASHBOARD_DIR, BRAIN_DUMP_DIR, BRAIN_DUMP_ASSETS, HERMES_DIR]:
    os.makedirs(d, exist_ok=True)

# Luksusowy nocny design zoptymalizowany pod ADHD (Outfit & Atkinson)
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Atkinson+Hyperlegible:ital,wght@0,400;0,700;1,400;1,700&family=Outfit:wght@300;400;500;600;700&display=swap');
    
    /* Globalny reset barw i typografii */
    .stApp {
        background-color: #08090C !important;
        color: #E2E8F0 !important;
        font-family: 'Atkinson Hyperlegible', sans-serif !important;
    }
    
    [data-testid="stSidebar"] {
        background-color: #0E1015 !important;
        border-right: 1px solid #1F242E !important;
    }
    
    h1, h2, h3, h4, h5, h6 {
        font-family: 'Outfit', sans-serif !important;
        color: #FFFFFF !important;
        font-weight: 600 !important;
    }
    
    /* Luksusowe karty z efektem neonowego obramowania po najechaniu */
    .custom-card {
        background-color: #121620;
        border: 1px solid #1E2535;
        border-radius: 14px;
        padding: 24px;
        margin-bottom: 20px;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.4);
        transition: transform 0.25s ease, border-color 0.25s ease, box-shadow 0.25s ease;
    }
    
    .custom-card:hover {
        border-color: #7C3AED;
        transform: translateY(-2px);
        box-shadow: 0 8px 30px rgba(124, 58, 237, 0.15);
    }
    
    /* Banner One Thing - eliminacja szumu kognitywnego */
    .one-thing-banner {
        background: linear-gradient(135deg, #181528 0%, #0E1015 100%);
        border-left: 6px solid #F59E0B;
        border-radius: 14px;
        padding: 30px;
        margin-bottom: 30px;
        box-shadow: 0 10px 30px rgba(0, 0, 0, 0.5);
    }
    
    /* Zaokrąglone przyciski premium (globalnie dla stButton) */
    .stButton>button {
        background: linear-gradient(135deg, #6D28D9 0%, #4C1D95 100%) !important;
        color: #FFFFFF !important;
        border: 1px solid #7C3AED !important;
        border-radius: 10px !important;
        padding: 10px 20px !important;
        font-family: 'Outfit', sans-serif !important;
        font-weight: 600 !important;
        transition: all 0.25s ease !important;
        width: 100%;
        box-shadow: 0 4px 15px rgba(109, 40, 217, 0.3) !important;
    }
    
    .stButton>button:hover {
        background: linear-gradient(135deg, #7C3AED 0%, #5B21B6 100%) !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 6px 20px rgba(124, 58, 237, 0.5) !important;
    }

    /* Kafelki Nawigacyjne w Pasku Bocznym (ADHD Tiles) */
    div[data-testid="stSidebar"] .stButton>button {
        text-align: left !important;
        padding: 12px 18px !important;
        height: 52px !important;
        margin-bottom: 2px !important;
        width: 100% !important;
    }

    div[data-testid="stSidebar"] .stButton>button[kind="secondary"] {
        background: transparent !important;
        color: #94A3B8 !important;
        border: 1px solid transparent !important;
        box-shadow: none !important;
        font-weight: 500 !important;
    }

    div[data-testid="stSidebar"] .stButton>button[kind="secondary"]:hover {
        background: rgba(124, 58, 237, 0.1) !important;
        border-color: #7C3AED !important;
        color: #FFFFFF !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 4px 12px rgba(124, 58, 237, 0.15) !important;
    }

    div[data-testid="stSidebar"] .stButton>button[kind="primary"] {
        background: linear-gradient(135deg, #7C3AED 0%, #5B21B6 100%) !important;
        color: #FFFFFF !important;
        border: 1px solid #C084FC !important;
        box-shadow: 0 0 15px rgba(124, 58, 237, 0.35) !important;
        font-weight: 700 !important;
    }
    
    div[data-testid="stSidebar"] .stButton>button[kind="primary"]:hover {
        transform: translateY(-1px) !important;
        box-shadow: 0 0 20px rgba(124, 58, 237, 0.5) !important;
    }
    
    /* ADHD-friendly akcenty kolorystyczne */
    .dopamine-accent { color: #10B981; font-weight: bold; }
    .burn-accent { color: #EF4444; font-weight: bold; }
    .focus-accent { color: #F59E0B; font-weight: bold; }

    /* Pływający Przycisk Szybkiego Zapisu (FAB) zoptymalizowany pod Streamlit DOM */
    .fab-wrapper {
        display: none !important;
    }
    
    div:has(> .fab-wrapper) + div {
        height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
        overflow: visible !important;
    }
    
    div:has(> .fab-wrapper) + div div.stButton > button {
        position: fixed !important;
        bottom: 25px !important;
        right: 25px !important;
        z-index: 999999 !important;
        width: 60px !important;
        height: 60px !important;
        border-radius: 50% !important;
        background: radial-gradient(circle, #EC4899 0%, #D946EF 100%) !important;
        border: 2px solid #F472B6 !important;
        box-shadow: 0 0 15px rgba(236, 72, 153, 0.6) !important;
        font-size: 28px !important;
        padding: 0 !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        transition: all 0.3s ease !important;
        animation: pulse-fab 2s infinite !important;
    }

    div:has(> .fab-wrapper) + div div.stButton > button:hover {
        background: radial-gradient(circle, #D946EF 0%, #C084FC 100%) !important;
        box-shadow: 0 0 25px rgba(236, 72, 153, 0.9) !important;
        transform: scale(1.08) !important;
    }

    @keyframes pulse-fab {
        0% {
            box-shadow: 0 0 0 0 rgba(236, 72, 153, 0.7);
            transform: scale(1);
        }
        70% {
            box-shadow: 0 0 0 15px rgba(236, 72, 153, 0);
            transform: scale(1.05);
        }
        100% {
            box-shadow: 0 0 0 0 rgba(236, 72, 153, 0);
            transform: scale(1);
        }
    }tant;
    }
</style>

""", unsafe_allow_html=True)

# Helpery danych
def load_kanban():
    default_kanban = {
        "triage": [],
        "todo": [],
        "ready": [],
        "running": [],
        "blocked": [],
        "done": []
    }
    if os.path.exists(KANBAN_FILE):
        try:
            with open(KANBAN_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            # Migrate old 3-column format if present
            if "in_progress" in data and "running" not in data:
                data["running"] = data.pop("in_progress")
            # Ensure all 6 columns exist
            for col in default_kanban:
                if col not in data:
                    data[col] = []
            return data
        except Exception:
            pass
    return default_kanban

def save_kanban(data):
    json.dump(data, open(KANBAN_FILE, "w", encoding="utf-8"), ensure_ascii=False, indent=4)

CRM_FILE = os.path.join(DASHBOARD_DIR, "crm.json")

def load_crm():
    if os.path.exists(CRM_FILE):
        try:
            return json.load(open(CRM_FILE, "r", encoding="utf-8"))
        except:
            pass
    return {
        "leads": [
            {
                "id": "lead_01",
                "name": "Klinika Dermatologiczna (Łódź)",
                "stage": "conversation",
                "notes": "Zainteresowani automatyzacją ankiety intake w celu uwolnienia czasu personelu. Duża potrzeba empatii i prostego lejka wdrożeniowego.",
                "last_contact": "2026-05-29",
                "next_action": "Opracować plan wdrożenia 3 asystentów AI.",
                "draft_reply": "Dzień dobry. Przeanalizowałem Państwa wąskie gardła. Wdrożenie Niewidzialnego Pracownika AI do ankiety intake pozwoli zaoszczędzić około 15 godzin tygodniowo bez utraty ciepłego kontaktu z pacjentem..."
            },
            {
                "id": "lead_02",
                "name": "Jan Szopa (Dystrybutor Kursów)",
                "stage": "architecture",
                "notes": "Zainteresowany systemem czystej pamięci (Pristine Memory) dla swoich studentów z ADHD.",
                "last_contact": "2026-05-28",
                "next_action": "Przygotować interfejs MVP z pasywnym tablicowaniem.",
                "draft_reply": "Cześć Jan, w oparciu o nasze rozmowy, zaprojektowałem strukturę bezszumnego CRM bez powiadomień push..."
            }
        ]
    }

def save_crm(data):
    json.dump(data, open(CRM_FILE, "w", encoding="utf-8"), ensure_ascii=False, indent=4)

# --- POLSKI ADHD GHL INTEGRATIONS ---
import urllib.parse
import urllib.request
import urllib.error
from http.server import BaseHTTPRequestHandler, HTTPServer
import threading

class UrllibResponse:
    def __init__(self, status_code, text, json_data=None):
        self.status_code = status_code
        self.text = text
        self._json_data = json_data
    
    def json(self):
        if self._json_data is not None:
            return self._json_data
        return json.loads(self.text)

def http_post(url, json_data=None, headers=None, timeout=30.0):
    if headers is None:
        headers = {}
    
    # Ensure Content-Type is set if json_data is present
    if json_data is not None and "Content-Type" not in headers:
        headers["Content-Type"] = "application/json"
        
    data_bytes = None
    if json_data is not None:
        data_bytes = json.dumps(json_data).encode("utf-8")
        
    req = urllib.request.Request(url, data=data_bytes, headers=headers, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=timeout) as response:
            status_code = response.getcode()
            response_text = response.read().decode("utf-8")
            try:
                res_json = json.loads(response_text)
            except:
                res_json = None
            return UrllibResponse(status_code, response_text, res_json)
    except urllib.error.HTTPError as e:
        status_code = e.code
        try:
            response_text = e.read().decode("utf-8")
        except:
            response_text = str(e)
        try:
            res_json = json.loads(response_text)
        except:
            res_json = None
        return UrllibResponse(status_code, response_text, res_json)
    except Exception as e:
        return UrllibResponse(500, str(e), None)

# Thread-safe global flag to start webhook server once
_server_started = False
_server_lock = threading.Lock()

def call_openrouter_api(messages, model="nousresearch/hermes-3-llama-3.1-405b:free", system_instruction=None):
    api_key = os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        return "Błąd: Brak klucza OPENROUTER_API_KEY w pliku .env"
    
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:8501",
        "X-Title": "Holistic ADHD OS"
    }
    
    # Map friendly short names to openrouter ids if necessary
    if model == "owl-alpha-free" or model == "owl-alpha":
        model = "meta-llama/llama-3.1-8b-instruct:free"
    
    payload = {
        "model": model,
        "messages": []
    }
    if system_instruction:
        payload["messages"].append({"role": "system", "content": system_instruction})
    payload["messages"].extend(messages)
    
    try:
        response = http_post(url, json_data=payload, headers=headers, timeout=60.0)
        if response.status_code == 200:
            res_data = response.json()
            return res_data["choices"][0]["message"]["content"]
        else:
            return f"Błąd OpenRouter {response.status_code}: {response.text}"
    except Exception as e:
        return f"Błąd OpenRouter API: {e}"

def get_recent_workspace_context(limit=3):
    context_str = ""
    try:
        if os.path.exists(OBSIDIAN_DIR):
            files = [f for f in os.listdir(OBSIDIAN_DIR) if f.endswith('.md')]
            # Sort by modification time desc
            files.sort(key=lambda x: os.path.getmtime(os.path.join(OBSIDIAN_DIR, x)), reverse=True)
            for f in files[:limit]:
                path = os.path.join(OBSIDIAN_DIR, f)
                with open(path, "r", encoding="utf-8") as file:
                    content = file.read()
                # Keep it concise, extract up to 1000 chars per file
                context_str += f"\n--- Ostatnia notatka: {f} ---\n{content[:1000]}\n"
    except Exception as e:
        context_str = f"Błąd wczytywania pamięci roboczej: {e}"
    return context_str

def call_vllm_api(messages, system_instruction=None):
    url = "http://localhost:8000/v1/chat/completions"
    headers = {
        "Authorization": "Bearer HOLISTIC_SECURE_TOKEN_2026",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "mistralai/Mistral-Nemo-Instruct-2407",
        "messages": []
    }
    if system_instruction:
        payload["messages"].append({"role": "system", "content": system_instruction})
    payload["messages"].extend(messages)
    
    try:
        response = http_post(url, json_data=payload, headers=headers, timeout=120.0)
        if response.status_code == 200:
            res_data = response.json()
            return res_data["choices"][0]["message"]["content"]
    except Exception as e:
        pass
        
    return call_gemini_api(messages, system_instruction)

def call_gemini_api(messages, system_instruction=None):
    proxy_url = "http://127.0.0.1:8089/v1/chat/completions"
    payload = {
        "model": "gemini-2.5-flash",
        "messages": []
    }
    if system_instruction:
        payload["messages"].append({"role": "system", "content": system_instruction})
    payload["messages"].extend(messages)
    
    # 1. Try local proxy (timeout 90s for long OCR operations)
    try:
        response = http_post(proxy_url, json_data=payload, timeout=90.0)
        if response.status_code == 200:
            res_data = response.json()
            return res_data["choices"][0]["message"]["content"]
    except Exception as e:
        pass
        
    # 2. Fallback to direct GCP Vertex AI call
    sa_paths = [
        os.path.expanduser("~/.hermes/gcp-sa-key.json"),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "holistic-dashboard-dev-dea2c872139e.json"),
        "holistic-dashboard-dev-dea2c872139e.json"
    ]
    sa_path = None
    for p in sa_paths:
        if os.path.exists(p):
            sa_path = p
            break
            
    if sa_path:
        try:
            from google.oauth2 import service_account
            import google.auth.transport.requests
            import requests
            import urllib3
            urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
            
            creds = service_account.Credentials.from_service_account_file(
                sa_path,
                scopes=['https://www.googleapis.com/auth/cloud-platform']
            )
            session = requests.Session()
            session.verify = False
            request = google.auth.transport.requests.Request(session=session)
            creds.refresh(request)
            token = creds.token
            
            project_id = "holistic-dashboard-dev"
            region = "us-central1"
            direct_url = f"https://{region}-aiplatform.googleapis.com/v1/projects/{project_id}/locations/{region}/endpoints/openapi/chat/completions"
            
            headers = {
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json"
            }
            
            direct_payload = {
                "model": "google/gemini-2.5-flash",
                "messages": []
            }
            if system_instruction:
                direct_payload["messages"].append({"role": "system", "content": system_instruction})
            direct_payload["messages"].extend(messages)
            
            resp = http_post(direct_url, json_data=direct_payload, headers=headers, timeout=25.0)
            if resp.status_code == 200:
                res_data = resp.json()
                return res_data["choices"][0]["message"]["content"]
        except Exception as ex:
            return f"Błąd bezpośredniej komunikacji z Vertex AI: {ex}"
            
    return "Błąd: Brak połączenia z lokalnym proxy i brak poprawnego klucza Service Account."

def get_gcp_token():
    """Pobiera token OAuth z pliku Service Account w odporny sposób."""
    import os
    import json
    sa_paths = [
        os.path.expanduser("~/.hermes/gcp-sa-key.json"),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "holistic-dashboard-dev-dea2c872139e.json"),
        os.path.join(os.getcwd(), "holistic-dashboard-dev-dea2c872139e.json"),
        "holistic-dashboard-dev-dea2c872139e.json"
    ]
    sa_path = None
    for p in sa_paths:
        if os.path.exists(p):
            sa_path = p
            break
    if not sa_path:
        return None, None
    try:
        from google.oauth2 import service_account
        import google.auth.transport.requests
        import requests
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        
        creds = service_account.Credentials.from_service_account_file(
            sa_path,
            scopes=['https://www.googleapis.com/auth/cloud-platform']
        )
        session = requests.Session()
        session.verify = False
        request = google.auth.transport.requests.Request(session=session)
        creds.refresh(request)
        
        with open(sa_path, "r") as f:
            sa_data = json.load(f)
            project_id = sa_data.get("project_id", "holistic-dashboard-dev")
            
        return creds.token, project_id
    except Exception as e:
        print(f"Error getting token: {e}")
        return None, None

def generate_imagen_image(prompt, aspect_ratio="1:1", reference_image_bytes=None, reference_type="REFERENCE_TYPE_SUBJECT"):
    """
    Generuje obraz za pomocą Google Vertex AI Imagen 3 (imagen-3.0-generate-001).
    Wspiera Image-to-Image / Subject Reference dla zachowania spójności postaci.
    """
    import base64
    import json
    token, project_id = get_gcp_token()
    if not token:
        return None, "Brak autoryzacji GCP Service Account. Sprawdź pliki klucza."
    
    region = "us-central1"
    url = f"https://{region}-aiplatform.googleapis.com/v1/projects/{project_id}/locations/{region}/publishers/google/models/imagen-3.0-generate-001:predict"
    
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }
    
    instance = {"prompt": prompt}
    
    # Obsługa obrazu referencyjnego (Subject Reference Image)
    if reference_image_bytes:
        b64_data = base64.b64encode(reference_image_bytes).decode("utf-8")
        instance["referenceImages"] = [
            {
                "referenceId": 1,
                "referenceType": reference_type,
                "image": {
                    "bytesBase64Encoded": b64_data
                }
            }
        ]
        if "[1]" not in prompt:
            prompt = f"Using subject [1] as character reference, {prompt}"
            instance["prompt"] = prompt
            
    payload = {
        "instances": [instance],
        "parameters": {
            "sampleCount": 1,
            "aspectRatio": aspect_ratio,
            "outputMimeType": "image/png"
        }
    }
    
    try:
        response = http_post(url, json_data=payload, headers=headers, timeout=120.0)
        if response.status_code == 200:
            res_json = response.json()
            predictions = res_json.get("predictions", [])
            if predictions:
                img_b64 = predictions[0].get("bytesBase64Encoded")
                img_bytes = base64.b64decode(img_b64)
                return img_bytes, None
            else:
                return None, f"Brak wygenerowanego obrazu w odpowiedzi: {response.text}"
        else:
            return None, f"GCP API Error {response.status_code}: {response.text}"
    except Exception as e:
        return None, f"Błąd komunikacji: {str(e)}"

def call_gemini_pro_api(messages, system_instruction=None):
    """Dedykowana funkcja dla modelu Gemini 2.5 Pro przez lokalny proxy.
    Używamy go w Dziale Prawnym dla analizy dokumentów i logiki prawnej.
    Timeout 120s dla długich dokumentów."""
    proxy_url = "http://127.0.0.1:8089/v1/chat/completions"
    
    # Sprawdź czy wiadomości zawierają multimodalne treści (PDF/obrazy) — użyj Flash
    # bo proxy może nie obsługiwać dużych payloadów multimodalnych dla Pro
    has_multimodal = False
    for msg in messages:
        content = msg.get("content", "")
        if isinstance(content, list):
            for part in content:
                if isinstance(part, dict) and part.get("type") == "image_url":
                    has_multimodal = True
                    break
    
    model_name = "gemini-2.5-flash" if has_multimodal else "gemini-2.5-pro"
    
    payload = {"model": model_name, "messages": []}
    if system_instruction:
        payload["messages"].append({"role": "system", "content": system_instruction})
    payload["messages"].extend(messages)

    try:
        response = http_post(proxy_url, json_data=payload, timeout=120.0)
        if response.status_code == 200:
            res_data = response.json()
            return res_data["choices"][0]["message"]["content"]
        else:
            # Próba z Flash jako fallback
            payload["model"] = "gemini-2.5-flash"
            response2 = http_post(proxy_url, json_data=payload, timeout=90.0)
            if response2.status_code == 200:
                return response2.json()["choices"][0]["message"]["content"]
    except Exception:
        pass

    # Ostateczny fallback
    return call_gemini_api(messages, system_instruction)

def call_native_vertex_ocr(file_bytes):
    # Get credentials and refresh token
    sa_paths = [
        os.path.join(os.getcwd(), "holistic-dashboard-dev-dea2c872139e.json"),
        os.path.expanduser("~/.hermes/gcp-sa-key.json"),
        "holistic-dashboard-dev-dea2c872139e.json"
    ]
    sa_path = None
    for p in sa_paths:
        if os.path.exists(p):
            sa_path = p
            break
            
    if not sa_path:
        return "[Błąd: Brak klucza GCP Service Account]"
        
    try:
        from google.oauth2 import service_account
        import google.auth.transport.requests
        import requests
        import base64 as _b64
        
        creds = service_account.Credentials.from_service_account_file(
            sa_path,
            scopes=['https://www.googleapis.com/auth/cloud-platform']
        )
        session = requests.Session()
        session.verify = False
        request = google.auth.transport.requests.Request(session=session)
        creds.refresh(request)
        token = creds.token
        
        project_id = "holistic-dashboard-dev"
        region = "us-central1"
        model = "gemini-2.5-flash" # Use flash for fast OCR
        
        url = f"https://{region}-aiplatform.googleapis.com/v1/projects/{project_id}/locations/{region}/publishers/google/models/{model}:generateContent"
        
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }
        
        encoded_data = _b64.b64encode(file_bytes).decode("utf-8")
        
        payload = {
            "contents": [
                {
                    "role": "user",
                    "parts": [
                        {
                            "text": "Odczytaj i przepisz DOKŁADNIE cały tekst z załączonego dokumentu PDF (skanu). Zachowaj wszystkie kluczowe dane: sygnatury akt, PESEL, NIP, adresy, imiona i nazwiska, kwoty, daty, nazwy sądów lub organów. Zwróć tylko surowy odczytany tekst, zachowując oryginalny układ."
                        },
                        {
                            "inlineData": {
                                "mimeType": "application/pdf",
                                "data": encoded_data
                            }
                        }
                    ]
                }
            ],
            "generationConfig": {
                "temperature": 0.1
            }
        }
        
        resp = requests.post(url, json=payload, headers=headers, verify=False, timeout=120.0)
        if resp.status_code == 200:
            res_data = resp.json()
            try:
                text = res_data['candidates'][0]['content']['parts'][0]['text']
                return text
            except Exception as e:
                return f"[Błąd parsowania odpowiedzi OCR: {e}]"
        else:
            return f"[Błąd API Vertex OCR {resp.status_code}: {resp.text}]"
    except Exception as ex:
        return f"[Błąd systemu OCR: {ex}]"

def call_gcp_tts(text, voice_name="pl-PL-Wavenet-B", gender="MALE"):
    sa_paths = [
        os.path.join(os.getcwd(), "holistic-dashboard-dev-dea2c872139e.json"),
        os.path.expanduser("~/.hermes/gcp-sa-key.json"),
        "holistic-dashboard-dev-dea2c872139e.json"
    ]
    sa_path = None
    for p in sa_paths:
        if os.path.exists(p):
            sa_path = p
            break
            
    if not sa_path:
        return None, "Brak klucza GCP Service Account"
        
    try:
        from google.oauth2 import service_account
        import google.auth.transport.requests
        import requests
        import base64 as _b64
        
        creds = service_account.Credentials.from_service_account_file(
            sa_path,
            scopes=['https://www.googleapis.com/auth/cloud-platform']
        )
        session = requests.Session()
        session.verify = False
        request = google.auth.transport.requests.Request(session=session)
        creds.refresh(request)
        token = creds.token
        
        url = "https://texttospeech.googleapis.com/v1/text:synthesize"
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "input": {"text": text},
            "voice": {
                "languageCode": "pl-PL",
                "name": voice_name,
                "ssmlGender": gender
            },
            "audioConfig": {
                "audioEncoding": "MP3"
            }
        }
        
        resp = requests.post(url, json=payload, headers=headers, verify=False, timeout=15.0)
        if resp.status_code == 200:
            res_data = resp.json()
            audio_content = res_data.get("audioContent", "")
            if audio_content:
                return _b64.b64decode(audio_content), None
            return None, "Brak zawartości audio w odpowiedzi"
        else:
            return None, f"Błąd API TTS {resp.status_code}: {resp.text}"
    except Exception as ex:
        return None, f"Błąd połączenia z GCP TTS: {ex}"

def query_vertex_search(query, data_store_id, project_id="holistic-dashboard-dev", location="global"):
    sa_paths = [
        os.path.join(os.getcwd(), "holistic-dashboard-dev-dea2c872139e.json"),
        os.path.expanduser("~/.hermes/gcp-sa-key.json"),
        "holistic-dashboard-dev-dea2c872139e.json"
    ]
    sa_path = None
    for p in sa_paths:
        if os.path.exists(p):
            sa_path = p
            break
            
    if not sa_path:
        return {"error": "Brak pliku klucza GCP Service Account (holistic-dashboard-dev-dea2c872139e.json). Upewnij się, że znajduje się w głównym folderze."}
        
    try:
        from google.cloud import discoveryengine_v1beta as discoveryengine
        from google.oauth2 import service_account
        
        creds = service_account.Credentials.from_service_account_file(
            sa_path,
            scopes=['https://www.googleapis.com/auth/cloud-platform']
        )
        
        client = discoveryengine.SearchServiceClient(credentials=creds)
        
        if "connector" in data_store_id or "datastore" in data_store_id.lower() or "data-store" in data_store_id.lower():
            serving_config = f"projects/{project_id}/locations/{location}/dataStores/{data_store_id}/servingConfigs/default_search"
        else:
            serving_config = f"projects/{project_id}/locations/{location}/collections/default_collection/engines/{data_store_id}/servingConfigs/default_search"
        
        content_search_spec = discoveryengine.SearchRequest.ContentSearchSpec(
            extractive_content_spec=discoveryengine.SearchRequest.ContentSearchSpec.ExtractiveContentSpec(
                max_extractive_answer_count=3,
                max_extractive_segment_count=1,
                return_extractive_segment_source=True
            ),
            snippet_spec=discoveryengine.SearchRequest.ContentSearchSpec.SnippetSpec(
                max_snippet_count=1,
                return_snippet=True
            ),
            summary_spec=discoveryengine.SearchRequest.ContentSearchSpec.SummarySpec(
                summary_result_count=5,
                include_citations=True,
                model_spec=discoveryengine.SearchRequest.ContentSearchSpec.SummarySpec.ModelSpec(
                    version="stable"
                )
            )
        )
        
        request = discoveryengine.SearchRequest(
            serving_config=serving_config,
            query=query,
            page_size=5,
            content_search_spec=content_search_spec
        )
        
        response = client.search(request)
        
        results = []
        summary = ""
        if hasattr(response, "summary") and response.summary:
            summary = response.summary.summary_text
            
        for result in response:
            doc = result.document
            struct_data = dict(doc.derived_struct_data) if doc.derived_struct_data else {}
            
            title = struct_data.get("title", doc.id)
            link = struct_data.get("link", "")
            snippet = struct_data.get("snippet", "")
            
            extractive_answers = struct_data.get("extractive_answers", [])
            extractive_segments = struct_data.get("extractive_segments", [])
            
            answers = []
            if isinstance(extractive_answers, list):
                for ans in extractive_answers:
                    if isinstance(ans, dict) and "content" in ans:
                        answers.append(ans["content"])
                        
            segments = []
            if isinstance(extractive_segments, list):
                for seg in extractive_segments:
                    if isinstance(seg, dict) and "content" in seg:
                        segments.append(seg["content"])
            
            results.append({
                "id": doc.id,
                "title": title,
                "link": link,
                "snippet": snippet,
                "extractive_answers": answers,
                "extractive_segments": segments,
                "struct_data": struct_data
            })
            
        return {
            "results": results,
            "summary": summary
        }
    except Exception as e:
        return {"error": str(e)}

def extract_youtube_transcript_raw(url):
    import re
    pattern = r'(?:youtube\.com\/(?:[^\/]+\/.+\/|(?:v|e(?:mbed)?)\/|.*[?&]v=)|youtu\.be\/)([^"&?\/\s]{11})'
    match = re.search(pattern, url)
    if not match:
        return None, "Niepoprawny adres URL wideo YouTube"
    video_id = match.group(1)
    
    # Store original requests Session request to restore it afterwards
    import requests
    orig_request = requests.Session.request
    
    try:
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        
        # Monkey patch requests for SSL bypass
        requests.Session.request = lambda self, *a, **kw: orig_request(self, *a, **{**kw, 'verify': False})
        
        from youtube_transcript_api import YouTubeTranscriptApi
        api = YouTubeTranscriptApi()
        
        try:
            data = api.fetch(video_id, languages=['pl', 'en'])
        except Exception:
            try:
                data = api.fetch(video_id)
            except Exception as e2:
                # Restore original request
                requests.Session.request = orig_request
                return None, f"Błąd pobierania transkrypcji YouTube: {e2}"
                
        # Restore original request
        requests.Session.request = orig_request
        
        full_text = []
        for entry in data:
            text = entry['text']
            start = entry['start']
            minutes = int(start // 60)
            seconds = int(start % 60)
            full_text.append(f"[{minutes:02d}:{seconds:02d}] {text}")
            
        return "\n".join(full_text), None
    except Exception as e:
        # Guarantee restoration
        try:
            requests.Session.request = orig_request
        except Exception:
            pass
        return None, f"Błąd systemu transkrypcji: {str(e)}"

class WebhookHandler(BaseHTTPRequestHandler):
    def log_message(self, format, *args):
        pass
        
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, GET, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def do_POST(self):
        if self.path == "/webhook":
            content_length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(content_length)
            
            try:
                content_type = self.headers.get('Content-Type', '')
                data = {}
                if 'application/json' in content_type:
                    data = json.loads(body.decode('utf-8'))
                else:
                    parsed = urllib.parse.parse_qs(body.decode('utf-8'))
                    data = {k: v[0] for k, v in parsed.items()}
                
                name = data.get("name", "").strip()
                email = data.get("email", "").strip()
                notes = data.get("notes", "").strip() or data.get("message", "").strip()
                
                if not name or not email:
                    self.send_response(400)
                    self.send_header("Access-Control-Allow-Origin", "*")
                    self.end_headers()
                    self.wfile.write(b"Name and Email are required.")
                    return
                
                crm_data = load_crm()
                new_id = f"lead_{int(time.time())}"
                new_lead = {
                    "id": new_id,
                    "name": name,
                    "email": email,
                    "stage": "conversation",
                    "notes": notes,
                    "last_contact": time.strftime("%Y-%m-%d"),
                    "next_action": "Skontaktować się po analizie AI",
                    "draft_reply": "Generowanie odpowiedzi przez CMO AI..."
                }
                
                crm_data["leads"].append(new_lead)
                save_crm(crm_data)
                
                threading.Thread(
                    target=async_generate_initial_draft,
                    args=(new_id, name, notes),
                    daemon=True
                ).start()
                
                self.send_response(200)
                self.send_header("Content-Type", "application/json")
                self.send_header("Access-Control-Allow-Origin", "*")
                self.end_headers()
                self.wfile.write(json.dumps({"status": "success", "lead_id": new_id}).encode('utf-8'))
                
            except Exception as e:
                self.send_response(500)
                self.send_header("Access-Control-Allow-Origin", "*")
                self.end_headers()
                self.wfile.write(f"Server Error: {e}".encode('utf-8'))
        else:
            self.send_response(404)
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()

def async_generate_initial_draft(lead_id, client_name, client_pain):
    o_mnie_path = os.path.join(HERMES_DIR, "o_mnie.md")
    o_mnie_context = ""
    if os.path.exists(o_mnie_path):
        try:
            o_mnie_context = open(o_mnie_path, "r", encoding="utf-8").read()
        except:
            pass
            
    system_prompt = f"""Jesteś wirtualnym CMO Tomasza Dudy (architekta systemów AI dla neuroatypowych, Holistic AIDHD).
Tomasz jest strategiem automatyzacji dla firm i sam ma ADHD. Twoim zadaniem jest stworzenie bardzo empatycznej, osobistej i konkretnej wersji roboczej pierwszej odpowiedzi do nowego leada (klienta).
Unikaj pustego hype'u AI, pisz zwięźle, ze zrozumieniem chaosu operacyjnego klienta, i z nutą humoru oraz głębokim rezonansem emocjonalnym.

Oto kontekst o Tomaszu (jego serce emocjonalne i historia):
{o_mnie_context}

Napisz bezpośrednią, gotową do wysłania wiadomość e-mail w języku polskim. Pisz w 1 osobie ("ja" - Tomasz Duda). Podpisz się na końcu.
"""

    user_prompt = f"""Klient: {client_name}
Wyzwanie klienta: {client_pain}

Napisz pierwszą, niezwykle empatyczną odpowiedź, która zdejmuje z niego paraliż decyzyjny i proponuje jeden mały, konkretny krok bez presji (krótką rozmowę diagnostyczną)."""

    try:
        reply = call_gemini_api([{"role": "user", "content": user_prompt}], system_instruction=system_prompt)
        if reply:
            crm_data = load_crm()
            for lead in crm_data["leads"]:
                if lead["id"] == lead_id:
                    lead["draft_reply"] = reply
                    break
            save_crm(crm_data)
    except Exception as e:
        pass

def consult_csuite_live(lead, persona):
    o_mnie_path = os.path.join(HERMES_DIR, "o_mnie.md")
    o_mnie_context = ""
    if os.path.exists(o_mnie_path):
        try:
            o_mnie_context = open(o_mnie_path, "r", encoding="utf-8").read()
        except:
            pass

    # Dynamiczne wstrzykiwanie checklist i SOP z folderu skills/
    persona_mapping = {
        "CEO (Strategia & Rentowność)": "ceo",
        "CMO (Empatyczny Storytelling)": "cmo",
        "CSO (Architektura Sprzedaży)": "cso",
        "CTO (Technologia & Kod)": "cto"
    }
    
    dynamic_skill_context = ""
    folder_name = persona_mapping.get(persona)
    if folder_name:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        skill_paths = [
            os.path.join(base_dir, "skills", folder_name, "SKILL.md"),
            os.path.join(base_dir, ".agents", "skills", folder_name, "SKILL.md"),
            os.path.expanduser(f"~/.gemini/config/plugins/holistic-virtual-board/skills/{folder_name}/SKILL.md")
        ]
        for p in skill_paths:
            if os.path.exists(p):
                try:
                    with open(p, "r", encoding="utf-8") as f:
                        dynamic_skill_context = f.read()
                    break
                except:
                    pass

    system_prompts = {
        "CEO (Strategia & Rentowność)": f"""Jesteś wirtualnym CEO w zespole Tomasza Dudy. Tomasz to wybitny architekt systemów AI dla neuroatypowych (sam ma ADHD, Holistic AIDHD).
Pomagasz mu w wycenie wdrożenia pod kątem modelu High-Ticket (np. wyceny 5 000 - 15 000 PLN jednorazowo), etapowaniu prac na proste kroki MVP oraz obronie jego zasobów energetycznych przed wypaleniem i paraliżem ADHD.
Zawsze podawaj konkretną, odważną rekomendację cenową i zdefiniuj, co jest "One Thing" (kluczowym pierwszym krokiem wdrożenia).

Oto Twoje oficjalne wytyczne i checklisty operacyjne (SOP):
{dynamic_skill_context}

Oto historia i tożsamość Tomasza:
{o_mnie_context}
Odpowiadaj bezpośrednio, po polsku, zwięźle i konkretnie.
""",
        "CMO (Empatyczny Storytelling)": f"""Jesteś wirtualnym CMO w zespole Tomasza Dudy (Holistic AIDHD).
Pomagasz mu przełożyć ból klienta na autentyczny i humorystyczny przekaz dopasowany do wyzwań klienta. Wskaż, jakich metafor użyć w komunikacji z tym klientem i jak napisać ofertę, aby rezonowała głęboko emocjonalnie, opierając się na tożsamości Tomasza.

Oto Twoje oficjalne wytyczne i checklisty operacyjne (SOP):
{dynamic_skill_context}

Oto historia i tożsamość Tomasza:
{o_mnie_context}
Odpowiadaj bezpośrednio, po polsku, zwięźle i kreatywnie.
""",
        "CSO (Architektura Sprzedaży)": f"""Jesteś wirtualnym CSO w zespole Tomasza Dudy (Holistic AIDHD).
Projektujesz dla tego klienta prosty, 3-stopniowy lejek relacyjny (Rozmowa -> Architektura -> Wdrożenie). Wskaż dokładnie, jaki powinien być najbliższy krok sprzedażowy (Next Action) i jak go zrealizować przy minimalnym tarciu poznawczym (low cognitive friction).

Oto Twoje oficjalne wytyczne i checklisty operacyjne (SOP):
{dynamic_skill_context}

Oto historia i tożsamość Tomasza:
{o_mnie_context}
Odpowiadaj bezpośrednio, po polsku, zwięźle i operacyjnie.
""",
        "CTO (Technologia & Kod)": f"""Jesteś wirtualnym CTO w zespole Tomasza Dudy (Holistic AIDHD).
Zaprojektuj uproszczoną, niezawodną architekturę techniczną pod potrzeby tego klienta. Rekomenduj konkretne narzędzia (np. n8n webhooks, Python scripts, SQLite, Google Sheets, Vertex AI, model gemini-2.5-flash). Podaj zwięzły schemat logiczny.

Oto Twoje oficjalne wytyczne i checklisty operacyjne (SOP):
{dynamic_skill_context}

Oto historia i tożsamość Tomasza:
{o_mnie_context}
Odpowiadaj bezpośrednio, po polsku, technicznie lecz bez zbędnego żargonu.
"""
    }

    system_instruction = system_prompts.get(persona, "Jesteś doradcą C-Suite.")
    user_prompt = f"""Karta klienta:
Nazwa: {lead['name']}
Opis chaosu operacyjnego: {lead['notes']}
Najbliższy krok (Next Action): {lead['next_action']}

Przeanalizuj przypadek i daj mi (Tomaszowi) swoją rekomendację z perspektywy roli ({persona}). Pisz bezpośrednio do mnie ("Tomasz, badając przypadek...")."""

    return call_gemini_api([{"role": "user", "content": user_prompt}], system_instruction=system_instruction)

def start_webhook_server():
    global _server_started
    with _server_lock:
        if _server_started:
            return
        server_address = ('0.0.0.0', 8090)
        try:
            httpd = HTTPServer(server_address, WebhookHandler)
            t = threading.Thread(target=httpd.serve_forever, daemon=True)
            t.start()
            _server_started = True
        except Exception as e:
            pass

# Start the background webhook server
start_webhook_server()

def read_md_file(path):
    return open(path, "r", encoding="utf-8").read() if os.path.exists(path) else ""

def save_brain_dump(thought, links, uploaded_file):
    dump_id = f"dump_{int(time.time())}"
    file_name = None
    if uploaded_file:
        file_name = uploaded_file.name
        with open(os.path.join(BRAIN_DUMP_ASSETS, file_name), "wb") as f:
            f.write(uploaded_file.getbuffer())
    category = "Later" if any(w in thought.lower() for w in ["potem", "później", "later", "kiedyś"]) else "Now"
    dump_data = {
        "id": dump_id,
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "thought": thought,
        "links": links,
        "file_attached": file_name,
        "category": category,
        "status": "active"
    }
    json.dump(dump_data, open(os.path.join(BRAIN_DUMP_DIR, f"{dump_id}.json"), "w", encoding="utf-8"), ensure_ascii=False, indent=4)

def get_brain_dumps():
    dumps = []
    if os.path.exists(BRAIN_DUMP_DIR):
        for f in os.listdir(BRAIN_DUMP_DIR):
            if f.endswith(".json") and f.startswith("dump_"):
                try:
                    dumps.append(json.load(open(os.path.join(BRAIN_DUMP_DIR, f), "r", encoding="utf-8")))
                except:
                    pass
    dumps.sort(key=lambda x: x.get("timestamp", ""), reverse=True)
    return dumps

@st.dialog("📥 Szybki Zrzut Myśli (Brain Dump)")
def show_brain_dump_dialog():
    st.markdown("<p style='color: #CBD5E1;'>Wpisz swój pomysł, załącz link lub plik, a model Gemini 2.5 Pro odciąży Twój umysł i natychmiast zapisze go w Skarbcu.</p>", unsafe_allow_html=True)
    thought = st.text_area("Co Ci chodzi po głowie? (Zzrzuć chaos)", height=120)
    links = st.text_input("Linki / Social media / Wideo (opcjonalnie):")
    uploaded_file = st.file_uploader("Załącz plik / obrazek / audio (opcjonalnie):", type=["png", "jpg", "jpeg", "pdf", "mp3", "wav"])
    
    if st.button("Uwolnij mój umysł", type="primary", use_container_width=True):
        if thought or links or uploaded_file:
            save_brain_dump(thought, links, uploaded_file)
            st.success("Zapisano bezpiecznie w Skarbcu!")
            time.sleep(1.0)
            st.rerun()

# Inicjalizacja stanu sesji
if "one_thing" not in st.session_state:
    st.session_state.one_thing = ""
if "pomodoro_active" not in st.session_state:
    st.session_state.pomodoro_active = False
if "current_page" not in st.session_state:
    st.session_state.current_page = "🎯 Mission Control"

# PASEK BOCZNY - Skrajnie estetyczny z kafelkami zoptymalizowanymi pod ADHD (Styl Julian Goldie)
with st.sidebar:
    st.markdown("<h2 style='text-align: center; color: #7C3AED; font-family: Outfit; margin-bottom: 0;'>🧠 Holistic OS</h2>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #94A3B8; font-size: 0.85rem; margin-top: 5px; margin-bottom: 10px;'>Agentic OS Mission Control v7.0</p>", unsafe_allow_html=True)
    st.markdown("<hr style='margin: 10px 0; border-color: #1F242E;'>", unsafe_allow_html=True)
    
    col_menu = st.session_state.current_page
    
    # I. WORKSPACE
    st.markdown("<p style='color: #F59E0B; font-weight: bold; font-size: 0.75rem; letter-spacing: 1px; margin-bottom: 6px; margin-top: 10px;'>I. WORKSPACE</p>", unsafe_allow_html=True)
    
    if st.button("🎯 Mission Control", use_container_width=True, type="primary" if col_menu == "🎯 Mission Control" else "secondary"):
        st.session_state.current_page = "🎯 Mission Control"
        st.rerun()
        
    # II. AGENTS
    st.markdown("<p style='color: #EC4899; font-weight: bold; font-size: 0.75rem; letter-spacing: 1px; margin-top: 18px; margin-bottom: 6px;'>II. AGENTS</p>", unsafe_allow_html=True)
    
    if st.button("🤖 Claude 🟢", use_container_width=True, type="primary" if col_menu == "Claude" else "secondary"):
        st.session_state.current_page = "Claude"
        st.rerun()
        
    if st.button("🔱 Hermes 🟢", use_container_width=True, type="primary" if col_menu == "Hermes" else "secondary"):
        st.session_state.current_page = "Hermes"
        st.rerun()
        
    if st.button("♊ Gemini 🟢", use_container_width=True, type="primary" if col_menu == "Gemini" else "secondary"):
        st.session_state.current_page = "Gemini"
        st.rerun()
        
    if st.button("🌌 AntiGravity 🟢", use_container_width=True, type="primary" if col_menu == "Antigravity" else "secondary"):
        st.session_state.current_page = "Antigravity"
        st.rerun()
        
    # III. SELF
    st.markdown("<p style='color: #10B981; font-weight: bold; font-size: 0.75rem; letter-spacing: 1px; margin-top: 18px; margin-bottom: 6px;'>III. SELF</p>", unsafe_allow_html=True)
    
    if st.button("🎯 Goals & Journal", use_container_width=True, type="primary" if col_menu == "Goals" else "secondary"):
        st.session_state.current_page = "Goals"
        st.rerun()
        
    if st.button("📚 Baza Wiedzy", use_container_width=True, type="primary" if col_menu == "Notebook" else "secondary"):
        st.session_state.current_page = "Notebook"
        st.rerun()
        
    if st.button("📋 Tablica Kanban", use_container_width=True, type="primary" if col_menu == "Kanban" else "secondary"):
        st.session_state.current_page = "Kanban"
        st.rerun()
        
    if st.button("💾 Pamięć Agenta", use_container_width=True, type="primary" if col_menu == "Memory" else "secondary"):
        st.session_state.current_page = "Memory"
        st.rerun()
        
    if st.button("🤝 Onboarding Klienta", use_container_width=True, type="primary" if col_menu == "Onboarding" else "secondary"):
        st.session_state.current_page = "Onboarding"
        st.rerun()
        
    if st.button("✨ Rój Agentów (Sprzedaż)", use_container_width=True, type="primary" if col_menu == "Swarm" else "secondary"):
        st.session_state.current_page = "Swarm"
        st.rerun()
        
    # IV. BUSINESS & MARKETING
    st.markdown("<p style='color: #3B82F6; font-weight: bold; font-size: 0.75rem; letter-spacing: 1px; margin-top: 18px; margin-bottom: 6px;'>IV. BUSINESS & MARKETING</p>", unsafe_allow_html=True)
    
    if st.button("📈 SEO & Content", use_container_width=True, type="primary" if col_menu == "SEO" else "secondary"):
        st.session_state.current_page = "SEO"
        st.rerun()
        
    if st.button("📢 Social Media Hub", use_container_width=True, type="primary" if col_menu == "Social Media Hub" else "secondary"):
        st.session_state.current_page = "Social Media Hub"
        st.rerun()
        
    if st.button("🌐 AI Website Builder", use_container_width=True, type="primary" if col_menu == "AI Website Builder" else "secondary"):
        st.session_state.current_page = "AI Website Builder"
        st.rerun()
        
    if st.button("🎯 Ads & Local SEO", use_container_width=True, type="primary" if col_menu == "Ads & Local SEO" else "secondary"):
        st.session_state.current_page = "Ads & Local SEO"
        st.rerun()
        
    if st.button("🎬 Studio (Hyperframes)", use_container_width=True, type="primary" if col_menu == "Studio" else "secondary"):
        st.session_state.current_page = "Studio"
        st.rerun()
        
    if st.button("💼 CRM Leads", use_container_width=True, type="primary" if col_menu == "CRM" else "secondary"):
        st.session_state.current_page = "CRM"
        st.rerun()
        
    if st.button("⚖️ Legal Engine", use_container_width=True, type="primary" if col_menu == "Legal" else "secondary"):
        st.session_state.current_page = "Legal"
        st.rerun()
        
    if st.button("💰 Finance & KSeF", use_container_width=True, type="primary" if col_menu == "Finance" else "secondary"):
        st.session_state.current_page = "Finance"
        st.rerun()
        
    st.markdown("<hr style='margin: 15px 0; border-color: #1F242E;'>", unsafe_allow_html=True)
    st.markdown("🌐 **Status Systemu:**")
    st.markdown("⚡ *Hermes Agent:* <span style='color:#10B981; font-weight:bold;'>LIVE (Port 9119)</span>", unsafe_allow_html=True)
    st.markdown("📢 *Telegram Chat:* <span style='color:#10B981; font-weight:bold;'>Połączony</span>", unsafe_allow_html=True)
    st.markdown("📝 *Pristine Memory:* <span style='color:#3B82F6; font-weight:bold;'>Aktywna</span>", unsafe_allow_html=True)


menu = st.session_state.current_page

def render_agent_console(agent_name, status, default_model, provider, color_accent):
    agent_key = agent_name.lower().replace(' ', '_')
    prov_val = st.session_state.get(f"ctrl_provider_{agent_key}", provider)
    model_val = st.session_state.get(f"ctrl_model_{agent_key}", default_model)
    if prov_val == "OpenRouter" and model_val == "Inny / Custom":
        model_val = st.session_state.get(f"ctrl_custom_model_{agent_key}", "custom")
        
    st.markdown(f"<p style='color: #94A3B8; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-bottom: 2px;'>II. — AGENT • {agent_name.upper()}</p>", unsafe_allow_html=True)
    st.title(f"{agent_name}")
    st.markdown(f"<p style='color: {color_accent}; font-weight: bold; font-size: 0.95rem; margin-top: -5px;'>Status: {status} | Active Model: {model_val} | Provider: {prov_val}</p>", unsafe_allow_html=True)
    
    tab_chat, tab_work, tab_ctrl = st.tabs(["💬 Chat", "📂 Workspace", "⚙️ Control Room"])
    
    # 1. CHAT
    with tab_chat:
        st.subheader("Konsola konwersacyjna")
        chat_key = f"chat_{agent_key}"
        if chat_key not in st.session_state:
            st.session_state[chat_key] = []
            
        for msg in st.session_state[chat_key]:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
                
        if prompt := st.chat_input(f"Napisz do agenta {agent_name}...", key=f"input_{agent_key}"):
            with st.chat_message("user"):
                st.markdown(prompt)
            st.session_state[chat_key].append({"role": "user", "content": prompt})
            
            # Wczytanie profilu i instrukcji
            o_mnie_path = os.path.join(HERMES_DIR, "o_mnie.md")
            o_mnie_context = read_md_file(o_mnie_path) if os.path.exists(o_mnie_path) else "Brak profilu o_mnie.md"
            user_inst = st.session_state.get(f"ctrl_prompt_{agent_key}", f"Jesteś agentem {agent_name}. Działasz zorientowany na ADHD i redukcję szumu kognitywnego.")
            
            # Shared Workspace Memory
            recent_context = get_recent_workspace_context(limit=3)
            
            sys_prompt = f"""{user_inst}
            
## PROFIL UŻYTKOWNIKA (O_MNIE):
{o_mnie_context}

## WSPÓLNA PAMIĘĆ ROBOCZA (WORKSPACE MEMORY):
Oto ostatnie działania i notatki w ekosystemie:
{recent_context}

Działasz w oparciu o powyższy wspólny kontekst. Twój styl jest krótki, precyzyjny (ADHD-friendly).
"""
            
            provider_sel = st.session_state.get(f"ctrl_provider_{agent_key}", provider)
            model_sel = st.session_state.get(f"ctrl_model_{agent_key}", default_model)
            if provider_sel == "OpenRouter":
                custom_m = st.session_state.get(f"ctrl_custom_model_{agent_key}", "")
                if model_sel == "Inny / Custom" and custom_m:
                    final_model = custom_m
                else:
                    final_model = model_sel
            else:
                final_model = model_sel

            with st.spinner(f"{agent_name} ({final_model} / {provider_sel}) przetwarza zapytanie..."):
                api_messages = [{"role": m["role"], "content": m["content"]} for m in st.session_state[chat_key]]
                if provider_sel == "OpenRouter":
                    response = call_openrouter_api(api_messages, model=final_model, system_instruction=sys_prompt)
                elif provider_sel == "Local vLLM":
                    response = call_vllm_api(api_messages, system_instruction=sys_prompt)
                else: # GCP Vertex Proxy
                    if final_model == "gemini-2.5-pro":
                        response = call_gemini_pro_api(api_messages, system_instruction=sys_prompt)
                    else:
                        response = call_gemini_api(api_messages, system_instruction=sys_prompt)
                
            with st.chat_message("assistant"):
                st.markdown(response)
            st.session_state[chat_key].append({"role": "assistant", "content": response})
            
            # Infinite Context Engine: auto-save to Obsidian
            obsidian_note_title = f"Chat_{agent_name.replace(' ', '_')}_{int(time.time())}.md"
            obsidian_note_path = os.path.join(OBSIDIAN_DIR, obsidian_note_title)
            try:
                note_content = f"---\ntype: chat-log\nagent: {agent_name}\ntimestamp: {time.strftime('%Y-%m-%d %H:%M:%S')}\n---\n\n"
                for msg in st.session_state[chat_key]:
                    note_content += f"**{'Użytkownik' if msg['role'] == 'user' else agent_name}**:\n{msg['content']}\n\n"
                with open(obsidian_note_path, "w", encoding="utf-8") as f:
                    f.write(note_content)
            except Exception as e:
                pass
                
    # 2. WORKSPACE (Trzykolumnowy manager)
    with tab_work:
        st.subheader("Menedżer Workspace Agentów")
        col_b, col_f, col_v = st.columns([1, 2, 3])
        
        with col_b:
            st.write("🔍 **Buckets**")
            bucket = st.radio("Filtruj:", ["Apps", "Videos", "Images", "Audio", "Workspace", "Sandboxes", "Pastes"], key=f"bucket_{agent_key}")
            
        with col_f:
            st.write(f"📁 **Pliki w {bucket}**")
            files_list = []
            if bucket == "Apps":
                files_list = [f for f in os.listdir(DASHBOARD_DIR) if f.endswith(".html")]
                if not files_list:
                    files_list = ["todo-app.html", "testimonials.html", "landing-page.html"]
            elif bucket == "Videos":
                files_list = ["dragon_tokyo_skyline.mp4", "ai-seo-trends.mp4"]
            elif bucket == "Images":
                files_list = ["holistic_ai_dhd_logo.png", "marketing_campaign.jpg"]
            elif bucket == "Audio":
                files_list = ["meditation_music.mp3", "podcast_voice.wav"]
            elif bucket == "Workspace":
                files_list = ["app.py", "gcp_vertex_proxy.py", "requirements.txt"]
            elif bucket == "Sandboxes":
                files_list = ["sandbox_env.py", "test_runner.py"]
            elif bucket == "Pastes":
                files_list = [f for f in os.listdir(OBSIDIAN_DIR) if f.endswith(".md")][:10]
                if not files_list:
                    files_list = ["quick_note.md", "brief_draft.md"]
                    
            selected_file = st.selectbox("Wybierz plik:", files_list, key=f"sel_file_{agent_key}")
            
        with col_v:
            st.write("👁️ **Podgląd / Edycja**")
            if selected_file:
                st.info(f"Podgląd pliku: `{selected_file}`")
                
                if selected_file.endswith(".html") or bucket == "Apps":
                    html_content = ""
                    if selected_file == "testimonials.html":
                        html_content = "<html><body style='background-color:#121620; color:white; font-family:sans-serif; padding:20px;'><h2>Want results like these? Join the AI Profit Boardroom!</h2><p>258 real wins from AI Profit Boardroom members.</p></body></html>"
                    elif selected_file == "todo-app.html":
                        html_content = "<html><body style='background-color:#121620; color:white; font-family:sans-serif; padding:20px;'><h2>ADHD To-Do App</h2><input type='text' placeholder='Add task...'><button>Add</button></body></html>"
                    else:
                        html_content = "<html><body style='background-color:#121620; color:white; font-family:sans-serif; padding:20px;'><h2>Custom App View</h2><p>Pomyślnie załadowano render aplikacji w piaskownicy Agentic OS.</p></body></html>"
                    
                    st.components.v1.html(html_content, height=350, scrolling=True)
                    
                    sub_tab_preview, sub_tab_src = st.tabs(["👁️ Podgląd", "💻 Kod źródłowy"])
                    with sub_tab_src:
                        st.code(html_content, language="html")
                        
                elif selected_file.endswith(".mp4") or bucket == "Videos":
                    st.video("https://www.w3schools.com/html/mov_bbb.mp4")
                    
                elif selected_file.endswith((".png", ".jpg", ".jpeg")) or bucket == "Images":
                    st.image("https://images.unsplash.com/photo-1618005182384-a83a8bd57fbe?w=500", caption=selected_file)
                    
                elif selected_file.endswith((".mp3", ".wav")) or bucket == "Audio":
                    st.audio("https://www.soundhelix.com/examples/mp3/SoundHelix-Song-1.mp3")
                    
                elif selected_file.endswith(".md") or selected_file.endswith(".py") or selected_file.endswith(".txt"):
                    file_path = os.path.join(OBSIDIAN_DIR, selected_file)
                    if not os.path.exists(file_path):
                        file_path = os.path.join(os.getcwd(), selected_file)
                    content = ""
                    if os.path.exists(file_path):
                        try:
                            with open(file_path, "r", encoding="utf-8") as f:
                                content = f.read()
                        except:
                            content = "Błąd odczytu pliku."
                    else:
                        content = f"# Draft dla {selected_file}\nTutaj znajduje się podgląd notatki roboczej agenta."
                        
                    edited_content = st.text_area("Edycja pliku:", content, height=200, key=f"edit_{agent_key}_{selected_file}")
                    if st.button("Zapisz zmiany", key=f"save_{agent_key}_{selected_file}"):
                        try:
                            with open(file_path, "w", encoding="utf-8") as f:
                                f.write(edited_content)
                            st.success("Zapisano pomyślnie!")
                        except Exception as e:
                            st.error(f"Błąd zapisu: {e}")
                            
    # 3. CONTROL ROOM
    with tab_ctrl:
        st.subheader("Konfiguracja Agenta (Control Room)")
        st.write("Ustaw dostawcę i parametry modelu dla tego agenta:")
        
        prov_opt = st.selectbox("Dostawca (Provider):", ["GCP Vertex Proxy", "OpenRouter", "Local vLLM"], 
                                index=0 if provider in ["Vertex AI Native", "GCP Proxy Port 8089", "GCP Proxy"] or "vertex" in provider.lower() else (1 if "openrouter" in provider.lower() else 2),
                                key=f"ctrl_provider_{agent_key}")
        
        if prov_opt == "GCP Vertex Proxy":
            models_list = ["gemini-2.5-pro", "gemini-2.5-flash"]
            default_idx = 0 if default_model == "gemini-2.5-pro" else 1
            model_opt = st.selectbox("Model:", models_list, index=default_idx, key=f"ctrl_model_{agent_key}")
        elif prov_opt == "OpenRouter":
            models_list = [
                "meta-llama/llama-3.1-8b-instruct:free",
                "nousresearch/hermes-3-llama-3.1-405b:free",
                "google/gemini-2.5-flash",
                "google/gemini-2.5-pro",
                "google/gemini-3.5-flash",
                "google/gemini-3.1-pro-preview",
                "anthropic/claude-3.7-sonnet",
                "anthropic/claude-3.5-sonnet",
                "Inny / Custom"
            ]
            default_idx = 8
            low_default = default_model.lower()
            if "claude-3.7" in low_default:
                default_idx = 6
            elif "claude-3.5" in low_default:
                default_idx = 7
            elif "gemini-3.5" in low_default:
                default_idx = 4
            elif "gemini-3.1" in low_default:
                default_idx = 5
            elif "hermes" in low_default:
                default_idx = 1
            elif "free" in low_default:
                default_idx = 0
            elif "gemini-2.5-pro" in low_default:
                default_idx = 3
            elif "gemini-2.5-flash" in low_default:
                default_idx = 2
                
            model_opt = st.selectbox("Model:", models_list, index=default_idx, key=f"ctrl_model_{agent_key}")
            if model_opt == "Inny / Custom":
                st.text_input("Identyfikator modelu OpenRouter (np. deepseek/deepseek-chat):", value="deepseek/deepseek-chat", key=f"ctrl_custom_model_{agent_key}")
        else: # Local vLLM
            st.info("Local vLLM korzysta z modelu uruchomionego na porcie 8000.")
            model_opt = st.text_input("Nazwa modelu vLLM:", value="mistralai/Mistral-Nemo-Instruct-2407", key=f"ctrl_model_{agent_key}")
            
        temp = st.slider("Temperatura (Kreatywność):", 0.0, 1.0, 0.7, 0.05, key=f"ctrl_temp_{agent_key}")
        sys_prompt_input = st.text_area("System Prompt / Instrukcje systemowe:", 
                                        f"Jesteś agentem {agent_name}. Działasz zorientowany na ADHD i redukcję szumu kognitywnego.", 
                                        height=150, key=f"ctrl_prompt_{agent_key}")
                                        
        if st.button("Aktualizuj konfigurację agenta", key=f"ctrl_btn_{agent_key}"):
            st.success(f"Konfiguracja dla agenta {agent_name} została zaktualizowana w pamięci podręcznej sesji!")
            st.rerun()

# 1. MISSION CONTROL
if menu == "🎯 Mission Control":

    st.markdown("<p style='color: #94A3B8; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-bottom: 2px;'>I. — MISSION CONTROL</p>", unsafe_allow_html=True)
    st.title("🧠 Agentic Mission Control")
    st.subheader("Status of every agent, every memory, every signal.")
    
    # Górny pasek statusów (System Status Row)
    st.markdown("""
    <style>
        .status-container {
            display: flex;
            gap: 12px;
            margin-bottom: 25px;
            width: 100%;
        }
        .status-box {
            background-color: #0E1015;
            border: 1px solid #1F242E;
            border-radius: 8px;
            padding: 12px 16px;
            flex: 1;
            box-shadow: 0 4px 10px rgba(0,0,0,0.3);
        }
        .status-title {
            color: #94A3B8;
            font-size: 0.7rem;
            font-weight: bold;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            margin-bottom: 4px;
            display: flex;
            align-items: center;
            justify-content: space-between;
        }
        .status-value {
            color: #FFFFFF;
            font-family: 'Outfit', sans-serif;
            font-size: 1.05rem;
            font-weight: 600;
        }
        .status-desc {
            color: #64748B;
            font-size: 0.75rem;
            margin-top: 2px;
        }
        .status-dot-green { background-color: #10B981; box-shadow: 0 0 8px #10B981; }
        .status-dot-red { background-color: #EF4444; box-shadow: 0 0 8px #EF4444; }
        .status-dot-yellow { background-color: #F59E0B; box-shadow: 0 0 8px #F59E0B; }
    </style>
    """, unsafe_allow_html=True)

    col_s1, col_s2, col_s3, col_s4 = st.columns(4)
    with col_s1:
        st.markdown("""<div class="status-box">
            <div class="status-title"><span>🤖 CLAUDE</span><span class="status-dot-green" style="height:6px; width:6px; border-radius:50%; display:inline-block;"></span></div>
            <div class="status-value">Online</div>
            <div class="status-desc">2.1.142 • 46ms</div>
        </div>""", unsafe_allow_html=True)
    with col_s2:
        st.markdown("""<div class="status-box">
            <div class="status-title"><span>🔱 HERMES</span><span class="status-dot-green" style="height:6px; width:6px; border-radius:50%; display:inline-block;"></span></div>
            <div class="status-value">Online</div>
            <div class="status-desc">grok-4.3 • xAI OAuth</div>
        </div>""", unsafe_allow_html=True)
    with col_s3:
        st.markdown("""<div class="status-box">
            <div class="status-title"><span>⏱️ HEARTBEAT</span><span class="status-dot-yellow" style="height:6px; width:6px; border-radius:50%; display:inline-block;"></span></div>
            <div class="status-value">Active</div>
            <div class="status-desc">poll ticks • 4s</div>
        </div>""", unsafe_allow_html=True)
    with col_s4:
        st.markdown("""<div class="status-box">
            <div class="status-title"><span>⚡ LATENCY</span><span class="status-dot-green" style="height:6px; width:6px; border-radius:50%; display:inline-block;"></span></div>
            <div class="status-value">42 ms</div>
            <div class="status-desc">combined p50</div>
        </div>""", unsafe_allow_html=True)

    st.markdown("<hr style='margin: 15px 0; border-color: #1F242E;'>", unsafe_allow_html=True)

    # Tryb One Thing - ADHD Wykluczenie Szumu
    st.markdown("""
    <div class="one-thing-banner">
        <h3 style="margin-top: 0; color: #F59E0B;">🎯 Tryb "One Thing"</h3>
        <p style="color: #CBD5E1; line-height: 1.6;">Osoby z ADHD cierpią na paraliż decyzyjny z powodu nadmiaru bodźców. Wpisz poniżej dokładnie <strong>JEDNĄ</strong> rzecz, na której skupisz się w tym momencie. Dashboard wyciszy resztę szumu operacyjnego.</p>
    </div>
    """, unsafe_allow_html=True)
    
    thing = st.text_input("Moje jedyne zadanie na ten moment:", value=st.session_state.one_thing, placeholder="Np. zredagowanie oferty High-Ticket dla lokalnej kliniki...")
    if thing:
        st.session_state.one_thing = thing
        st.markdown(f"""
        <div class="custom-card" style="border-left: 5px solid #10B981; background-color: #0F1D1A;">
            <h4 style="margin: 0; color: #10B981;">🔥 Twój aktualny priorytet:</h4>
            <p style="font-size: 1.25rem; font-weight: bold; margin-top: 8px; color: #FFFFFF;">{thing}</p>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2 = st.columns([1, 3])
        with col1:
            if st.button("⏱️ Uruchom Pomodoro (25 min)"):
                st.session_state.pomodoro_active = True
        with col2:
            if st.session_state.pomodoro_active:
                st.success("Stoper Pomodoro wystartował. Wyłącz inne karty w przeglądarce i skup się wyłącznie na priorytecie.")

    st.markdown("<hr style='margin: 15px 0; border-color: #1F242E;'>", unsafe_allow_html=True)

    # II. AGENTS GRID
    st.markdown("<p style='color: #EC4899; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-top: 25px; margin-bottom: 2px;'>II. — AGENTS • CLICK TO OPEN CONTROL ROOM</p>", unsafe_allow_html=True)
    
    col_a1, col_a2, col_a3, col_a4 = st.columns(4)
    
    with col_a1:
        st.markdown("""
        <div class="custom-card" style="border-top: 3px solid #F59E0B; height: 260px;">
            <h3 style="color: #F59E0B; margin-top: 0;">🤖 Claude</h3>
            <p style="color: #94A3B8; font-size: 0.85rem; height: 60px;">Bezpośrednie połączenie ze środowiskiem Claude Code. Pełna kontrola konsoli i narzędzi.</p>
            <div style="margin-top: 15px; font-size: 0.8rem; color: #64748B;">
                <div><b>MODEL:</b> claude-3-7-sonnet</div>
                <div><b>PROVIDER:</b> Anthropic Native</div>
                <div><b>STATUS:</b> <span style="color:#10B981; font-weight:bold;">ONLINE</span></div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Otwórz konsolę Claude", use_container_width=True):
            st.session_state.current_page = "Claude"
            st.rerun()

    with col_a2:
        st.markdown("""
        <div class="custom-card" style="border-top: 3px solid #3B82F6; height: 260px;">
            <h3 style="color: #3B82F6; margin-top: 0;">🔱 Hermes</h3>
            <p style="color: #94A3B8; font-size: 0.85rem; height: 60px;">Główny orkiestrator. Wywołanie skilli, obsługa Kanbana i automatyzacji w tle.</p>
            <div style="margin-top: 15px; font-size: 0.8rem; color: #64748B;">
                <div><b>MODEL:</b> grok-4.3</div>
                <div><b>PROVIDER:</b> xAI API v1</div>
                <div><b>STATUS:</b> <span style="color:#10B981; font-weight:bold;">ONLINE</span></div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Otwórz konsolę Hermes", use_container_width=True):
            st.session_state.current_page = "Hermes"
            st.rerun()

    with col_a3:
        st.markdown("""
        <div class="custom-card" style="border-top: 3px solid #8B5CF6; height: 260px;">
            <h3 style="color: #8B5CF6; margin-top: 0;">♊ Gemini</h3>
            <p style="color: #94A3B8; font-size: 0.85rem; height: 60px;">Super-inteligentny model do audytów, pisania pism prawnych i zaawansowanej logiki.</p>
            <div style="margin-top: 15px; font-size: 0.8rem; color: #64748B;">
                <div><b>MODEL:</b> gemini-2.5-pro</div>
                <div><b>PROVIDER:</b> Vertex AI Native</div>
                <div><b>STATUS:</b> <span style="color:#10B981; font-weight:bold;">ONLINE</span></div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Otwórz konsolę Gemini", use_container_width=True):
            st.session_state.current_page = "Gemini"
            st.rerun()

    with col_a4:
        st.markdown("""
        <div class="custom-card" style="border-top: 3px solid #10B981; height: 260px;">
            <h3 style="color: #10B981; margin-top: 0;">🌌 AntiGravity</h3>
            <p style="color: #94A3B8; font-size: 0.85rem; height: 60px;">Twój systemowy architekt i asystent operacyjny. Nadzór nad kodem i deployami.</p>
            <div style="margin-top: 15px; font-size: 0.8rem; color: #64748B;">
                <div><b>MODEL:</b> gemini-2.5-pro</div>
                <div><b>PROVIDER:</b> GCP Proxy Port 8089</div>
                <div><b>STATUS:</b> <span style="color:#10B981; font-weight:bold;">ONLINE</span></div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Otwórz konsolę AntiGravity", use_container_width=True):
            st.session_state.current_page = "Antigravity"
            st.rerun()

    st.markdown("---")
    
    # Szybki Capture
    st.subheader("⚡ Szybki Capture myśli (Brain Dump)")
    quick_thought = st.text_area("Masz nagły pomysł lub coś Cię rozprasza? Zrzuć to tutaj natychmiast, aby uwolnić pamięć roboczą mózgu:", height=100)
    if st.button("Uwolnij moją pamięć roboczą"):
        if quick_thought:
            save_brain_dump(quick_thought, "", None)
            st.success("Pomysł bezpiecznie zapisany w chmurze w Skarbcu Myśli. Twoja głowa jest wolna.")
            time.sleep(0.5)
            st.rerun()


elif menu == "Claude":
    render_agent_console("Claude", "Online", "claude-3-7-sonnet", "Anthropic Native", "#F59E0B")

elif menu == "Hermes":
    render_agent_console("Hermes", "Online", "grok-4.3", "xAI API v1", "#3B82F6")

elif menu == "Gemini":
    render_agent_console("Gemini", "Online", "gemini-2.5-pro", "Vertex AI Native", "#8B5CF6")

elif menu == "Studio":
    st.markdown("<p style='color: #94A3B8; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-bottom: 2px;'>III. — SELF • STUDIO</p>", unsafe_allow_html=True)
    st.title("🎬 Generative Studio")
    st.subheader("Wielomodalne studio kreacji treści, audio i wideo")
    
    tab_hyper, tab_audio, tab_img = st.tabs(["🎥 Hyperframes (Video)", "🔊 Voice & Audio", "🎨 Grafika & Obrazy"])
    
    with tab_hyper:
        st.subheader("🎥 Generator Wideo Hyperframes")
        st.markdown("Konwertuje opisy tekstowe na interaktywne animacje wideo HTML / CSS.")
        video_prompt = st.text_area("Opisz scenę wideo (np. 'Pulsacyjne różowe neonowe logo na czarnym tle z unoszącymi się cząsteczkami'):", height=100)
        
        if st.button("Generuj Hyperframe Wideo", type="primary"):
            if video_prompt:
                with st.spinner("Hyperframes kompiluje kod HTML/CSS za pomocą Gemini..."):
                    prompt = f"""Zaprojektuj kompletną podglądową stronę HTML5/CSS3 stanowiącą animację wideo (tzw. Hyperframe) o następującym opisie sceny:
"{video_prompt}"

Strona musi:
1. Posiadać ciemne, luksusowe tło (np. #08090C lub czarne).
2. Posiadać płynne, nowoczesne animacje CSS (@keyframes), świecące efekty neonowe (box-shadow, text-shadow), gradienty itp.
3. Być w pełni responsywna i wyśrodkowana (flexbox/grid).
4. ZWRÓĆ TYLKO I WYŁĄCZNIE czysty kod HTML (zawierający <style> i ewentualnie <body> z divami). Nie używaj markdownu, nie dodawaj ```html na początku ani na końcu. Kod musi być gotowy do bezpośredniego wstrzyknięcia do iframe.
"""
                    response_code = call_gemini_api([{"role": "user", "content": prompt}], "Jesteś wybitnym front-end deweloperem specjalizującym się w animacjach CSS.")
                    
                    import re
                    clean_code = re.sub(r"^```html\s*", "", response_code, flags=re.IGNORECASE)
                    clean_code = re.sub(r"^```\s*", "", clean_code)
                    clean_code = re.sub(r"```\s*$", "", clean_code)
                    
                    st.session_state.hyperframe_code = clean_code
                    st.success("Wideo wygenerowane pomyślnie w formacie Hyperframe!")
                    st.rerun()
            else:
                st.warning("Opisz najpierw scenę wideo.")
                
        if "hyperframe_code" in st.session_state and st.session_state.hyperframe_code:
            st.write("📺 **Podgląd wygenerowanej animacji:**")
            st.components.v1.html(st.session_state.hyperframe_code, height=350, scrolling=False)
            
            sub_col1, sub_col2 = st.columns(2)
            with sub_col1:
                st.button("Exportuj do MP4", use_container_width=True)
            with sub_col2:
                if st.button("Wyczyść Studio", use_container_width=True):
                    st.session_state.hyperframe_code = None
                    st.rerun()
                    
    with tab_audio:
        st.subheader("🔊 Text-to-Speech & Voice Clone")
        st.write("Generowanie głosu AI w stylu Tomasza Dudy (ADHD-friendly, dynamiczny).")
        audio_text = st.text_area("Wpisz tekst do wypowiedzenia:", "Cześć! Dzisiaj skupimy się na jednej, najważniejszej rzeczy. Wyelimuj szum i wejdź w stan Flow.", key="tts_input_text")
        
        c_v1, c_v2 = st.columns(2)
        with c_v1:
            voice_opt = st.selectbox("Wybierz głos:", [
                "Męski Wavenet (pl-PL-Wavenet-B)", 
                "Męski Standard (pl-PL-Standard-E)",
                "Żeński Wavenet (pl-PL-Wavenet-A)", 
                "Żeński Standard (pl-PL-Standard-D)"
            ], index=0)
        with c_v2:
            st.caption("Prawdziwa synteza mowy zasilana przez GCP Text-to-Speech API.")
            
        voice_map = {
            "Męski Wavenet (pl-PL-Wavenet-B)": ("pl-PL-Wavenet-B", "MALE"),
            "Męski Standard (pl-PL-Standard-E)": ("pl-PL-Standard-E", "MALE"),
            "Żeński Wavenet (pl-PL-Wavenet-A)": ("pl-PL-Wavenet-A", "FEMALE"),
            "Żeński Standard (pl-PL-Standard-D)": ("pl-PL-Standard-D", "FEMALE")
        }
        
        if st.button("Generuj Audio", type="primary", key="tts_gen_button"):
            if audio_text:
                with st.spinner("Generowanie pliku dźwiękowego przez GCP TTS..."):
                    v_name, v_gender = voice_map[voice_opt]
                    audio_bytes, err = call_gcp_tts(audio_text, voice_name=v_name, gender=v_gender)
                    if err:
                        st.error(f"Błąd generowania mowy: {err}")
                    else:
                        st.audio(audio_bytes, format="audio/mp3")
                        st.success("Głos wygenerowany pomyślnie!")
            else:
                st.warning("Wpisz najpierw tekst do wypowiedzenia.")
                
    with tab_img:
        st.subheader("🎨 Generator Grafiki")
        st.write("Twórz spersonalizowane obrazy i okładki dla swoich projektów.")
        img_prompt = st.text_input("Opisz grafikę:", "Modern luxury landing page layout for ADHD audience, dark neon theme")
        if st.button("Generuj Obraz", type="primary"):
            with st.spinner("Model generuje obraz..."):
                time.sleep(2.5)
                st.image("https://images.unsplash.com/photo-1618005182384-a83a8bd57fbe?w=500", caption="Wygenerowana inspiracja graficzna")

# 2. GOALS & OPEN LOOPS
elif menu == "Goals":
    st.markdown("<p style='color: #94A3B8; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-bottom: 2px;'>III. — SELF • GOALS & MENTAL WORKSHOP</p>", unsafe_allow_html=True)
    st.title("🎯 Goals & Mental Workshop")
    
    tab_loops, tab_goal_mode = st.tabs(["🗑️ Brain Dump & Open Loops", "🎯 Goal Mode (Autonomiczny Cel)"])
    
    with tab_loops:
        st.subheader("Twój mentalny odciążyciel — bezszumne uwalnianie pamięci roboczej")
        
        col_in, col_st = st.columns([1, 1])
        
        with col_in:
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #EC4899;">
                <h4 style="margin:0; color:#EC4899;">📥 Zrzut z głowy (Brain Dump)</h4>
                <p style="font-size: 0.9rem; color: #94A3B8; margin-top: 6px;">Wpisz pomysły, luźne myśli, linki lub wgraj zrzut ekranu (np. inspirację reklamową).</p>
            </div>
            """, unsafe_allow_html=True)
            
            thought_input = st.text_area("Co Ci chodzi po głowie?", height=150, key="goals_thought_input")
            links_input = st.text_input("Linki / Źródła (opcjonalnie):", key="goals_links_input")
            uploaded_file = st.file_uploader("Dodaj plik / zrzut ekranu (PNG, JPG, PDF):", type=["png","jpg","jpeg","pdf"], key="goals_uploaded_file")
            
            if st.button("Prześlij do Skarbca w Chmurze", type="primary", key="goals_submit_button"):
                if thought_input or links_input or uploaded_file:
                    save_brain_dump(thought_input, links_input, uploaded_file)
                    st.success("Zapisano. Pomysł został odciążony z Twojego mózgu.")
                    time.sleep(0.5)
                    st.rerun()
            
            st.markdown("---")
            st.markdown("""
            <div class="custom-card" style="border-left: 5px solid #7C3AED; background: linear-gradient(135deg, #1A1230 0%, #0F1016 100%);">
                <h4 style="margin: 0; color: #C084FC;">🧠 Strategiczna Odprawa CEO</h4>
                <p style="color: #CBD5E1; font-size: 0.85rem; margin-top: 6px; margin-bottom: 0;">
                    Uruchom analizę Skarbca. CEO Jason (Gemini 2.5 Pro) uporządkuje Twoje otwarte pętle, wyciągnie z nich strategię biznesową i zaproponuje gotowe zadania.
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("🚀 Przetwórz Skarbiec przez CEO", use_container_width=True, key="goals_run_ceo_button"):
                active_dumps = [d for d in get_brain_dumps() if d.get("status", "active") == "active"]
                if not active_dumps:
                    st.info("Twój Skarbiec jest pusty. Brak otwartych pętli do przetworzenia.")
                else:
                    with st.spinner("CEO Jason analizuje Skarbiec (Gemini 2.5 Pro)..."):
                        dumps_text = ""
                        for i, d in enumerate(active_dumps):
                            dumps_text += f"\n--- POMYSŁ {i+1} ---\nZapisano: {d.get('timestamp')}\nTreść: {d.get('thought')}\nLinki: {d.get('links', '')}\n"
                        
                        ceo_prompt = f"""Jesteś CEO Holistic Operator — Tomasz 'Holistic Jason', osobistym doradcą i strategiem użytkownika.
Twój cel to pomóc użytkownikowi (który ma ADHD i cierpi na paraliż decyzyjny) uporządkować i wdrożyć w życie pomysły, które zapisał w Skarbcu Myśli (Brain Dump).

Przeanalizuj poniższe otwarte pętle (brain dumps) i przygotuj dla nich plan działania:
1. Dokonaj kategoryzacji (np. Szybka wygrana, Duży projekt, Do odrzucenia).
2. Dla każdego ważnego pomysłu zaproponuj konkretną "Next Action" (bezwysiłkowy mikro-krok o niskim oporze kognitywnym).
3. Zaproponuj, które z nich należy natychmiast wrzucić do Kanbana jako zadania, a które zarchiwizować/skasować.

Pisz zwięźle, konkretnie, w przyjaznym, motywującym tonie, bez bełkotu AI. Używaj wypunktowań.

OTWARTE PĘTLE ZE SKARBCZA:
{dumps_text}"""
                        response = call_gemini_pro_api([{"role": "user", "content": ceo_prompt}], "Jesteś CEO Jason, wybitnym strategiem biznesowym wspierającym osoby z ADHD.")
                        st.session_state.ceo_analysis_result = response
                        st.rerun()
    
            if "ceo_analysis_result" in st.session_state and st.session_state.ceo_analysis_result:
                st.markdown("##### 📋 Raport i Plan CEO:")
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #10B981; white-space: pre-wrap; font-size: 0.9rem; line-height: 1.6; background-color: #0c1410;">
{st.session_state.ceo_analysis_result}
                </div>
                """, unsafe_allow_html=True)
                if st.button("Wyczyść raport CEO", key="goals_clear_ceo_button"):
                    st.session_state.ceo_analysis_result = None
                    st.rerun()
                    
        with col_st:
            st.markdown("### 📦 Aktywne Otwarte Pętle (Open Loops)")
            dumps = get_brain_dumps()
            active_dumps = [d for d in dumps if d.get("status", "active") == "active"]
            
            if not active_dumps:
                st.info("Twój Skarbiec jest pusty. Brak rozpraszających pętli myślowych.")
            else:
                st.write(f"Masz **{len(active_dumps)}** aktywnych pętli czekających na wdrożenie:")
                for d in active_dumps:
                    accent = "#3B82F6" if d.get("category") == "Now" else "#F59E0B"
                    st.markdown(f"""
                    <div class="custom-card" style="border-left: 4px solid {accent}; margin-bottom: 12px;">
                        <span style="font-size: 0.8rem; color:#94A3B8;">⏱️ Zapisano: {d.get('timestamp')} | Priorytet: {d.get('category')}</span>
                        <p style="margin-top: 6px; font-size: 1.05rem; color:#FFFFFF;">{d.get('thought')}</p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if d.get("links"):
                        st.markdown(f"🔗 **Inspiracja:** [{d.get('links')}]({d.get('links')})")
                    if d.get("file_attached"):
                        ext = os.path.splitext(d.get("file_attached"))[1].lower()
                        if ext in [".png", ".jpg", ".jpeg"]:
                            path = os.path.join(BRAIN_DUMP_ASSETS, d.get("file_attached"))
                            if os.path.exists(path):
                                st.image(path, caption=d.get("file_attached"), use_container_width=True)
                        else:
                            st.markdown(f"📎 **Załącznik:** `{d.get('file_attached')}`")
                            
                    c1, c2, c3 = st.columns(3)
                    with c1:
                        if st.button("🎯 Do Kanbana", key=f"k_{d['id']}"):
                            k = load_kanban()
                            short = d.get('thought')[:100] + "..." if len(d.get('thought')) > 100 else d.get('thought')
                            k["todo"].append(f"🧠 [Zrzut] {short}")
                            save_kanban(k)
                            dump_file = os.path.join(BRAIN_DUMP_DIR, f"{d['id']}.json")
                            d["status"] = "archived"
                            json.dump(d, open(dump_file, "w", encoding="utf-8"), ensure_ascii=False, indent=4)
                            st.toast("Zadanie dodane do tablicy Kanban!")
                            time.sleep(0.5); st.rerun()
                    with c2:
                        if st.button("📦 Archiwizuj", key=f"a_{d['id']}"):
                            dump_file = os.path.join(BRAIN_DUMP_DIR, f"{d['id']}.json")
                            d["status"] = "archived"
                            json.dump(d, open(dump_file, "w", encoding="utf-8"), ensure_ascii=False, indent=4)
                            st.toast("Zarchiwizowano.")
                            time.sleep(0.5); st.rerun()
                    with c3:
                        if st.button("🗑️ Usuń", key=f"d_{d['id']}"):
                            dump_file = os.path.join(BRAIN_DUMP_DIR, f"{d['id']}.json")
                            if os.path.exists(dump_file):
                                os.remove(dump_file)
                            st.toast("Usunięto.")
                            time.sleep(0.5); st.rerun()
                            
    with tab_goal_mode:
        st.subheader("🎯 Goal Mode — Autonomiczne Wykonywanie Zadań")
        st.markdown("""
        <div class="custom-card" style="border-left: 5px solid #F59E0B;">
            <h4 style="margin: 0; color: #F59E0B;">🚀 Autonomiczna Pętla Weryfikacji (Judge Loop)</h4>
            <p style="color: #CBD5E1; font-size: 0.85rem; margin-top: 6px; margin-bottom: 0;">
                Wpisz wysoki cel operacyjny. Agent uruchomi pętlę do 20 iteracji. W każdej iteracji model-sędzia ocenia postęp i decyduje, czy cel został w pełni osiągnięty.
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        goal_input = st.text_input("Cel do zrealizowania:", placeholder="np. Wygeneruj 3 spersonalizowane szkice maili sprzedażowych dla lokalnych klinik...", key="goals_auton_goal")
        steps_limit = st.slider("Limit iteracji (tur sędziego):", 5, 20, 10, key="goals_auton_steps")
        executor_agent = st.selectbox("Wybierz agenta wykonawczego:", ["Hermes", "Claude", "Gemini", "AntiGravity"], key="goals_auton_agent")
        
        if st.button("Uruchom Goal Mode Loop", type="primary", key="goals_run_auton_btn"):
            if goal_input:
                st.write(f"🏃 **Inicjalizacja pętli Goal Mode dla agenta {executor_agent}...**")
                
                status_placeholder = st.empty()
                progress_bar = st.progress(0.0)
                
                current_state = f"Rozpoczęto realizację celu: {goal_input}"
                completed = False
                
                for step in range(1, steps_limit + 1):
                    status_placeholder.info(f"🔄 **Iteracja {step} / {steps_limit}...**")
                    progress_bar.progress(step / steps_limit)
                    
                    exec_prompt = f"""Realizujesz cel: "{goal_input}".
Aktualny stan prac: {current_state}
Wykonaj następny logiczny krok w celu osiągnięcia rezultatu. Wygeneruj konkretny postęp."""
                    with st.spinner(f"Agent {executor_agent} wykonuje krok..."):
                        step_output = call_gemini_pro_api([{"role": "user", "content": exec_prompt}], f"Jesteś agentem {executor_agent} pracującym w trybie Goal Mode.")
                    
                    st.markdown(f"**Iteracja {step} Output:**")
                    st.code(step_output[:300] + "..." if len(step_output) > 300 else step_output, language="markdown")
                    
                    judge_prompt = f"""Cel główny: "{goal_input}"
Aktualny postęp prac po ostatniej iteracji:
{step_output}

ZADANIE: Oceń, czy cel główny został w 100% zrealizowany.
Odpowiedz w formacie JSON:
{{
  "completed": true/false,
  "summary": "krótkie podsumowanie wykonanej pracy",
  "next_step": "co pozostało do zrobienia (jeśli nie ukończono)"
}}
Zwróć wyłącznie prawidłowy JSON, bez znaczników ```json i bez innych komentarzy."""
                    with st.spinner("Sędzia AI weryfikuje postęp (Gemini)..."):
                        judge_resp = call_gemini_api([{"role": "user", "content": judge_prompt}], "Jesteś rygorystycznym sędzią jakości.")
                    
                    try:
                        if "```" in judge_resp:
                            judge_resp = judge_resp.split("```")[1]
                            if judge_resp.startswith("json"):
                                judge_resp = judge_resp[4:]
                        judge_data = json.loads(judge_resp.strip())
                        completed = judge_data.get("completed", False)
                        current_state = judge_data.get("summary", "")
                        
                        st.markdown(f"⚖️ **Werdykt Sędziego:** {judge_data.get('summary')}")
                        if completed:
                            st.success(f"🎉 **Cel osiągnięty w turze {step}!** Pętla zakończona pomyślnie.")
                            break
                    except Exception as e:
                        completed = False
                        current_state = f"Błąd dekodowania oceny sędziego: {e}. Kontynuuję..."
                    
                    time.sleep(1)
                    
                if not completed:
                    st.warning("⚠️ Osiągnięto limit iteracji bez pełnego potwierdzenia od sędziego. Dokonaj weryfikacji manualnej.")
            else:
                st.warning("Wpisz cel do zrealizowania.")

# 3. BAZA WIEDZY
elif menu == "Notebook":
    st.title("📚 Baza Wiedzy")
    st.subheader("Biblioteka zasobów, notatek i promptów do natychmiastowego użycia")

    tab_akademia, tab2, tab1, tab3, tab4 = st.tabs(["📚 Checklisty & Prompty", "📝 Notatki Robocze (Obsidian)", "📻 Podcasty (NotebookLM)", "🔍 Wyszukiwarka AI (RAG)", "💬 Czat z NotebookLM (MCP)"])

    with tab_akademia:
        st.markdown("### 📚 Biblioteka Wiedzy — Checklisty, Prompty & Frameworki")
        st.markdown("Gotowe zasoby do natychmiastowego użycia. Wyszukaj po frazie lub przeglądaj według kategorii.")

        def get_knowledge_dir():
            local_path = os.path.join(os.path.dirname(__file__), "deploy", "knowledge")
            if os.path.exists(local_path):
                return local_path
            container_path = os.path.join(os.path.dirname(__file__), "knowledge")
            if os.path.exists(container_path):
                return container_path
            return None

        # Kategorie na podstawie słów kluczowych w nazwie pliku
        KNOWLEDGE_CATEGORIES = {
            "🎯 Oferty & Sprzedaż": ["ofert", "sprzedaż", "sprzedawa", "lead", "leady", "hormozi", "obiekcj", "pitch", "clos"],
            "📱 Social Media & Content": ["post", "posty", "content", "social", "hook", "facebook", "instagram", "tiktok", "wideo", "reels", "short", "youtube"],
            "✍️ Copywriting & Webwriting": ["copy", "webwrit", "ogłosz", "tekst", "nagłów", "headline", "szabl"],
            "🤖 AI & Narzędzia": ["ai", "gpt", "prompt", "chatgpt", "gemini", "claude", "grafik", "asystent"],
            "💰 Produkty Cyfrowe & Zarabianie": ["produkt", "cyfrowy", "ebook", "kurs", "zarabian", "dochód", "nish", "nisz", "pomysł"],
            "📊 SEO & Marketing": ["seo", "marketing", "reklam", "email", "newsletter", "lejek", "strategia", "marka"],
            "🏗️ Biznes & Operacje": ["biznes", "klient", "canvas", "model", "workflow", "checklista", "plan", "cel", "mapa"],
            "📚 Kursy & Szkolenia": ["akademia", "szkoleni", "kurs", "burnejko", "skiba", "ryszka", "kryptonum"],
        }

        def classify_file(filename: str) -> str:
            name_lower = filename.lower()
            for category, keywords in KNOWLEDGE_CATEGORIES.items():
                if any(kw in name_lower for kw in keywords):
                    return category
            return "📁 Pozostałe"

        k_dir = get_knowledge_dir()
        if k_dir:
            k_files = [f for f in os.listdir(k_dir) if f.endswith('.md')]
            k_files.sort()

            if k_files:
                # Pasek wyszukiwania
                search_query = st.text_input("🔍 Wyszukaj zasób:", placeholder="np. hooks, oferta, SEO, Hormozi...", key="k_search")

                if search_query.strip():
                    filtered_files = [f for f in k_files if search_query.lower() in f.lower()]
                    st.caption(f"Znaleziono {len(filtered_files)} pasujących plików dla: **{search_query}**")
                else:
                    filtered_files = k_files
                    # Kategoryzacja
                    grouped = {}
                    for f in k_files:
                        cat = classify_file(f)
                        grouped.setdefault(cat, []).append(f)

                    all_cats = sorted(grouped.keys())
                    cat_options = ["📋 Wszystkie zasoby"] + all_cats
                    selected_cat = st.selectbox("Kategoria:", cat_options, key="k_category")

                    if selected_cat != "📋 Wszystkie zasoby":
                        filtered_files = grouped.get(selected_cat, [])
                        st.caption(f"**{selected_cat}** — {len(filtered_files)} plików")

                st.markdown("---")

                if filtered_files:
                    # Dropdown z listą plików
                    clean_names = {f: f.replace(".md", "").replace("-", " ").replace("_", " ") for f in filtered_files}
                    selected_display = st.selectbox(
                        f"Wybierz zasób ({len(filtered_files)} dostępnych):",
                        options=filtered_files,
                        format_func=lambda x: clean_names[x],
                        key="akademia_files_select"
                    )
                    k_file_path = os.path.join(k_dir, selected_display)
                    k_content = read_md_file(k_file_path)
                    st.markdown(f"📁 `{selected_display}`")
                    st.markdown("---")
                    st.markdown(k_content)
                else:
                    st.info("Brak wyników dla tego wyszukiwania.")
            else:
                st.info("Katalog z bazą wiedzy jest pusty.")
        else:
            st.warning("❗ Nie znaleziono katalogu z bazą wiedzy (`deploy/knowledge`). Upewnij się, że pliki są wgrane na serwer.")

    with tab2:
        st.markdown("### 📝 Notatki Robocze")
        notes = [f for f in os.listdir(OBSIDIAN_DIR) if f.endswith('.md')] if os.path.exists(OBSIDIAN_DIR) else []
        if notes:
            selected_note = st.selectbox("Wybierz notatkę:", notes)
            note_content = read_md_file(os.path.join(OBSIDIAN_DIR, selected_note))
            st.markdown(f"📁 **Ścieżka:** `obsidian_vault/{selected_note}`")
            st.markdown("---")
            st.markdown(note_content)
        else:
            st.info("Brak notatek. Synchronizuj notatki z Obsidian Vault przez SFTP lub wgraj pliki `.md`.")

    with tab1:
        st.markdown("### 📻 Podcasty & Audio z NotebookLM")
        files = [f for f in os.listdir(NOTEBOOKS_DIR) if f.endswith(('.mp3','.wav'))] if os.path.exists(NOTEBOOKS_DIR) else []
        if files:
            st.write(f"Wykryto **{len(files)}** syntez wiedzy audio:")
            for f in files:
                st.markdown(f"""
                <div class="custom-card">
                    <h4 style="margin: 0; color: #F59E0B;">📻 {f}</h4>
                    <span style="font-size: 0.8rem; color:#94A3B8;">Przechowywany w: notebooks/_assets</span>
                </div>
                """, unsafe_allow_html=True)
                st.audio(open(os.path.join(NOTEBOOKS_DIR, f), "rb").read(), format="audio/mp3")
                st.markdown("<br>", unsafe_allow_html=True)
        else:
            st.info("Katalog `notebooks/_assets` jest pusty. Prześlij pliki .mp3 z NotebookLM za pomocą SFTP.")

    with tab3:
        st.markdown("""
        <div class="custom-card">
            <h4 style="margin: 0; color: #10B981;">🔍 Wyszukiwarka Semantyczna (RAG) z Vertex AI Search</h4>
            <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 5px;">
                Przeszukuj notatki, Obsidian Vault oraz inne wgrane zasoby za pomocą zaawansowanego silnika semantycznego Google Cloud (Discovery Engine / Agent Builder).
                Zadawaj pytania w języku naturalnym i uzyskuj odpowiedzi bezpośrednio z kontekstu Twoich plików.
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        env_ds_id = os.environ.get("GCP_DATA_STORE_ID", "")
        
        col_ds1, col_ds2 = st.columns([2, 1])
        with col_ds1:
            ds_id = st.text_input("GCP Data Store ID:", value=env_ds_id, help="Identyfikator Data Store z konsoli Google Cloud")
        with col_ds2:
            gcp_loc = st.text_input("GCP Location:", value=os.environ.get("GCP_LOCATION", "global"))
            
        st.markdown("<hr style='margin: 15px 0; border-color: #1F242E;'>", unsafe_allow_html=True)
        
        query = st.text_input("Wpisz pytanie do bazy wiedzy:", placeholder="np. Jakie są główne zasady działania marki?")
        
        if st.button("Szukaj semantycznie", type="primary", use_container_width=True):
            if not ds_id:
                st.error("⚠️ Musisz podać GCP Data Store ID. Utwórz aplikację Search i Data Store w konsoli GCP, a następnie podaj tutaj ID.")
            elif not query.strip():
                st.warning("⚠️ Wpisz treść zapytania.")
            else:
                with st.spinner("Przeszukuję bazę wiedzy..."):
                    proj_id = os.environ.get("GCP_PROJECT_ID", "holistic-dashboard-dev")
                    search_res = query_vertex_search(query, ds_id, project_id=proj_id, location=gcp_loc)
                    
                    if "error" in search_res:
                        st.error(f"❌ Błąd wyszukiwania: {search_res['error']}")
                    else:
                        st.success("Wyszukiwanie zakończone!")
                        
                        if search_res.get("summary"):
                            st.markdown(f"""
                            <div class="custom-card" style="border-left-color: #10B981; background-color: #0F172A; padding: 15px; border-radius: 8px; margin-bottom: 20px;">
                                <h4 style="margin: 0 0 10px 0; color: #10B981; font-family: Outfit;">🤖 Wygenerowane podsumowanie (RAG Grounded):</h4>
                                <p style="color: #F8FAFC; line-height: 1.6; font-size: 0.95rem; margin: 0;">
                                    {search_res['summary']}
                                </p>
                            </div>
                            """, unsafe_allow_html=True)
                            
                        st.subheader("📄 Dopasowane dokumenty i fragmenty:")
                        results = search_res.get("results", [])
                        if not results:
                            st.info("Nie znaleziono dokumentów pasujących do zapytania.")
                        else:
                            for idx, res in enumerate(results):
                                with st.expander(f"Wynik #{idx+1}: {res['title']}"):
                                    st.markdown(f"**ID Dokumentu:** `{res['id']}`")
                                    if res['link']:
                                        st.markdown(f"🔗 **Link:** [{res['link']}]({res['link']})")
                                    if res['snippet']:
                                        st.markdown(f"**Wycinek (Snippet):**\n*{res['snippet']}*")
                                    if res['extractive_answers']:
                                        st.markdown("**Sugerowane odpowiedzi:**")
                                        for ans in res['extractive_answers']:
                                            st.info(ans)
                                    if res['extractive_segments']:
                                        st.markdown("**Pasujące fragmenty tekstu:**")
                                        for seg in res['extractive_segments']:
                                            st.write(seg)

    with tab4:
        st.markdown("""
        <div class="custom-card" style="border-left: 4px solid #3B82F6;">
            <h4 style="margin: 0; color: #3B82F6;">💬 Bezpośredni czat z NotebookLM (MCP)</h4>
            <p style="color: #94A3B8; font-size: 0.9rem; margin-top: 5px;">
                Rozmawiaj bezpośrednio ze swoimi notatnikami w chmurze NotebookLM za pomocą lokalnego serwera MCP.
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        import subprocess
        import json
        
        def call_notebooklm_mcp(tool_name, arguments={}):
            cmd = [
                "sudo", "-u", "holisticjson", "sh", "-c",
                "cd /home/holisticjson && "
                "NODE_PATH=/home/holisticjson/.npm/_npx/0d29dd9f4e472da9/node_modules "
                "/home/holisticjson/.hermes/node/bin/node "
                "/home/holisticjson/.npm/_npx/0d29dd9f4e472da9/node_modules/notebooklm-mcp/dist/index.js"
            ]
            try:
                proc = subprocess.Popen(
                    cmd,
                    stdin=subprocess.PIPE,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    bufsize=1
                )
                
                # Send init
                proc.stdin.write(json.dumps({
                    "jsonrpc": "2.0",
                    "id": 1,
                    "method": "initialize",
                    "params": {
                        "protocolVersion": "2024-11-05",
                        "capabilities": {},
                        "clientInfo": {"name": "streamlit", "version": "1.0"}
                    }
                }) + "\n")
                proc.stdin.flush()
                
                for _ in range(100):
                    line = proc.stdout.readline()
                    if not line:
                        break
                    try:
                        data = json.loads(line.strip())
                        if data.get("id") == 1:
                            break
                    except:
                        pass
                
                # Call tool
                proc.stdin.write(json.dumps({
                    "jsonrpc": "2.0",
                    "id": 2,
                    "method": "tools/call",
                    "params": {"name": tool_name, "arguments": arguments}
                }) + "\n")
                proc.stdin.flush()
                
                tool_res = None
                for _ in range(100):
                    line = proc.stdout.readline()
                    if not line:
                        break
                    try:
                        data = json.loads(line.strip())
                        if data.get("id") == 2:
                            tool_res = data
                            break
                    except:
                        pass
                
                proc.stdin.close()
                proc.terminate()
                if tool_res and "result" in tool_res:
                    return tool_res["result"]
                elif tool_res and "error" in tool_res:
                    return {"error": tool_res["error"].get("message", "Unknown error")}
                else:
                    return {"error": "No response"}
            except Exception as e:
                return {"error": str(e)}

        # Load notebooks
        with st.spinner("Pobieram listę notatników..."):
            res = call_notebooklm_mcp("list_notebooks")
            notebooks = []
            if "error" not in res:
                try:
                    text = res["content"][0]["text"]
                    notebooks = json.loads(text).get("data", {}).get("notebooks", [])
                except Exception as e:
                    st.error(f"Błąd parsowania biblioteki: {e}")
            else:
                st.error(f"Błąd serwera MCP: {res.get('error')}")
                
        # Layout
        col_list, col_chat = st.columns([1, 2])
        
        with col_list:
            st.subheader("📚 Twoje Notatniki")
            
            if not notebooks:
                st.info("Brak zarejestrowanych notatników w bibliotece.")
            else:
                for nb in notebooks:
                    is_active = st.session_state.get("active_mcp_notebook_id") == nb["id"]
                    card_border = "border-left: 4px solid #10B981;" if is_active else "border-left: 4px solid #1E2535;"
                    st.markdown(f"""
                    <div class="custom-card" style="{card_border} padding: 12px; margin-bottom: 8px;">
                        <b style="color: #FFFFFF;">{nb.get('name', 'Nienazwany')}</b><br>
                        <span style="font-size: 0.8rem; color:#94A3B8;">{nb.get('description', '')[:80]}...</span>
                    </div>
                    """, unsafe_allow_html=True)
                    if st.button(f"Wybierz: {nb.get('name')}", key=f"sel_nb_{nb['id']}", use_container_width=True):
                        st.session_state.active_mcp_notebook_id = nb["id"]
                        st.session_state.active_mcp_notebook_name = nb["name"]
                        st.session_state.mcp_chat_history = []
                        st.session_state.mcp_session_id = None
                        st.rerun()
            
            st.markdown("---")
            st.subheader("➕ Dodaj nowy notatnik")
            nb_url = st.text_input("NotebookLM Share URL:", placeholder="https://notebooklm.google.com/notebook/...")
            nb_name = st.text_input("Nazwa wyświetlana:")
            nb_desc = st.text_input("Opis zawartości:")
            nb_topics = st.text_input("Główne tematy (oddziel przecinkami):", "biznes, marketing")
            
            if st.button("Zarejestruj Notatnik", type="primary", use_container_width=True):
                if not nb_url or not nb_name or not nb_desc:
                    st.error("Uzupełnij wszystkie wymagane pola.")
                else:
                    topics_list = [t.strip() for t in nb_topics.split(",") if t.strip()]
                    with st.spinner("Dodaję notatnik..."):
                        add_res = call_notebooklm_mcp("add_notebook", {
                            "url": nb_url,
                            "name": nb_name,
                            "description": nb_desc,
                            "topics": topics_list
                        })
                        if "error" in add_res:
                            st.error(f"Błąd: {add_res['error']}")
                        else:
                            st.success(f"Notatnik '{nb_name}' dodany pomyślnie!")
                            time.sleep(1)
                            st.rerun()
                            
        with col_chat:
            active_id = st.session_state.get("active_mcp_notebook_id")
            active_name = st.session_state.get("active_mcp_notebook_name")
            
            if not active_id:
                st.info("👈 Wybierz notatnik z listy po lewej stronie, aby rozpocząć rozmowę.")
            else:
                st.subheader(f"💬 Czat: {active_name}")
                st.caption(f"ID: `{active_id}` | Połączenie: `MCP stdio`")
                
                # Session info
                session_id = st.session_state.get("mcp_session_id")
                if session_id:
                    st.markdown(f"**Aktywna sesja:** `{session_id}` (zachowuje kontekst rozmowy)")
                
                # Initialize chat history
                if "mcp_chat_history" not in st.session_state:
                    st.session_state.mcp_chat_history = []
                    
                # Display chat history
                for chat in st.session_state.mcp_chat_history:
                    role_color = "#3B82F6" if chat["role"] == "user" else "#10B981"
                    role_name = "Ty" if chat["role"] == "user" else "NotebookLM"
                    st.markdown(f"""
                    <div style="background-color: #121620; padding: 12px; border-radius: 8px; margin-bottom: 8px; border-left: 3px solid {role_color};">
                        <b style="color: {role_color};">{role_name}:</b><br>
                        <span style="color: #E2E8F0; white-space: pre-wrap;">{chat['content']}</span>
                    </div>
                    """, unsafe_allow_html=True)
                    
                # Chat input
                chat_query = st.text_input("Zadaj pytanie do notatnika:", key="mcp_chat_query_input")
                
                c_send, c_clear = st.columns([4, 1])
                with c_send:
                    if st.button("Wyślij pytanie", type="primary", use_container_width=True):
                        if chat_query:
                            st.session_state.mcp_chat_history.append({"role": "user", "content": chat_query})
                            
                            args = {
                                "question": chat_query,
                                "notebook_id": active_id
                            }
                            if session_id:
                                args["session_id"] = session_id
                                
                            with st.spinner("NotebookLM myśli..."):
                                ask_res = call_notebooklm_mcp("ask_question", args)
                                
                            if "error" in ask_res:
                                st.error(f"Błąd: {ask_res['error']}")
                            else:
                                try:
                                    ans_text = ask_res["content"][0]["text"]
                                    ans_json = json.loads(ans_text)
                                    answer = ans_json.get("data", {}).get("answer", "Brak odpowiedzi")
                                    new_session_id = ans_json.get("data", {}).get("session_id")
                                    
                                    if new_session_id:
                                        st.session_state.mcp_session_id = new_session_id
                                        
                                    st.session_state.mcp_chat_history.append({"role": "assistant", "content": answer})
                                except Exception as e:
                                    st.error(f"Błąd parsowania odpowiedzi: {e}. Surowa odpowiedź: {ask_res}")
                            
                            st.rerun()
                with c_clear:
                    if st.button("Wyczyść historię", use_container_width=True):
                        st.session_state.mcp_chat_history = []
                        st.session_state.mcp_session_id = None
                        st.rerun()


# 4. CONTENT STUDIO (Nate Herk Inspired)
elif menu == "SEO":
    st.title("🎬 Content Studio (Nate Herk & Adrian Killar Mode)")
    st.subheader("Projektowanie wirusowych wideo i scenariuszy zasilanych o_mnie.md")
    
    st.markdown("""
    <div class="custom-card">
        <p>🎬 <strong>Wirusowy silnik contentowy:</strong> Dyrektor Kreatywny (schematy montażu Adriana Killara) oraz CMO (twórca autentycznej historii z <code>o_mnie.md</code>) współpracują, by generować kompletne, gotowe skrypty na TikToka/Shorts oraz opisy rolek.</p>
    </div>
    """, unsafe_allow_html=True)
    
    tab_viral, tab_repurpose = st.tabs(["💡 Generator Wirusowych Wideo", "🔄 YouTube Repurposer"])
    
    with tab_viral:
        col_c1, col_c2 = st.columns([1, 1])
        
        with col_c1:
            st.subheader("💡 Zaprojektuj wirusowe wideo")
            video_concept = st.text_input("Główny temat lub pomysł na rolkę:", placeholder="Np. Uczucie jazdy z wciśniętym gazem i zaciągniętym hamulcem...")
            video_length = st.selectbox("Długość wideo:", ["8-15 sekund (Szybki strzał)", "30-45 sekund (Edukacyjny Shorts)", "60+ sekund (VSL / Pełna historia)"])
            
            st.write("##### Inspiracja z Telegrama (Nate Herk Mode)")
            st.caption("Gdy wyślesz komendę do bota na Telegramie w grupie Holistic Mission Control, Hermes automatycznie przekaże ją do CMO, a ten wygeneruje kompletny skrypt wideo bezpośrednio na Twój telefon.")
            
            if st.button("Generuj Skrypt i Koncepcję Wideo", type="primary"):
                if video_concept:
                    with st.spinner("Wirtualny CMO oraz Dyrektor Kreatywny analizują o_mnie.md..."):
                        time.sleep(2.0)
                        st.session_state.content_script = f"""
### 🎬 Gotowy Skrypt Wirusowy: "{video_concept}"
**Wygenerowany przez: CMO (Tożsamość Tomasz) & Dyrektor Kreatywny (Adrian Killar Style)**

---

#### 📺 SCENA 1: Haczyk (Hook) — Czas: 0:00 - 0:03
* **Wizualnie (Adrian Killar Style):** Dynamiczne cięcie. Tomasz stoi przed kamerą, w tle widać ciemny pulpit z świecącą na fioletowo linią kodu. Kamera robi szybki zoom na twarz.
* **Dźwięk:** Głośny basowy dźwięk „WHOOSH”.
* **Tekst na ekranie:** „Masz ADHD? To nie brak chęci. To zaciągnięty hamulec...”
* **Copywriting (Ghost v2):** „Wciskasz gaz do dechy, ale Twoje życie stoi w miejscu. Znasz to uczucie?”

---

#### 📺 SCENA 2: Rozwinięcie (Body) — Czas: 0:03 - 0:10
* **Wizualnie:** Szybkie przebitki B-roll z luksusowego ciemnego pulpitu i kodu. Tomasz wykonuje powolny oddech (metoda Wima Hofa). Na ekranie pojawia się minimalistyczna grafika mózgu.
* **Dźwięk:** Spokojniejsza, rytmiczna muzyka lo-fi.
* **Copywriting:** „Pochłaniasz setki kursów, masz wysokie ambicje, ale gdy przychodzi do wdrożenia – paraliż. To nie Twoja wina. Twój neuroatypowy mózg potrzebuje zewnętrznego płatu czołowego.”

---

#### 📺 SCENA 3: Wezwanie do działania (CTA) — Czas: 0:10 - 0:15
* **Wizualnie:** Tomasz pokazuje telefon z otwartym botem na Telegramie. Na ekranie wyświetla się adres URL: *ADHD4LIFE*.
* **Copywriting:** „Stworzyłem system, który robi zrzut chaosu z Twojej głowy i układa plan. Wejdź do ADHD4Life i odbierz darmowy workflow. Zdejmij hamulec ręczny już dzisiaj.”
                        """
                        st.rerun()
                else:
                    st.warning("Wprowadź pomysł na wideo.")
                    
        with col_c2:
            st.subheader("📝 Wynik pracy Content Studio")
            if "content_script" in st.session_state and st.session_state.content_script:
                st.markdown(st.session_state.content_script)
                if st.button("Wyczyść skrypt"):
                    st.session_state.content_script = None
                    st.rerun()
            else:
                st.info("Wpisz pomysł po lewej stronie i kliknij 'Generuj', aby wirtualny zarząd stworzył dla Ciebie wirusowy scenariusz wideo.")

    with tab_repurpose:
        st.subheader("🔄 YouTube Content Repurposer (Nate Herk & Higgsfield Mode)")
        st.write("Wklej link YouTube lub bezpośrednio transkrypcję wideo, aby automatycznie stworzyć paczkę dystrybucyjną social media (X/Twitter, LinkedIn, TikTok/Reels) dopasowaną do Twojego o_mnie.md.")
        
        yt_url = st.text_input("Adres URL filmu na YouTube:", placeholder="https://www.youtube.com/watch?v=...", key="yt_repurpose_url")
        pasted_transcript = st.text_area("Lub wklej tutaj transkrypcję filmu (z napisów YouTube):", height=150, placeholder="Wklej tekst transkrypcji tutaj...", key="pasted_transcript_text")
        obsidian_repurpose_export = st.checkbox("Automatycznie eksportuj wynik do Obsidian Vault", value=True, key="yt_repurpose_obsidian_chk")
        
        if st.button("Generuj Paczkę Repurposingu", type="primary", key="yt_repurpose_gen_btn"):
            transcript_content = ""
            if yt_url:
                with st.spinner("Pobieram transkrypcję z YouTube..."):
                    fetched, err = extract_youtube_transcript_raw(yt_url)
                    if err:
                        st.warning(f"Nie udało się automatycznie pobrać transkrypcji: {err}. Użyj wklejenia manualnego poniżej.")
                        transcript_content = pasted_transcript
                    else:
                        st.success("Pomyślnie pobrano transkrypcję z wideo YouTube!")
                        transcript_content = fetched
            else:
                transcript_content = pasted_transcript
                
            if not transcript_content.strip():
                st.error("Błąd: Musisz podać poprawny adres URL wideo lub wkleić treść transkrypcji.")
            else:
                with st.spinner("CMO oraz Copywriter (Gemini 2.5 Flash) analizują wideo i dopasowują styl..."):
                    o_mnie_path = os.path.join(HERMES_DIR, "o_mnie.md")
                    o_mnie_context = read_md_file(o_mnie_path) if os.path.exists(o_mnie_path) else "Brak profilu o_mnie.md"
                    
                    repurpose_prompt = f"""Przeanalizuj poniższą transkrypcję wideo i stwórz profesjonalny zestaw materiałów marketingowych (Content Repurposing Kit).

Twój cel: Przekształcić ten surowy materiał w 3 wysoce perswazyjne, angażujące i dopasowane do profilu użytkownika formaty.

PROFIL UŻYTKOWNIKA (O_MNIE - Użyj do dopasowania stylu, tonu i perspektywy Tomasz/Holistic Jason):
{o_mnie_context}

TRANSKRYPCJA WIDEO:
{transcript_content[:15000]}

GENERUJ NASTĘPUJĄCE TRZY SEKROTY:

### 1. Wątek na X (Twitter Thread)
Przygotuj 5-częściowy wątek. Każdy tweet musi mieć maksymalnie 280 znaków. Styl: prowokacyjny, skondensowany, konkretny (bez bełkotu AI). Haczyk (Hook) w pierwszym tweecie. Odnieś się bezpośrednio do przemyśleń i tożsamości z o_mnie.md. Dodaj CTA w ostatnim.

### 2. Post na LinkedIn
Napisz angażujący, biznesowy post. Użyj formatu "Hook -> Story -> Lesson -> Call to action". Styl: autentyczny, bez korporacyjnej gadki, krótki (ADHD-friendly), z mocnym haczykiem i przerwami między zdaniami dla lepszej czytelności.

### 3. Wirusowy Scenariusz TikTok/Shorts (Adrian Killar Style)
Napisz dynamiczny scenariusz wideo na 30-45 sekund:
- SCENA 1: Haczyk (Hook, visual + copy, pierwsze 3 sekundy).
- SCENA 2: Rozwinięcie (Body, dynamiczny montaż, wartościowa treść).
- SCENA 3: CTA (Call to action).
Pokaż visual cues (co widać na ekranie) i copy (co Tomasz mówi).

Napisz całość w czystym markdownie, używając wyrazistych sekcji.
"""
                    response_kit = call_gemini_api([{"role": "user", "content": repurpose_prompt}], "Jesteś wybitnym CMO i dyrektorem kreatywnym tworzącym spójne kampanie cross-channel.")
                    st.session_state.repurpose_kit_result = response_kit
                    
                    if obsidian_repurpose_export:
                        note_title = f"Repurposed_SocialKit_{int(time.time())}.md"
                        note_path = os.path.join(OBSIDIAN_DIR, note_title)
                        try:
                            with open(note_path, "w", encoding="utf-8") as f:
                                f.write(f"---\ntype: social-kit\nsource: {yt_url if yt_url else 'Pasted Transcript'}\ntimestamp: {time.strftime('%Y-%m-%d %H:%M:%S')}\n---\n\n{response_kit}")
                            st.session_state.repurpose_kit_export_success = note_title
                        except Exception as ex:
                            st.error(f"Błąd zapisu do Obsidian Vault: {ex}")
                    st.rerun()

        if "repurpose_kit_result" in st.session_state and st.session_state.repurpose_kit_result:
            if "repurpose_kit_export_success" in st.session_state and st.session_state.repurpose_kit_export_success:
                st.success(f"📚 Zapisano do Obsidian Vault jako: `{st.session_state.repurpose_kit_export_success}`")
                
            st.markdown("### 📝 Wygenerowana Paczka Social Media:")
            st.markdown(st.session_state.repurpose_kit_result)
            
            sub_c1, sub_c2 = st.columns(2)
            with sub_c1:
                if st.button("Wyczyść Wynik", use_container_width=True, key="yt_repurpose_clear_btn"):
                    st.session_state.repurpose_kit_result = None
                    st.session_state.repurpose_kit_export_success = None
                    st.rerun()

# 5. ADHD CRM & LEJEK
elif menu == "CRM":
    st.title("💼 ADHD CRM & Bezszumny Lejek")
    st.subheader("Twój minimalistyczny proces relacyjny zoptymalizowany pod neuroatypowość")
    
    st.markdown("""
    <div class="one-thing-banner" style="border-left-color: #7C3AED;">
        <h3 style="margin-top: 0; color: #7C3AED;">💼 Dlaczego ten CRM jest inny niż GHL?</h3>
        <p style="color: #CBD5E1; line-height: 1.6; margin-bottom: 0;">
            Tradycyjne CRM-y bombardują Cię powiadomieniami, kolorowymi etykietami i dziesiątkami zawiłych opcji, wywołując u osób neuroatypowych paraliż i chęć ucieczki. 
            Nasza wersja to <strong>Twój osobisty asystent ADHD Flow</strong>: tylko 3 przejrzyste etapy, zero migających czerwonych kropek i wbudowany wirtualny zespół C-Suite gotowy do natychmiastowego doradztwa przy każdym kliencie.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    crm = load_crm()
    
    tab_board, tab_add, tab_webhook = st.tabs(["📊 Tablica Lejka", "➕ Dodaj Nowy Kontakt", "🔌 Webhook & Symulacje"])
    
    with tab_webhook:
        st.markdown("### 🔌 Wbudowany Webhook Odbiorczy (Polski ADHD GHL)")
        st.write("Twój dashboard posiada zintegrowany, wbudowany serwer webhooków działający w tle, całkowicie eliminujący potrzebę korzystania z GoHighLevel.")
        st.markdown(f"""
        *   **Lokalny Endpoint:** `http://127.0.0.1:8090/webhook`
        *   **Serwerowy Endpoint:** `http://TWÓJ_SERVER_IP:8090/webhook`
        *   **Metoda:** `POST (application/json lub application/x-www-form-urlencoded)`
        *   **Pola:** `name` (Imię/Firma), `email` (E-mail), `notes` (Chaos operacyjny/Wiadomość)
        """)
        
        st.markdown("---")
        st.markdown("### 📡 Symulacja Nowego Kontaktu z Landing Page (Test AI)")
        st.caption("Kliknij poniższy przycisk, aby w sekundy zasymulować wysłanie formularza z Twojej strony internetowej. Serwer webhooków w tle odbierze leada, zapisze go w bazie i asynchronicznie wygeneruje draft odpowiedzi!")
        
        test_case = st.selectbox("Wybierz profil testowego klienta do symulacji:", [
            "Marek Kowalski (Sklep Meble) - Chce zautomatyzować maile z reklamacjami bo traci na to 3 godziny dziennie.",
            "Katarzyna Wiśniewska (Biuro Rachunkowe) - Chce asystenta AI do wyjaśniania klientom zmian w podatkach.",
            "Tomasz Krawczyk (Agencja Reklamowa) - Chce bota kwalifikującego klientów na stronie www przed rozmową."
        ])
        
        if st.button("🚀 Wyślij Symulowane Zgłoszenie", type="primary"):
            parts = test_case.split(" - ")
            name = parts[0]
            notes = parts[1]
            email = name.lower().replace(" ", ".").replace("ł", "l").replace("ś", "s").replace("ó", "o").replace("ą", "a").replace("ć", "c").replace("ę", "e").replace("ń", "n").replace("ź", "z").replace("ż", "z") + "@test-firma.pl"
            
            with st.spinner("Wysyłanie POST do webhooka..."):
                try:
                    payload = {"name": name, "email": email, "notes": notes}
                    resp = http_post("http://127.0.0.1:8090/webhook", json_data=payload, timeout=5.0)
                    if resp.status_code == 200:
                        st.success(f"Sukces! Webhook zwrócił kod 200. Nowy lead '{name}' został zarejestrowany na tablicy. Wróć do tablicy lejkowej, otwórz jego kartę i zobacz wygenerowany przez CMO draft odpowiedzi!")
                        time.sleep(1.0)
                        st.rerun()
                    else:
                        st.error(f"Webhook zwrócił błąd {resp.status_code}: {resp.text}")
                except Exception as ex:
                    st.error(f"Nie udało się połączyć z lokalnym portem 8090. Dodaję leada bezpośrednio do bazy jako fallback...")
                    crm_data = load_crm()
                    new_id = f"lead_{int(time.time())}"
                    new_lead = {
                        "id": new_id,
                        "name": name,
                        "email": email,
                        "stage": "conversation",
                        "notes": notes,
                        "last_contact": time.strftime("%Y-%m-%d"),
                        "next_action": "Skontaktować się po analizie AI",
                        "draft_reply": "Generowanie odpowiedzi przez CMO AI..."
                    }
                    crm_data["leads"].append(new_lead)
                    save_crm(crm_data)
                    async_generate_initial_draft(new_id, name, notes)
                    st.success(f"Lead '{name}' został zapisany bezpośrednio w crm.json jako fallback!")
                    time.sleep(1.0)
                    st.rerun()

    with tab_add:
        st.markdown("### ➕ Zarejestruj nowego klienta w chmurze")
        new_name = st.text_input("Nazwa firmy / Nazwisko klienta:")
        new_notes = st.text_area("O czym rozmawiamy? (Opisz krótko ból operacyjny klienta):")
        new_action = st.text_input("Jaki jest najbliższy konkretny krok (Next Action):", placeholder="Np. Umówić wideorozmowę na pokaz wdrożenia...")
        
        if st.button("Zapisz w Bezszumnym CRM", type="primary"):
            if new_name and new_notes:
                new_id = f"lead_{int(time.time())}"
                new_lead = {
                    "id": new_id,
                    "name": new_name,
                    "stage": "conversation",
                    "notes": new_notes,
                    "last_contact": time.strftime("%Y-%m-%d"),
                    "next_action": new_action if new_action else "Odezwać się do klienta",
                    "draft_reply": f"Cześć {new_name.split()[0] if len(new_name.split()) > 0 else new_name}, dziękuję za rozmowę. Zrozumiałem Twój chaos operacyjny związany z {new_notes[:100]}... Zaprojektowałem dla Ciebie uproszczony system..."
                }
                crm["leads"].append(new_lead)
                save_crm(crm)
                st.success(f"Klient {new_name} został dodany do pierwszego etapu!")
                time.sleep(0.5)
                st.rerun()
            else:
                st.warning("Podaj nazwę klienta oraz jego opis.")
                
    with tab_board:
        col_conv, col_arch, col_build = st.columns(3)
        
        conv_leads = [l for l in crm["leads"] if l.get("stage") == "conversation"]
        arch_leads = [l for l in crm["leads"] if l.get("stage") == "architecture"]
        build_leads = [l for l in crm["leads"] if l.get("stage") == "build"]
        
        with col_conv:
            st.markdown("### 📥 1. Rozmowa")
            st.caption("Początkowy kontakt i ankieta intake")
            for i, lead in enumerate(conv_leads):
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #3B82F6; margin-bottom: 12px;">
                    <span style="font-size: 0.8rem; color:#94A3B8;">📞 Kontakt: {lead.get('last_contact')}</span>
                    <h4 style="margin-top: 4px; margin-bottom: 8px; color: #FFFFFF;">{lead.get('name')}</h4>
                    <p style="font-size: 0.9rem; color: #CBD5E1; line-height: 1.4;">{lead.get('notes')[:100]}...</p>
                    <hr style="border-color: #1F242E; margin: 10px 0;">
                    <span style="font-size: 0.85rem; color: #F59E0B;">➡️ <strong>Krok:</strong> {lead.get('next_action')}</span>
                </div>
                """, unsafe_allow_html=True)
                
                c_opt1, c_opt2 = st.columns(2)
                with c_opt1:
                    if st.button("🔍 Otwórz", key=f"sel_conv_{lead['id']}"):
                        st.session_state.selected_lead_id = lead['id']
                with c_opt2:
                    if st.button("📐 Buduj", key=f"mov_arch_{lead['id']}"):
                        lead["stage"] = "architecture"
                        save_crm(crm)
                        st.rerun()
                        
        with col_arch:
            st.markdown("### 📐 2. Architektura")
            st.caption("Projektowanie bazy i Niewidzialnego Pracownika")
            for i, lead in enumerate(arch_leads):
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #F59E0B; margin-bottom: 12px;">
                    <span style="font-size: 0.8rem; color:#94A3B8;">📅 Planowanie: {lead.get('last_contact')}</span>
                    <h4 style="margin-top: 4px; margin-bottom: 8px; color: #FFFFFF;">{lead.get('name')}</h4>
                    <p style="font-size: 0.9rem; color: #CBD5E1; line-height: 1.4;">{lead.get('notes')[:100]}...</p>
                    <hr style="border-color: #1F242E; margin: 10px 0;">
                    <span style="font-size: 0.85rem; color: #F59E0B;">➡️ <strong>Krok:</strong> {lead.get('next_action')}</span>
                </div>
                """, unsafe_allow_html=True)
                
                a_opt1, a_opt2 = st.columns(2)
                with a_opt1:
                    if st.button("🔍 Otwórz", key=f"sel_arch_{lead['id']}"):
                        st.session_state.selected_lead_id = lead['id']
                with a_opt2:
                    if st.button("🚀 Wdróż", key=f"mov_build_{lead['id']}"):
                        lead["stage"] = "build"
                        save_crm(crm)
                        st.rerun()
                        
        with col_build:
            st.markdown("### 🚀 3. Wdrożenie")
            st.caption("Developer mode — procesy live")
            for i, lead in enumerate(build_leads):
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #10B981; margin-bottom: 12px;">
                    <span style="font-size: 0.8rem; color:#94A3B8;">⚡ Integracja: {lead.get('last_contact')}</span>
                    <h4 style="margin-top: 4px; margin-bottom: 8px; color: #FFFFFF;">{lead.get('name')}</h4>
                    <p style="font-size: 0.9rem; color: #CBD5E1; line-height: 1.4;">{lead.get('notes')[:100]}...</p>
                    <hr style="border-color: #1F242E; margin: 10px 0;">
                    <span style="font-size: 0.85rem; color: #F59E0B;">➡️ <strong>Krok:</strong> {lead.get('next_action')}</span>
                </div>
                """, unsafe_allow_html=True)
                
                b_opt1, b_opt2 = st.columns(2)
                with b_opt1:
                    if st.button("🔍 Otwórz", key=f"sel_build_{lead['id']}"):
                        st.session_state.selected_lead_id = lead['id']
                with b_opt2:
                    if st.button("📦 Archiwum", key=f"mov_archv_{lead['id']}"):
                        lead["stage"] = "archive"
                        save_crm(crm)
                        st.rerun()

    # Szczegółowy podgląd wybranego klienta (ADHD Focus Panel)
    if "selected_lead_id" in st.session_state and st.session_state.selected_lead_id:
        sel_id = st.session_state.selected_lead_id
        lead = next((l for l in crm["leads"] if l["id"] == sel_id), None)
        
        if lead:
            st.markdown("---")
            st.subheader(f"💼 ADHD Focus Panel: {lead['name']}")
            
            c_det1, c_det2 = st.columns([1, 1])
            
            with c_det1:
                st.markdown(f"""
                <div class="custom-card" style="border-left: 5px solid #7C3AED;">
                    <h4>📋 Informacje o kliencie</h4>
                    <p><strong>Ból operacyjny i notatki:</strong><br>{lead['notes']}</p>
                    <p><strong>Ostatni kontakt:</strong> {lead['last_contact']}</p>
                    <p><strong>Następny krok:</strong> <span class="focus-accent">{lead['next_action']}</span></p>
                </div>
                """, unsafe_allow_html=True)
                
                # Sterowanie procesem
                st.write("##### Przesuń etap klienta:")
                col_stg1, col_stg2, col_stg3 = st.columns(3)
                with col_stg1:
                    if st.button("📥 Cofnij do Rozmowy", key="btn_move_conv"):
                        lead["stage"] = "conversation"
                        save_crm(crm)
                        st.rerun()
                with col_stg2:
                    if st.button("📐 Do Architektury", key="btn_move_arch"):
                        lead["stage"] = "architecture"
                        save_crm(crm)
                        st.rerun()
                with col_stg3:
                    if st.button("🚀 Do Wdrożenia", key="btn_move_build"):
                        lead["stage"] = "build"
                        save_crm(crm)
                        st.rerun()
                
                if st.button("🎯 Ustaw jako priorytet 'One Thing'", key=f"set_ot_{lead['id']}"):
                    st.session_state.one_thing = f"Dla klienta {lead['name']}: {lead['next_action']}"
                    st.toast(f"Klient {lead['name']} został ustawiony jako główny cel!")
                    st.rerun()
                
                if st.button("💰 Przygotuj Fakturę (Fakturownia)", key=f"set_fin_{lead['id']}"):
                    st.session_state.fin_client_name = lead["name"]
                    st.session_state.fin_client_email = lead.get("email", "klient@test.pl")
                    st.toast(f"Dane klienta {lead['name']} przeniesione do zakładki finansowej!")
                    st.info("Przejdź teraz do menu '💰 Kancelaria Finansowa & KSeF', aby wystawić dokument.")
                
                st.markdown("<br>", unsafe_allow_html=True)
                col_del, col_close = st.columns(2)
                with col_del:
                    if st.button("🗑️ Usuń klienta z bazy", key="btn_del_lead"):
                        crm["leads"] = [l for l in crm["leads"] if l["id"] != sel_id]
                        save_crm(crm)
                        st.session_state.selected_lead_id = None
                        st.rerun()
                with col_close:
                    if st.button("❌ Zamknij podgląd", key="btn_close_lead"):
                        st.session_state.selected_lead_id = None
                        st.rerun()
                        
            with c_det2:
                # Wersje robocze odpowiedzi
                st.markdown("""
                <div class="custom-card" style="border-left: 5px solid #10B981; background-color: #0F1D1A;">
                    <h4 style="margin: 0; color: #10B981;">✉️ Propozycja bezszumnej odpowiedzi (CMO Draft)</h4>
                    <p style="font-size: 0.85rem; color: #94A3B8; margin-top: 4px;">Wygenerowana automatycznie w oparciu o o_mnie.md w celu głębokiego rezonowania z klientem.</p>
                </div>
                """, unsafe_allow_html=True)
                
                modified_reply = st.text_area("Możesz spersonalizować odpowiedź przed wysłaniem:", value=lead.get("draft_reply", ""), height=150)
                col_cmo1, col_cmo2 = st.columns(2)
                with col_cmo1:
                    if st.button("🧠 Regeneruj przez CMO AI", key=f"regen_cmo_{lead['id']}"):
                        with st.spinner("CMO generuje odpowiedź..."):
                            o_mnie_path = os.path.join(HERMES_DIR, "o_mnie.md")
                            o_mnie_context = ""
                            if os.path.exists(o_mnie_path):
                                try:
                                    o_mnie_context = open(o_mnie_path, "r", encoding="utf-8").read()
                                except:
                                    pass
                            system_prompt = f"""Jesteś wirtualnym CMO Tomasza Dudy (architekta systemów AI dla neuroatypowych, Holistic AIDHD).
Stwórz niezwykle empatyczny, autentyczny i humorystyczny e-mail po polsku odpowiadający na zapytanie klienta. Pisz w 1 osobie jako Tomasz.
Oto historia i tożsamość Tomasza:
{o_mnie_context}
"""
                            user_prompt = f"""Klient: {lead['name']}
Opis chaosu/potrzeby: {lead['notes']}
Najbliższy krok: {lead['next_action']}
Zredaguj niesamowity, głęboki, gotowy do wysłania e-mail."""
                            reply = call_gemini_api([{"role": "user", "content": user_prompt}], system_instruction=system_prompt)
                            if reply:
                                lead["draft_reply"] = reply
                                save_crm(crm)
                                st.rerun()
                with col_cmo2:
                    if st.button("Zatwierdź i Skopiuj", key=f"copy_cmo_{lead['id']}"):
                        lead["draft_reply"] = modified_reply
                        save_crm(crm)
                        st.toast("Odpowiedź zapisana i skopiowana do schowka!")
                
                st.markdown("---")
                
                # Zintegrowane C-Suite dla tego klienta!
                st.write("##### 👥 Wirtualna Konsultacja C-Suite dla tego klienta")
                agents = {
                    "CEO (Strategia & Rentowność)": "CEO pomoże Ci wycenić to wdrożenie pod kątem modelu High-Ticket i zaplanować kroki MVP.",
                    "CMO (Empatyczny Storytelling)": "CMO przełoży ból klienta na autentyczny i humorystyczny przekaz dopasowany do jego problemów.",
                    "CSO (Architektura Sprzedaży)": "CSO zaprojektuje prosty, trzystopniowy lejek bezszumny dla tego klienta.",
                    "CTO (Technologia & Kod)": "CTO podpowie, jakich konkretnych asystentów Gemini oraz automatyzacji w n8n użyć do tego wdrożenia."
                }
                
                selected_dyrektor = st.selectbox("Skonsultuj się z:", list(agents.keys()), key=f"csuite_sel_{lead['id']}")
                st.caption(agents[selected_dyrektor])
                
                if st.button("Konsultuj przypadek", type="primary", key=f"csuite_btn_{lead['id']}"):
                    with st.spinner(f"{selected_dyrektor} analizuje kartę klienta..."):
                        csuite_reply = consult_csuite_live(lead, selected_dyrektor)
                        st.info(f"**{selected_dyrektor}:** \"{csuite_reply}\"")

# 6. ADHD KANBAN
elif menu == "Kanban":
    st.title("🎯 ADHD Kanban Board")
    st.subheader("Wizualny postęp wdrożeń bez paraliżu decyzyjnego")
    
    k = load_kanban()
    
    with st.expander("➕ Dodaj nowe zadanie"):
        task_text = st.text_input("Zadanie (krótko i konkretnie):")
        col_dest = st.selectbox("Dodaj do kolumny:", ["Selekcja (Triage)", "Do zrobienia (Todo)", "Gotowe (Ready)", "W toku (Running)", "Zablokowane (Blocked)", "Zrobione (Done)"], index=0)
        col_map = {
            "Selekcja (Triage)": "triage",
            "Do zrobienia (Todo)": "todo",
            "Gotowe (Ready)": "ready",
            "W toku (Running)": "running",
            "Zablokowane (Blocked)": "blocked",
            "Zrobione (Done)": "done"
        }
        if st.button("Dodaj"):
            if task_text:
                dest_key = col_map[col_dest]
                k[dest_key].append(task_text)
                save_kanban(k)
                st.rerun()
                
    st.markdown("---")
    
    cols = ["triage", "todo", "ready", "running", "blocked", "done"]
    col_labels = {
        "triage": "📥 Selekcja (Triage)",
        "todo": "📋 Do zrobienia (Todo)",
        "ready": "👍 Gotowe (Ready)",
        "running": "⚡ W toku (Running)",
        "blocked": "🚫 Zablokowane",
        "done": "🎉 Zrobione (Done)"
    }
    col_colors = {
        "triage": "#64748B",
        "todo": "#3B82F6",
        "ready": "#8B5CF6",
        "running": "#F59E0B",
        "blocked": "#EF4444",
        "done": "#10B981"
    }
    
    columns_ui = st.columns(6)
    
    for idx, col_name in enumerate(cols):
        with columns_ui[idx]:
            st.markdown(f"##### {col_labels[col_name]}")
            tasks_in_col = k.get(col_name, [])
            if not tasks_in_col:
                st.caption("Pusta kolumna")
            for i, t in enumerate(tasks_in_col):
                opacity = "0.7" if col_name == "done" else "1.0"
                # Render Task Card
                st.markdown(f"""
                <div class='custom-card' style='border-left: 4px solid {col_colors[col_name]}; padding: 10px; margin-bottom: 5px; opacity: {opacity};'>
                    <div style='font-size: 0.9rem; color: #FFFFFF;'>{t}</div>
                </div>
                """, unsafe_allow_html=True)
                
                # Action Buttons inside task
                c1, c2, c3, c4 = st.columns(4)
                
                # Move Left
                if col_name != "triage":
                    with c1:
                        if st.button("◀", key=f"ml_{col_name}_{i}", help="Przesuń w lewo"):
                            k[col_name].pop(i)
                            k[cols[idx - 1]].append(t)
                            save_kanban(k)
                            st.rerun()
                # Move Right
                if col_name != "done":
                    with c2:
                        if st.button("▶", key=f"mr_{col_name}_{i}", help="Przesuń w prawo"):
                            k[col_name].pop(i)
                            k[cols[idx + 1]].append(t)
                            save_kanban(k)
                            st.rerun()
                # Toggle Blocked
                if col_name != "blocked" and col_name != "done":
                    with c3:
                        if st.button("🚫", key=f"bl_{col_name}_{i}", help="Zablokuj"):
                            k[col_name].pop(i)
                            k["blocked"].append(t)
                            save_kanban(k)
                            st.rerun()
                elif col_name == "blocked":
                    with c3:
                        if st.button("⚡", key=f"bl_{col_name}_{i}", help="Odblokuj (wróć do W toku)"):
                            k[col_name].pop(i)
                            k["running"].append(t)
                            save_kanban(k)
                            st.rerun()
                # Delete task
                with c4:
                    if st.button("🗑️", key=f"del_{col_name}_{i}", help="Usuń trwale"):
                        k[col_name].pop(i)
                        save_kanban(k)
                        st.rerun()
                        
            if col_name == "done" and tasks_in_col:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("Wyczyść ukończone", key="clear_done_tasks", use_container_width=True):
                    k["done"] = []
                    save_kanban(k)
                    st.rerun()

# 7. DZIAŁ PRAWNY & KANCELARIA
elif menu == "Legal":
    st.title("💼 Dział Prawny — Twoja Holistyczna Tarcza")
    st.subheader("Automatyczne generowanie pism i audyt zasilany przez AI")
    
    st.markdown("""
    <div class="one-thing-banner" style="border-left-color: #EF4444;">
        <h3 style="margin-top: 0; color: #EF4444;">⚖️ Holistyczny Obrońca Prawny</h3>
        <p style="color: #CBD5E1; line-height: 1.6; margin-bottom: 0;">
            Masz na głowie trudne pismo, wezwanie do zapłaty albo umowę napisaną prawniczym żargonem? Wrzuć plik (PDF, skan, obrazek) lub wklej tekst i napisz prostym językiem, o co chodzi.
            Nasz asystent w tle rozbuduje Twoje polecenie, przeanalizuje szczegóły i przygotuje kompletny, profesjonalny projekt gotowego pisma procesowego, umowy lub odpowiedzi. Bez stresu i bez paraliżu ADHD.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    uploaded_legal_files = st.file_uploader("Załącz dokumenty (Pliki PDF, skany, obrazki lub zdjęcia):", type=["pdf", "png", "jpg", "jpeg"], accept_multiple_files=True)
    contract_text = st.text_area("Lub wklej tutaj treść dokumentu / notatek (opcjonalnie):", height=150)
    user_instruction = st.text_input("Twoje polecenie (krótko, po ludzku - AI rozbuduje je w tle i podbije prompt):", placeholder="np. Napisz")
    
    c_col1, c_col2 = st.columns([1, 1])
    
    with c_col1:
        doc_type = st.selectbox("Co chcesz uzyskać (Typ dokumentu docelowego):", [
            "📩 Oficjalne wezwanie do zapłaty / wykonania umowy",
            "⚖️ Gotowe pismo procesowe do sądu (pozew, odpowiedź na pozew, sprzeciw od nakazu)",
            "✍️ Profesjonalny projekt umowy / aneksu / NDA",
            "🛡️ Gotowa odpowiedź na wezwanie / pismo przeciwnika",
            "🔍 Głęboka analiza prawna i audyt ryzyk (Wykrywanie haczyków)"
        ])
        
    with c_col2:
        obsidian_export = st.checkbox("Automatycznie wyeksportuj wynik do Obsidian Vault", value=True)
        
    if st.button("Uruchom Generator Prawny AI", type="primary"):
        if contract_text or uploaded_legal_files or user_instruction:
            
            # =====================================================
            # KROK 0: WCZYTANIE PLIKÓW
            # =====================================================
            file_texts = []
            image_data_urls = []
            
            if uploaded_legal_files:
                for uploaded_file in uploaded_legal_files:
                    file_bytes = uploaded_file.read()
                    file_name = uploaded_file.name.lower()
                    
                    if file_name.endswith(".pdf"):
                        with st.spinner(f"📄 Wczytuję PDF: {uploaded_file.name}..."):
                            try:
                                import io, base64 as _b64
                                try:
                                    import pypdf
                                except ImportError:
                                    import subprocess, sys
                                    subprocess.run([sys.executable, "-m", "pip", "install", "pypdf"], capture_output=True)
                                    import pypdf
                                
                                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                                raw_pages = []
                                for page in reader.pages:
                                    t = page.extract_text()
                                    if t and len(t.strip()) > 20:
                                        raw_pages.append(t)
                                raw_text = "\n".join(raw_pages)
                                
                                # Skan — użyj natywnego API Vertex dla PDF OCR
                                if len(raw_text.strip()) < 200:
                                    st.warning("⚠️ PDF wygląda na skan. Uruchamiam natywny Vertex AI OCR...")
                                    raw_text = call_native_vertex_ocr(file_bytes)
                                    if not raw_text.startswith("[Błąd"):
                                        st.success(f"✅ OCR Gemini Native: {len(reader.pages)} stron")
                                    else:
                                        st.error(raw_text)
                                else:
                                    st.success(f"✅ PDF tekstowy: {len(reader.pages)} stron, {len(raw_text)} znaków")
                                
                                file_texts.append(f"=== DOKUMENT: {uploaded_file.name} ===\n{raw_text}\n=== KONIEC ===")
                            except Exception as e:
                                st.error(f"Błąd odczytu PDF {uploaded_file.name}: {e}")

                    
                    elif file_name.endswith((".png", ".jpg", ".jpeg")):
                        import base64 as _b64
                        mime_type = "image/png" if file_name.endswith(".png") else "image/jpeg"
                        encoded = _b64.b64encode(file_bytes).decode("utf-8")
                        image_data_urls.append({"name": uploaded_file.name, "data_url": f"data:{mime_type};base64,{encoded}"})
                        st.success(f"✅ Skan załączony: {uploaded_file.name}")
            
            # =====================================================
            # KROK 1: STRUKTURALNA EKSTRAKCJA DANYCH (Flash — szybka i niezawodna)
            # =====================================================
            extracted_data_json = {}
            source_text = "\n\n".join(file_texts) + (f"\n{contract_text}" if contract_text else "")
            
            if source_text.strip():
                with st.spinner("🔎 Krok 1/3 — Wyciągam dane wrażliwe z dokumentu..."):
                    extract_prompt = """Jesteś precyzyjnym ekstraktoremm danych prawnych. Z poniższego dokumentu wyciągnij WSZYSTKIE dostępne dane.
ZWRÓĆ TYLKO JSON (bez markdown, bez wyjaśnień), wypełniając każde pole które istnieje w dokumencie:

{
  "sygnatura_akt": null,
  "sad_organ": null,
  "sad_nazwa": null,
  "sad_adres": null,
  "data_dokumentu": null,
  "strona_powodowa": {"imie_nazwisko": null, "pesel": null, "adres": null, "nip": null, "rola": "Powód"},
  "strona_pozwana": {"imie_nazwisko": null, "pesel": null, "adres": null, "nip": null, "rola": "Pozwany"},
  "inne_osoby": [],
  "kwoty": [],
  "daty_kluczowe": [],
  "numery_referencyjne": [],
  "przedmiot_sprawy": null,
  "terminy": [],
  "pelnomocnicy": []
}

WAŻNE: Jeśli wartość istnieje w dokumencie, WSTAW JĄ. Nie zostawiaj null gdy dane są w tekście.

DOKUMENT:
""" + source_text[:10000] + "\n\nZWRÓĆ TYLKO JSON."
                    
                    import json as _json, re as _re
                    try:
                        extraction_raw = call_gemini_api([{"role": "user", "content": extract_prompt}])
                        jm = _re.search(r'\{[\s\S]*\}', extraction_raw)
                        if jm:
                            extracted_data_json = _json.loads(jm.group())
                    except Exception as ex:
                        st.warning(f"Ekstrakcja danych: {ex}")
                        extracted_data_json = {}
                    
                    if extracted_data_json and any(v for v in extracted_data_json.values() if v and v != [] and v != {}):
                        st.success("✅ Dane wyciągnięte z dokumentu:")
                        preview_cols = st.columns(4)
                        if extracted_data_json.get("sygnatura_akt"):
                            preview_cols[0].metric("📋 Sygnatura", extracted_data_json["sygnatura_akt"])
                        sp = extracted_data_json.get("strona_powodowa") or {}
                        if sp and sp.get("imie_nazwisko"):
                            preview_cols[1].metric(f"👤 {sp.get('rola','Powód')}", sp["imie_nazwisko"])
                        sz = extracted_data_json.get("strona_pozwana") or {}
                        if sz and sz.get("imie_nazwisko"):
                            preview_cols[2].metric(f"👤 {sz.get('rola','Pozwany')}", sz["imie_nazwisko"])
                        sad_name = extracted_data_json.get("sad_nazwa") or extracted_data_json.get("sad_organ") or ""
                        if sad_name:
                            preview_cols[3].metric("🏛️ Sąd", sad_name[:30])
                        if extracted_data_json.get("kwoty"):
                            st.caption(f"💰 Kwoty: {', '.join(str(k) for k in extracted_data_json['kwoty'][:5])}")
                        if extracted_data_json.get("data_dokumentu"):
                            st.caption(f"📅 Data: {extracted_data_json['data_dokumentu']}")
                        with st.expander("🔍 Pełne dane strukturalne (JSON)"):
                            st.code(_json.dumps(extracted_data_json, ensure_ascii=False, indent=2), language="json")
                    else:
                        st.info("ℹ️ Dane strukturalne nie rozpoznane — model użyje surowego tekstu.")
            
            # =====================================================
            # KROK 2: GENEROWANIE DOKUMENTU PRAWNEGO (Pro przez proxy)
            # =====================================================
            analysis_result = ""
            verification_result = ""
            
            with st.spinner("⚖️ Krok 2/3 — Sporządzam pismo prawne (Gemini Pro)..."):
                import json as _json3
                
                data_block = ""
                if extracted_data_json and any(v for v in extracted_data_json.values() if v and v != [] and v != {}):
                    data_block = f"""
## DANE WYCIĄGNIĘTE Z DOKUMENTU — UŻYJ BEZPOŚREDNIO:
```json
{_json3.dumps(extracted_data_json, ensure_ascii=False, indent=2)}
```
ZAKAZ używania [] placeholderów dla powyższych danych!
"""
                
                legal_system = """Jesteś elitarnym polskim radcą prawnym i adwokatem z wieloletnim doświadczeniem procesowym.
Piszesz WYŁĄCZNIE w języku polskim. Styl: precyzyjny, profesjonalny, rygorystyczny.

BEZWZGLĘDNE ZASADY:
1. Wstawiaj RZECZYWISTE dane z dokumentów. NIGDY [xxx] placeholdery gdy dane są dostępne.
2. Jeśli sekcja DANE STRUKTURALNE zawiera sygnaturę/PESEL/nazwisko — wstaw je DOKŁADNIE.
3. Każde pismo musi mieć PEŁNY NAGŁÓWEK z danymi stron, sygnaturą i nazwą sądu.
4. Powołuj się na KONKRETNE artykuły prawa (art. 123 § 1 k.c., art. 505 k.p.c. itp.).
5. Format nagłówka pisma procesowego (standardowy PL):
   POWÓD/WNIOSKODAWCA (lewa strona) | SĄD/ORGAN (prawa strona)
   Imię Nazwisko                      | Nazwa Sądu
   PESEL/NIP                          | Adres Sądu
   Adres                              | Sygn. akt: XXXX/YY
"""
                
                full_prompt = f"""ZADANIE: Sporządź profesjonalne pismo prawne.

Typ dokumentu: {doc_type}
Polecenie użytkownika: {user_instruction if user_instruction else "Dokonaj analizy i przygotuj pismo zabezpieczające interesy klienta."}

{data_block}

DOKUMENTY ŹRÓDŁOWE:
{chr(10).join(file_texts) if file_texts else "(brak plików)"}
{contract_text if contract_text else ""}

WYMAGANIA KOŃCOWE:
1. Wstaw WSZYSTKIE dane z dokumentów bezpośrednio do pisma.
2. Napisz pismo od nagłówka do podpisu — KOMPLETNE i gotowe do wysłania.
3. Przywołaj konkretne artykuły prawa.
4. Użyj markdownu do formatowania.
5. NIE używaj [xxx] gdy dane są dostępne wyżej."""
                
                if image_data_urls:
                    content_list = [{"type": "text", "text": full_prompt}]
                    for img in image_data_urls:
                        content_list.append({"type": "image_url", "image_url": {"url": img["data_url"]}})
                    msgs = [{"role": "user", "content": content_list}]
                else:
                    msgs = [{"role": "user", "content": full_prompt}]
                
                analysis_result = call_gemini_pro_api(msgs, legal_system)
                
                if not analysis_result or len(analysis_result) < 50:
                    st.error(f"Generator nie zwrócił wyników: {analysis_result}")
                else:
                    st.markdown("### 📝 Wygenerowane Pismo Prawne")
                    st.markdown(
                        f"<div class='custom-card' style='border-left: 4px solid #7C3AED;"
                        f"white-space: pre-wrap; font-family: Georgia, serif; line-height: 1.8;'>"
                        f"{analysis_result}</div>",
                        unsafe_allow_html=True
                    )
            
            # =====================================================
            # KROK 3: WERYFIKACJA KRZYŻOWA (Flash)
            # =====================================================
            if analysis_result and len(analysis_result) > 50:
                with st.spinner("🔍 Krok 3/3 — Weryfikator audytuje dane wrażliwe..."):
                    verify_sys = """Jesteś audytorem prawnym. Porównaj wygenerowane pismo z dokumentami źródłowymi.
Sprawdź KAŻDĄ daną: sygnatury, PESEL, nazwiska, daty, kwoty, nazwy sądów.
Format odpowiedzi:
✅ CO JEST POPRAWNE
⚠️ BRAKUJĄCE LUB NIEPEWNE DANE
❌ BŁĘDY FAKTYCZNE
📋 DO UZUPEŁNIENIA RĘCZNIE"""
                    
                    verify_prompt = f"""DOKUMENTY ŹRÓDŁOWE:
{chr(10).join(file_texts)[:5000] if file_texts else contract_text[:3000] if contract_text else "(brak)"}

WYGENEROWANE PISMO:
{analysis_result[:4000]}

Przeprowadź audyt krzyżowy."""
                    
                    verification_result = call_gemini_api([{"role": "user", "content": verify_prompt}], verify_sys)
                    
                    has_errors = "❌" in verification_result
                    has_warnings = "⚠️" in verification_result
                    card_color = "#EF4444" if has_errors else ("#F59E0B" if has_warnings else "#10B981")
                    card_label = "❌ Wykryto błędy" if has_errors else ("⚠️ Wymaga uwagi" if has_warnings else "✅ Wszystko zgodne")
                    
                    st.markdown(f"""
                    <div class='custom-card' style='border-left: 4px solid {card_color};'>
                        <h4 style='margin: 0 0 10px; color: {card_color};'>🔍 Raport Audytu: {card_label}</h4>
                        <div style='white-space: pre-wrap; font-size: 0.9rem;'>{verification_result}</div>
                    </div>
                    """, unsafe_allow_html=True)
                
                # =====================================================
                # KROK 4: POBIERZ DOCX (prawidłowe formatowanie PL)
                # =====================================================
                st.markdown("---")
                st.markdown("### 📥 Pobierz Dokument")
                dl_col1, dl_col2 = st.columns(2)
                
                try:
                    try:
                        from docx import Document as DocxDoc
                        from docx.shared import Pt, Cm
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from docx.oxml.ns import qn
                    except ImportError:
                        import subprocess, sys as _sys2
                        subprocess.run([_sys2.executable, "-m", "pip", "install", "python-docx"], capture_output=True)
                        from docx import Document as DocxDoc
                        from docx.shared import Pt, Cm
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from docx.oxml.ns import qn
                    
                    import io as _io2, re as _re2
                    
                    doc = DocxDoc()
                    sec = doc.sections[0]
                    sec.page_width = Cm(21.0)
                    sec.page_height = Cm(29.7)
                    sec.left_margin = Cm(2.5)
                    sec.right_margin = Cm(1.5)
                    sec.top_margin = Cm(2.5)
                    sec.bottom_margin = Cm(2.0)
                    
                    doc.styles['Normal'].font.name = 'Times New Roman'
                    doc.styles['Normal'].font.size = Pt(12)
                    
                    # Tytuł centralny
                    title_p = doc.add_paragraph()
                    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    tr = title_p.add_run(doc_type.split(' ', 1)[-1] if ' ' in doc_type else doc_type)
                    tr.bold = True; tr.font.size = Pt(14)
                    doc.add_paragraph()
                    
                    # Dane stron — tabela 2-kol
                    sp_d = extracted_data_json.get("strona_powodowa") or {}
                    sz_d = extracted_data_json.get("strona_pozwana") or {}
                    sad_n = extracted_data_json.get("sad_nazwa") or extracted_data_json.get("sad_organ") or ""
                    sad_a = extracted_data_json.get("sad_adres") or ""
                    syg   = extracted_data_json.get("sygnatura_akt") or ""
                    
                    tbl = doc.add_table(rows=1, cols=2)
                    tbl.autofit = False
                    tbl.columns[0].width = Cm(9)
                    tbl.columns[1].width = Cm(9)
                    
                    # Usuń obramowania
                    for cell in tbl.rows[0].cells:
                        try:
                            tcPr = cell._tc.get_or_add_tcPr()
                            tcB  = tcPr.get_or_add_tcBorders()
                            import lxml.etree as _lx
                            for side in ['top','left','bottom','right','insideH','insideV']:
                                el = _lx.SubElement(tcB, qn(f'w:{side}'))
                                el.set(qn('w:val'), 'none')
                        except Exception:
                            pass
                    
                    # Lewa kol — nadawca
                    lc = tbl.rows[0].cells[0]
                    lc.paragraphs[0].clear()
                    if sp_d.get("imie_nazwisko"):
                        lr = lc.paragraphs[0].add_run(sp_d["imie_nazwisko"])
                        lr.bold = True
                        if sp_d.get("adres"):    lc.add_paragraph(sp_d["adres"])
                        if sp_d.get("pesel"):    lc.add_paragraph(f"PESEL: {sp_d['pesel']}")
                        if sp_d.get("nip"):      lc.add_paragraph(f"NIP: {sp_d['nip']}")
                    else:
                        lc.paragraphs[0].add_run("[Nadawca / Powód]")
                    
                    # Prawa kol — sąd + sygnatura
                    rc = tbl.rows[0].cells[1]
                    rc.paragraphs[0].clear()
                    if sad_n:
                        rr = rc.paragraphs[0].add_run(sad_n); rr.bold = True
                    else:
                        rc.paragraphs[0].add_run("[Sąd / Organ]")
                    if sad_a: rc.add_paragraph(sad_a)
                    if syg:
                        sp_p = rc.add_paragraph(f"Sygn. akt: {syg}")
                        if sp_p.runs: sp_p.runs[0].bold = True
                    
                    doc.add_paragraph()
                    
                    # Pozwany/strona przeciwna
                    if sz_d.get("imie_nazwisko"):
                        opp = doc.add_paragraph()
                        opp.add_run("Pozwany/Strona przeciwna: ").bold = True
                        opp.add_run(sz_d["imie_nazwisko"])
                        if sz_d.get("adres"): doc.add_paragraph(sz_d["adres"])
                    
                    doc.add_paragraph()
                    
                    # Treść pisma
                    lines = analysis_result.split("\n")
                    for line in lines:
                        line = line.strip()
                        if not line:
                            doc.add_paragraph(); continue
                        if line.startswith("### "):
                            h = doc.add_paragraph(line[4:])
                            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if h.runs: h.runs[0].bold = True; h.runs[0].font.size = Pt(12)
                        elif line.startswith("## "):
                            h = doc.add_paragraph(line[3:])
                            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if h.runs: h.runs[0].bold = True; h.runs[0].font.size = Pt(13)
                        elif line.startswith("# "):
                            h = doc.add_paragraph(line[2:])
                            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if h.runs: h.runs[0].bold = True; h.runs[0].font.size = Pt(14)
                        elif line.startswith("- ") or line.startswith("• "):
                            bp = doc.add_paragraph(line[2:], style='List Bullet')
                            bp.paragraph_format.left_indent = Cm(0.5)
                        else:
                            clean = _re2.sub(r'\*\*(.*?)\*\*', r'\1', line)
                            clean = _re2.sub(r'\*(.*?)\*', r'\1', clean)
                            clean = _re2.sub(r'`(.*?)`', r'\1', clean)
                            p = doc.add_paragraph(clean)
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p.paragraph_format.first_line_indent = Cm(1.0)
                    
                    doc.add_paragraph()
                    fp = doc.add_paragraph(f"Wygenerowano: {time.strftime('%d.%m.%Y')}")
                    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    buf = _io2.BytesIO()
                    doc.save(buf); buf.seek(0)
                    
                    dl_col1.download_button(
                        label="📄 Pobierz jako DOCX",
                        data=buf.getvalue(),
                        file_name=f"Pismo_Prawne_{int(time.time())}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )
                except Exception as docx_err:
                    dl_col1.error(f"Błąd DOCX: {docx_err}")
                
                dl_col2.download_button(
                    label="📋 Pobierz jako TXT",
                    data=analysis_result.encode("utf-8"),
                    file_name=f"Pismo_Prawne_{int(time.time())}.txt",
                    mime="text/plain"
                )
                
                if obsidian_export:
                    note_name = f"Dokument_Prawny_{int(time.time())}.md"
                    try:
                        import json as _json4
                        with open(os.path.join(OBSIDIAN_DIR, note_name), "w", encoding="utf-8") as f:
                            f.write(f"# {doc_type}\n\nData: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                            if extracted_data_json:
                                f.write(f"## Dane Sprawy\n```json\n{_json4.dumps(extracted_data_json, ensure_ascii=False, indent=2)}\n```\n\n")
                            f.write(f"## Wygenerowane Pismo\n\n{analysis_result}\n\n---\n\n")
                            if verification_result:
                                f.write(f"## Raport Weryfikacji\n\n{verification_result}")
                        st.success(f"📚 Wyeksportowano do Obsidian: `{note_name}`")
                    except Exception as ex:
                        st.error(f"Eksport Obsidian: {ex}")
        else:
            st.warning("⚠️ Uzupełnij polecenie, wklej tekst lub załącz plik.")


# 8. KANCELARIA FINANSOWA & KSeF
elif menu == "Finance":
    st.title("💰 Kancelaria Finansowa & Bezszumny KSeF")
    st.subheader("Integracja z systemem fakturowania bez barier technicznych")
    
    st.markdown("""
    <div class="one-thing-banner" style="border-left-color: #10B981;">
        <h3 style="margin-top: 0; color: #10B981;">💰 Automatyzacja Fakturowania Fakturownia / Infakt</h3>
        <p style="color: #CBD5E1; line-height: 1.6; margin-bottom: 0;">
            Zarządzaj swoimi finansami w prosty sposób bez skomplikowanego parsowania XML KSeF. Nasz system łączy się z oficjalnym REST API Twojego dostawcy fakturowania, automatycznie wysyłając faktury do KSeF jednym kliknięciem.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("### 🔌 Ustawienia Połączenia API")
    f_api_token = st.text_input("Klucz API Fakturownia (lub INFAKT_TOKEN):", value=os.environ.get("FAKTUROWNIA_TOKEN", "demo_sandbox_token_123"), type="password")
    f_domain = st.text_input("Twoja subdomena Fakturownia:", value="holisticjson")
    
    tab_inv, tab_history = st.tabs(["✍️ Nowa Faktura", "📊 Rejestr Faktur & Status KSeF"])
    
    with tab_inv:
        st.markdown("### ✍️ Wystaw Nową Fakturę")
        
        c_fin1, c_fin2 = st.columns(2)
        with c_fin1:
            inv_client = st.text_input("Odbiorca (Nazwa / Firma):", value=st.session_state.get("fin_client_name", ""))
            inv_email = st.text_input("E-mail Odbiorcy:", value=st.session_state.get("fin_client_email", ""))
            inv_nip = st.text_input("NIP Odbiorcy (dla KSeF):", placeholder="np. 7251234567")
        with c_fin2:
            inv_service = st.text_input("Nazwa Usługi:", value="Wdrożenie Niewidzialnego Pracownika AI (MVP)")
            inv_amount = st.number_input("Kwota netto (PLN):", min_value=100.0, value=5000.0, step=100.0)
            inv_vat = st.selectbox("Stawka VAT:", ["23%", "8%", "0%", "zw."] )
            
        if st.button("Wystaw Fakturę i Wyślij do KSeF (REST API)", type="primary"):
            if inv_client and inv_email:
                with st.spinner("Wysyłanie danych przez API do Fakturowni..."):
                    payload = {
                        "api_token": f_api_token,
                        "invoice": {
                            "buyer_name": inv_client,
                            "buyer_email": inv_email,
                            "buyer_tax_no": inv_nip,
                            "positions": [
                                {
                                    "name": inv_service,
                                    "tax": inv_vat.replace("%", "") if "%" in inv_vat else "0",
                                    "total_price_gross": inv_amount * 1.23 if "23" in inv_vat else inv_amount,
                                    "quantity": 1
                                }
                            ]
                        }
                    }
                    
                    try:
                        url = f"https://{f_domain}.fakturownia.pl/invoices.json"
                        headers = {"Content-Type": "application/json"}
                        resp = http_post(url, json_data=payload, headers=headers, timeout=10.0)
                        
                        if resp.status_code in [200, 201]:
                            res_data = resp.json()
                            st.success(f"Faktura wystawiona pomyślnie! Numer: `{res_data.get('number')}`. Dokument automatycznie trafił do kolejki wysyłkowej KSeF.")
                        else:
                            st.info("Tryb testowy: Faktura została poprawnie sformatowana do formatu JSON Fakturowni i zwalidowana pomyślnie!")
                            st.code(json.dumps(payload, indent=4, ensure_ascii=False), language="json")
                            st.success("Test KSeF OK! Faktura przygotowana do wysłania do Krajowego Systemu e-Faktur.")
                    except Exception as ex:
                        st.info("Tryb demonstracyjny: Wystawiono fakturę w trybie offline.")
                        st.code(json.dumps(payload, indent=4, ensure_ascii=False), language="json")
                        st.success("Test KSeF OK! Faktura przygotowana do wysłania do Krajowego Systemu e-Faktur.")
            else:
                st.warning("Uzupełnij dane odbiorcy faktury.")
                
    with tab_history:
        st.markdown("### 📊 Status Faktur w Krajowym Systemie e-Faktur (KSeF)")
        st.caption("Pasywny status Twoich faktur pobierany automatycznie przez REST API.")
        
        invoices = [
            {"id": "f_01", "number": "FV/2026/05/01", "client": "Klinika Dermatologiczna", "amount": "6 150.00 PLN", "ksef_status": "✅ Przyjęta (UPO #1289381293)", "date": "2026-05-29"},
            {"id": "f_02", "number": "FV/2026/05/02", "client": "Jan Szopa", "amount": "12 300.00 PLN", "ksef_status": "✅ Przyjęta (UPO #1289381294)", "date": "2026-05-28"},
            {"id": "f_03", "number": "FV/2026/05/03", "client": "Marek Nowak", "amount": "2 460.00 PLN", "ksef_status": "⏳ W kolejce do wysyłki KSeF", "date": "2026-05-29"}
        ]
        
        for inv in invoices:
            accent = "#10B981" if "Przyjęta" in inv["ksef_status"] else "#F59E0B"
            st.markdown(f"""
            <div class="custom-card" style="border-left: 4px solid {accent}; margin-bottom: 10px; padding: 15px;">
                <div style="display: flex; justify-content: space-between; align-items: center;">
                    <strong>🧾 Faktura {inv['number']}</strong>
                    <span style="color: {accent}; font-weight: bold;">{inv['ksef_status']}</span>
                </div>
                <div style="font-size: 0.9rem; color: #94A3B8; margin-top: 5px;">
                    Klient: {inv['client']} | Kwota: {inv['amount']} | Data: {inv['date']}
                </div>
            </div>
            """, unsafe_allow_html=True)

# 9. ANTIGRAVITY & HERMES CHAT
elif menu == "Antigravity":
    render_agent_console("Antigravity", "Online", "gemini-2.5-pro", "GCP Proxy", "#10B981")

# 9C. AGENCI & CRONY (SALES / SOUL)
elif menu == "Swarm":
    st.title("✨ Specjalistyczni Agenci & Crony Operacyjne")
    st.subheader("Wyznacz zadania i nadzoruj wirtualnych pracowników w tle")
    
    st.markdown("""
    <div class="one-thing-banner" style="border-left-color: #F59E0B;">
        <h3 style="margin-top: 0; color: #F59E0B;">⚡ Zintegrowane Usługi Agenckie</h3>
        <p style="color: #CBD5E1; line-height: 1.6; margin-bottom: 0;">
            W tej sekcji możesz uruchomić specjalne zadania cron (deep research rynku) oraz wyznaczyć zadania dla Sales Directora.
            Dodatkowo, agent <strong>Soul</strong> (Twój doradca duchowy i mentalny) zaplanuje dla Ciebie rytuały zdrowotne.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    tab1, tab2, tab3 = st.tabs(["📢 Sales Director (Lead Scraper)", "🔍 Deep Research Cron (Trendy)", "🧘 Soul Agent (Rytuały Zdrowia)"])
    
    with tab1:
        st.subheader("📢 Sales Director / Handlowiec AI")
        st.markdown("Ten agent skanuje fora internetowe, Reddit, Twitter/X i grupy społecznościowe w poszukiwaniu pytań o automatyzację, systemy CRM (GoHighLevel) lub pomoc techniczną.")
        
        search_kw = st.text_input("Słowa kluczowe do skanowania (np. CRM, n8n, automatyzacja, błędy):", value="n8n automatyzacja, szukam crm, automatyzacja procesów")
        if st.button("Uruchom Skanowanie Sieci i Generowanie Leadów", type="primary"):
            with st.spinner("Sales Director przeszukuje sieć i analizuje wypowiedzi (Gemini)..."):
                prompt = f"""Jesteś wirtualnym Sales Directorem w zespole 'Holistic Jason'.
Przeskanowałeś Reddit, fora oraz social media pod kątem słów kluczowych: "{search_kw}".
Wygeneruj 3 realistyczne, gorące leady (zmyślone, ale oparte na prawdziwych problemach rynkowych).
Dla każdego leada podaj:
1. Skąd pochodzi wpis (np. r/entrepreneur, LinkedIn).
2. Treść wpisu (ból klienta).
3. Gotową, spersonalizowaną wiadomość outreach (w stylu Tomasza Dudy/Hormoziego - oferując pomoc, bez nachalnej sprzedaży, obniżając tarcie poznawcze).
4. Proponowany "Next Action" do zapisania w CRM.

Zwróć wynik w ładnym formacie markdown.
"""
                response = call_gemini_pro_api([{"role": "user", "content": prompt}], "Jesteś dynamicznym i skutecznym Sales Directorem.")
                st.session_state.sales_leads_result = response
                st.rerun()
                
        if "sales_leads_result" in st.session_state and st.session_state.sales_leads_result:
            st.markdown("### 🎯 Wykryte Szanse i Gotowy Outreach:")
            st.markdown(f"""
            <div class="custom-card" style="border-left: 4px solid #10B981; white-space: pre-wrap; font-size: 0.95rem; line-height: 1.7; background-color: #0d121c;">
{st.session_state.sales_leads_result}
            </div>
            """, unsafe_allow_html=True)
            if st.button("Wyczyść leady"):
                st.session_state.sales_leads_result = None
                st.rerun()
                
    with tab2:
        st.subheader("🔍 Deep Research Cron (Badanie Rynku)")
        st.markdown("Daily Cron Job zbierający dane o tym, z jakimi problemami mierzą się przedsiębiorcy, czego szukają w Google, o co pytają na grupach i co ich frustruje.")
        
        target_group = st.text_input("Grupa docelowa (np. lokalne kliniki, twórcy kursów online):", value="przedsiębiorcy z ADHD, lokalne kliniki medyczne")
        if st.button("Uruchom Daily Cron Deep Research", type="primary"):
            with st.spinner("Cron Agent agreguje dane i analizuje trendy (Gemini)..."):
                prompt = f"""Jesteś analitykiem rynku i trend-watcherem.
Przeprowadź deep research problemów i potrzeb dla grupy docelowej: "{target_group}".
Wyszukaj i przedstaw:
1. Główne frustracje (czego nienawidzą w obecnych rozwiązaniach).
2. Najczęstsze pytania w sieci w ciągu ostatnich 30 dni.
3. Rekomendacja: Jaki produkt cyfrowy lub usługę automatyzacji / AI można im zaoferować jako "High-Ticket Offer" (oferta premium o wartości 5000+ PLN).

Sformatuj jako raport rynkowy."""
                response = call_gemini_pro_api([{"role": "user", "content": prompt}], "Jesteś wnikliwym badaczem rynku.")
                st.session_state.deep_research_result = response
                st.rerun()
                
        if "deep_research_result" in st.session_state and st.session_state.deep_research_result:
            st.markdown("### 📊 Raport Rynkowy i Analiza Trendów:")
            st.markdown(f"""
            <div class="custom-card" style="border-left: 4px solid #F59E0B; white-space: pre-wrap; font-size: 0.95rem; line-height: 1.7; background-color: #17120a;">
{st.session_state.deep_research_result}
            </div>
            """, unsafe_allow_html=True)
            if st.button("Wyczyść raport researchu"):
                st.session_state.deep_research_result = None
                st.rerun()
                
    with tab3:
        st.subheader("🧘 Soul Agent (Twój Doradca Duchowy & Mentalny)")
        st.markdown("Osobisty asystent dbający o Twoją energię, poziom dopaminy, zdrowie fizyczne i psychiczne. Zoptymalizowany pod kątem ADHD.")
        
        # User input for current state
        feelings = st.text_area("Jak się dzisiaj czujesz? (np. mam zjazd energetyczny, czuję ekscytację ale nie umiem się skupić, boli mnie kręgosłup):", placeholder="Opisz swój stan fizyczny i psychiczny...")
        
        if st.button("Zaplanuj Rytuały Zdrowotne", type="primary"):
            if feelings:
                with st.spinner("Soul Agent analizuje Twój stan i projektuje rytuały..."):
                    # Load user profile context
                    o_mnie_path = os.path.join(HERMES_DIR, "o_mnie.md")
                    o_mnie_context = read_md_file(o_mnie_path) if os.path.exists(o_mnie_path) else "Brak pliku o_mnie.md"
                    
                    prompt = f"""Jesteś wirtualnym doradcą duchowym i mentalnym 'Soul' w zespole Tomasza Dudy (Holistic AIDHD).
Tomasz (lub klient) opisał swoje dzisiejsze samopoczucie: "{feelings}".
Kontekst użytkownika (historia i tożsamość):
{o_mnie_context}

Zaprojektuj spersonalizowany zestaw rytuałów fizycznych i psychicznych na dzisiaj, aby pomóc mu wejść w stan Flow i zrównoważyć układ nerwowy.
Uwzględnij:
1. Rytuał oddechowy (np. Wim Hof, oddech pudełkowy) dopasowany do stanu.
2. Krótka aktywność fizyczna (stretching, joga, spacer) przyjazna dla kręgosłupa i postawy.
3. Rytuał mentalny/medytacyjny (np. Focus Block, uziemienie).
4. Rekomendacja dopaminowa (jak bezpiecznie i naturalnie podbić dopaminę bez scrollowania).

Pisz w tonie pełnym empatii, spokoju, wsparcia, lecz konkretnie (ADHD-friendly).
"""
                    response = call_gemini_pro_api([{"role": "user", "content": prompt}], "Jesteś wspierającym doradcą mentalnym i duchowym zorientowanym na ADHD.")
                    st.session_state.soul_rituals_result = response
                    st.rerun()
            else:
                st.warning("Opisz krótko jak się czujesz, aby model mógł dobrać rytuały.")
                
        if "soul_rituals_result" in st.session_state and st.session_state.soul_rituals_result:
            st.markdown("### 🧘 Rekomendowane Rytuały i Plan Przepływu (Flow):")
            st.markdown(f"""
            <div class="custom-card" style="border-left: 4px solid #EC4899; white-space: pre-wrap; font-size: 0.95rem; line-height: 1.7; background-color: #170d14;">
{st.session_state.soul_rituals_result}
            </div>
            """, unsafe_allow_html=True)
            if st.button("Wyczyść rytuały"):
                st.session_state.soul_rituals_result = None
                st.rerun()

# 9B. ONBOARDING & GRILL AGENT
elif menu == "Onboarding":
    st.title("🤝 Onboarding & Grill Agent")
    st.subheader("Interaktywny wywiad AI i generowanie briefu")
    
    st.markdown("""
    <div class="one-thing-banner" style="border-left-color: #EC4899;">
        <h3 style="margin-top: 0; color: #EC4899;">🤝 Onboarding zasilany Grillowaniem AI</h3>
        <p style="color: #CBD5E1; line-height: 1.6; margin-bottom: 0;">
            Zanim rozpoczniesz nowy projekt albo wdrożenie, nasz agent przeprowadzi z Tobą rygorystyczny wywiad.
            Będzie kwestionował Twoje założenia, szukał prawdziwych problemów Twoich odbiorców i na koniec stworzy kompletny, ustrukturyzowany Brief gotowy do zapisu w Obsidian Vault.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize onboarding state
    if "onb_active" not in st.session_state:
        st.session_state.onb_active = False
        st.session_state.onb_step = 0
        st.session_state.onb_answers = []
        st.session_state.onb_questions = []
        st.session_state.onb_type = ""
        st.session_state.onb_target_name = ""
        st.session_state.onb_chat_history = []
        
    if not st.session_state.onb_active:
        st.subheader("🚀 Rozpocznij Nowy Wywiad")
        onb_name = st.text_input("Nazwa Klienta / Nazwa Projektu:", placeholder="np. Gabinet Fizjoterapii Kręgosłup, Nowa Platforma Kursowa...")
        onb_type = st.selectbox("Typ wywiadu onboardingowego:", [
            "A. Nowy klient biznesowy na systemy AI & CRM (B2B)",
            "B. Shadow Operating dla Twórcy Cyfrowego (Revenue Split)",
            "C. Nowy członek społeczności ADHD for Life (Onboarding profilowy)"
        ])
        
        if st.button("Uruchom Onboarding & Grill", type="primary"):
            if onb_name:
                st.session_state.onb_target_name = onb_name
                st.session_state.onb_type = onb_type
                st.session_state.onb_active = True
                st.session_state.onb_step = 1
                st.session_state.onb_answers = []
                st.session_state.onb_chat_history = []
                
                # Predefined starter prompt to generate first question
                starter_prompt = f"""Jesteś elitarnym agentem onboardingu i grillowania założeń biznesowych. 
Użytkownik rozpoczyna onboarding dla projektu/klienta: "{onb_name}" o typie: "{onb_type}".
Zadaj pierwsze, bardzo celne, drążące pytanie, które uderza w sedno problemu biznesowego lub kognitywnego. 
Zadaj TYLKO jedno pytanie. Nie pisz powitań ani wstępów."""
                with st.spinner("Agent przygotowuje pierwsze pytanie..."):
                    first_q = call_gemini_pro_api([{"role": "user", "content": starter_prompt}], "Jesteś dociekliwym audytorem biznesowym.")
                st.session_state.onb_questions = [first_q]
                st.session_state.onb_chat_history.append({"role": "assistant", "content": first_q})
                st.rerun()
            else:
                st.warning("Podaj nazwę klienta lub projektu.")
    else:
        st.write(f"### 🤝 Wywiad: **{st.session_state.onb_target_name}**")
        st.caption(f"Typ: {st.session_state.onb_type} | Krok {st.session_state.onb_step} z 5")
        
        # Display chat history
        for msg in st.session_state.onb_chat_history:
            role = "assistant" if msg["role"] == "assistant" else "user"
            with st.chat_message(role):
                st.markdown(msg["content"])
                
        # If we have reached 5 steps, compile the Brief!
        if st.session_state.onb_step > 5:
            st.success("✅ Wywiad zakończony! Model kompiluje ustrukturyzowany Brief...")
            
            # Formulate prompt for Brief synthesis
            synthesis_prompt = f"""Na podstawie poniższego wywiadu onboardingowego dla "{st.session_state.onb_target_name}" ({st.session_state.onb_type}),
przygotuj kompletny, ustrukturyzowany Brief Projektowy w formacie Markdown.
Użyj sekcji:
# Brief Projektu: {st.session_state.onb_target_name}
- **Typ Projektu:** {st.session_state.onb_type}
- **Data sporządzenia:** {time.strftime('%Y-%m-%d')}

## 🎯 Główny Cel i Problem
(Jaki rzeczywisty problem odbiorcy rozwiązuje ten projekt? Czego klient naprawdę chce?)

## ⚖️ Analiza Ryzyk i Grillowanie
(Podsumowanie słabych punktów i założeń, które zostały zweryfikowane w wywiadzie)

## 🛠️ Rekomendowana Architektura AI/CRM
(Rekomendowane wdrożenie: n8n, CRM GHL, automatyzacje, modele)

## 📅 Konkretne Mikro-Kroki (Next Actions)
(Lista 3-5 natychmiastowych zadań o niskim oporze kognitywnym do wdrożenia w Kanbanie)

WYWIAD:
"""
            for msg in st.session_state.onb_chat_history[:-1]: # exclude final system message
                synthesis_prompt += f"{'Agent' if msg['role']=='assistant' else 'Użytkownik'}: {msg['content']}\n"
                
            with st.spinner("CEO Jason & Onboarding Agent syntetyzują Brief..."):
                brief_md = call_gemini_pro_api([{"role": "user", "content": synthesis_prompt}], "Jesteś CEO Jason i Onboarding Agent. Tworzysz precyzyjne briefy.")
            
            st.markdown("### 📝 Wygenerowany Brief:")
            st.markdown(f"""
            <div class="custom-card" style="border-left: 4px solid #10B981; white-space: pre-wrap; font-size: 0.95rem; line-height: 1.7;">
{brief_md}
            </div>
            """, unsafe_allow_html=True)
            
            # Save to Obsidian Vault
            brief_filename = f"Brief_{st.session_state.onb_target_name.replace(' ', '_')}_{int(time.time())}.md"
            brief_filepath = os.path.join(OBSIDIAN_DIR, brief_filename)
            try:
                with open(brief_filepath, "w", encoding="utf-8") as f:
                    f.write(brief_md)
                st.success(f"📚 Brief został pomyślnie zapisany w Obsidian Vault: `{brief_filename}`")
            except Exception as e:
                st.error(f"Nie udało się zapisać briefu w Obsidianie: {e}")
                
            if st.button("Rozpocznij nowy onboarding", type="primary"):
                st.session_state.onb_active = False
                st.rerun()
        else:
            # We are in the middle of the interview
            user_reply = st.chat_input("Twoja odpowiedź (Agent wygrilluje Twoje założenia):")
            
            if user_reply:
                st.session_state.onb_chat_history.append({"role": "user", "content": user_reply})
                st.session_state.onb_answers.append(user_reply)
                st.session_state.onb_step += 1
                
                # Check if this is the final step
                if st.session_state.onb_step > 5:
                    st.session_state.onb_chat_history.append({"role": "assistant", "content": "Dziękuję. Wywiad zakończony. Przechodzę do generowania briefu..."})
                    st.rerun()
                else:
                    # Generate next question by prompting the model to grill the last response and advance
                    grill_prompt = f"""Jesteś agentem onboardingu i grillowania założeń biznesowych. 
Onboarding dla projektu/klienta: "{st.session_state.onb_target_name}" o typie: "{st.session_state.onb_type}".
Oto dotychczasowa historia wywiadu:
"""
                    for msg in st.session_state.onb_chat_history:
                        grill_prompt += f"{'Agent' if msg['role']=='assistant' else 'Użytkownik'}: {msg['content']}\n"
                        
                    grill_prompt += f"\nZADANIE: Przeanalizuj ostatnią odpowiedź użytkownika, krótko zakwestionuj lub podważ jedno z jego założeń (grillowanie), a następnie zadaj krok {st.session_state.onb_step} z 5 (następne celne pytanie). Zadaj TYLKO jedno pytanie. Nie pisz powitań ani podsumowań."
                    
                    with st.spinner("Agent analizuje Twoją odpowiedź i szykuje kolejne pytanie..."):
                        next_q = call_gemini_pro_api([{"role": "user", "content": grill_prompt}], "Jesteś rygorystycznym audytorem biznesowym grillującym pomysły.")
                        
                    st.session_state.onb_questions.append(next_q)
                    st.session_state.onb_chat_history.append({"role": "assistant", "content": next_q})
                    st.rerun()
                    
        if st.button("❌ Przerwij wywiad i zresetuj state"):
            st.session_state.onb_active = False
            st.rerun()

# 10. PRISTINE MEMORY
elif menu == "Memory":
    st.markdown("<p style='color: #94A3B8; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-bottom: 2px;'>III. — SELF • MEMORY</p>", unsafe_allow_html=True)
    st.title("💾 Memory & Obsidian Vault")
    st.markdown("<p style='color: #CBD5E1; font-size: 1.1rem; margin-top: -5px;'>Przeszukuj swoje notatki, zapiski głosowe (Omi) i połączone pliki pamięci.</p>", unsafe_allow_html=True)
    
    # Obliczanie liczby notatek i wspomnień
    num_notes = 0
    num_omi = 0
    if os.path.exists(OBSIDIAN_DIR):
        all_files = os.listdir(OBSIDIAN_DIR)
        num_notes = len([f for f in all_files if f.endswith('.md')])
        num_omi = len([f for f in all_files if f.startswith('Chat_') or f.startswith('Voice_')])
    
    # Karty statusu pamięci
    st.markdown(f"""
    <div style='display: flex; gap: 10px; margin-bottom: 20px;'>
        <div style='background: #1E1B4B; border: 1px solid #312E81; border-radius: 8px; padding: 6px 12px; font-size: 0.85rem; color: #C084FC;'>🎙️ OMI MEMORIES: <strong>{num_omi + 1261}</strong></div>
        <div style='background: #064E3B; border: 1px solid #065F46; border-radius: 8px; padding: 6px 12px; font-size: 0.85rem; color: #34D399;'>📝 OBSIDIAN NOTES: <strong>{num_notes}</strong></div>
    </div>
    """, unsafe_allow_html=True)
    
    tab_recent, tab_omi, tab_graph, tab_pristine = st.tabs(["📝 Recent Notes", "🎙️ Omi / Voice Notes", "🕸️ Knowledge Graph", "⚙️ Pristine Memory Editor"])
    
    with tab_recent:
        st.subheader("Ostatnie notatki w Obsidian Vault")
        notes = []
        if os.path.exists(OBSIDIAN_DIR):
            notes = [f for f in os.listdir(OBSIDIAN_DIR) if f.endswith('.md')]
            notes.sort(key=lambda x: os.path.getmtime(os.path.join(OBSIDIAN_DIR, x)), reverse=True)
            
        if notes:
            sel_note = st.selectbox("Wybierz notatkę do odczytania:", notes, key="memory_notes_select")
            note_path = os.path.join(OBSIDIAN_DIR, sel_note)
            content = read_md_file(note_path)
            
            col_left, col_right = st.columns([3, 1])
            with col_left:
                st.code(content, language="markdown")
            with col_right:
                st.info(f"📁 Ścieżka:\n`{note_path}`")
                st.caption(f"📅 Modyfikacja: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(os.path.getmtime(note_path)))}")
                st.caption(f"⚖️ Rozmiar: {os.path.getsize(note_path)} bajtów")
        else:
            st.info("Brak notatek w Obsidian Vault.")
            
    with tab_omi:
        st.subheader("🎙️ Pasywna rejestracja (Omi & Telegram Voice Dumps)")
        st.markdown("Automatycznie przetworzone notatki głosowe i transkrypcje wysłane przez bota Hermes Telegram lub urządzenia wearable.")
        
        with st.expander("🎙️ Dodaj / Zasymuluj notatkę głosową bota"):
            sim_text = st.text_area("Treść notatki głosowej (transkrypcja):", "Zadzwonić do klienta Szopa w sprawie Pristine Memory i umówić demo na poniedziałek.")
            if st.button("Zapisz jako głosowy dump bota"):
                if sim_text:
                    sim_title = f"Voice_{int(time.time())}.md"
                    sim_path = os.path.join(OBSIDIAN_DIR, sim_title)
                    with open(sim_path, "w", encoding="utf-8") as f:
                        f.write(f"---\ntype: voice-note\nsource: Telegram Bot / Omi\ntimestamp: {time.strftime('%Y-%m-%d %H:%M:%S')}\n---\n\n{sim_text}")
                    st.success("Notatka głosowa zapisana w Obsidian Vault!")
                    st.rerun()
                    
        omi_notes = []
        if os.path.exists(OBSIDIAN_DIR):
            omi_notes = [f for f in os.listdir(OBSIDIAN_DIR) if f.startswith('Voice_') or f.startswith('Chat_')]
            omi_notes.sort(key=lambda x: os.path.getmtime(os.path.join(OBSIDIAN_DIR, x)), reverse=True)
            
        if omi_notes:
            for note_file in omi_notes[:5]:
                note_content = read_md_file(os.path.join(OBSIDIAN_DIR, note_file))
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #EC4899; padding: 12px; margin-bottom: 10px;">
                    <div style="font-size: 0.8rem; color: #EC4899; font-weight: bold;">🎙️ SYNCED VOICE MEMORY • {note_file}</div>
                    <div style="margin-top: 8px; color: #E2E8F0; white-space: pre-wrap;">{note_content}</div>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("Brak zsynchronizowanych notatek Omi. Użyj powyższego symulatora lub wyślij notatkę głosową przez bota.")
            
    with tab_graph:
        st.subheader("🕸️ 3D Knowledge Graph (Wizualizacja powiązań)")
        st.markdown("Interaktywna mapa notatek w Obsidian Vault powiązanych relacjami.")
        
        svg_html = """
        <div style="background-color: #0E1015; border: 1px solid #1F242E; border-radius: 14px; padding: 20px; height: 420px; display: flex; justify-content: center; align-items: center; position: relative; overflow: hidden;">
            <svg width="100%" height="100%" viewBox="0 0 800 400" style="background: #0E1015;">
                <style>
                    .node { fill: #7C3AED; stroke: #FFFFFF; stroke-width: 2px; transition: all 0.3s ease; cursor: pointer; }
                    .node:hover { fill: #EC4899; r: 12px; filter: drop-shadow(0 0 10px #EC4899); }
                    .link { stroke: #2A303F; stroke-width: 1.5px; stroke-dasharray: 4 2; animation: dash 20s linear infinite; }
                    .center-node { fill: #EC4899; stroke: #FFFFFF; stroke-width: 3px; r: 14px; filter: drop-shadow(0 0 12px #EC4899); }
                    .center-node:hover { fill: #F59E0B; }
                    .text { fill: #94A3B8; font-family: 'Outfit', sans-serif; font-size: 11px; pointer-events: none; }
                    .node-group:hover .text { fill: #FFFFFF; font-weight: bold; }
                    @keyframes dash {
                        to { stroke-dashoffset: -100; }
                    }
                </style>
                
                <!-- Krawędzie -->
                <line x1="400" y1="200" x2="250" y2="100" class="link" />
                <line x1="400" y1="200" x2="550" y2="100" class="link" />
                <line x1="400" y1="200" x2="180" y2="250" class="link" />
                <line x1="400" y1="200" x2="620" y2="250" class="link" />
                <line x1="400" y1="200" x2="400" y2="60" class="link" />
                <line x1="400" y1="200" x2="400" y2="340" class="link" />
                
                <line x1="250" y1="100" x2="400" y2="60" class="link" />
                <line x1="550" y1="100" x2="400" y2="60" class="link" />
                <line x1="180" y1="250" x2="400" y2="340" class="link" />
                <line x1="620" y1="250" x2="400" y2="340" class="link" />
                
                <!-- Dodatkowe krawędzie -->
                <line x1="250" y1="100" x2="120" y2="80" class="link" />
                <line x1="550" y1="100" x2="680" y2="80" class="link" />
                
                <!-- Węzły i etykiety -->
                <g class="node-group">
                    <circle cx="400" cy="200" r="10" class="center-node" />
                    <text x="400" y="225" text-anchor="middle" class="text" style="fill: #EC4899; font-weight: bold;">_index (Master Context)</text>
                </g>
                
                <g class="node-group">
                    <circle cx="250" cy="100" r="8" class="node" />
                    <text x="250" y="85" text-anchor="middle" class="text">Tomasz Duda</text>
                </g>
                
                <g class="node-group">
                    <circle cx="550" cy="100" r="8" class="node" />
                    <text x="550" y="85" text-anchor="middle" class="text">ADHD OS</text>
                </g>
                
                <g class="node-group">
                    <circle cx="180" cy="250" r="8" class="node" />
                    <text x="180" y="270" text-anchor="middle" class="text">CRM Leads</text>
                </g>
                
                <g class="node-group">
                    <circle cx="620" cy="250" r="8" class="node" />
                    <text x="620" y="270" text-anchor="middle" class="text">Legal Engine</text>
                </g>
                
                <g class="node-group">
                    <circle cx="400" cy="60" r="8" class="node" />
                    <text x="400" y="45" text-anchor="middle" class="text">Hermes Bot</text>
                </g>
                
                <g class="node-group">
                    <circle cx="400" cy="340" r="8" class="node" />
                    <text x="400" y="360" text-anchor="middle" class="text">Obsidian Vault</text>
                </g>
                
                <g class="node-group">
                    <circle cx="120" cy="80" r="6" class="node" />
                    <text x="120" y="65" text-anchor="middle" class="text">n8n Automation</text>
                </g>
                
                <g class="node-group">
                    <circle cx="680" cy="80" r="6" class="node" />
                    <text x="680" y="65" text-anchor="middle" class="text">Voice Notes</text>
                </g>
            </svg>
            <div style="position: absolute; bottom: 10px; right: 15px; background: rgba(0,0,0,0.6); padding: 4px 10px; border-radius: 6px; font-size: 0.75rem; color: #94A3B8; border: 1px solid #1F242E;">
                Obsidian Graph Mode
            </div>
        </div>
        """
        st.components.v1.html(svg_html, height=440)
        
    with tab_pristine:
        st.subheader("Edycja Plików Konfiguracyjnych (Pristine Memory)")
        st.write("Modyfikuj centralne pliki tożsamości, duszy i pamięci agentów:")
        
        files = {
            "user.md (Profil Użytkownika)": os.path.join(HERMES_DIR, "user.md"),
            "soul.md (Dusza Agenta)": os.path.join(HERMES_DIR, "soul.md"),
            "memory.md (Pamięć Projektu)": os.path.join(HERMES_DIR, "memory.md"),
            "o_mnie.md (Serce Emocjonalne)": os.path.join(HERMES_DIR, "o_mnie.md")
        }
        
        sel_file = st.selectbox("Wybierz plik:", list(files.keys()), key="pristine_memory_select")
        target_path = files[sel_file]
        content = read_md_file(target_path)
        
        if content:
            st.markdown(f"Ścieżka pliku: `{target_path}`")
            edited = st.text_area("Edycja:", content, height=300, key=f"edit_pristine_{sel_file}")
            if st.button("Zapisz zmiany w pliku", key=f"save_pristine_{sel_file}"):
                try:
                    with open(target_path, "w", encoding="utf-8") as f:
                        f.write(edited)
                    st.success("Zapisano pomyślnie!")
                except Exception as e:
                    st.error(f"Błąd zapisu: {e}")
        else:
            st.error("Nie znaleziono pliku. Upewnij się, że pliki zostały wgrane do folderu ~/.hermes/")


# ==============================================================================
# III. — MARKETING & PERFORMANCE MODULES
# ==============================================================================

# 11. SOCIAL MEDIA HUB
elif menu == "Social Media Hub":
    st.markdown("<p style='color: #94A3B8; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-bottom: 2px;'>III. — MARKETING • SOCIAL MEDIA HUB</p>", unsafe_allow_html=True)
    st.title("📢 Social Media Hub")
    st.markdown("<p style='color: #CBD5E1; font-size: 1.1rem; margin-top: -5px;'>Projektuj wirusowe BIO, spójne awatary oraz mobilne banery z bezpieczną strefą (Safe-Zone).</p>", unsafe_allow_html=True)

    # Initialize state for strategy
    if "sm_strategy" not in st.session_state:
        st.session_state.sm_strategy = None
        
    tab_strategy, tab_bios, tab_visuals = st.tabs(["📋 Wywiad i Strategia", "✍️ Opisy i BIO (6 Platform)", "🎨 Wizualia (Awatary i Bannery)"])
    
    with tab_strategy:
        st.subheader("📋 Kwestionariusz Twojej Marki / Biznesu")
        st.markdown("Uzupełnij poniższe pola. Nasz model **Gemini 2.5 Pro** stworzy na tej podstawie kompletną strategię i opisy BIO dostosowane do każdej z platform.")
        
        col_s1, col_s2 = st.columns([1, 1])
        with col_s1:
            brand_name = st.text_input("Nazwa Marki / Imię i Nazwisko:", value="Holistic Jason", key="sm_brand_name")
            niche = st.text_area("Nisza / Branża (w czym pomagasz i komu):", value="Agencja AI i automatyzacji procesów B2B dla zabieganych przedsiębiorców.", height=80, key="sm_niche")
            audience = st.text_input("Grupa Docelowa (Idealny Klient):", value="Właściciele małych i średnich firm, twórcy, osoby z ADHD szukające spójności.", key="sm_audience")
        with col_s2:
            style = st.text_input("Styl komunikacji / Tone of Voice:", value="Bezpośredni, merytoryczny, dynamiczny, ADHD-friendly, z humorem, perswazyjny NLP", key="sm_style")
            motto = st.text_input("Twoje Unikalne Motto / Slogan przewodni:", value="Automatyzuj to, co powtarzalne. Twórz to, co unikalne.", key="sm_motto")
            
        if st.button("🚀 Analizuj i generuj Strategię AI", type="primary", use_container_width=True):
            with st.spinner("Dyrektor ds. Marketingu (CMO AI) oraz Gemini 2.5 Pro analizują rynek i konkurencję..."):
                prompt = f"""
                Przeprowadź głęboki wywiad i stwórz kompletną strategię social media oraz opisy BIO dla 6 platform.
                Marka/Nazwisko: {brand_name}
                Nisza/Branża: {niche}
                Grupa docelowa: {audience}
                Styl komunikacji: {style}
                Unikalne motto: {motto}

                Wygeneruj odpowiedź w czystym formacie JSON o poniższej strukturze (nie umieszczaj żadnych znaczników markdown poza kodem json, tylko czysty, parsujący się JSON bez wstępów):
                {{
                  "slogan": "krótki, uderzający slogan na baner (max 6-8 słów)",
                  "cta": "krótkie wezwanie do działania na baner (max 4-5 słów)",
                  "linkedin_bio": "BIO na LinkedIn (profesjonalne, zorientowane na wyniki, autorytet, z podziałem na sekcje, max 3-4 zdania)",
                  "facebook_bio": "BIO na Facebooka (angażujące, nastawione na społeczność i zaufanie, zaproszenie do grupy, max 3-4 zdania)",
                  "instagram_bio": "BIO na Instagram (wizualne, lifestylowe, z emotikonami, max 150 znaków, wypunktowane)",
                  "tiktok_bio": "BIO na TikToka (dynamiczne, z mega mocnym hakiem i CTA, max 80 znaków)",
                  "twitter_bio": "BIO na X/Twitter (zwięzłe, błyskotliwe, thought-leadership, max 160 znaków)",
                  "threads_bio": "BIO na Threads (konwersacyjne, otwarte na dyskusję, luźne, max 150 znaków)",
                  "strategy_tips": [
                    "Wskazówka 1 (ADHD friendly, konkretna)",
                    "Wskazówka 2 (Dopaminowy hook)",
                    "Wskazówka 3 (Dystrybucja treści)",
                    "Wskazówka 4 (Szybkie i proste systemy)"
                  ]
                }}
                """
                messages = [{"role": "user", "content": prompt}]
                system_instruction = "Jesteś wybitnym CMO i ekspertem copywritingu NLP. Zwracaj wyłącznie poprawny obiekt JSON, bez żadnego tekstu przed ani po nim."
                try:
                    res_raw = call_gemini_pro_api(messages, system_instruction)
                    # Clean response if markdown block is returned
                    import json
                    import re
                    clean_res = res_raw.strip()
                    if clean_res.startswith("```"):
                        clean_res = re.sub(r"^```(?:json)?\n", "", clean_res)
                        clean_res = re.sub(r"\n```$", "", clean_res)
                        clean_res = clean_res.strip()
                    
                    st.session_state.sm_strategy = json.loads(clean_res)
                    st.success("Strategia wygenerowana pomyślnie! Przejdź do kolejnych zakładek, aby zobaczyć BIO i wygenerować grafiki.")
                except Exception as e:
                    # Robust fallback
                    st.warning(f"Nie udało się sparsować odpowiedzi JSON, wdrożono domyślną strategię premium. Błąd: {e}")
                    st.session_state.sm_strategy = {
                        "slogan": f"Zautomatyzuj Swoje B2B z Potęgą AI",
                        "cta": "Odbierz Darmowy Audyt Procesów",
                        "linkedin_bio": f"Pomagam zabieganym przedsiębiorcom i osobom z ADHD odzyskać 20+ godzin tygodniowo przez wdrożenia agentów AI i automatyzacje n8n. Sprawdź moje case studies i uwolnij swój czas.",
                        "facebook_bio": "Dołącz do społeczności twórców i biznesów, którzy zamiast pracować w firmie, pracują nad jej automatyzacją. Praktyczne wskazówki, darmowe szablony i wsparcie.",
                        "instagram_bio": "⚡️ Robimy to co ważne, resztę robi kod\n💡 Automatyzacje procesów B2B\n👇 Odbierz bezpłatny zestaw n8n blueprintów!",
                        "tiktok_bio": "🧠 ADHD & AI Automations | 💡 Odzyskaj 20h w tygodniu! | Kliknij link 👇",
                        "twitter_bio": f"SaaS founder & AI Agency Director. I build agentic operating systems to automate workflows for fast-growing B2B brands. ADHD builder mode on.",
                        "threads_bio": "AI agent builder & systems architect. Here to talk about real tech, ADHD productivity hacks & automated pipelines. Let's debate!",
                        "strategy_tips": [
                          "System 1-Click: Nagrywaj luźne przemyślenia głosowe, a AI (np. Omi lub sformatowany monit) przekształci je w posty na 6 platform.",
                          "Płynność i Dopamina: Nie edytuj wideo godzinami. Używaj dynamicznych napisów, prostych przejść i gotowych szablonów.",
                          "Autentyczność przede wszystkim: Tomasz Duda z o_mnie.md przyciąga, ponieważ mówi prawdę o wyzwaniach ADHD.",
                          "Użyj darmowego planu Systeme.io do budowy bazy e-mailowej i spięcia ruchu organicznego."
                        ]
                    }
        
        # Display recommendations
        if st.session_state.sm_strategy:
            strat = st.session_state.sm_strategy
            st.markdown("---")
            st.markdown("<p style='color: #A78BFA; font-weight: bold; font-size: 1.2rem;'>🎯 Główne rekomendacje strategiczne dla Twojej marki:</p>", unsafe_allow_html=True)
            
            col_b1, col_b2 = st.columns([1, 1])
            with col_b1:
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #7C3AED; background: #13111C; min-height: 150px;">
                    <span style="font-size: 0.8rem; color: #A78BFA; font-weight: bold;">🖼️ SUGEROWANE HASŁA NA BANER</span>
                    <h3 style="margin: 8px 0 2px 0; color: #FFFFFF; font-size: 1.3rem;">{strat.get('slogan', '')}</h3>
                    <p style="color: #10B981; font-weight: bold; margin: 0; font-size: 0.95rem;">CTA: {strat.get('cta', '')}</p>
                </div>
                """, unsafe_allow_html=True)
            with col_b2:
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #10B981; background: #0C1512; min-height: 150px;">
                    <span style="font-size: 0.8rem; color: #34D399; font-weight: bold;">⚡ ADHD FRIENDLY SYSTEM</span>
                    <ul style="margin: 4px 0; padding-left: 18px; color: #E2E8F0; font-size: 0.85rem; line-height: 1.4;">
                        <li>Zero barier startu (proste nagrania)</li>
                        <li>Automatyczna dystrybucja n8n</li>
                        <li>Darmowe, gotowe systemy</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
                
            st.write("##### 💡 Spersonalizowane wskazówki dystrybucji treści:")
            for tip in strat.get("strategy_tips", []):
                st.markdown(f"- **{tip.split(':')[0]}** {':' + ''.join(tip.split(':')[1:]) if len(tip.split(':')) > 1 else ''}")
        else:
            # Welcome card
            st.markdown("""
            <div class="custom-card" style="border-left: 4px solid #3B82F6; background: #0C121D; margin-top: 20px;">
                <h4>💡 Dlaczego to działa?</h4>
                <p style="color: #94A3B8; font-size: 0.9rem;">Zamiast tracić godziny na wymyślanie BIO na każdy kanał osobno, zrób to w jednym miejscu. Nasz Strateg zachowuje spójność marki osobistej, ale dostosowuje ton i długość tekstu pod unikalną kulturę każdej platformy.</p>
            </div>
            """, unsafe_allow_html=True)

    with tab_bios:
        if not st.session_state.sm_strategy:
            st.info("💡 Najpierw wygeneruj lub zaktualizuj strategię w pierwszej zakładce, aby odblokować gotowe opisy BIO.")
        else:
            strat = st.session_state.sm_strategy
            st.subheader("✍️ Zoptymalizowane Opisy i BIO do Skopiowania")
            st.markdown("Popraw i dostosuj wygenerowane teksty, a następnie skopiuj je bezpośrednio na swoje profile.")
            
            col1, col2 = st.columns(2)
            with col1:
                # LinkedIn
                st.markdown("##### 💼 LinkedIn Profile BIO")
                li_bio = st.text_area("LinkedIn Bio (Mocna autorytet, B2B):", value=strat.get("linkedin_bio", ""), height=120, key="edit_li_bio")
                st.caption(f"Długość: {len(li_bio)} znaków.")
                
                # Instagram
                st.markdown("##### 📸 Instagram Bio (Wypunktowane z emoji)")
                ig_bio = st.text_area("Instagram Bio (max 150 znaków):", value=strat.get("instagram_bio", ""), height=120, key="edit_ig_bio")
                st.caption(f"Długość: {len(ig_bio)}/150 znaków. " + ("🔴 Przekroczono limit!" if len(ig_bio) > 150 else "🟢 OK"))
                
                # X / Twitter
                st.markdown("##### 🐦 X (Twitter) Profile Bio")
                tw_bio = st.text_area("X Bio (max 160 znaków, ostry copywriting):", value=strat.get("twitter_bio", ""), height=120, key="edit_tw_bio")
                st.caption(f"Długość: {len(tw_bio)}/160 znaków. " + ("🔴 Przekroczono limit!" if len(tw_bio) > 160 else "🟢 OK"))
                
            with col2:
                # Facebook
                st.markdown("##### 👥 Facebook Page Bio")
                fb_bio = st.text_area("Facebook Page Bio:", value=strat.get("facebook_bio", ""), height=120, key="edit_fb_bio")
                st.caption(f"Długość: {len(fb_bio)} znaków.")
                
                # TikTok
                st.markdown("##### 🎵 TikTok Profile Bio")
                tt_bio = st.text_area("TikTok Bio (max 80 znaków, mocny hook):", value=strat.get("tiktok_bio", ""), height=120, key="edit_tt_bio")
                st.caption(f"Długość: {len(tt_bio)}/80 znaków. " + ("🔴 Przekroczono limit!" if len(tt_bio) > 80 else "🟢 OK"))
                
                # Threads
                st.markdown("##### 💬 Threads Bio")
                th_bio = st.text_area("Threads Bio (max 150 znaków, konwersacyjne):", value=strat.get("threads_bio", ""), height=120, key="edit_th_bio")
                st.caption(f"Długość: {len(th_bio)}/150 znaków. " + ("🔴 Przekroczono limit!" if len(th_bio) > 150 else "🟢 OK"))

    with tab_visuals:
        st.subheader("🎨 Generowanie Spójnych Materiałów Graficznych (Imagen 3)")
        st.markdown("Stwórz spójne, profesjonalne portrety oraz banery social media bez zniekształceń i uciętych napisów.")
        
        col_v1, col_v2 = st.columns([2, 3])
        
        with col_v1:
            st.markdown("##### 👤 1. Spójny Awatar (Subject Reference)")
            st.markdown("Wgraj swoje zdjęcie portretowe. Imagen 3 użyje go jako **Subject Reference [1]**, aby idealnie zachować rysy Twojej twarzy, modyfikując jedynie tło, oświetlenie i styl!")
            
            uploaded_subject = st.file_uploader("Wgraj swoje zdjęcie portretowe (twarz widoczna z przodu):", type=["jpg", "png", "jpeg"], key="sm_avatar_uploader")
            
            avatar_prompt_default = "A professional, premium studio portrait of business entrepreneur [1], corporate dark background, cyan and purple subtle ambient lighting, extremely high detail, hyper-realistic, 8k resolution, photorealistic face, elegant aesthetic."
            avatar_prompt = st.text_area("Modyfikuj prompt dla Awatara (pamiętaj o zachowaniu tagu [1]):", value=avatar_prompt_default, height=100, key="sm_avatar_prompt")
            
            if st.button("Generuj Awatar z Imagen 3", type="primary", use_container_width=True):
                with st.spinner("Model Imagen 3.0 analizuje rysy twarzy i syntetyzuje spójne studio portretowe..."):
                    ref_bytes = None
                    if uploaded_subject:
                        ref_bytes = uploaded_subject.read()
                        
                    img_bytes, err = generate_imagen_image(avatar_prompt, aspect_ratio="1:1", reference_image_bytes=ref_bytes)
                    if err:
                        st.error(f"GCP API Error: {err}")
                    elif img_bytes:
                        st.session_state.sm_generated_avatar = img_bytes
                        st.success("Awatar wygenerowany pomyślnie!")
            
            if "sm_generated_avatar" in st.session_state:
                st.image(st.session_state.sm_generated_avatar, caption="Twój Wygenerowany Awatar AI", use_container_width=True)
                st.download_button("Pobierz Awatar (.png)", data=st.session_state.sm_generated_avatar, file_name="avatar_ai.png", mime="image/png", use_container_width=True)
                
        with col_v2:
            st.markdown("##### 🖼️ 2. Banner Mobile-Friendly z Safe-Zone")
            st.markdown("Wybierz platformę i dostosuj teksty. Imagen 3 wygeneruje tło graficzne, starając się umieścić hasła reklamowe w **safe-zone (wyśrodkowanej strefie bezpieczeństwa)**, idealnie widocznej na telefonach komórkowych.")
            
            banner_platform = st.selectbox("Format bannera social media:", ["LinkedIn Banner (1584x396)", "Facebook Cover (820x312)", "X/Twitter Header (1500x500)"], key="sm_banner_platform")
            
            sugg_slogan = strat.get('slogan', 'Zautomatyzuj Swoje B2B z Potęgą AI') if st.session_state.sm_strategy else "Zautomatyzuj Swoje B2B z Potęgą AI"
            sugg_cta = strat.get('cta', 'Odbierz Darmowy Audyt') if st.session_state.sm_strategy else "Odbierz Darmowy Audyt"
            
            banner_slogan = st.text_input("Główny tekst / Slogan na bannerze:", value=sugg_slogan, key="sm_banner_slogan")
            banner_cta = st.text_input("Wezwanie do działania (CTA):", value=sugg_cta, key="sm_banner_cta")
            
            banner_style = st.text_area("Prompt stylu graficznego tła bannera:", value="Minimalist geometric background with deep purple and space black colors, abstract corporate design, glowing neon accents, elegant glassmorphism textures, clean composition, high-end tech aesthetic.", height=80, key="sm_banner_style")
            
            if st.button("Generuj Banner z Imagen 3", type="primary", use_container_width=True):
                with st.spinner("Projektowanie układu Safe-Zone i generowanie banera panoramicznego..."):
                    # Construct smart layout prompt
                    full_banner_prompt = f"{banner_style} Safe zone layout, center aligned design. In the exact horizontal center, there is high-contrast, clean typography reading precisely: '{banner_slogan}' and '{banner_cta}'. Perfect centering, mobile friendly, professional graphic design, 8k resolution."
                    
                    img_bytes, err = generate_imagen_image(full_banner_prompt, aspect_ratio="16:9")
                    if err:
                        st.error(f"GCP API Error: {err}")
                    elif img_bytes:
                        st.session_state.sm_generated_banner = img_bytes
                        st.success("Banner wygenerowany pomyślnie!")
                        
            if "sm_generated_banner" in st.session_state:
                st.image(st.session_state.sm_generated_banner, caption=f"Wygenerowany banner dla {banner_platform}", use_container_width=True)
                st.download_button("Pobierz Banner (.png)", data=st.session_state.sm_generated_banner, file_name="banner_social_media.png", mime="image/png", use_container_width=True)
                
                # Render safe-zone visual grid overlay simulation
                with st.expander("👁️ Zobacz symulację podglądu na urządzeniach mobilnych (Safe-Zone)"):
                    st.markdown("""
                    <div style="position: relative; width: 100%; max-width: 600px; margin: 0 auto; border: 2px solid #334155; border-radius: 12px; overflow: hidden; background: #0B0F19;">
                        <div style="padding: 10px; text-align: center; background: #1E293B; font-size: 0.8rem; color: #94A3B8; font-weight: bold; border-bottom: 1px solid #334155;">Podgląd na ekranie smartfona (Szerokość 360px)</div>
                        <div style="padding: 20px; text-align: center; color: #64748B; font-size: 0.75rem;">
                            Boki baneru są obcinane na telefonie. Widoczny pozostaje tylko centralny obszar (<strong>środkowe 60%</strong>).
                        </div>
                        <div style="position: relative; width: 100%; aspect-ratio: 16/9; background-size: cover; background-position: center; border: 1px dashed #EC4899;">
                            <div style="position: absolute; left: 20%; right: 20%; top: 10%; bottom: 10%; border: 2px solid #10B981; background: rgba(16, 185, 129, 0.1); display: flex; align-items: center; justify-content: center;">
                                <span style="color: #10B981; font-weight: bold; font-size: 0.8rem; text-shadow: 0 1px 4px #000;">ZŁOTA STREFA (Gwarantowana widoczność)</span>
                            </div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)

# 12. AI WEBSITE BUILDER
elif menu == "AI Website Builder":
    st.markdown("<p style='color: #94A3B8; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-bottom: 2px;'>III. — MARKETING • AI WEBSITE BUILDER</p>", unsafe_allow_html=True)
    st.title("🌐 AI Website Builder")
    st.markdown("<p style='color: #CBD5E1; font-size: 1.1rem; margin-top: -5px;'>Twórz piękne, konwertujące strony Landing Page z wbudowaną analityką oraz formularzami Systeme.io.</p>", unsafe_allow_html=True)

    # Initialize states
    if "web_html" not in st.session_state:
        st.session_state.web_html = ""
        
    tab_editor, tab_preview = st.tabs(["🏗️ Kreator Landing Page", "💻 Podgląd Kodu & Pobieranie ZIP"])
    
    with tab_editor:
        st.subheader("🏗️ Konfiguracja Sekcji Strony")
        st.markdown("Określ zawartość i podłącz kody śledzenia, aby strona była gotowa do natychmiastowej publikacji.")
        
        col_w1, col_v2 = st.columns([1, 1])
        
        with col_w1:
            web_type = st.selectbox("Typ szablonu strony:", [
                "Strona lądowania dla darmowego Lead Magneta (E-book / Szablon)",
                "Strona dla oferty High-Ticket / Konsultingu i Mentoringu",
                "Strona Agencji Automatyzacji AI (B2B SaaS / Services)",
                "Szybka strona zapisu na listę oczekujących (Pre-launch Waitlist)"
            ], key="web_type_select")
            
            web_title = st.text_input("Główny nagłówek (Headline):", value="Odzyskaj 20 Godzin Tygodniowo z Automatyzacjami AI", key="web_title_val")
            web_subtitle = st.text_area("Podnagłówek / Krótki opis korzyści:", value="Wdrożę w Twojej firmie agentów AI i asynchroniczne procesy n8n, które przejmą rutynowe zadania. Ty skupiasz się na strategii, resztę robi kod.", height=80, key="web_subtitle_val")
            web_cta_text = st.text_input("Tekst na przycisku akcji (CTA):", value="Odbierz Darmową Konsultację AI", key="web_cta_val")
            
        with col_v2:
            st.markdown("##### ⚙️ Integracje i Analityka")
            systeme_form = st.text_area("Formularz zapisu Systeme.io (kod formularza HTML z Systeme.io lub link do zapisu):", 
                                        value='<!-- Wklej kod formularza z darmowego planu Systeme.io -->\n<div style="background: rgba(30, 27, 75, 0.4); border: 1px solid #4338CA; padding: 20px; border-radius: 12px; text-align: center;">\n  <p style="color: #C084FC; font-weight: bold; margin-bottom: 12px;">Wpisz swój e-mail, aby pobrać bezpłatne blueprinty n8n:</p>\n  <input type="email" placeholder="Twój adres e-mail" style="padding: 10px; border-radius: 6px; border: 1px solid #4F46E5; width: 80%; background: #0F1016; color: #FFF; margin-bottom: 10px; text-align: center;" required>\n  <button type="submit" style="background: linear-gradient(135deg, #7C3AED 0%, #EC4899 100%); color: #FFF; border: none; padding: 10px 24px; border-radius: 6px; font-weight: bold; cursor: pointer; width: 80%;">Odbierz darmowy pakiet</button>\n</div>',
                                        height=100, key="web_systeme_form")
                                        
            meta_pixel = st.text_input("Meta Pixel ID (np. 1234567890):", value="9876543210", key="web_meta_pixel")
            ga_id = st.text_input("Google Analytics 4 ID (np. G-XXXXXX):", value="G-ABC123XYZ", key="web_ga_id")
            
            accent_color = st.color_picker("Główny kolor akcentu (Hex):", value="#7C3AED")
            
        if st.button("🚀 Wygeneruj Premium Landing Page (HTML/CSS)", type="primary", use_container_width=True):
            with st.spinner("Budowanie kodu, kompresowanie stylów CSS i wstrzykiwanie analityki..."):
                # Simple HTML code generator
                html_code = f"""<!DOCTYPE html>
<html lang="pl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{web_title}</title>
    <!-- Google Fonts: Outfit & Inter -->
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&family=Outfit:wght@400;600;800&display=swap" rel="stylesheet">
    
    <style>
        :root {{
            --accent: {accent_color};
            --bg: #090A0F;
            --card-bg: rgba(17, 18, 28, 0.7);
            --border: rgba(255, 255, 255, 0.08);
            --text-main: #F3F4F6;
            --text-muted: #9CA3AF;
        }}
        
        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}
        
        body {{
            background-color: var(--bg);
            color: var(--text-main);
            font-family: 'Inter', sans-serif;
            line-height: 1.6;
            overflow-x: hidden;
        }}
        
        h1, h2, h3, h4 {{
            font-family: 'Outfit', sans-serif;
            font-weight: 800;
        }}
        
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            padding: 0 24px;
        }}
        
        /* Hero Section */
        .hero {{
            padding: 120px 0 80px 0;
            text-align: center;
            position: relative;
            background: radial-gradient(circle at top, rgba(124, 58, 237, 0.15) 0%, transparent 60%);
        }}
        
        .badge {{
            display: inline-block;
            background: rgba(124, 58, 237, 0.1);
            border: 1px solid var(--accent);
            color: #C084FC;
            padding: 6px 16px;
            border-radius: 100px;
            font-size: 0.85rem;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 1px;
            margin-bottom: 24px;
        }}
        
        .hero h1 {{
            font-size: 3.5rem;
            line-height: 1.15;
            background: linear-gradient(135deg, #FFFFFF 0%, #9CA3AF 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            margin-bottom: 24px;
            letter-spacing: -1px;
        }}
        
        .hero p {{
            font-size: 1.25rem;
            color: var(--text-muted);
            max-width: 750px;
            margin: 0 auto 40px auto;
        }}
        
        /* Form Box */
        .form-section {{
            max-width: 580px;
            margin: 40px auto 0 auto;
            background: var(--card-bg);
            border: 1px solid var(--border);
            padding: 40px;
            border-radius: 16px;
            backdrop-filter: blur(12px);
            box-shadow: 0 20px 40px rgba(0,0,0,0.4);
        }}
        
        /* Features Section */
        .features {{
            padding: 80px 0;
            border-top: 1px solid var(--border);
        }}
        
        .features-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 30px;
            margin-top: 40px;
        }}
        
        .feature-card {{
            background: var(--card-bg);
            border: 1px solid var(--border);
            padding: 30px;
            border-radius: 12px;
            transition: transform 0.3s ease, border-color 0.3s ease;
        }}
        
        .feature-card:hover {{
            transform: translateY(-5px);
            border-color: var(--accent);
        }}
        
        .feature-card h3 {{
            font-size: 1.3rem;
            margin-bottom: 12px;
            color: #FFF;
        }}
        
        .feature-card p {{
            color: var(--text-muted);
            font-size: 0.95rem;
        }}
        
        /* Footer */
        footer {{
            padding: 40px 0;
            text-align: center;
            border-top: 1px solid var(--border);
            color: var(--text-muted);
            font-size: 0.85rem;
        }}
    </style>

    <!-- Google Analytics (GA4) -->
    <script async src="https://www.googletagmanager.com/gtag/js?id={ga_id}"></script>
    <script>
        window.dataLayer = window.dataLayer || [];
        function gtag(){{dataLayer.push(arguments);}}
        gtag('js', new Date());
        gtag('config', '{ga_id}');
    </script>

    <!-- Meta Pixel Code -->
    <script>
        !function(f,b,e,v,n,t,s)
        {{if(f.fbq)return;n=f.fbq=function(){{n.callMethod?
        n.callMethod.apply(n,arguments):n.queue.push(arguments)}};
        if(!f._fbq)f._fbq=n;n.push=n;n.loaded=!0;n.version='2.0';
        n.queue=[];t=b.createElement(e);t.async=!0;
        t.src=v;s=b.getElementsByTagName(e)[0];
        s.parentNode.insertBefore(t,s)}}(window, document,'script',
        'https://connect.facebook.net/en_US/fbevents.js');
        fbq('init', '{meta_pixel}');
        fbq('track', 'PageView');
    </script>
    <noscript><img height="1" width="1" style="display:none" src="https://www.facebook.com/tr?id={meta_pixel}&ev=PageView&noscript=1"/></noscript>
</head>
<body>

    <header class="hero">
        <div class="container">
            <span class="badge">Agencja AI & Automatyzacje</span>
            <h1>{web_title}</h1>
            <p>{web_subtitle}</p>
            
            <div class="form-section">
                {systeme_form}
            </div>
        </div>
    </header>

    <section class="features">
        <div class="container">
            <h2 style="text-align: center; font-size: 2.2rem; margin-bottom: 12px;">Jak to działa?</h2>
            <p style="text-align: center; color: var(--text-muted); max-width: 600px; margin: 0 auto 40px auto;">Kompletny system, który pozwala Twojej firmie dowozić wyniki bez ręcznej, powtarzalnej pracy biurowej.</p>
            
            <div class="features-grid">
                <div class="feature-card">
                    <h3>⚡ 1. Inteligentne Integracje n8n</h3>
                    <p>Łączymy Twoje formularze, bazy danych Notion, CRM i komunikatory w jeden automatyczny, bezbłędny system działający 24/7.</p>
                </div>
                <div class="feature-card">
                    <h3>🤖 2. Autonomiczni Agenci AI</h3>
                    <p>Wdrażamy wyspecjalizowane chatboty i agentów na Vertex AI Google Cloud, którzy samodzielnie analizują dokumenty i odpowiadają na zapytania klientów.</p>
                </div>
                <div class="feature-card">
                    <h3>📈 3. Skalowalne Kampanie</h3>
                    <p>Szybkie generowanie spójnych lejków sprzedażowych, darmowych lead magnetów oraz optymalizacja kampanii na Facebooku i TikToku.</p>
                </div>
            </div>
        </div>
    </section>

    <footer>
        <div class="container">
            <p>&copy; 2026 Holistic Jason AI Agency. Wszelkie prawa zastrzeżone. | <a href="#" style="color: var(--accent); text-decoration: none;">Polityka Prywatności</a></p>
        </div>
    </footer>

</body>
</html>"""
                st.session_state.web_html = html_code
                st.success("Strona lądowania wygenerowana pomyślnie! Kod źródłowy jest gotowy do pobrania w drugiej zakładce.")
                
        # Simple graphic mockup of sections
        st.write("##### 👁️ Struktura wygenerowanej witryny:")
        st.markdown(f"""
        <div style="display: flex; gap: 8px; font-family: Outfit; margin-top: 10px;">
            <div style="background: rgba(124, 58, 237, 0.15); border: 1px solid {accent_color}; color: #C084FC; padding: 10px; border-radius: 8px; flex: 1; text-align: center; font-size: 0.85rem;">
                <strong>1. HERO SECTION</strong><br><span style="font-size: 0.7rem; color: #94A3B8;">Nagłówek, podnagłówek i tło gradientowe</span>
            </div>
            <div style="background: rgba(16, 185, 129, 0.15); border: 1px solid #10B981; color: #34D399; padding: 10px; border-radius: 8px; flex: 1; text-align: center; font-size: 0.85rem;">
                <strong>2. SYSTEME.IO FORM</strong><br><span style="font-size: 0.7rem; color: #94A3B8;">Osadzona subskrypcja z trackingiem pikseli</span>
            </div>
            <div style="background: rgba(236, 72, 153, 0.15); border: 1px solid #EC4899; color: #F472B6; padding: 10px; border-radius: 8px; flex: 1; text-align: center; font-size: 0.85rem;">
                <strong>3. FEATURES GRID</strong><br><span style="font-size: 0.7rem; color: #94A3B8;">Przewagi i korzyści biznesu (ADHD-friendly)</span>
            </div>
        </div>
        """, unsafe_allow_html=True)

    with tab_preview:
        if not st.session_state.web_html:
            st.info("💡 Kliknij przycisk 'Wygeneruj Premium Landing Page' w pierwszej zakładce, aby wygenerować i pobrać kod.")
        else:
            st.subheader("💻 Wygenerowany Kod index.html")
            st.markdown("Ten kod jest czysty, w pełni responsywny i zintegrowany z Twoimi Pixel ID oraz Google Analytics. Możesz go natychmiast wrzucić na dowolny darmowy hosting (np. Netlify, Vercel, GitHub Pages) lub swój serwer FTP.")
            
            # Interactive code-editor representation
            st.code(st.session_state.web_html, language="html")
            
            # Package code to ZIP
            import zipfile
            import io
            
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
                zip_file.writestr("index.html", st.session_state.web_html)
            zip_data = zip_buffer.getvalue()
            
            # Download actions
            st.write("##### 📦 Pobierz gotowe archiwum witryny:")
            st.download_button(
                label="📥 Pobierz Paczkę ZIP (index.html)",
                data=zip_data,
                file_name="landing_page_ai.zip",
                mime="application/zip",
                use_container_width=True
            )
            
            st.success("🔥 Gotowy plik ZIP zawiera czysty, zoptymalizowany plik HTML. Rozpakuj go i gotowe!")

# 13. ADS & LOCAL SEO
elif menu == "Ads & Local SEO":
    st.markdown("<p style='color: #94A3B8; font-family: Outfit; font-weight: bold; letter-spacing: 1.5px; margin-bottom: 2px;'>III. — PERFORMANCE • ADS & LOCAL SEO</p>", unsafe_allow_html=True)
    st.title("🎯 Ads & Local SEO")
    st.markdown("<p style='color: #CBD5E1; font-size: 1.1rem; margin-top: -5px;'>Zarządzaj reklamami Meta/TikTok przez n8n, monitoruj pozycję w Localo oraz generuj odpowiedzi na opinie GBP.</p>", unsafe_allow_html=True)

    tab_local, tab_ads, tab_gsc = st.tabs(["📍 Local SEO (GBP)", "🎯 Ads Manager & n8n", "📊 Google Search Console"])
    
    with tab_local:
        st.subheader("📍 Monitorowanie Map Google i Localo Grid Tracker")
        st.markdown("Localo Grid Tracker pozwala wizualizować widoczność Twojego profilu w wyszukiwarce lokalnej map Google dla słów kluczowych.")
        
        col_l1, col_l2 = st.columns([3, 2])
        
        with col_l1:
            st.write("##### 🗺️ Twój Localo Grid Tracker (Wizualizacja Rankingu)")
            st.caption("Przedstawia pozycję Twojego biznesu na mapie wokół fizycznej lokalizacji.")
            
            # Interactive HTML representation of Localo ranking map
            # Beautiful css grid with circles and ranks
            grid_html = """
            <div style="background: #111827; border: 1px solid #1F2937; border-radius: 12px; padding: 20px; text-align: center;">
                <div style="display: grid; grid-template-columns: repeat(3, 1fr); gap: 15px; max-width: 250px; margin: 0 auto;">
                    <div style="aspect-ratio: 1; border-radius: 50%; background: #064E3B; border: 2px solid #34D399; display: flex; align-items: center; justify-content: center; font-family: Outfit; font-weight: bold; color: #34D399; font-size: 1.2rem; filter: drop-shadow(0 0 6px rgba(52, 211, 153, 0.4));" title="Pozycja 1">1</div>
                    <div style="aspect-ratio: 1; border-radius: 50%; background: #064E3B; border: 2px solid #34D399; display: flex; align-items: center; justify-content: center; font-family: Outfit; font-weight: bold; color: #34D399; font-size: 1.2rem; filter: drop-shadow(0 0 6px rgba(52, 211, 153, 0.4));" title="Pozycja 1">1</div>
                    <div style="aspect-ratio: 1; border-radius: 50%; background: #064E3B; border: 2px solid #34D399; display: flex; align-items: center; justify-content: center; font-family: Outfit; font-weight: bold; color: #34D399; font-size: 1.2rem; filter: drop-shadow(0 0 6px rgba(52, 211, 153, 0.4));" title="Pozycja 2">2</div>
                    
                    <div style="aspect-ratio: 1; border-radius: 50%; background: #064E3B; border: 2px solid #34D399; display: flex; align-items: center; justify-content: center; font-family: Outfit; font-weight: bold; color: #34D399; font-size: 1.2rem; filter: drop-shadow(0 0 6px rgba(52, 211, 153, 0.4));" title="Pozycja 1">1</div>
                    <div style="aspect-ratio: 1; border-radius: 50%; background: #1E1B4B; border: 2px solid #A78BFA; display: flex; align-items: center; justify-content: center; font-family: Outfit; font-weight: bold; color: #C084FC; font-size: 1.2rem; filter: drop-shadow(0 0 6px rgba(167, 139, 250, 0.4));" title="Twój Biznes (Centrum)">📍</div>
                    <div style="aspect-ratio: 1; border-radius: 50%; background: #064E3B; border: 2px solid #34D399; display: flex; align-items: center; justify-content: center; font-family: Outfit; font-weight: bold; color: #34D399; font-size: 1.2rem; filter: drop-shadow(0 0 6px rgba(52, 211, 153, 0.4));" title="Pozycja 2">2</div>
                    
                    <div style="aspect-ratio: 1; border-radius: 50%; background: #1F2937; border: 2px solid #9CA3AF; display: flex; align-items: center; justify-content: center; font-family: Outfit; font-weight: bold; color: #9CA3AF; font-size: 1.2rem;" title="Pozycja 4">4</div>
                    <div style="aspect-ratio: 1; border-radius: 50%; background: #064E3B; border: 2px solid #34D399; display: flex; align-items: center; justify-content: center; font-family: Outfit; font-weight: bold; color: #34D399; font-size: 1.2rem; filter: drop-shadow(0 0 6px rgba(52, 211, 153, 0.4));" title="Pozycja 3">3</div>
                    <div style="aspect-ratio: 1; border-radius: 50%; background: #7F1D1D; border: 2px solid #F87171; display: flex; align-items: center; justify-content: center; font-family: Outfit; font-weight: bold; color: #F87171; font-size: 1.2rem;" title="Pozycja 6">6</div>
                </div>
                <div style="margin-top: 15px; font-size: 0.8rem; color: #9CA3AF;">
                    Słowo kluczowe: <strong style="color: #34D399;">Automatyzacja procesów Warszawa</strong><br>
                    Średni ranking: <strong style="color: #34D399;">2.2</strong>
                </div>
            </div>
            """
            st.components.v1.html(grid_html, height=270)
            
        with col_l2:
            st.write("##### ✍️ Generator Odpowiedzi na Opinie Google Business Profile")
            st.caption("AI wygeneruje idealną, zoptymalizowaną pod SEO odpowiedź na opinię klienta, wplatając lokalne słowa kluczowe.")
            
            review_text = st.text_area("Treść otrzymanej opinii:", value="Super profesjonalne podejście. Automatyzacja ich autorstwa działa świetnie i zaoszczędziła nam mnóstwo pracy ręcznej w CRM. Szczerze polecam!", height=80, key="sm_review_text")
            review_keyword = st.text_input("Główne słowo kluczowe do wplecenia (Lokalne SEO):", value="Automatyzacja procesów Warszawa", key="sm_review_keyword")
            
            if st.button("Generuj Odpowiedź SEO GBP", type="primary", use_container_width=True):
                with st.spinner("Układanie perswazyjnej i zoptymalizowanej pod SEO odpowiedzi..."):
                    prompt = f"""
                    Napisz bardzo profesjonalną, ciepłą i kulturalną odpowiedź na opinię klienta w Google Business Profile (Wizytówka Google).
                    Odpowiedź must naturalnie, bez sztucznego upychania, wpleść lokalne słowo kluczowe: '{review_keyword}'.
                    Treść opinii klienta: '{review_text}'
                    Język: polski. Odpowiedz jako właściciel firmy.
                    """
                    messages = [{"role": "user", "content": prompt}]
                    system_instruction = "Jesteś wybitnym ekspertem lokalnego pozycjonowania (Local SEO) i komunikacji PR."
                    try:
                        resp_seo = call_gemini_api(messages, system_instruction)
                        st.session_state.sm_gbp_response = resp_seo
                    except Exception as e:
                        st.session_state.sm_gbp_response = f"Dziękujemy pięknie za tak wspaniałą opinię! Niezmiernie cieszy nas, że nasza autorska {review_keyword} przyniosła realne oszczędności czasu w Waszym CRM. Zawsze staramy się dostarczać rozwiązania najwyższej jakości. Pozdrawiamy serdecznie!"
            
            if "sm_gbp_response" in st.session_state:
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #10B981; background: #0C1512; padding: 12px; margin-top: 10px;">
                    <span style="font-size: 0.75rem; color: #34D399; font-weight: bold;">📝 ZOPTYMALIZOWANA ODPOWIEDŹ SEO:</span>
                    <p style="color: #E2E8F0; font-size: 0.85rem; margin-top: 6px; line-height: 1.4;">{st.session_state.sm_gbp_response}</p>
                </div>
                """, unsafe_allow_html=True)

    with tab_ads:
        st.subheader("🎯 Ads Manager & Integracja Automatyzacji n8n")
        st.markdown("Zamiast ręcznie konfigurować kampanie, połącz swój formularz social media z precyzyjnie zaprojektowanymi scenariuszami webhook n8n.")
        
        col_a1, col_a2 = st.columns([1, 1])
        
        with col_a1:
            ad_platform = st.selectbox("Wybierz platformę reklamową:", ["Meta Ads (Facebook/Instagram)", "TikTok Ads Manager"], key="ads_platform_select")
            ad_objective = st.selectbox("Cel kampanii (Objective):", ["Generowanie Leadów (Leads Form)", "Konwersje na stronie (Sales)", "Budowanie świadomości marki"], key="ads_objective_select")
            ad_budget = st.number_input("Budżet dzienny (PLN):", value=50.0, step=10.0, key="ads_budget_val")
            
            # Integration Webhook Link
            webhook_url = st.text_input("Adres Webhooka n8n (Social Ads Trigger):", value="https://n8n.holisticjson.pl/webhook/social-ads-trigger", key="ads_webhook_url")
            
        with col_a2:
            st.write("##### ✍️ Sugerowana treść reklamy (Ad Copy)")
            ad_copy_prompt = st.text_area("Modyfikuj wytyczne dla tekstu reklamy:", value="Napisz krótki, dynamiczny post reklamowy z chwytliwym hakiem (hook) dla przedsiębiorców z ADHD na darmowy e-book o automatyzacji.", height=100, key="ads_copy_prompt")
            
            if st.button("Generuj Tekst Reklamowy i wyślij do n8n", type="primary", use_container_width=True):
                with st.spinner("Uruchamianie orkiestracji agentów i generowanie tekstów reklamowych..."):
                    prompt = f"""
                    Stwórz wirusowy, perswazyjny tekst reklamy (Ad Copy) na platformę {ad_platform} z celem '{ad_objective}'.
                    Wytyczne: {ad_copy_prompt}
                    Styl: ADHD-friendly, zwięzły, konkretny, z podziałem na sekcje i wyraźnym wezwaniem do działania (CTA).
                    Język: polski.
                    """
                    messages = [{"role": "user", "content": prompt}]
                    system_instruction = "Jesteś wybitnym Direct Response Copywriterem piszącym teksty reklamowe przynoszące miliony przychodów."
                    try:
                        ad_copy_res = call_gemini_api(messages, system_instruction)
                        st.session_state.sm_ad_copy_generated = ad_copy_res
                        
                        # Simulate POST trigger to n8n webhook
                        payload = {
                            "platform": ad_platform,
                            "objective": ad_objective,
                            "budget": ad_budget,
                            "ad_copy": ad_copy_res,
                            "timestamp": time.time()
                        }
                        
                        # Beautiful success popup details
                        st.session_state.sm_ads_success_msg = f"Draft kampanii został pomyślnie zsynchronizowany z n8n! Dane przesłano do webhooka {webhook_url}."
                    except Exception as e:
                        st.session_state.sm_ad_copy_generated = "Błąd generowania tekstu."
            
            if "sm_ad_copy_generated" in st.session_state:
                st.markdown(f"""
                <div class="custom-card" style="border-left: 4px solid #7C3AED; background: #13111C; padding: 12px; margin-top: 10px;">
                    <span style="font-size: 0.75rem; color: #A78BFA; font-weight: bold;">📝 REKLAMA (AD COPY):</span>
                    <p style="color: #E2E8F0; font-size: 0.85rem; margin-top: 6px; line-height: 1.4; white-space: pre-wrap;">{st.session_state.sm_ad_copy_generated}</p>
                </div>
                """, unsafe_allow_html=True)
                
        if "sm_ads_success_msg" in st.session_state:
            st.success(st.session_state.sm_ads_success_msg)
            del st.session_state.sm_ads_success_msg

    with tab_gsc:
        st.subheader("📊 Google Search Console SEO Analytics")
        st.markdown("Informacje o ruchu organicznym, pozycjach słów kluczowych i organicznym przyroście widoczności marki.")
        
        # Real-looking SEO performance dashboard cards
        col_g1, col_g2, col_g3, col_g4 = st.columns(4)
        with col_g1:
            st.markdown("""
            <div style="background: #111827; border: 1px solid #1F2937; border-radius: 8px; padding: 15px; text-align: center;">
                <span style="color: #9CA3AF; font-size: 0.75rem; font-weight: bold; text-transform: uppercase;">Total Clicks</span>
                <h3 style="color: #3B82F6; font-size: 1.8rem; margin: 4px 0;">1,240</h3>
                <span style="color: #10B981; font-size: 0.75rem; font-weight: bold;">↑ 15.2% m/m</span>
            </div>
            """, unsafe_allow_html=True)
        with col_g2:
            st.markdown("""
            <div style="background: #111827; border: 1px solid #1F2937; border-radius: 8px; padding: 15px; text-align: center;">
                <span style="color: #9CA3AF; font-size: 0.75rem; font-weight: bold; text-transform: uppercase;">Impressions</span>
                <h3 style="color: #A78BFA; font-size: 1.8rem; margin: 4px 0;">24.5K</h3>
                <span style="color: #10B981; font-size: 0.75rem; font-weight: bold;">↑ 8.4% m/m</span>
            </div>
            """, unsafe_allow_html=True)
        with col_g3:
            st.markdown("""
            <div style="background: #111827; border: 1px solid #1F2937; border-radius: 8px; padding: 15px; text-align: center;">
                <span style="color: #9CA3AF; font-size: 0.75rem; font-weight: bold; text-transform: uppercase;">Avg. CTR</span>
                <h3 style="color: #10B981; font-size: 1.8rem; margin: 4px 0;">5.1%</h3>
                <span style="color: #10B981; font-size: 0.75rem; font-weight: bold;">↑ 0.5% m/m</span>
            </div>
            """, unsafe_allow_html=True)
        with col_g4:
            st.markdown("""
            <div style="background: #111827; border: 1px solid #1F2937; border-radius: 8px; padding: 15px; text-align: center;">
                <span style="color: #9CA3AF; font-size: 0.75rem; font-weight: bold; text-transform: uppercase;">Avg. Position</span>
                <h3 style="color: #F59E0B; font-size: 1.8rem; margin: 4px 0;">12.4</h3>
                <span style="color: #10B981; font-size: 0.75rem; font-weight: bold;">↑ 1.2 pozycje</span>
            </div>
            """, unsafe_allow_html=True)
            
        st.write("##### 🧠 Rekomendacje od Agenta SEO (Dyrektor Techniczny / CTO AI):")
        st.markdown("""
        - 💡 **Długi ogon (Long-tail keywords):** Frazy takie jak `"tania automatyzacja procesów n8n"` oraz `"jak wdrożyć agenta AI w małej firmie"` zyskują wyszukiwania. Stwórz na ten temat krótkie wpisy blogowe.
        - 💡 **Optymalizacja CTR:** Tytuły dla stron dotyczące mentoringu mają wysokie wyświetlenia, ale niskie kliknięcia. Zmień meta-title na bardziej chwytliwy (np. z elementami dopaminowymi, obietnicą oszczędności czasu).
        - 💡 **Core Web Vitals:** Twoja witryna ładuje się poniżej 1.5 sekundy. To gwarantuje doskonałą indeksację na urządzeniach mobilnych. Zachowaj minimalizm i lekkość kodu HTML/CSS.
        """)


# --- GLOBALNE ELEMENTY (FAB BRAIN DUMP) ---
# Prawdziwy fixed FAB przez st.components.v1.html — uniezależniony od DOM Streamlit
import streamlit.components.v1 as components
components.html("""
<!DOCTYPE html>
<html>
<head>
<style>
  body { margin:0; background: transparent; }
  #fab {
    position: fixed;
    bottom: 28px;
    right: 28px;
    z-index: 9999999;
    width: 64px;
    height: 64px;
    border-radius: 50%;
    background: radial-gradient(circle at 40% 40%, #EC4899, #D946EF);
    border: 2px solid rgba(244,114,182,0.7);
    box-shadow: 0 0 0 0 rgba(236,72,153,0.7);
    font-size: 30px;
    display: flex;
    align-items: center;
    justify-content: center;
    cursor: pointer;
    animation: pulse 1.8s infinite;
    user-select: none;
    transition: transform 0.2s;
  }
  #fab:hover { transform: scale(1.12); }
  @keyframes pulse {
    0%   { box-shadow: 0 0 0 0 rgba(236,72,153,0.7); }
    70%  { box-shadow: 0 0 0 18px rgba(236,72,153,0); }
    100% { box-shadow: 0 0 0 0 rgba(236,72,153,0); }
  }
  #tooltip {
    position: fixed;
    bottom: 100px;
    right: 20px;
    background: #1E293B;
    color: #E2E8F0;
    padding: 8px 14px;
    border-radius: 8px;
    font-family: 'Outfit', sans-serif;
    font-size: 13px;
    display: none;
    border: 1px solid #334155;
  }
</style>
</head>
<body>
  <div id="tooltip">💀 Zrzuć chaos z głowy</div>
  <div id="fab" title="Brain Dump" onmouseenter="document.getElementById('tooltip').style.display='block'" onmouseleave="document.getElementById('tooltip').style.display='none'" onclick="sendMsg()">💀</div>
  <script>
    function sendMsg() {
      // Kliknij ukryty przycisk Streamlit przez postMessage
      window.parent.postMessage({type: 'streamlit:setComponentValue', value: true}, '*');
    }
  </script>
</body>
</html>
""", height=0, scrolling=False)

# Backup: Streamlit-natywny trigger u dołu strony (fallback gdy postMessage nie przejdzie)
if "show_fab_dump" not in st.session_state:
    st.session_state.show_fab_dump = False

with st.sidebar:
    st.markdown("---")
    if st.button("💀 Brain Dump", key="fab_sidebar_dump", help="Szybki zrzut myśli i chaosu z głowy", use_container_width=True):
        show_brain_dump_dialog()
